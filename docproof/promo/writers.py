"""Turning the model's answer into the two documents the press ships.

Unlike the review and prep writers, which edit a manuscript that already exists,
these write from nothing — so they use python-docx (already a dependency, for the
change log) rather than hand-rolling OOXML. Both files are regenerated from
`promo.json` whenever a human edits the copy in the panel, so they are outputs,
never a source of truth.
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document

from .model import SocialPost


def write_teaser(path: str | Path, teaser: str, *, title: str) -> Path:
    doc = Document()
    doc.add_heading(f"{title} — Teaser", level=1)
    # Keep the author's paragraphing: a teaser is occasionally two short paras.
    for block in teaser.split("\n\n"):
        block = block.strip()
        if block:
            doc.add_paragraph(block)
    doc.save(str(path))
    return Path(path)


def write_posts(path: str | Path, posts: list[SocialPost], *,
                title: str) -> Path:
    doc = Document()
    doc.add_heading(f"{title} — Social Posts", level=1)
    for i, post in enumerate(posts, start=1):
        head = f"Post {i}"
        if post.platform:
            head += f" · {post.platform}"
        doc.add_heading(head, level=2)
        doc.add_paragraph(post.text)
        if post.hashtags:
            tags = " ".join(h if h.startswith("#") else f"#{h}"
                            for h in post.hashtags)
            doc.add_paragraph(tags)
    doc.save(str(path))
    return Path(path)


_BOLD = re.compile(r"\*\*(.+?)\*\*")


def _emit_inline(paragraph, text: str) -> None:
    """Add `text` to a paragraph, turning `**bold**` spans into bold runs.

    The plan prompt uses bold only for the label at the head of a bullet, so
    this is all the inline Markdown the writer needs — anything else is written
    as plain text, never as stray asterisks."""
    pos = 0
    for match in _BOLD.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        paragraph.add_run(match.group(1)).bold = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def write_plan(path: str | Path, plan: str, *, title: str) -> Path:
    """Render the marketing plan's Markdown to the delivered .docx.

    The model is asked for a fixed, small subset — `#` title, `##` section
    headings, `-` bullets, `**bold**` labels — so this parses exactly that and
    treats every other line as a paragraph. It never fails on unexpected input:
    a plan that arrives as plain prose (an older model, a hand edit that
    stripped the marks) still writes, as headingless paragraphs, rather than
    losing the copy. `title` is the fallback document heading when the plan text
    carries no `#` line of its own."""
    doc = Document()
    seen_title = False
    for raw in plan.splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            seen_title = True
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.lstrip().startswith(("- ", "* ")):
            body = line.lstrip()[2:].strip()
            _emit_inline(doc.add_paragraph(style="List Bullet"), body)
        else:
            _emit_inline(doc.add_paragraph(), line.strip())
    if not seen_title:
        # No `#` line came back: give the document a heading so it is not a wall
        # of unlabelled paragraphs. Inserted first, before everything parsed.
        head = doc.add_heading(f"{title} — Marketing Plan", level=1)
        if doc.paragraphs and doc.paragraphs[0]._p is not head._p:
            doc.paragraphs[0]._p.addprevious(head._p)
    doc.save(str(path))
    return Path(path)
