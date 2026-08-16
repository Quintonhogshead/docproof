"""The change log the press hands to an author alongside the manuscript.

`summary.md` is for whoever ran the review. This is for the author and the
proofreading team: what was changed, what was only asked about, what was
deliberately left alone, and — the part the house brief cares most about —
what was *not* covered.

The brief is blunt about that last one: "Never overstate the depth of the
pass." So the honesty sections here are not padding. A change log that lists
four hundred corrections and says nothing about what it did not look at
invites the reader to assume it looked at everything.

Every number in this document comes from the run that produced it. Nothing is
asserted that was not measured.
"""
from __future__ import annotations

import logging
from datetime import datetime, timezone
from pathlib import Path

from .models import Finding
from .queries import shows_margin_comment

log = logging.getLogger("docproof.changelog")

# Rows past this and Word starts to struggle, and nobody reads them anyway.
# The table says so rather than truncating silently.
_MAX_TABLE_ROWS = 600


def _style_basis(cfg, variant) -> list[str]:
    lines = [f"This manuscript was proofread as {variant.name}."]
    lines += list(variant.authorities)
    ellipsis_clause = {
        "closed": "the ellipsis character closed up to the word before it",
        "nbsp": "the ellipsis character with a non-breaking space before it",
        "space": "the ellipsis character with a space on each side",
    }[cfg.style.ellipsis]
    lines.append(
        "The Atmosphere Press House Style Guide applies throughout, and "
        "overrides the above where they differ: serial comma always, "
        "spelled-out numbers to one hundred, spelled-out centuries, unspaced "
        f"em dashes, and {ellipsis_clause}.")
    if variant.confirm:
        # The brief asks for this by name. Canadian and Australian English are
        # hybrids — Canadian takes U.S. punctuation with Canadian spelling —
        # so the choice is a judgment the press should ratify, not an
        # assumption buried in a config file.
        lines.append(
            f"{variant.name} is a hybrid of conventions, so please confirm "
            f"this was the right choice before treating the pass as final. "
            f"If the manuscript is meant to be a different variant, the rules "
            f"that differ — quotation-mark nesting, date and time formats, "
            f"decade apostrophes, percent versus per cent, and the that/which "
            f"distinction — were applied in the wrong form throughout.")
    return lines


def _covered_locations(doc) -> set[str]:
    return {p.location for p in doc.paragraphs}


def _notes_statement(doc) -> str:
    """Whether the notes were read. The brief singles this out because the
    usual failure is silent: most tools read only the body, and an academic
    book can carry thousands of words of unproofed citations."""
    locations = _covered_locations(doc)
    notes = {"footnote", "endnote"} & locations
    if not notes:
        return ("This manuscript has no footnotes or endnotes. Nothing was "
                "skipped on that account.")
    counted = sum(1 for p in doc.paragraphs if p.location in notes)
    return (f"This manuscript has {', '.join(sorted(notes))}s, and they were "
            f"read: {counted} note paragraph(s) went through the same sweeps, "
            f"the same passes and the same reject-all audit as the body. Word "
            f"keeps them in separate parts of the file, which is where tools "
            f"that read only the main document lose them.")


def _limits(cfg, doc, findings, spell) -> list[str]:
    heading_clause = (
        "Heading-styled paragraphs got the scripted checks and the "
        "title-case pass, but no model read"
        if getattr(cfg.style, "heading_title_case", False) else
        "Heading-styled paragraphs got the scripted checks but no model read")
    lines = [
        "This was a mechanical and grammatical pass. It does not read for "
        "plot, character, pacing, structure, or fact — nothing here should be "
        "taken as a comment on the writing.",
        "Every paragraph of running text was read, and the scripted checks "
        f"above ran over all of them exhaustively. {heading_clause}; "
        "title pages and table-of-contents lines were not touched at all, "
        f"though they were included in the audit ({len(doc.skipped)} "
        "skipped).",
        "Text that is not paragraph text was not reviewed: words inside "
        "images, charts, embedded objects, and text boxes.",
        "Corrections are proposed, not imposed. Everything in the table above "
        "is a tracked change to accept or reject, and the queries are "
        "questions with no answer attached.",
    ]
    if spell is not None and not spell.available:
        which = f" ({spell.dictionary} is not installed)" if spell.dictionary else ""
        lines.append(
            f"The dictionary scan did not run for this pass{which}, so the "
            f"model was not given this manuscript's own vocabulary as a "
            f"do-not-flag list. Invented names and place names may be flagged "
            f"as misspellings more often than they otherwise would be.")
    if cfg.min_confidence != "low":
        lines.append(
            f"Only findings of {cfg.min_confidence} confidence or better were "
            f"applied. The rest were raised as questions rather than dropped, "
            f"so nothing was discarded silently — but a real error the model "
            f"was unsure about will be a comment here, not a correction.")
    return lines


def _reason(f: Finding) -> str:
    kind = f.error_type.replace("_", " ")
    if f.error_type.startswith("sweep_"):
        kind = kind.replace("sweep ", "") + " (scripted check)"
    return f"{kind}. {f.explanation}".strip()


def write_change_log(path: Path, *, doc, findings: list[Finding], cfg,
                     applied_ids, sweeps=None, spell=None, normalization=None,
                     audit=None, usage=None, variant=None, fmt=None) -> None:
    """Write the change log. Imports python-docx lazily so that a run which
    does not want one never pays for the import."""
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    from .reporting import scripted_check_rows

    applied_set = set(applied_ids)
    # Document order, not the order the passes happened to produce them: a
    # reader works through this table alongside the manuscript, and grouping
    # by which sweep found what would make them jump around the book.
    order = {p.para_id: i for i, p in enumerate(doc.paragraphs)}
    applied = sorted(
        (f for f in findings if f.finding_id in applied_set),
        key=lambda f: (order.get(f.para_id, len(order)),
                       f.anchor.start if f.anchor else 0))
    queries = [f for f in findings if f.status == "query"]
    low = [f for f in findings if f.status == "skipped_low_confidence"]
    rejected = [f for f in findings if f.status.startswith("rejected")]
    oversized = [f for f in findings if f.status == "rejected_oversized"]
    # What the manuscript actually carries, decided by the rule the reassembler
    # applies rather than by a second count that can drift from it — and named
    # in the format's own word, because an InDesign file holds notes and this
    # log is read beside the file it describes.
    in_margin = [f for f in queries + low + oversized
                 if f.anchor and shows_margin_comment(
                     f, query_comments=cfg.query_comments,
                     not_applied_comments=cfg.not_applied_comments)]
    noun = fmt.comment_noun if fmt is not None else "margin comment"

    d = docx.Document()
    name = Path(doc.source_path).stem
    d.add_heading(f"{name} — Atmosphere Press Proofreader Change Log", level=0)

    stamp = datetime.now(timezone.utc).strftime("%d %B %Y, %H:%M UTC")
    sub = d.add_paragraph(f"Atmosphere Press · prepared {stamp}")
    sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in sub.runs:
        run.italic = True
        run.font.size = Pt(10)

    d.add_paragraph(
        f"{len(applied)} correction(s) were made as tracked changes, and "
        f"{len(in_margin)} question(s) were raised as {noun}s that change "
        f"nothing. Revisions and {noun}s are authored "
        f"by “{cfg.revision_author}”.")

    # --- style basis ---------------------------------------------------------
    d.add_heading("Style basis", level=1)
    for line in _style_basis(cfg, variant):
        d.add_paragraph(line, style="List Bullet")

    # --- corrections ---------------------------------------------------------
    d.add_heading("Corrections", level=1)
    if not applied:
        d.add_paragraph("No corrections were applied.")
    else:
        shown = applied[:_MAX_TABLE_ROWS]
        d.add_paragraph(
            f"Every tracked change in the manuscript, in document order."
            + (f" The first {len(shown)} of {len(applied)} are listed here; "
               f"the rest are in the accompanying findings.json."
               if len(applied) > len(shown) else ""))
        table = d.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        for cell, text in zip(table.rows[0].cells,
                              ("Where", "Original", "Corrected", "Why")):
            cell.text = text
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True
        paras = {p.para_id: p for p in doc.paragraphs}
        for f in shown:
            row = table.add_row().cells
            loc = paras[f.para_id].location if f.para_id in paras else "body"
            row[0].text = f"{f.para_id} ({loc})"
            row[1].text = f.original_text
            row[2].text = f.corrected_text
            row[3].text = _reason(f)

    # --- scripted checks -----------------------------------------------------
    rows = scripted_check_rows(sweeps, findings, applied_ids)
    if rows:
        d.add_heading("Scripted checks", level=1)
        d.add_paragraph(
            "Each rule below was run as a pattern sweep over every paragraph, "
            "not looked for by eye, so these counts are exact. “Remaining” is "
            "what the sweep's own patterns still match after its fixes: zero "
            "means nothing of that kind is left in the manuscript.")
        for r in rows:
            line = (f"{r['name']}: {r['flagged']} flagged, {r['applied']} "
                    f"applied, {r['remaining']} remaining")
            if r["applied"] < r["flagged"]:
                line += (f" ({r['flagged'] - r['applied']} overlapped an "
                         f"earlier change and were left to it)")
            d.add_paragraph(line, style="List Bullet")

    # --- silent normalizations ----------------------------------------------
    if normalization is not None and normalization.ran and normalization.total:
        d.add_heading("Applied without tracked changes", level=1)
        d.add_paragraph(
            f"House style allows two edits outside the tracked-changes system, "
            f"because neither changes what the book says: "
            f"{normalization.quotes} straight quote(s) were curled and "
            f"{normalization.spaces} run(s) of two or more spaces were "
            f"collapsed to one, across {normalization.paragraphs} paragraph(s). "
            f"These are not in the table above and cannot be rejected in Word.")
        if normalization.ambiguous:
            d.add_paragraph(
                f"{normalization.ambiguous} straight mark(s) were left as they "
                f"were, because the surrounding text does not settle which way "
                f"they should curl — usually a single quote that could be a "
                f"nested quotation or a dialect elision. Curling one the wrong "
                f"way would not be visible as a tracked change, so these were "
                f"left for a person to set.")

    # --- queries -------------------------------------------------------------
    d.add_heading("Queries", level=1)
    if not (queries or low):
        d.add_paragraph("No questions were raised.")
    else:
        carried = (f"{len(in_margin)} of them a {noun} in the manuscript"
                   if len(in_margin) != len(queries) + len(low)
                   else f"each a {noun} in the manuscript")
        d.add_paragraph(
            f"{len(queries) + len(low)} question(s), {carried}. Nothing here "
            f"was changed. Where a query has a suggestion attached, it is a "
            f"suggestion and not a correction.")
        for f in sorted(queries + low,
                        key=lambda f: (order.get(f.para_id, len(order)),
                                       f.anchor.start if f.anchor else 0)):
            p = d.add_paragraph(style="List Bullet")
            p.add_run(f"{f.para_id}: ").bold = True
            # Italic rather than wrapped in quotes: the quoted text usually
            # carries its own, and doubling them reads as a mistake.
            p.add_run(f.original_text).italic = True
            p.add_run(f" — {f.explanation}")

    # --- deliberately left unchanged ----------------------------------------
    d.add_heading("Deliberately left unchanged", level=1)
    d.add_paragraph(
        "Author voice wins. Dialect in dialogue, intentional fragments, "
        "stylized punctuation and genre conventions are preserved rather than "
        "corrected, and anything that might be one of those is raised as a "
        "question instead of an edit.")
    if low:
        d.add_paragraph(
            f"{len(low)} finding(s) were anchored but not applied, because "
            f"the model read the passage as possibly deliberate. They are "
            f"listed under Queries above.")
    if rejected:
        d.add_paragraph(
            f"{len(rejected)} finding(s) were discarded by the validator "
            f"before reaching the manuscript — a quote that did not match the "
            f"text, an edit overlapping one already made, a duplicate, or an "
            f"edit too large to be a proofreading fix. None of them changed "
            f"anything.")
    if oversized:
        d.add_paragraph(
            f"Of those, {len(oversized)} were refused specifically for "
            f"overstepping: an edit that re-typed a long passage wholesale or "
            f"added new text the original did not contain. A proofreading pass "
            f"corrects; it does not rewrite or invent, so these were held back "
            f"rather than applied.")
    # --- the words taken on trust -------------------------------------------
    # The full protected list, never a truncated aside. This is the one part
    # of the pass that shields text from checking, so it is rendered as a
    # checklist a person can actually review — alphabetical, so near-identical
    # spellings sit next to each other, with each word's occurrence count.
    if spell is not None and spell.available and (
            spell.lexicon or spell.near_duplicates or spell.name_pairs):
        d.add_heading("Words taken on trust — please verify", level=1)
        if spell.lexicon:
            counts = dict(zip(spell.lexicon, spell.lexicon_counts))
            listed = ", ".join(
                f"{w} ({counts[w]}×)" if counts.get(w) else w
                for w in sorted(spell.lexicon, key=str.lower))
            d.add_paragraph(
                f"{len(spell.lexicon)} word(s) not in the dictionary are "
                f"written as names — coined terms, invented places, "
                f"characters — and were protected from spell-checking on that "
                f"evidence. Nothing flagged them; nothing ever would have. If "
                f"any word below is NOT spelled as the author intends, it is "
                f"wrong everywhere it appears and this pass could not have "
                f"caught it — that is what this list is for:")
            d.add_paragraph(listed + ".")
        if spell.near_duplicates:
            d.add_paragraph(
                f"{len(spell.near_duplicates)} near-duplicate name(s) were "
                f"NOT taken on trust, because a protected name one edit away "
                f"decisively owns the book. Each occurrence was ruled on in "
                f"context instead: " + "; ".join(
                    f"{nd.word} ({nd.count}×), one edit from {nd.of} "
                    f"({nd.of_count}×)" for nd in spell.near_duplicates)
                + ".")
        if spell.name_pairs:
            d.add_paragraph(
                f"{len(spell.name_pairs)} near-identical name pair(s) were "
                f"too close in usage for this pass to call, so both spellings "
                f"stand and a query marks each pair: " + "; ".join(
                    f"{p.a} ({p.a_count}×) beside {p.b} ({p.b_count}×)"
                    for p in spell.name_pairs) + ".")
    if spell is not None and spell.available and spell.recurring:
        d.add_paragraph(
            f"{len(spell.recurring)} word(s) not in the dictionary were used "
            f"more than once but are not written as names. These were shown to "
            f"the model as the book's likely vocabulary — evidence, not "
            f"protection — and read in context rather than corrected, because a "
            f"misspelling an author repeats reads the same as a coinage: "
            f"{', '.join(sorted((c.word for c in spell.recurring),
                                key=str.lower))}.")

    # --- coverage and honesty ------------------------------------------------
    d.add_heading("Footnotes and endnotes", level=1)
    d.add_paragraph(_notes_statement(doc))

    d.add_heading("Final audit", level=1)
    if audit is not None and audit.ran and audit.passed:
        d.add_paragraph(
            f"Passed. Rejecting every tracked change in the manuscript "
            f"reproduces all {audit.checked} paragraph(s) exactly as they were "
            f"received, so nothing reached the text without a revision mark "
            f"around it. The two silent normalizations above are applied "
            f"before that baseline is taken and are counted separately.")
    elif audit is not None and audit.ran:
        d.add_paragraph(
            f"FAILED. {audit.summary()}. This document should not be treated "
            f"as final.")
    else:
        d.add_paragraph(
            "Not run for this pass, so there is no mechanical proof that "
            "rejecting the tracked changes returns the original text.")

    d.add_heading("Limits of this pass", level=1)
    for line in _limits(cfg, doc, findings, spell):
        d.add_paragraph(line, style="List Bullet")

    d.add_heading("How this pass was made", level=1)
    passes = "; ".join(" + ".join(g) for g in cfg.error_type_groups)
    d.add_paragraph(
        f"Model: {cfg.api.model}. Confidence gate: {cfg.min_confidence}. "
        f"{len(doc.paragraphs)} paragraph(s) reviewed in "
        f"{len(cfg.error_type_groups)} pass(es) over "
        f"{len(cfg.error_type_keys)} error type(s): {passes}.")
    if spell is not None and spell.available:
        d.add_paragraph(
            f"Dictionary scan: {spell.tokens:,} words read, {spell.unknown:,} "
            f"not in the dictionary, {len(spell.lexicon)} protected as names, "
            f"{len(spell.candidates)} raised for a look, and "
            f"{len(spell.recurring)} noted as repeated vocabulary.")

    d.save(path)
    log.info("Wrote %s", path)
