"""The whole-book continuity read: the model's contradictions turned into
correctly-anchored margin queries, the two-quote precision guardrail, and the
deterministic date->weekday tier."""
from __future__ import annotations

import itertools
import logging

from docproof.continuity import (ContinuityFinding, ContinuityReport,
                                 build_continuity, calendar_findings,
                                 default_continuity_prompt, report_to_findings)
from docproof.models import DocumentModel, ParagraphRef, Usage
from docproof.providers.base import ProviderResult
from docproof.validator import validate_findings

from .fakes import FakeProvider


def _para(pid: str, text: str) -> ParagraphRef:
    return ParagraphRef(para_id=pid, part="word/document.xml", location="body",
                        text=text, style="Normal")


def _cf(**kw) -> ContinuityFinding:
    base = dict(category="attribute",
                earlier_quote="Mira's eyes were deep brown.",
                quote="Mira turned, her blue eyes bright.",
                question="Are Mira's eyes brown or blue?",
                confidence="high")
    base.update(kw)
    return ContinuityFinding(**base)


# --- report -> query findings -------------------------------------------------

def test_a_grounded_contradiction_becomes_one_query_at_the_later_site():
    paras = [_para("body-0", "Mira's eyes were deep brown."),
             _para("body-1", "Mira turned, her blue eyes bright.")]
    out = report_to_findings(ContinuityReport(findings=[_cf()]), paras,
                             itertools.count(1))
    assert len(out) == 1
    f = out[0]
    assert f.force_query is True                       # never an edit
    assert f.error_type == "continuity"
    assert f.para_id == "body-1"                        # the later, breaking site
    assert f.original_text == "Mira turned, her blue eyes bright."
    assert f.corrected_text == f.original_text          # a question edits nothing
    assert f.explanation.startswith("attribute:")
    assert 'Earlier: "Mira\'s eyes were deep brown."' in f.explanation


def test_a_hallucinated_earlier_quote_is_dropped():
    # The later quote is real; the establishing one is invented — exactly what a
    # fabricated contradiction looks like, and what the two-quote rule kills.
    paras = [_para("body-0", "Mira turned, her blue eyes bright.")]
    cf = _cf(earlier_quote="Mira had piercing green eyes.")
    assert report_to_findings(ContinuityReport(findings=[cf]), paras,
                              itertools.count(1)) == []


def test_a_hallucinated_later_quote_is_dropped():
    paras = [_para("body-0", "Mira's eyes were deep brown.")]
    cf = _cf(quote="Mira turned, her blue eyes bright.")   # nowhere in the text
    assert report_to_findings(ContinuityReport(findings=[cf]), paras,
                              itertools.count(1)) == []


def test_min_confidence_drops_soft_calls():
    paras = [_para("body-0", "Mira's eyes were deep brown."),
             _para("body-1", "Mira turned, her blue eyes bright.")]
    rep = ContinuityReport(findings=[_cf(confidence="low")])
    assert report_to_findings(rep, paras, itertools.count(1),
                              min_confidence="medium") == []
    # ...but the same finding survives when the floor is lowered.
    assert len(report_to_findings(rep, paras, itertools.count(1),
                                  min_confidence="low")) == 1


def test_max_queries_caps_the_output():
    paras = [_para("body-0", "Mira's eyes were deep brown."),
             _para("body-1", "Mira turned, her blue eyes bright.")]
    rep = ContinuityReport(findings=[_cf() for _ in range(5)])
    out = report_to_findings(rep, paras, itertools.count(1), max_queries=2)
    assert len(out) == 2


def test_a_query_routes_through_the_validator_to_the_margin():
    # End to end: a force_query finding anchors on the whole quoted sentence and
    # comes out of the validator as a "query", not an edit — even at a high
    # confidence floor, because it never enters the edit machinery.
    paras = (_para("body-0", "Mira's eyes were deep brown."),
             _para("body-1", "Mira turned, her blue eyes bright."))
    doc = DocumentModel(source_path="x.docx", paragraphs=paras)
    findings = report_to_findings(ContinuityReport(findings=[_cf()]),
                                  list(paras), itertools.count(1))
    validated = validate_findings(findings, doc, min_confidence="high")
    assert [f.status for f in validated] == ["query"]
    anchor = validated[0].anchor
    assert doc.paragraphs[1].text[anchor.start:anchor.end] == \
        "Mira turned, her blue eyes bright."
    assert anchor.insert_text == ""                    # a comment inserts nothing


# --- punctuation-tolerant locating ---------------------------------------------
#
# The ingest normalizer curls straight quotes before review, so the canonical
# text of a real manuscript holds “ ” ‘ ’ (and house-style nbsp) in ~98% of its
# sentences — and a model re-typing those marks as ASCII is the routine habit
# validator.py folds for. Before the fold, that habit silently zeroed the whole
# model tier: every finding with a retyped mark in either quote was dropped.

def test_a_straight_retyped_quote_survives_the_curly_manuscript():
    paras = [_para("body-0", "“Mira’s eyes were deep brown,” her mother said."),
             _para("body-1", "Mira turned, and the girl’s blue eyes flashed.")]
    cf = _cf(earlier_quote='"Mira\'s eyes were deep brown," her mother said.',
             quote="Mira turned, and the girl's blue eyes flashed.")
    out = report_to_findings(ContinuityReport(findings=[cf]), paras,
                             itertools.count(1))
    assert len(out) == 1
    f = out[0]
    # The anchor is the manuscript's own characters, not the model's ASCII...
    assert f.original_text == "Mira turned, and the girl’s blue eyes flashed."
    assert f.corrected_text == f.original_text
    # ...and so is the earlier sentence the margin comment shows the author.
    assert "“Mira’s eyes were deep brown,” her mother said." in f.explanation


def test_an_nbsp_retyped_as_a_space_still_locates():
    # House style sets nbsp before an ellipsis; a model re-types a plain
    # space. The anchor that comes back is the manuscript's own characters,
    # nbsp included — the comment's display copy may collapse it, the anchor
    # never does.
    real = "He waited\u00a0… then walked through the locked door."
    paras = [_para("body-0", "The door stayed locked all night."),
             _para("body-1", real)]
    cf = _cf(category="object",
             earlier_quote="The door stayed locked all night.",
             quote="He waited … then walked through the locked door.",
             question="Who unlocked the door?")
    out = report_to_findings(ContinuityReport(findings=[cf]), paras,
                             itertools.count(1))
    assert len(out) == 1
    assert out[0].original_text == real          # nbsp, not the model's space


def test_a_fold_recovered_query_anchors_through_the_validator():
    # End to end: the finding _locate recovered through the fold must be one
    # the validator can anchor exactly — original_text is a real slice and the
    # occurrence is counted on it, so the exact path lands without the fold.
    paras = (_para("body-0", "“Mira’s eyes were deep brown,” her mother said."),
             _para("body-1", "Mira turned, and the girl’s blue eyes flashed."))
    doc = DocumentModel(source_path="x.docx", paragraphs=paras)
    cf = _cf(earlier_quote='"Mira\'s eyes were deep brown," her mother said.',
             quote="Mira turned, and the girl's blue eyes flashed.")
    findings = report_to_findings(ContinuityReport(findings=[cf]),
                                  list(paras), itertools.count(1))
    validated = validate_findings(findings, doc, min_confidence="high")
    assert [f.status for f in validated] == ["query"]
    anchor = validated[0].anchor
    assert doc.paragraphs[1].text[anchor.start:anchor.end] == \
        "Mira turned, and the girl’s blue eyes flashed."


def test_fold_recovery_counts_occurrences_on_the_real_text():
    # The same sentence twice in one paragraph: the folded search matches the
    # first copy, and the occurrence stored must be the REAL slice's — what the
    # validator's exact anchor path counts — or the query lands on the wrong
    # copy.
    paras = [_para("body-0", "Petra locked the door at nine."),
             _para("body-1", "She said it’s time. She said it’s time.")]
    cf = _cf(earlier_quote="Petra locked the door at nine.",
             quote="She said it's time.")
    findings = report_to_findings(ContinuityReport(findings=[cf]), paras,
                                  itertools.count(1))
    assert len(findings) == 1
    assert findings[0].occurrence == 1
    doc = DocumentModel(source_path="x.docx", paragraphs=tuple(paras))
    validated = validate_findings(findings, doc, min_confidence="high")
    assert validated[0].status == "query"
    assert validated[0].anchor.start == 0               # the first copy


def test_the_fold_forgives_typography_never_wording():
    # A quote that differs in wording — a dropped apostrophe-s changes the
    # word, not just the mark — still fails to locate and is still dropped:
    # the guardrail's point survives the fold.
    paras = [_para("body-0", "Mira’s eyes were deep brown."),
             _para("body-1", "Mira turned, and the girl’s blue eyes flashed.")]
    cf = _cf(earlier_quote="Miras eyes were deep brown.",     # not the text
             quote="Mira turned, and the girl's blue eyes flashed.")
    assert report_to_findings(ContinuityReport(findings=[cf]), paras,
                              itertools.count(1)) == []


def test_kept_and_dropped_counts_reach_the_log(caplog):
    # The one visibility line: kept-of-flagged with every drop reason counted,
    # INFO on a normal run...
    paras = [_para("body-0", "Mira's eyes were deep brown."),
             _para("body-1", "Mira turned, her blue eyes bright.")]
    rep = ContinuityReport(findings=[
        _cf(),                                          # kept
        _cf(confidence="low"),                          # below the floor
        _cf(quote="A sentence found nowhere at all.")])  # unlocatable
    with caplog.at_level(logging.INFO, logger="docproof.continuity"):
        out = report_to_findings(rep, paras, itertools.count(1))
    assert len(out) == 1
    kept = [r for r in caplog.records if "kept 1 of 3" in r.message]
    assert kept and kept[0].levelno == logging.INFO
    assert "1 below medium confidence" in kept[0].message
    assert "1 with a quote that did not locate" in kept[0].message


def test_a_read_whose_findings_all_die_warns(caplog):
    # ...and a WARNING when a non-empty report kept nothing: the hosted app
    # records nothing below WARNING, and "flagged but none reached the margin"
    # is the outcome an operator must be able to tell apart from "found
    # nothing" — before this line, both were the same silence.
    paras = [_para("body-0", "Mira's eyes were deep brown.")]
    rep = ContinuityReport(findings=[
        _cf(quote="A sentence found nowhere at all.")])
    with caplog.at_level(logging.INFO, logger="docproof.continuity"):
        assert report_to_findings(rep, paras, itertools.count(1)) == []
    warned = [r for r in caplog.records if "kept 0 of 1" in r.message]
    assert warned and warned[0].levelno == logging.WARNING


# --- deterministic calendar tier ----------------------------------------------

def test_a_wrong_weekday_for_a_real_date_is_flagged():
    # June 3, 2019 was a Monday, not a Tuesday.
    paras = [_para("body-0", "They met on Tuesday, June 3, 2019, in the rain.")]
    out = calendar_findings(paras, itertools.count(1))
    assert len(out) == 1
    f = out[0]
    assert f.error_type == "continuity" and f.force_query is True
    assert "was a Monday" in f.explanation and "Tuesday" in f.explanation


def test_a_correct_weekday_is_not_flagged():
    # June 3, 2019 really was a Monday.
    paras = [_para("body-0", "They met on Monday, June 3, 2019, in the rain.")]
    assert calendar_findings(paras, itertools.count(1)) == []


def test_a_weekday_and_date_in_different_sentences_are_not_paired():
    paras = [_para("body-0", "It was a Tuesday. Much later, June 3, 2019 arrived.")]
    assert calendar_findings(paras, itertools.count(1)) == []


def test_an_impossible_date_is_skipped_not_crashed():
    paras = [_para("body-0", "On Monday, February 30, 2019 nothing happened.")]
    assert calendar_findings(paras, itertools.count(1)) == []


def test_a_date_with_no_weekday_claim_is_ignored():
    paras = [_para("body-0", "It happened on June 3, 2019, without ceremony.")]
    assert calendar_findings(paras, itertools.count(1)) == []


# --- build_continuity: additive and best-effort -------------------------------

def test_build_reads_once_and_parses_the_report():
    from tests.fakes import USAGE
    result = ProviderResult(parsed={"findings": [{
        "category": "timeline", "earlier_quote": "It was Monday.",
        "quote": "The next day, Wednesday, they left.",
        "question": "Where did Tuesday go?", "confidence": "high"}]},
        usage=USAGE)
    prov = FakeProvider([result])
    usage = Usage()
    rep = build_continuity([_para("body-0", "It was Monday.")], prov,
                           model="claude-fable-5", max_tokens=8000, usage=usage)
    assert len(rep.findings) == 1 and rep.findings[0].category == "timeline"
    assert usage.api_calls == 1                         # billed once, whole book


def test_build_degrades_to_empty_on_a_bad_response():
    from tests.fakes import USAGE
    prov = FakeProvider([ProviderResult(stop_reason="error", error="refused",
                                        usage=USAGE)])
    rep = build_continuity([_para("body-0", "text")], prov,
                           model="claude-fable-5", max_tokens=8000, usage=Usage())
    assert rep.findings == []                           # review proceeds regardless


def test_a_truncated_read_retries_once_at_double_the_ceiling():
    # The ceiling covers thinking, and reasoning spend varies run to run — one
    # truncation is a coin flip lost, not proof the book is too big. This is
    # one indivisible call (no window to split), so the recovery is headroom:
    # retry once at double, and both calls are billed as the two calls they
    # were.
    from tests.fakes import USAGE
    good = ProviderResult(parsed={"findings": [{
        "category": "timeline", "earlier_quote": "It was Monday.",
        "quote": "The next day, Wednesday, they left.",
        "question": "Where did Tuesday go?", "confidence": "high"}]},
        usage=USAGE)
    prov = FakeProvider([
        ProviderResult(stop_reason="max_tokens", error="output truncated",
                       usage=USAGE),
        good])
    usage = Usage()
    rep = build_continuity([_para("body-0", "It was Monday.")], prov,
                           model="claude-fable-5", max_tokens=8000, usage=usage)
    assert len(rep.findings) == 1                       # the retry's answer
    assert [c["max_tokens"] for c in prov.calls] == [8000, 16000]
    assert usage.api_calls == 2                         # the truncation was billed


def test_a_read_that_truncates_twice_degrades_to_empty():
    # The retry is once, not a loop: a second truncation lands on the ordinary
    # error path — empty report, review proceeds — exactly where one
    # truncation landed before the retry existed.
    from tests.fakes import USAGE
    truncated = ProviderResult(stop_reason="max_tokens",
                               error="output truncated", usage=USAGE)
    prov = FakeProvider([truncated, truncated])
    usage = Usage()
    rep = build_continuity([_para("body-0", "text")], prov,
                           model="claude-fable-5", max_tokens=8000, usage=usage)
    assert rep.findings == []
    assert len(prov.calls) == 2                         # no third attempt
    assert usage.api_calls == 2


# --- pipeline wiring ----------------------------------------------------------

def test_the_pass_runs_in_the_pipeline_and_emits_a_query(tmp_path):
    # End to end through run_sync: with continuity on, the deterministic calendar
    # tier surfaces a date->weekday mismatch as a force_query finding, proving the
    # post-loop block is wired and appends to the run's findings. The model read
    # returns an empty report from the fake (like every other structured call),
    # so this isolates the wiring, not the model.
    import docx
    from docproof.config import load_config
    from docproof.pipeline import prepare, run_sync

    d = docx.Document()
    d.add_paragraph("They met on Tuesday, June 3, 2019, in the rain.")
    src = tmp_path / "m.docx"
    d.save(src)

    cfg = load_config("config/default.yaml")
    cfg.error_types = [["spelling"]]
    cfg.glossary.enabled = False
    cfg.adjudicate.enabled = False
    cfg.rewrite.enabled = False
    cfg.languagetool.enabled = False
    cfg.continuity.enabled = True                       # calendar_check on by default

    prepared = prepare(cfg, str(src), "config/error_types")
    findings, _ = run_sync(cfg, prepared,
                           provider_factory=lambda dcfg: FakeProvider())

    conts = [f for f in findings if f.error_type == "continuity"]
    assert len(conts) == 1
    assert conts[0].force_query is True
    assert "was a Monday" in conts[0].explanation


def _batch_cfg():
    from docproof.config import load_config
    cfg = load_config("config/default.yaml")
    cfg.error_types = [["spelling"]]
    cfg.glossary.enabled = False
    cfg.adjudicate.enabled = False
    cfg.rewrite.enabled = False
    cfg.languagetool.enabled = False
    cfg.continuity.enabled = True                       # calendar_check on by default
    return cfg


def _mismatch_docx(tmp_path):
    """A manuscript whose stated weekday is wrong — June 3, 2019 was a Monday."""
    import docx
    d = docx.Document()
    d.add_paragraph("They met on Tuesday, June 3, 2019, in the rain.")
    src = tmp_path / "m.docx"
    d.save(src)
    return src


def test_the_pass_runs_on_the_batch_path_too(tmp_path, monkeypatch):
    # Batch is the app's default submission mode, so this is the common path, not
    # an edge case: the read cannot ride the review batch (its own model, one
    # request), so — like the glossary read — it runs synchronously at collect.
    # Without it a batch review would report and price a read it never made.
    from docproof import batch as batchlib
    from docproof import providers as providers_mod

    cfg = _batch_cfg()
    reader = FakeProvider()                 # the continuity model's own client
    monkeypatch.setattr(providers_mod, "build_provider", lambda c, **kw: reader)

    provider = FakeProvider()
    job = batchlib.submit(cfg, str(_mismatch_docx(tmp_path)),
                          "config/error_types", provider, tmp_path / "ws")
    result = batchlib.collect_findings(job, provider, "config/error_types")

    conts = [f for f in result.findings if f.error_type == "continuity"]
    assert len(conts) == 1                              # the calendar tier's query
    assert conts[0].force_query is True
    assert "was a Monday" in conts[0].explanation
    # And the model read itself happened, not only the free calendar check.
    assert [c["schema_name"] for c in reader.calls] == ["continuity"]


def test_a_collected_batch_review_reports_only_the_read_it_made(tmp_path,
                                                                monkeypatch):
    # The deliverable's own claim: summary.md prints the pass list straight off
    # the config, so "Continuity read on" has to mean a continuity finding could
    # have come out of this run. The two must not disagree.
    import json

    from docproof import batch as batchlib
    from docproof import providers as providers_mod

    cfg = _batch_cfg()
    monkeypatch.setattr(providers_mod, "build_provider",
                        lambda c, **kw: FakeProvider())

    provider = FakeProvider()
    ws = tmp_path / "ws"
    job = batchlib.submit(cfg, str(_mismatch_docx(tmp_path)),
                          "config/error_types", provider, ws)
    outputs = batchlib.collect(job, provider, "config/error_types", ws)

    assert "Continuity read on" in outputs.summary_md.read_text("utf-8")
    findings = json.loads(outputs.findings_json.read_text("utf-8"))["findings"]
    assert [f["error_type"] for f in findings if f["error_type"] == "continuity"]


# --- editable prompt ----------------------------------------------------------

def test_default_continuity_prompt_is_the_builtin():
    from docproof.continuity import _SYSTEM
    assert default_continuity_prompt() == _SYSTEM


def test_a_custom_prompt_replaces_the_builtin_and_reaches_the_provider():
    from tests.fakes import USAGE
    prov = FakeProvider([ProviderResult(parsed={"findings": []}, usage=USAGE)])
    build_continuity([_para("body-0", "Some prose.")], prov,
                     model="claude-fable-5", max_tokens=8000, usage=Usage(),
                     prompt="  Only look for eye-colour drift.  ")
    assert prov.calls[0]["system"] == "Only look for eye-colour drift."  # stripped


def test_an_empty_prompt_uses_the_builtin_default():
    from tests.fakes import USAGE
    from docproof.continuity import _SYSTEM
    prov = FakeProvider([ProviderResult(parsed={"findings": []}, usage=USAGE)])
    build_continuity([_para("body-0", "Some prose.")], prov,
                     model="claude-fable-5", max_tokens=8000, usage=Usage(),
                     prompt="   ")                       # whitespace = untouched
    assert prov.calls[0]["system"] == _SYSTEM


def test_the_prompt_is_part_of_the_cache_key(tmp_path):
    from tests.fakes import USAGE

    def _read(prompt):
        prov = FakeProvider([ProviderResult(parsed={"findings": []}, usage=USAGE)])
        build_continuity([_para("body-0", "Same prose, same model.")], prov,
                         model="claude-fable-5", max_tokens=8000, usage=Usage(),
                         prompt=prompt, cache_dir=str(tmp_path))
        return prov

    a = _read("look for timeline slips")
    b = _read("look for timeline slips")               # identical → cache hit
    c = _read("look for eye-colour drift")             # changed → re-read
    assert len(a.calls) == 1                            # first read: a real call
    assert len(b.calls) == 0                            # served from cache
    assert len(c.calls) == 1                            # different prompt, new key
    assert len(list(tmp_path.glob("continuity-*.json"))) == 2   # two distinct keys


# --- continuity-only shape ----------------------------------------------------

def test_continuity_only_runs_with_no_detector_passes(tmp_path):
    # The shape JobStore.config_for produces for a "continuity only" run: no
    # error types, no sweeps, other passes off, continuity on. The review still
    # completes and yields only continuity queries.
    import docx
    from docproof.config import load_config
    from docproof.pipeline import prepare, run_sync

    d = docx.Document()
    d.add_paragraph("They met on Tuesday, June 3, 2019, in the rain.")
    src = tmp_path / "m.docx"
    d.save(src)

    cfg = load_config("config/default.yaml")
    cfg.error_types = []
    cfg.sweeps = []
    for p in ("glossary", "adjudicate", "rewrite", "languagetool",
              "consistency", "spellcheck"):
        getattr(cfg, p).enabled = False
    cfg.continuity.enabled = True

    prepared = prepare(cfg, str(src), "config/error_types")
    assert prepared.request_count == 0                 # nothing sent per chunk
    findings, _ = run_sync(cfg, prepared,
                           provider_factory=lambda dcfg: FakeProvider())
    assert [f.error_type for f in findings] == ["continuity"]
