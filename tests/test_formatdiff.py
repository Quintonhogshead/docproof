"""The InDesign-formatting comparator: each paragraph's style NAME is read and
resolved, paragraphs align by reject-all content, and two docs' styling buckets
into same / different / only-A / only-B."""
from __future__ import annotations

from pathlib import Path

import docx
import pytest

from docproof.formatdiff import (DocStyles, FormatReport, TrackDiffError,
                                 compare_files, compare_styles, extract_styles,
                                 open_docx, render_markdown, report_json)


# --- the core comparison, on hand-built style maps ---------------------------

def _doc(styles: dict[str, str], base: dict[str, str] | None = None) -> DocStyles:
    base = base or {pid: f"paragraph {pid}" for pid in styles}
    return DocStyles(base=base, style=styles)


def test_same_style_agrees_different_style_does_not():
    a = _doc({"body-0000": "Chapter Title", "body-0001": "Body"})
    b = _doc({"body-0000": "Heading", "body-0001": "Body"})
    r = compare_styles(a, b, label_a="human", label_b="docproof")
    assert (r.agree, r.different) == (1, 1)
    assert r.agreement == 0.5
    assert not r.only_a and not r.only_b


def test_a_paragraph_in_only_one_file_has_no_partner():
    a = _doc({"body-0000": "Heading", "body-0001": "Body"},
             base={"body-0000": "Chapter One", "body-0001": "An extra line."})
    b = _doc({"body-0000": "Heading"}, base={"body-0000": "Chapter One"})
    r = compare_styles(a, b)
    assert r.agree == 1
    assert r.only_a == ["body-0001"] and r.only_b == []


def test_agreement_is_none_when_nothing_aligns():
    a = _doc({"body-0000": "Heading"}, base={"body-0000": "The cat sat."})
    b = _doc({"body-0000": "Heading"}, base={"body-0000": "A dog ran."})
    r = compare_styles(a, b)
    assert r.aligned_paras == 0
    assert r.agreement is None
    assert r.only_a == ["body-0000"] and r.only_b == ["body-0000"]


def test_blank_paragraphs_are_not_compared():
    # An empty paragraph carries no InDesign meaning; the prep pass removes it.
    a = _doc({"body-0000": "Heading", "body-0001": "Spacer"},
             base={"body-0000": "Chapter One", "body-0001": "   "})
    b = _doc({"body-0000": "Heading"}, base={"body-0000": "Chapter One"})
    r = compare_styles(a, b)
    assert r.aligned_paras == 1 and r.agree == 1
    assert not r.only_a and not r.only_b


# --- reading styles out of real .docx files ----------------------------------

def _styled(tmp_path: Path, name: str, specs) -> Path:
    """A .docx whose paragraphs carry the given (text, style-name) pairs. A
    None style leaves the paragraph on the default (Normal)."""
    d = docx.Document()
    for text, style in specs:
        d.add_paragraph(text, style=style) if style else d.add_paragraph(text)
    path = tmp_path / f"{name}.docx"
    d.save(path)
    return path


def test_extract_resolves_the_style_name(tmp_path):
    doc = _styled(tmp_path, "one",
                  [("Chapter One", "Heading 1"), ("It was dark.", None)])
    styles = extract_styles(open_docx(doc))
    # Word stores an id; the comparison wants the human-facing NAME.
    assert styles.style["body-0000"] == "heading 1"   # builtin's own casing
    assert styles.style["body-0001"] == "Normal"
    assert styles.styled_count == 2


def test_compare_files_end_to_end(tmp_path):
    a = _styled(tmp_path, "human",
                [("Chapter One", "Heading 1"), ("It was dark.", None)])
    b = _styled(tmp_path, "docproof",
                [("Chapter One", "Title"), ("It was dark.", None)])
    r = compare_files(a, b, label_a="human", label_b="docproof")
    assert r.agree == 1          # the body paragraph, Normal in both
    assert r.different == 1      # the heading: "heading 1" vs "Title"
    diffs = report_json(r)["differences"]
    assert len(diffs) == 1
    assert diffs[0]["style_a"] != diffs[0]["style_b"]
    assert diffs[0]["text"] == "Chapter One"


def test_report_json_shape():
    a = _doc({"body-0000": "Title", "body-0001": "Body"})
    b = _doc({"body-0000": "Heading", "body-0001": "Body"})
    j = report_json(compare_styles(a, b, label_a="human", label_b="docproof"))
    assert j["mode"] == "formatting"
    assert j["totals"] == {"agree": 1, "different": 1, "only_a": 0, "only_b": 0}
    assert j["agreement"] == 0.5


def test_render_markdown_smoke():
    a = _doc({"body-0000": "Title", "body-0001": "Body"})
    b = _doc({"body-0000": "Heading", "body-0001": "Body"})
    md = render_markdown(compare_styles(a, b, label_a="human", label_b="docproof"))
    assert "InDesign-formatting comparison" in md
    assert "Style agreement" in md


def test_open_docx_rejects_non_docx(tmp_path):
    bad = tmp_path / "x.docx"
    bad.write_bytes(b"not a zip")
    with pytest.raises(TrackDiffError, match="not a .docx"):
        open_docx(bad)
