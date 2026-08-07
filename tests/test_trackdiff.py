"""The tracked-changes comparator: edits extract into reject-coordinates, a
replacement coalesces into one anchor, and two docs' edit sets bucket into
agree / different-fix / only-A / only-B against a shared base."""
from __future__ import annotations

from pathlib import Path

import docx
import pytest

from docproof.config import load_config
from docproof.models import Anchor
from docproof.pipeline import finish, prepare, run_sync
from docproof.trackdiff import (DocEdits, TrackDiffError, compare_edits,
                                compare_files, extract_edits, open_docx,
                                render_markdown)
from docproof.analyzer import MockAnalyzer

ROOT = Path(__file__).parent.parent
ERROR_DIR = ROOT / "config" / "error_types"
CONFIG = str(ROOT / "config" / "default.yaml")

SPLICE_0 = "The manuscript was finished, nobody wanted to read it."
SPLICE_1 = "She opened the letter, her hands would not stay still."


def _reviewed(tmp_path, name: str, canned: list[dict]) -> Path:
    """Build a base doc with two comma splices and run docproof (mock findings)
    to produce a real tracked-changes .docx."""
    import json
    base = tmp_path / f"{name}-base.docx"
    d = docx.Document()
    d.add_paragraph(SPLICE_0)
    d.add_paragraph(SPLICE_1)
    d.save(base)

    cfg = load_config(CONFIG)
    cfg.error_types = [["comma_splice"]]
    cfg.output_dir = str(tmp_path / name)
    prepared = prepare(cfg, str(base), ERROR_DIR)
    ids = __import__("itertools").count(1)
    from docproof.models import Usage
    usage = Usage()
    findings = []
    for group in prepared.groups:
        mock = MockAnalyzer(group, canned, ids)
        for chunk in prepared.chunks:
            found, _ = mock.analyze_chunk(chunk, usage)
            findings.extend(found)
    out = finish(prepared, findings, usage, cfg,
                 out_dir=cfg.output_dir, source_path=str(base))
    return out.reviewed_path


def _fix(para_id: str, original: str, corrected: str) -> dict:
    return {"para_id": para_id, "error_type": "comma_splice",
            "original_text": original, "corrected_text": corrected,
            "confidence": "high", "explanation": "x"}


def test_extract_coalesces_a_replacement_into_one_anchor(tmp_path):
    doc = _reviewed(tmp_path, "one",
                    [_fix("body-0000", SPLICE_0,
                          SPLICE_0.replace(", nobody", "; nobody"))])
    edits = extract_edits(open_docx(doc))
    anchors = edits.edits["body-0000"]
    assert len(anchors) == 1
    a = anchors[0]
    # A comma → semicolon replacement: both sides present, one anchor.
    assert a.delete_text == "," and a.insert_text == ";"
    # Offsets index the reject (base) text, which is the original splice.
    assert edits.base["body-0000"] == SPLICE_0
    assert SPLICE_0[a.start:a.end] == ","


def test_agree_when_both_make_the_same_fix(tmp_path):
    fix0 = _fix("body-0000", SPLICE_0, SPLICE_0.replace(", nobody", "; nobody"))
    a = _reviewed(tmp_path, "a", [fix0])
    b = _reviewed(tmp_path, "b", [fix0])
    r = compare_files(a, b)
    assert (r.agree, r.diff_fix, r.only_a, r.only_b) == (1, 0, 0, 0)
    assert r.located_recall == 1.0 and r.precision == 1.0


def test_only_b_is_a_docproof_extra(tmp_path):
    fix0 = _fix("body-0000", SPLICE_0, SPLICE_0.replace(", nobody", "; nobody"))
    fix1 = _fix("body-0001", SPLICE_1, SPLICE_1.replace(", her", "; her"))
    a = _reviewed(tmp_path, "a", [fix0])
    b = _reviewed(tmp_path, "b", [fix0, fix1])
    r = compare_files(a, b, label_a="human", label_b="docproof")
    assert (r.agree, r.only_a, r.only_b) == (1, 0, 1)
    assert r.located_recall == 1.0        # human's one edit was matched
    assert r.precision == 0.5             # docproof made two, one unmatched


def test_only_a_is_a_docproof_miss(tmp_path):
    fix0 = _fix("body-0000", SPLICE_0, SPLICE_0.replace(", nobody", "; nobody"))
    fix1 = _fix("body-0001", SPLICE_1, SPLICE_1.replace(", her", "; her"))
    a = _reviewed(tmp_path, "a", [fix0, fix1])   # human caught both
    b = _reviewed(tmp_path, "b", [fix0])         # docproof caught one
    r = compare_files(a, b)
    assert (r.agree, r.only_a, r.only_b) == (1, 1, 0)
    assert r.located_recall == 0.5        # missed one of the human's two


def test_different_fix_same_place(tmp_path):
    # Human uses a semicolon; docproof splits into two sentences. Same span,
    # different net text → different fix, not agreement.
    a = _reviewed(tmp_path, "a",
                  [_fix("body-0000", SPLICE_0,
                        SPLICE_0.replace(", nobody", "; nobody"))])
    b = _reviewed(tmp_path, "b",
                  [_fix("body-0000", SPLICE_0,
                        SPLICE_0.replace(", nobody", ". Nobody"))])
    r = compare_files(a, b)
    assert (r.agree, r.diff_fix, r.only_a, r.only_b) == (0, 1, 0, 0)


def test_word_level_and_char_level_edits_agree():
    """The report pathology: one document records a change as a single wide
    edit (a whole-word retype) and the other as the minimal per-character fix.
    Same resulting text, so they must agree — not land in 'different fix' with a
    phantom 'only in docproof' left over, which is what pairwise anchor matching
    produced before the two sides were canonicalized to one granularity."""
    base = "the middle city"
    a = DocEdits(base={"body-0000": base},
                 edits={"body-0000": [Anchor(4, 15, "middle city", "Middle City")]})
    b = DocEdits(base={"body-0000": base},
                 edits={"body-0000": [Anchor(4, 5, "m", "M"),
                                      Anchor(11, 12, "c", "C")]})
    r = compare_edits(a, b, label_a="human", label_b="docproof")
    assert (r.agree, r.diff_fix, r.only_a, r.only_b) == (2, 0, 0, 0)
    assert r.precision == 1.0 and r.exact_recall == 1.0


def test_partial_overlap_is_one_clean_miss_not_a_double_count():
    """Human capitalizes two words as one wide edit; docproof caught only the
    first. The second is a single clean miss, not a phantom smeared across
    'different fix' and 'only human'."""
    base = "the middle city"
    a = DocEdits(base={"body-0000": base},
                 edits={"body-0000": [Anchor(4, 15, "middle city", "Middle City")]})
    b = DocEdits(base={"body-0000": base},
                 edits={"body-0000": [Anchor(4, 5, "m", "M")]})
    r = compare_edits(a, b, label_a="human", label_b="docproof")
    assert (r.agree, r.diff_fix, r.only_a, r.only_b) == (1, 0, 1, 0)


def test_canonical_anchors_merges_a_close_fix_but_not_distant_ones():
    from docproof.trackdiff import _canonical_anchors
    # A dialogue-tag fix — period→comma and the pronoun it lowercases, with the
    # closing quote and a space between — is one logical edit, so one anchor.
    assert len(_canonical_anchors('stop." He ran', 'stop," he ran')) == 1
    # Two capitalizations six characters apart are two separate edits.
    assert len(_canonical_anchors("the middle city", "the Middle City")) == 2


def test_unaligned_paragraph_is_reported_not_matched():
    # Two synthetic docs whose bases differ: nothing should match, and the
    # paragraph is reported as unaligned.
    a = DocEdits(base={"body-0000": "The cat sat."},
                 edits={"body-0000": [Anchor(3, 4, " ", "")]})
    b = DocEdits(base={"body-0000": "A dog ran fast today."},
                 edits={"body-0000": [Anchor(1, 2, " ", "")]})
    r = compare_edits(a, b)
    assert r.aligned_paras == 0
    assert r.unaligned == ["body-0000"]
    assert r.agree == 0 and r.only_a == 0 and r.only_b == 0


def test_open_docx_rejects_non_docx(tmp_path):
    bad = tmp_path / "x.docx"
    bad.write_bytes(b"not a zip")
    with pytest.raises(TrackDiffError, match="not a .docx"):
        open_docx(bad)


def test_render_markdown_smoke(tmp_path):
    fix0 = _fix("body-0000", SPLICE_0, SPLICE_0.replace(", nobody", "; nobody"))
    fix1 = _fix("body-0001", SPLICE_1, SPLICE_1.replace(", her", "; her"))
    a = _reviewed(tmp_path, "a", [fix0])
    b = _reviewed(tmp_path, "b", [fix0, fix1])
    ea = extract_edits(open_docx(a))
    r = compare_files(a, b, label_a="human", label_b="docproof")
    md = render_markdown(r, ea.base)
    assert "Tracked-changes comparison" in md
    assert "ground truth" in md
