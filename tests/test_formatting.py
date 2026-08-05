"""Findings that change how text is set rather than what it says.

A long-work title in roman is a house-style error that alters no characters,
so it cannot be an insertion and a deletion. Word records it as w:rPrChange:
the run carries its new properties and the change element holds the old ones.
The invariant that matters is that the text is untouched in both views.
"""
from __future__ import annotations

import itertools
import re
import zipfile

import docx
import pytest

from docproof.analyzer import build_system_prompt
from docproof.config import load_config
from docproof.error_registry import load_error_types
from docproof.models import DocumentModel, Finding, ParagraphRef, Usage
from docproof.pipeline import finish, prepare
from docproof.reassembler import paragraph_view_text
from docproof.utils.xml_helpers import DocxPackage, qn, walk_package
from docproof.validator import validate_findings

from .test_error_types import ERROR_DIR

FORMAT = {"title_italics": "italic"}
GATSBY = "The Great Gatsby sat unopened on the nightstand."


def _doc(*texts):
    return DocumentModel(source_path="x.docx", paragraphs=tuple(
        ParagraphRef(f"body-{i:04d}", "word/document.xml", "body", t, "Normal")
        for i, t in enumerate(texts)))


def _finding(**kw):
    base = dict(finding_id="f-0001", chunk_id="chunk-000", para_id="body-0000",
                error_type="title_italics", original_text=GATSBY, occurrence=1,
                corrected_text="The Great Gatsby", explanation="Book title.",
                confidence="high")
    return Finding(**{**base, **kw})


# --- the error type -----------------------------------------------------------

def test_title_italics_is_a_format_type():
    et = load_error_types(ERROR_DIR, ["title_italics"])["title_italics"]
    assert et.is_format and et.channel == "format" and et.format == "italic"
    assert not et.is_query


def test_the_prompt_explains_what_corrected_text_means_here():
    types = list(load_error_types(ERROR_DIR, ["title_italics"]).values())
    prompt = build_system_prompt(types)
    assert "CHANNEL: FORMAT" in prompt
    assert "exact span to mark" in prompt
    # It must also say the model cannot see existing italics, or it would try
    # to skip titles that are already correct and miss the ones that are not.
    assert "cannot see existing formatting" in prompt


# --- the validator ------------------------------------------------------------

def test_a_format_finding_anchors_the_span_not_the_sentence():
    out = validate_findings([_finding()], _doc(GATSBY), "medium",
                            format_types=FORMAT)
    assert out[0].status == "validated" and out[0].format == "italic"
    assert out[0].anchor.start == 0
    assert out[0].anchor.end == len("The Great Gatsby")
    # Nothing is deleted or inserted: the text does not change.
    assert out[0].anchor.delete_text == out[0].anchor.insert_text


def test_a_span_that_is_not_in_the_sentence_is_rejected():
    out = validate_findings([_finding(corrected_text="Moby-Dick")],
                            _doc(GATSBY), "medium", format_types=FORMAT)
    assert out[0].status == "rejected_no_anchor"


def test_a_format_finding_is_not_rejected_as_a_noop():
    """The text is identical by design, which for a text edit would mean
    there was nothing to do."""
    out = validate_findings([_finding(corrected_text=GATSBY)], _doc(GATSBY),
                            "medium", format_types=FORMAT)
    assert out[0].status == "validated"


def test_formatting_does_not_block_a_text_edit_on_the_same_words():
    """Italicising a title and fixing a typo inside it are not in conflict, so
    they claim spans in separate registers."""
    text = "The Great Gatsy sat unopened, nobody had read it."
    out = validate_findings([
        _finding(original_text=text, corrected_text="The Great Gatsy"),
        Finding(finding_id="f-0002", chunk_id="c", para_id="body-0000",
                error_type="spelling", original_text=text, occurrence=1,
                corrected_text="The Great Gatsby sat unopened, nobody had read it.",
                explanation="", confidence="high"),
    ], _doc(text), "medium", format_types=FORMAT)
    assert [f.status for f in out] == ["validated", "validated"]


def test_the_same_title_marked_twice_is_a_duplicate():
    out = validate_findings([_finding(), _finding(finding_id="f-0002")],
                            _doc(GATSBY), "medium", format_types=FORMAT)
    assert [f.status for f in out] == ["validated", "rejected_duplicate"]


def test_a_doubtful_title_falls_below_the_gate_like_anything_else():
    out = validate_findings([_finding(confidence="low")], _doc(GATSBY),
                            "medium", format_types=FORMAT)
    assert out[0].status == "skipped_low_confidence"


# --- into the document --------------------------------------------------------

def _run(tmp_path, build, mock):
    d = docx.Document()
    build(d)
    src = tmp_path / "Titles.docx"
    d.save(src)
    cfg = load_config("config/default.yaml")
    cfg.error_types = ["title_italics"]
    prepared = prepare(cfg, src, "config/error_types")

    from docproof.analyzer import MockAnalyzer
    findings, ids = [], itertools.count(1)
    for group in prepared.groups:
        for chunk in prepared.chunks:
            found, _ = MockAnalyzer(group, mock, ids).analyze_chunk(chunk, Usage())
            findings += found
    return finish(prepared, findings, Usage(), cfg, out_dir=tmp_path / "out",
                  source_path=src)


MOCK_GATSBY = [{"para_id": "body-0000", "error_type": "title_italics",
                "original_text": GATSBY, "corrected_text": "The Great Gatsby",
                "explanation": "Book title.", "confidence": "high"}]


def test_a_roman_title_becomes_a_tracked_formatting_revision(tmp_path):
    out = _run(tmp_path, lambda d: d.add_paragraph(GATSBY), MOCK_GATSBY)
    assert out.applied == 1

    body = zipfile.ZipFile(out.reviewed_path).read("word/document.xml").decode()
    run = re.search(r"<w:r>(?:(?!</w:r>).)*?rPrChange.*?</w:r>", body, re.S)
    assert run, "no formatting revision was written"
    assert "The Great Gatsby" in run.group(0)
    # w:rPr's children are a schema-enforced sequence: w:i before w:iCs, and
    # w:rPrChange last.
    assert re.search(r"<w:i/><w:iCs/><w:rPrChange", run.group(0))


def test_the_text_is_untouched_in_both_views(tmp_path):
    """The whole point: a formatting change alters no characters, so accept
    and reject read identically and the audit has nothing to object to."""
    out = _run(tmp_path, lambda d: d.add_paragraph(GATSBY), MOCK_GATSBY)
    para = next(iter(walk_package(DocxPackage(out.reviewed_path)))).element
    assert paragraph_view_text(para, "accept") == GATSBY
    assert paragraph_view_text(para, "reject") == GATSBY


def test_rejecting_restores_the_original_run_properties(tmp_path):
    """w:rPrChange holds the OLD properties, which is what reject applies."""
    out = _run(tmp_path, lambda d: d.add_paragraph(GATSBY), MOCK_GATSBY)
    body = zipfile.ZipFile(out.reviewed_path).read("word/document.xml").decode()
    change = re.search(r"<w:rPrChange[^>]*>(.*?)</w:rPrChange>", body, re.S)
    # The run was plain, so what reject restores is a plain run.
    assert change and "<w:i/>" not in change.group(1)


def test_a_title_already_italic_is_left_alone(tmp_path):
    """The model reads plain text and cannot see italics, so it reports every
    title. Marking a correct one would be a revision to click through for
    nothing."""
    def build(d):
        p = d.add_paragraph("He reread ")
        p.add_run("Hamlet").italic = True
        p.add_run(" every winter without fail.")

    out = _run(tmp_path, build, [
        {"para_id": "body-0000", "error_type": "title_italics",
         "original_text": "He reread Hamlet every winter without fail.",
         "corrected_text": "Hamlet", "explanation": "Play title.",
         "confidence": "high"}])
    assert out.applied == 0
    body = zipfile.ZipFile(out.reviewed_path).read("word/document.xml").decode()
    assert "rPrChange" not in body


def test_italics_turned_explicitly_off_still_counts_as_not_italic():
    """A toggle present with w:val="0" is the document turning italics OFF,
    which is not the same as the property being absent."""
    from lxml import etree
    from docproof.reassembler import _is_set
    rpr = etree.Element(qn("w:rPr"))
    assert not _is_set(rpr, "w:i")
    etree.SubElement(rpr, qn("w:i")).set(qn("w:val"), "0")
    assert not _is_set(rpr, "w:i")
    rpr.find(qn("w:i")).set(qn("w:val"), "1")
    assert _is_set(rpr, "w:i")


def test_the_audit_still_passes_over_a_formatting_change(tmp_path):
    out = _run(tmp_path, lambda d: d.add_paragraph(GATSBY), MOCK_GATSBY)
    import json
    payload = json.loads(out.findings_json.read_text())
    assert payload["audit"]["passed"] is True


def test_the_shipped_config_enables_it():
    assert "title_italics" in load_config("config/default.yaml").error_type_keys
