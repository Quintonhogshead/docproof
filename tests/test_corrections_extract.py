"""Extracting a corrections edit list from the two sources buildable without a
sample PDF: an author's redlined Word file (deterministic) and a free-form list
read by a model (with a fake provider — no vendor is touched).

Both feed the same deterministic core, so each test ends where the engine does:
an anchored edit list that applies to the real InDesign fixture.
"""
from __future__ import annotations

import zipfile
from pathlib import Path

import docx
import pytest

from docproof.corrections.apply import apply_edits
from docproof.corrections.extract import ExtractionError, extract_edits
from docproof.corrections.from_word import edits_from_docx
from docproof.corrections.idml import parse_story
from docproof.corrections.model import DESIGN, JUDGMENT, MECHANICAL
from docproof.models import Usage
from docproof.providers import NormalizedUsage, ProviderResult

from .conftest import FIXTURES
from .fakes import FakeProvider

LAYOUT = FIXTURES / "layout.idml"
W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _w(tag: str) -> str:
    return f"{{{W}}}{tag}"


def make_tracked_docx(path: Path, paragraphs: list[list[tuple[str, str]]]) -> Path:
    """Write a .docx whose paragraphs carry tracked changes. Each paragraph is a
    list of (mode, text) segments: mode "" is plain text, "ins" a tracked
    insertion, "del" a tracked deletion."""
    d = docx.Document()
    for segs in paragraphs:
        p = d.add_paragraph()
        for mode, text in segs:
            run = p.add_run(text)
            if mode in ("ins", "del"):
                wrap = run._r.makeelement(_w(mode), {
                    _w("id"): "1", _w("author"): "Author",
                    _w("date"): "2026-01-01T00:00:00Z"})
                run._r.addprevious(wrap)
                if mode == "del":
                    t = run._r.find(_w("t"))
                    if t is not None:
                        t.tag = _w("delText")
                wrap.append(run._r)
    d.save(str(path))
    return path


def make_commented_pdf(path: Path, lines: list[tuple[float, float, str]],
                       annots: list[dict],
                       font_widths: tuple[int, list[float]] | None = None) -> Path:
    """Write a one-page PDF with text at given positions and comment/highlight
    annotations, so the PDF reader can be tested without a real proof. Each line
    is (x, y, text); each annot is {subtype, rect, contents, quad?}.

    `font_widths` is an optional (first_char, widths) pair that embeds a `/Widths`
    table, so a test can pin exactly how wide each glyph is and thus where a word
    on a multi-word line renders — the geometry the highlight reader walks."""
    from pypdf import PdfWriter
    from pypdf.generic import (ArrayObject, DecodedStreamObject, DictionaryObject,
                               FloatObject, NameObject, NumberObject,
                               TextStringObject)
    w = PdfWriter()
    page = w.add_blank_page(width=612, height=792)
    ops = [f"BT /F1 12 Tf {x} {y} Td ({t}) Tj ET" for x, y, t in lines]
    stream = DecodedStreamObject()
    stream.set_data("\n".join(ops).encode("latin-1"))
    page[NameObject("/Contents")] = w._add_object(stream)
    font = DictionaryObject({NameObject("/Type"): NameObject("/Font"),
                             NameObject("/Subtype"): NameObject("/Type1"),
                             NameObject("/BaseFont"): NameObject("/Helvetica")})
    if font_widths is not None:
        first, widths = font_widths
        font[NameObject("/FirstChar")] = NumberObject(first)
        font[NameObject("/LastChar")] = NumberObject(first + len(widths) - 1)
        font[NameObject("/Widths")] = ArrayObject(
            [NumberObject(v) for v in widths])
    page[NameObject("/Resources")] = DictionaryObject({NameObject("/Font"):
        DictionaryObject({NameObject("/F1"): w._add_object(font)})})
    arr = ArrayObject()
    for spec in annots:
        d = DictionaryObject({
            NameObject("/Type"): NameObject("/Annot"),
            NameObject("/Subtype"): NameObject(spec["subtype"]),
            NameObject("/Rect"): ArrayObject(
                [FloatObject(v) for v in spec["rect"]]),
            NameObject("/Contents"): TextStringObject(spec["contents"])})
        if "quad" in spec:
            d[NameObject("/QuadPoints")] = ArrayObject(
                [FloatObject(v) for v in spec["quad"]])
        arr.append(w._add_object(d))
    page[NameObject("/Annots")] = arr
    with path.open("wb") as fh:
        w.write(fh)
    return path


# A proof with a highlight (over "fish oil.") and a sticky note (on the tobacco
# line, dropped just below its baseline as real notes are).
def _proof(path: Path) -> Path:
    return make_commented_pdf(path,
        lines=[(72, 700, "were slick with "), (150, 700, "fish oil."),
               (72, 680, "he carried a pouch of tobacco here")],
        annots=[
            {"subtype": "/Highlight", "rect": [150, 698, 195, 712],
             "quad": [150, 712, 195, 712, 150, 698, 195, 698],
             "contents": "Slick with petroleum jelly"},
            {"subtype": "/Text", "rect": [72, 677, 90, 695],
             "contents": "Replace tobacco with candlestick"}])


def story_text(idml: Path, story_id: str = "ue0") -> list[str]:
    with zipfile.ZipFile(idml) as z:
        s = parse_story(z.read(f"Stories/Story_{story_id}.xml"), story_id)
    return [p.text for p in s.paragraphs]


# --- reading a redlined Word file ---------------------------------------------

def test_a_replacement_becomes_one_anchored_edit(tmp_path):
    doc = make_tracked_docx(tmp_path / "r.docx", [
        [("del", "Their"), ("ins", "There"), ("", " were three ships.")]])
    result = edits_from_docx(doc)
    assert result.ok and len(result.edits) == 1
    e = result.edits[0]
    # The anchor is widened past the bare "ir"/"re" so it locates uniquely.
    assert "Their" in e.find and "There" in e.replace
    assert e.kind == MECHANICAL


def test_a_deletion_removes_the_phrase(tmp_path):
    """A mid-paragraph deletion anchors with context on both sides, so it is a
    replacement whose net effect drops the phrase — the removed words are in the
    find and gone from the replace."""
    doc = make_tracked_docx(tmp_path / "d.docx", [
        [("", "A third paragraph with plain text"),
         ("del", " for good measure"), ("", ".")]])
    result = edits_from_docx(doc)
    assert result.ok and len(result.edits) == 1
    e = result.edits[0]
    assert "for good measure" in e.find and "for good measure" not in e.replace


def test_unchanged_paragraphs_produce_no_edits(tmp_path):
    doc = make_tracked_docx(tmp_path / "u.docx", [
        [("", "Nothing was touched here.")],
        [("del", "Their"), ("ins", "There"), ("", " it is.")]])
    result = edits_from_docx(doc)
    assert len(result.edits) == 1        # only the changed paragraph


def test_a_wholly_inserted_paragraph_is_flagged_not_guessed(tmp_path):
    doc = make_tracked_docx(tmp_path / "i.docx", [
        [("ins", "A brand new paragraph the author added.")]])
    result = edits_from_docx(doc)
    assert not result.edits
    assert result.issues and "anchor" in result.issues[0].reason


def test_a_redline_applies_to_the_real_indesign_file(tmp_path):
    """End to end: a Word redline of the book's text extracts to an edit that
    anchors and applies to the InDesign fixture, and verifies clean."""
    # layout.idml paragraph 4 reads "Their were several mistakes here to find."
    doc = make_tracked_docx(tmp_path / "book.docx", [
        [("del", "Their"), ("ins", "There"),
         ("", " were several mistakes here to find.")]])
    result = edits_from_docx(doc)
    assert result.ok and len(result.edits) == 1

    out = tmp_path / "out.idml"
    report = apply_edits(LAYOUT, out, list(result.edits))
    assert report.applied == 1
    assert story_text(out)[4] == "There were several mistakes here to find."


# --- reading a free-form list with a model ------------------------------------

def _canned(edits: list[dict]) -> FakeProvider:
    return FakeProvider([ProviderResult(
        parsed={"edits": edits},
        usage=NormalizedUsage(input_tokens=200, output_tokens=40))])


def test_a_prose_list_is_extracted_into_edits():
    provider = _canned([
        {"find": "Their were", "replace": "There were", "instruction": "",
         "kind": "mechanical", "occurrence": 0},
        {"find": "Chapter One", "replace": "", "instruction": "move to a recto",
         "kind": "design", "occurrence": 0}])
    usage = Usage()
    result = extract_edits("...an author's notes...", provider,
                           model="claude-haiku-4-5", usage=usage)
    assert result.ok and len(result.edits) == 2
    assert result.edits[0].find == "Their were"
    assert result.edits[1].kind == DESIGN
    # The source text and the corrections schema went to the model, and its
    # tokens were counted.
    assert provider.calls[0]["user"].startswith("...an author")
    assert usage.output_tokens == 40


def test_an_extracted_context_scopes_the_edit(tmp_path):
    """The model can return a context anchor for a repeated word; it flows
    through to the edit and lands it on the one place the anchor names."""
    provider = _canned([
        {"find": "runs on", "replace": "rambles on", "context": "it also runs on",
         "instruction": "", "kind": "mechanical", "occurrence": 0}])
    result = extract_edits("a comment on the footnote line", provider,
                           model="m", usage=Usage())
    assert result.edits[0].context == "it also runs on"
    out = tmp_path / "out.idml"
    report = apply_edits(LAYOUT, out, list(result.edits))
    assert report.applied == 1
    assert "rambles on" in story_text(out)[5]                 # footnote, not table
    assert story_text(out, "uff")[0].endswith("it runs on.")


def test_the_extractor_reuses_the_parse_contract_for_bad_entries():
    """A model that returns an unanchorable entry has it flagged, not dropped —
    the same 'flag, never guess' contract a typed list gets."""
    provider = _canned([
        {"find": "Their were", "replace": "There were", "instruction": "",
         "kind": "mechanical", "occurrence": 0},
        {"find": "", "replace": "orphan", "instruction": "", "kind": "mechanical",
         "occurrence": 0}])
    result = extract_edits("notes", provider, model="m", usage=Usage())
    assert len(result.edits) == 1
    assert result.issues and "no find text" in result.issues[0].reason


def test_a_refusal_or_truncation_raises():
    provider = FakeProvider([ProviderResult(parsed=None, stop_reason="refusal")])
    with pytest.raises(ExtractionError):
        extract_edits("notes", provider, model="m", usage=Usage())


def test_an_empty_source_makes_no_model_call():
    provider = FakeProvider([])          # would raise if a call popped nothing
    result = extract_edits("   ", provider, model="m", usage=Usage())
    assert result.ok and not result.edits
    assert not provider.calls


def test_extracted_edits_apply_to_the_real_indesign_file(tmp_path):
    provider = _canned([
        {"find": "Their were", "replace": "There were", "instruction": "",
         "kind": "mechanical", "occurrence": 0}])
    result = extract_edits("change 'Their were' to 'There were'", provider,
                           model="m", usage=Usage())
    out = tmp_path / "out.idml"
    report = apply_edits(LAYOUT, out, list(result.edits))
    assert report.applied == 1
    assert story_text(out)[4] == "There were several mistakes here to find."


# --- reading a commented PDF proof --------------------------------------------

def test_a_highlight_is_read_with_its_span_and_line(tmp_path):
    from docproof.corrections.from_pdf import read_pdf_comments
    comments = read_pdf_comments(_proof(tmp_path / "p.pdf"))
    hl = [c for c in comments if c.kind == "highlight"]
    assert len(hl) == 1
    assert hl[0].instruction == "Slick with petroleum jelly"
    assert "fish oil" in hl[0].anchor                 # the highlighted span
    assert "slick with fish oil" in hl[0].context     # the whole line for context


def test_a_highlight_on_a_word_inside_a_run_reads_that_word(tmp_path):
    """A highlight over a word that is not first in its text run. pypdf hands back
    a whole run at a single position, so this once read back empty — the run's
    start x sat left of the mark. The reader works from pdfplumber's exact
    per-word boxes now, so an interior word is found and comes back whole.

    At Helvetica 12pt "charlie" occupies x 138–173; the mark covers its middle and
    must come back as the whole word."""
    from docproof.corrections.from_pdf import read_pdf_comments
    pdf = make_commented_pdf(
        tmp_path / "run.pdf",
        lines=[(72, 700, "alpha bravo charlie delta echo"),
               (72, 680, "cal"), (87, 680, "ibrate")],   # 2 chunks -> scale 0.01
        annots=[{"subtype": "/Highlight",
                 "rect": [134, 698, 165, 712],
                 "quad": [134, 712, 165, 712, 134, 698, 165, 698],
                 "contents": "swap this word"}],
        font_widths=(32, [500] * 96))                     # codes 32..127
    hl = [c for c in read_pdf_comments(pdf) if c.kind == "highlight"]
    assert len(hl) == 1
    assert hl[0].anchor == "charlie"                      # whole interior word
    assert "alpha" not in hl[0].anchor                    # not the whole run
    assert "alpha bravo charlie delta echo" in hl[0].context


def test_a_note_is_anchored_to_the_line_it_sits_on(tmp_path):
    """The note is dropped just below the tobacco line's baseline; it must map to
    that line, not the (empty) space below it."""
    from docproof.corrections.from_pdf import read_pdf_comments
    comments = read_pdf_comments(_proof(tmp_path / "p.pdf"))
    note = [c for c in comments if c.kind == "note"]
    assert len(note) == 1
    assert note[0].instruction == "Replace tobacco with candlestick"
    assert "tobacco" in note[0].anchor


@pytest.mark.parametrize("rect_y", [
    (651, 671),      # a ~20pt icon box, its top at the click point
    (662, 662),      # a zero-height Rect — the anchor lands INSIDE the line's band
    (666, 666),      # the same, high on the line
    (659, 659),      # the same, on the baseline
    (658, 670),      # a box the height of the line itself
])
def test_a_note_lands_on_its_own_line_whatever_box_the_tool_wrote(tmp_path, rect_y):
    """The same click, stored five ways, must reach the same line.

    A sticky note's `/Rect` is not a reliable shape: the viewer draws a fixed-size
    icon and ignores the stored one, so tools write anything from a 20pt box to a
    zero-height point. The reader used to ask only whether a line's *bottom* edge
    was above the anchor, which meant a short box put the anchor inside the band of
    the line it marked — and that line then failed the test against its own bottom
    edge, so the line ABOVE was returned. Every note stored that way reached the
    extractor quoting the line before the one it was written about."""
    from docproof.corrections.from_pdf import read_pdf_comments
    lines = [(72, 700 - 20 * i, f"line {100 + i} of the proof here")
             for i in range(5)]
    pdf = make_commented_pdf(
        tmp_path / "n.pdf", lines,
        [{"subtype": "/Text", "rect": [72, rect_y[0], 90, rect_y[1]],
          "contents": "fix this"}])
    note = read_pdf_comments(pdf)[0]
    assert "line 102" in note.anchor          # the line the click was on


def test_a_note_in_the_leading_belongs_to_the_line_above_it(tmp_path):
    """The rule the box-overlap test replaces, still doing its job where it should:
    an icon that sits wholly in the gap between two lines hangs under the one it
    marks, so it belongs to the line above, not the one below."""
    from docproof.corrections.from_pdf import read_pdf_comments
    lines = [(72, 700 - 20 * i, f"line {100 + i} of the proof here")
             for i in range(3)]
    pdf = make_commented_pdf(
        tmp_path / "g.pdf", lines,
        [{"subtype": "/Text", "rect": [72, 692, 90, 695],
          "contents": "fix this"}])
    assert "line 100" in read_pdf_comments(pdf)[0].anchor


def test_a_pdf_with_no_comments_yields_nothing(tmp_path):
    from docproof.corrections.from_pdf import (pdf_corrections_source,
                                               read_pdf_comments)
    plain = make_commented_pdf(tmp_path / "plain.pdf",
                               lines=[(72, 700, "Just some text, no comments.")],
                               annots=[])
    assert read_pdf_comments(plain) == []
    source, comments = pdf_corrections_source(plain)
    assert source == "" and comments == []


def test_the_pdf_source_feeds_the_extractor_end_to_end(tmp_path):
    """The reader's output is exactly what the model extractor consumes: read a
    proof, hand its source to a (fake) model, get edits back."""
    from docproof.corrections.from_pdf import pdf_corrections_source
    source, comments = pdf_corrections_source(_proof(tmp_path / "p.pdf"))
    assert len(comments) == 2 and "petroleum jelly" in source
    provider = _canned([
        {"find": "fish oil", "replace": "petroleum jelly", "instruction": "",
         "kind": "mechanical", "occurrence": 0},
        {"find": "tobacco", "replace": "candlestick", "instruction": "",
         "kind": "mechanical", "occurrence": 0}])
    result = extract_edits(source, provider, model="m", usage=Usage())
    assert result.ok and len(result.edits) == 2


def test_a_highlight_on_a_single_chunk_line_is_read_exactly(tmp_path):
    """The failure that cost a real proof half its anchors.

    A line of typeset body text is drawn as one run. The old reader estimated a
    page's points-per-em from the gaps *between* runs, so a single-run line offered
    it nothing to measure and it fell back to a constant — roughly half the truth
    here. Every glyph was then placed at half its real advance, packed into the
    left of the measure, and a highlight over a word late in the line covered no
    glyph at all: the comment arrived with an empty anchor and the model had
    nothing to work from.

    At Helvetica 12pt the last word of this line, "juliet", renders at x 353–377.
    The mark is drawn exactly there; the old reader placed that glyph run around
    x 72–258 and so covered none of it.
    """
    from docproof.corrections.from_pdf import read_pdf_comments
    line = "alpha bravo charlie delta echo foxtrot golf hotel india juliet"
    pdf = make_commented_pdf(
        tmp_path / "single.pdf",
        lines=[(72, 700, line)],                          # ONE run, no gap to measure
        annots=[{"subtype": "/Highlight",
                 "rect": [353, 698, 377, 712],
                 "quad": [353, 712, 377, 712, 353, 698, 377, 698],
                 "contents": "Previously capitalized"}])
    hl = [c for c in read_pdf_comments(pdf) if c.kind == "highlight"]
    assert len(hl) == 1
    assert hl[0].anchor == "juliet"
    assert hl[0].context == line


def test_a_highlight_across_a_line_break_reads_as_one_span(tmp_path):
    """A mark that runs over the end of one line and onto the next comes back whole,
    with both lines as its context."""
    from docproof.corrections.from_pdf import read_pdf_comments
    pdf = make_commented_pdf(
        tmp_path / "wrap.pdf",
        lines=[(72, 700, "he won the player of the"), (72, 680, "tournament that year")],
        annots=[{"subtype": "/Highlight",
                 "rect": [72, 678, 201, 710],
                 # one rectangle per line, as a real wrapped highlight is drawn:
                 # "the" at x 183–200 on the upper line, "tournament" at x 72–133
                 # on the lower one
                 "quad": [183, 710, 201, 710, 183, 697, 201, 697,
                          72, 690, 133, 690, 72, 677, 133, 677],
                 "contents": "Previously capitalized"}])
    hl = [c for c in read_pdf_comments(pdf) if c.kind == "highlight"]
    assert len(hl) == 1
    # The line break collapses to a space: the book sets these words as one run,
    # so that is the quotation an edit has to carry.
    assert hl[0].anchor == "the tournament"
    assert "player of the" in hl[0].context
    assert "that year" in hl[0].context


def test_the_proof_hands_back_the_text_of_every_page(tmp_path):
    """The page texts are what the page map aligns against the book, so they come
    back from the same read — including for a page that carries no comment."""
    from docproof.corrections.from_pdf import read_pdf
    proof = read_pdf(_proof(tmp_path / "p.pdf"))
    assert len(proof.page_texts) == 1
    assert proof.page_texts[0] == ("were slick with fish oil.\n"
                                   "he carried a pouch of tobacco here")
    assert len(proof.comments) == 2


def test_a_comment_records_where_it_sat_on_its_page(tmp_path):
    """The offset into the page text is the closest thing a page-anchored mark has
    to an exact address, so it is carried rather than recomputed later."""
    from docproof.corrections.from_pdf import read_pdf
    proof = read_pdf(_proof(tmp_path / "p.pdf"))
    hl = [c for c in proof.comments if c.kind == "highlight"][0]
    page = proof.page_texts[0]
    assert hl.offset >= 0
    assert page[hl.offset:].startswith(hl.anchor)


def test_a_highlight_over_no_words_falls_back_to_its_line(tmp_path):
    """A mark dragged over blank space (or an image) resolves to no word. It still
    has to arrive anchored to something real, or the comment reaches the model with
    nothing at all — which is how a mark used to become an unusable edit."""
    from docproof.corrections.from_pdf import read_pdf_comments
    pdf = make_commented_pdf(
        tmp_path / "blank.pdf",
        lines=[(72, 700, "the line that carries the mark")],
        annots=[{"subtype": "/Highlight",
                 # to the right of where the text ends, on the same line
                 "rect": [400, 698, 460, 712],
                 "quad": [400, 712, 460, 712, 400, 698, 460, 698],
                 "contents": "Remove blank line break?"}])
    hl = [c for c in read_pdf_comments(pdf) if c.kind == "highlight"]
    assert len(hl) == 1
    assert hl[0].anchor == "the line that carries the mark"
