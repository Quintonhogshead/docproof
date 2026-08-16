"""Soft line breaks and tabs ARE characters of the canonical text.

The failure this file pins: a Shift+Enter inside a paragraph (w:br) rendered
as nothing, so "The Ripple Effect" over "QX Countdown" read as
"...EffectQX...". The fused token then poisoned everything word-shaped — the
spell scan harvested "EffectQX" into the protected lexicon, and the model was
shown a run-on no author wrote and proposed text to glue it back together.

The contract now: w:br and w:cr render as "\n", an in-run w:tab as "\t", in
paragraph_text AND in every offset the reassembler edits by AND in both audit
views. One source (iter_content_elements), so they cannot drift.
"""
from __future__ import annotations

import docx
from lxml import etree

from docproof.config import Config
from docproof.ingest import build_document_model, preflight
from docproof.models import Anchor, Finding
from docproof.normalize import _apply_untracked, _quote_edits
from docproof.reassembler import apply_tracked_changes, paragraph_view_text
from docproof.spellscan import scan
from docproof.utils.xml_helpers import paragraph_text, walk_package

_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _w(t):
    return f"{{{_W}}}{t}"


def _p_with(*runs) -> etree._Element:
    """A bare w:p whose runs hold the given content: a str becomes w:t,
    "BR"/"TAB" the corresponding empty element."""
    p = etree.Element(_w("p"))
    r = etree.SubElement(p, _w("r"))
    for item in runs:
        if item == "BR":
            etree.SubElement(r, _w("br"))
        elif item == "TAB":
            etree.SubElement(r, _w("tab"))
        else:
            etree.SubElement(r, _w("t")).text = item
    return p


def test_breaks_and_tabs_render_as_their_characters():
    p = _p_with("The Ripple Effect", "BR", "QX Countdown", "TAB", "page 9")
    assert paragraph_text(p) == "The Ripple Effect\nQX Countdown\tpage 9"


def test_tab_stop_definitions_are_not_characters():
    """w:tab under w:pPr/w:tabs positions text; only an in-run w:tab IS text."""
    p = _p_with("one", "TAB", "two")
    ppr = etree.Element(_w("pPr"))
    tabs = etree.SubElement(ppr, _w("tabs"))
    etree.SubElement(tabs, _w("tab"))            # a tab STOP, not a character
    p.insert(0, ppr)
    assert paragraph_text(p) == "one\ttwo"


def test_spell_scan_no_longer_harvests_fused_line_junk():
    """The exact DP-007 artifact: without the break, "EffectQX" is an unknown
    name-cased word and lands in the protected lexicon."""
    from docproof.models import ParagraphRef

    text = ("The Ripple Effect\nQX Countdown\n"
            "Home\nSteve finds the letter and reads it slowly.")
    para = ParagraphRef(para_id="body-0000", part="word/document.xml",
                        location="body", text=text, style="Normal",
                        reviewable=True)
    result = scan([para])
    assert "EffectQX" not in result.lexicon
    assert "HomeSteve" not in result.lexicon
    fused = [w for w in result.lexicon if "QX" in w and w != "QX"]
    assert not fused


def _breaks_docx(tmp_path):
    d = docx.Document()
    p = d.add_paragraph("The Ripple Effect")
    p.runs[0].add_break()                        # w:br — a Shift+Enter
    p.add_run("QX Countdown begins here.")
    d.add_paragraph("A second paragraph, long enough to review.")
    path = tmp_path / "breaks.docx"
    d.save(path)
    return path


def _first_para(doc):
    return next(p for p in doc.paragraphs if "Ripple" in p.text)


def test_ingest_carries_the_break(tmp_path):
    pkg = preflight(_breaks_docx(tmp_path), "abort")
    doc = build_document_model(pkg, Config())
    para = _first_para(doc)
    assert para.text == "The Ripple Effect\nQX Countdown begins here."


def _finding(pid, para_text, start, end, insert, i=1):
    return Finding(f"f-{i}", "chunk-000", pid, "spelling", para_text, 1, "",
                   "test", "high", status="validated",
                   anchor=Anchor(start=start, end=end,
                                 delete_text=para_text[start:end],
                                 insert_text=insert))


def test_edit_after_a_break_lands_where_it_says(tmp_path):
    """The regression the offset map exists for: before the fix, every offset
    past a break was one character early."""
    pkg = preflight(_breaks_docx(tmp_path), "abort")
    doc = build_document_model(pkg, Config())
    para = _first_para(doc)
    start = para.text.index("begins")
    f = _finding(para.para_id, para.text, start, start + len("begins"), "began")

    stats = apply_tracked_changes(pkg, doc, [f], Config())
    assert stats.applied == (f.finding_id,)
    elem = {wp.para_id: wp.element for wp in walk_package(pkg)}[para.para_id]
    assert paragraph_view_text(elem, "accept") == \
        "The Ripple Effect\nQX Countdown began here."
    assert paragraph_view_text(elem, "reject") == para.text


def test_deleting_a_break_is_a_tracked_change_that_rejects_clean(tmp_path):
    """Joining two soft-broken lines: the deletion span IS the break."""
    pkg = preflight(_breaks_docx(tmp_path), "abort")
    doc = build_document_model(pkg, Config())
    para = _first_para(doc)
    nl = para.text.index("\n")
    f = _finding(para.para_id, para.text, nl, nl + 1, " ")

    stats = apply_tracked_changes(pkg, doc, [f], Config())
    assert stats.applied == (f.finding_id,)
    elem = {wp.para_id: wp.element for wp in walk_package(pkg)}[para.para_id]
    assert paragraph_view_text(elem, "accept") == \
        "The Ripple Effect QX Countdown begins here."
    assert paragraph_view_text(elem, "reject") == para.text


def test_inserted_newline_is_written_as_a_real_break(tmp_path):
    pkg = preflight(_breaks_docx(tmp_path), "abort")
    doc = build_document_model(pkg, Config())
    para = next(p for p in doc.paragraphs if "second" in p.text)
    at = para.text.index("long")
    f = _finding(para.para_id, para.text, at, at, "very\n")

    stats = apply_tracked_changes(pkg, doc, [f], Config())
    assert stats.applied == (f.finding_id,)
    elem = {wp.para_id: wp.element for wp in walk_package(pkg)}[para.para_id]
    accepted = paragraph_view_text(elem, "accept")
    assert "very\nlong" in accepted
    assert len(elem.findall(f".//{_w('ins')}/{_w('r')}/{_w('br')}")) == 1
    assert paragraph_view_text(elem, "reject") == para.text


def test_normalize_offsets_stay_aligned_past_a_break():
    """A straight quote after a soft break curls as an opener — and lands on
    the quote, not one character early."""
    p = _p_with('He looked up.', "BR", '"Morning," she said.')
    text = paragraph_text(p)
    edits, _ambiguous = _quote_edits(text, single_primary=False)
    _apply_untracked(p, edits)
    assert paragraph_text(p) == "He looked up.\n“Morning,” she said."
