"""The seed -> review -> score loop, closed end-to-end (deliverable 5):
`galley seed` now writes a real, reviewable .docx alongside its JSON, so a
plain `docproof review` can consume it and `galley score` can grade the
result — with no API call anywhere in the chain (the review step uses
--mock-findings, standing in for a real fleet's output).
"""
from __future__ import annotations

import json

import docx

from docproof.__main__ import main
from docproof.config import load_config
from docproof.ingest import build_document_model, preflight

PARAGRAPHS = [
    "The lamp on the desk flickered, then finally caught, and Maria sighed.",
    "Its light was thin, but it was enough to read the letter by tonight.",
    "David counted the coins twice, sure the total would not change again.",
    "The garden gate creaked open, and the old dog limped out to greet her.",
    "Their car idled at the curb while the rain fell in long grey ropes.",
    "The captain said the storm would pass before it reached the harbor.",
]


def _manuscript(tmp_path):
    d = docx.Document()
    for p in PARAGRAPHS:
        d.add_paragraph(p)
    src = tmp_path / "book.docx"
    d.save(src)
    return src


def test_seed_writes_a_reviewable_docx(tmp_path):
    src = _manuscript(tmp_path)
    seed_dir = tmp_path / "seed"
    rc = main(["galley", "seed", str(src), "-n", "3", "--seed", "0",
              "--out", str(seed_dir), "--json"])
    assert rc == 0
    seeded_docx = seed_dir / "seeded_manuscript.docx"
    assert seeded_docx.exists()

    # It really is a plain, reviewable .docx: re-ingest it the way `docproof
    # review` would and check the planted mutations actually landed in the
    # paragraph text docproof itself would read.
    key = json.loads((seed_dir / "answer_key.json").read_text("utf-8"))
    assert key["planted"], "nothing planted — nothing to verify"
    cfg = load_config("config/default.yaml")
    pkg = preflight(str(seeded_docx), cfg.tracked_changes_policy)
    doc = build_document_model(pkg, cfg)
    seeded_text = {p.para_id: p.text for p in doc.paragraphs}
    for pe in key["planted"]:
        assert pe["mutated"] in seeded_text[pe["para_id"]] or not pe["mutated"]


def test_seed_review_score_loop_catches_every_planted_error(tmp_path):
    src = _manuscript(tmp_path)
    seed_dir = tmp_path / "seed"
    rc = main(["galley", "seed", str(src), "-n", "3", "--seed", "0",
              "--out", str(seed_dir)])
    assert rc == 0

    seeded = json.loads((seed_dir / "seeded_manuscript.json").read_text("utf-8"))
    key = json.loads((seed_dir / "answer_key.json").read_text("utf-8"))
    assert key["planted"]

    # The original (pre-seed) paragraph text, read the same way docproof
    # itself would — this is what a "catch" fixes the seeded text back to.
    cfg = load_config("config/default.yaml")
    pkg = preflight(str(src), cfg.tracked_changes_policy)
    doc = build_document_model(pkg, cfg)
    original_text = {p.para_id: p.text for p in doc.paragraphs}

    mocks = [
        {"para_id": pe["para_id"],
         "original_text": seeded["paragraphs"][pe["para_id"]],
         "corrected_text": original_text[pe["para_id"]],
         "confidence": "high"}
        for pe in key["planted"]
    ]
    mocks_path = tmp_path / "mocks.json"
    mocks_path.write_text(json.dumps(mocks), encoding="utf-8")

    review_dir = tmp_path / "review"
    rc = main(["review", str(seed_dir / "seeded_manuscript.docx"),
              "--error-types", "spelling", "--mock-findings", str(mocks_path),
              "--out", str(review_dir)])
    assert rc == 0
    findings_json = review_dir / "findings.json"
    assert findings_json.exists()

    payload = json.loads(findings_json.read_text("utf-8"))
    validated_paras = {f["para_id"] for f in payload["findings"]
                       if f["status"] == "validated"}
    planted_paras = {pe["para_id"] for pe in key["planted"]}
    assert planted_paras <= validated_paras, (
        "every planted error's paragraph must carry a validated fix")

    # galley score expects GFindings; the docproof-ladder adapter is the
    # existing converter from a findings.json to that shape.
    from galley.adapters.docproof_ladder import gfindings_from_json
    gfindings, dropped = gfindings_from_json(findings_json, wave=1, model="mock")
    assert dropped == 0
    gfindings_path = tmp_path / "gfindings.json"
    gfindings_path.write_text(
        json.dumps({"findings": [g.to_json() for g in gfindings]}),
        encoding="utf-8")

    score_dir = tmp_path / "score"
    rc = main(["galley", "score", str(gfindings_path), "--answer-key",
              str(seed_dir / "answer_key.json"), "--out", str(score_dir),
              "--json"])
    assert rc == 0
    recall = json.loads((score_dir / "recall.json").read_text("utf-8"))
    assert recall["caught"] == recall["planted"] == len(key["planted"])
    assert recall["rate"] == 1.0
