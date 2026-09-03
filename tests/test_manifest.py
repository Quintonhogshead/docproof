"""Approval manifest, model-route visibility, and the delivery certificate
(galley/manifest.py)."""
from __future__ import annotations

import json
from pathlib import Path

import pytest

from docproof.genre import materialize_genre_pack
from galley.manifest import (build_manifest, certify_run, config_hash,
                             model_routes, providers_in_use, verify_plan)

CONFIG = Path(__file__).parent.parent / "config" / "default.yaml"


def _mech_cfg():
    cfg, _ = materialize_genre_pack(CONFIG, "general_fiction",
                                    stage="mechanical-wave")
    return cfg


def _source(tmp_path) -> Path:
    p = tmp_path / "book.docx"
    p.write_bytes(b"fake manuscript bytes")
    return p


# --- routes ------------------------------------------------------------------

def test_model_routes_report_active_ensemble_and_provider():
    routes = model_routes(_mech_cfg())
    active = {r.role: r for r in routes if r.active}
    assert "ensemble.detector[gpt-5.6-luna]" in active
    assert active["ensemble.detector[gpt-5.6-luna]"].provider == "openai"
    assert any(r.role == "ensemble.verifier" and r.active for r in routes)


def test_providers_in_use_is_the_active_provider_set():
    assert set(providers_in_use(_mech_cfg())) == {"anthropic", "openai"}


def test_disabled_lane_routes_are_marked_inactive():
    routes = {r.role: r for r in model_routes(_mech_cfg())}
    # The mechanical-wave stage turns the two gates ON (it IS the "both gates"
    # recipe) -> active routes; factcheck stays off -> inactive route.
    if "meaning_check.model" in routes:
        assert routes["meaning_check.model"].active is True
    if "factcheck.model" in routes:
        assert routes["factcheck.model"].active is False


# --- manifest + verify -------------------------------------------------------

def test_build_manifest_captures_hashes_models_and_lanes(tmp_path):
    cfg = _mech_cfg()
    src = _source(tmp_path)
    m = build_manifest(source=src, config_path=str(CONFIG), cfg=cfg,
                       max_spend_usd=20.0, stage="mechanical-wave")
    assert m["source_sha256"] and m["config_sha256"] == config_hash(cfg)
    assert "gpt-5.6-luna" in m["allowed_models"]
    assert set(m["allowed_providers"]) == {"anthropic", "openai"}
    assert "ensemble" in m["enabled_lanes"]
    assert m["max_spend_usd"] == 20.0


def test_verify_plan_clean_run_has_no_deviations(tmp_path):
    cfg = _mech_cfg()
    src = _source(tmp_path)
    m = build_manifest(source=src, config_path=str(CONFIG), cfg=cfg,
                       max_spend_usd=20.0)
    assert verify_plan(m, source=src, cfg=cfg) == []


def test_verify_plan_flags_a_changed_manuscript(tmp_path):
    cfg = _mech_cfg()
    src = _source(tmp_path)
    m = build_manifest(source=src, config_path=str(CONFIG), cfg=cfg,
                       max_spend_usd=20.0)
    src.write_bytes(b"a different manuscript")
    kinds = {d.kind for d in verify_plan(m, source=src, cfg=cfg)}
    assert "source_changed" in kinds


def test_verify_plan_flags_a_model_outside_the_approved_set(tmp_path):
    cfg = _mech_cfg()
    src = _source(tmp_path)
    m = build_manifest(source=src, config_path=str(CONFIG), cfg=cfg,
                       max_spend_usd=20.0)
    cfg.api.model = "claude-opus-5"
    kinds = {d.kind for d in verify_plan(m, source=src, cfg=cfg)}
    assert "model_not_approved" in kinds


def test_verify_plan_flags_budget_over_cap(tmp_path):
    cfg = _mech_cfg()
    src = _source(tmp_path)
    m = build_manifest(source=src, config_path=str(CONFIG), cfg=cfg,
                       max_spend_usd=20.0)
    kinds = {d.kind for d in verify_plan(m, source=src, cfg=cfg,
                                         budget_usd=25.0)}
    assert "budget_over_cap" in kinds


# --- certify -----------------------------------------------------------------

def _run_dir(tmp_path, envelope: dict) -> Path:
    run = tmp_path / "run"
    run.mkdir()
    (run / "findings.json").write_text(json.dumps(envelope))
    return run


def test_certify_passes_a_clean_run(tmp_path):
    cfg = _mech_cfg()
    src = _source(tmp_path)
    m = build_manifest(source=src, config_path=str(CONFIG), cfg=cfg,
                       max_spend_usd=20.0)
    run = _run_dir(tmp_path, {"findings": [{"a": 1}],
                              "cost": {"total_usd": 4.2},
                              "checkpoint": {"x": 1}})
    cert = certify_run(run, manifest=m, cfg=cfg, source=src)
    assert cert.passed
    assert {c.name for c in cert.checks} >= {
        "source hash", "config hash", "approved model routes",
        "budget within approval", "artifact scan"}


def test_certify_fails_zero_cost_anomaly(tmp_path):
    run = _run_dir(tmp_path, {"findings": [{"a": 1}, {"b": 2}],
                              "cost": {"total_usd": 0.0}})
    cert = certify_run(run)
    anomaly = [c for c in cert.checks if c.name == "zero-cost anomaly"]
    assert anomaly and anomaly[0].status == "fail"
    assert not cert.passed


def test_certify_fails_budget_over_cap(tmp_path):
    cfg = _mech_cfg()
    src = _source(tmp_path)
    m = build_manifest(source=src, config_path=str(CONFIG), cfg=cfg,
                       max_spend_usd=3.0)
    run = _run_dir(tmp_path, {"findings": [{"a": 1}],
                              "cost": {"total_usd": 4.2},
                              "checkpoint": {"x": 1}})
    cert = certify_run(run, manifest=m, cfg=cfg, source=src)
    budget = [c for c in cert.checks if c.name == "budget within approval"]
    assert budget and budget[0].status == "fail"


def test_certify_fails_paid_run_missing_checkpoint(tmp_path):
    run = _run_dir(tmp_path, {"findings": [{"a": 1}],
                              "cost": {"total_usd": 4.2}})  # no checkpoint
    cert = certify_run(run)
    ckpt = [c for c in cert.checks if c.name == "checkpoint present"]
    assert ckpt and ckpt[0].status == "fail"


def test_certify_flags_a_merge_artifact_in_corrected_text(tmp_path):
    run = tmp_path / "run"
    run.mkdir()
    (run / "findings.json").write_text(json.dumps(
        {"findings": [{"finding_id": "f-0001", "para_id": "body-0001",
                       "original_text": "She paused then left.",
                       "corrected_text": "She paused,, then left."}],
         "cost": {"total_usd": 0.0}}))
    cert = certify_run(run)
    scan = [c for c in cert.checks if c.name == "artifact scan"]
    assert scan and scan[0].status == "fail"


def test_certify_does_not_flag_artifacts_quoted_from_the_original(tmp_path):
    """The change log and findings faithfully quote pre-fix text; the artifacts
    IN those quotes are what the run fixed, not defects of the deliverable."""
    run = tmp_path / "run"
    run.mkdir()
    (run / "findings.json").write_text(json.dumps(
        {"findings": [{"finding_id": "f-0001", "para_id": "body-0001",
                       "original_text": "She paused,,  then left. ",
                       "corrected_text": "She paused, then left."}],
         "cost": {"total_usd": 0.0}}))
    (run / "change_log.md").write_text("> She paused,,  then left. ")
    cert = certify_run(run)
    scan = [c for c in cert.checks if c.name == "artifact scan"]
    assert scan and scan[0].status == "pass"


def test_import_judgments_run_is_not_a_zero_cost_anomaly(tmp_path):
    """A model-free import legitimately costs $0 with findings — it must not
    trip the zero-cost anomaly check."""
    run = _run_dir(tmp_path, {"findings": [{"a": 1}],
                              "cost": {"total_usd": 0.0},
                              "judge_model": "external:import-judgments"})
    cert = certify_run(run)
    anomaly = [c for c in cert.checks if c.name == "zero-cost anomaly"]
    assert anomaly and anomaly[0].status == "pass"


# --- delivered text hygiene ---------------------------------------------------

def _docx_in(run, name, *texts):
    import docx as _docx
    d = _docx.Document()
    for t in texts:
        d.add_paragraph(t)
    d.save(run / name)


def test_certify_text_hygiene_skips_without_a_docx(tmp_path):
    run = _run_dir(tmp_path, {"findings": [], "cost": {"total_usd": 0.0}})
    cert = certify_run(run)
    check = next(c for c in cert.checks if c.name == "delivered text hygiene")
    assert check.status == "skip"


def test_certify_text_hygiene_fails_on_double_spaces(tmp_path):
    run = _run_dir(tmp_path, {"findings": [], "cost": {"total_usd": 0.0}})
    _docx_in(run, "book - Atmosphere Press Proofreader.docx",
             "A clean paragraph.", "Two  spaces survived here.",
             "Trailing spaces here.  ")
    cert = certify_run(run)
    check = next(c for c in cert.checks if c.name == "delivered text hygiene")
    assert check.status == "fail"
    assert "double space" in check.detail


def test_certify_text_hygiene_passes_clean_text_and_dividers(tmp_path):
    run = _run_dir(tmp_path, {"findings": [], "cost": {"total_usd": 0.0}})
    _docx_in(run, "book.docx", "A clean paragraph.", "*   *   *", "Another.")
    cert = certify_run(run)
    check = next(c for c in cert.checks if c.name == "delivered text hygiene")
    assert check.status == "pass"


# --- pre-existing vs. introduced hygiene faults (HARNESS_NOTES item 10) -----
#
# certify's canonical text strips paragraph-trailing whitespace, so a fault
# already in the SOURCE manuscript can never be expressed as a finding row —
# no lane could ever have fixed it. Without --source that is unprovable and
# every hit still fails (today's behaviour); with a matching --source, a
# hit that is byte-for-byte the same in the source's same paragraph passes
# instead, while a hit the run actually introduced still fails.

def _docx_with_internal_break(path, before, trailing_ws, after):
    """A single paragraph whose text has a mid-paragraph line break (w:br) —
    the Redding body-0734 shape: 'gone. \\n\\n', a pre-existing space before a
    paragraph-INTERNAL break, not the paragraph's own trailing whitespace."""
    import docx as _docx
    from docx.enum.text import WD_BREAK
    d = _docx.Document()
    p = d.add_paragraph()
    p.add_run(before + trailing_ws)
    p.add_run().add_break(WD_BREAK.LINE)
    if after:
        p.add_run(after)
    d.save(path)


def test_certify_text_hygiene_pre_existing_double_space_is_not_a_fail(tmp_path):
    run = _run_dir(tmp_path, {"findings": [], "cost": {"total_usd": 0.0}})
    _docx_in(run, "book.docx", "A clean paragraph.",
             "Two  spaces survived here.")
    _docx_in(tmp_path, "source.docx", "A clean paragraph.",
             "Two  spaces survived here.")
    cert = certify_run(run, source=tmp_path / "source.docx")
    check = next(c for c in cert.checks if c.name == "delivered text hygiene")
    assert check.status == "pass"
    assert "pre-existing" in check.detail


def test_certify_text_hygiene_introduced_double_space_still_fails(tmp_path):
    run = _run_dir(tmp_path, {"findings": [], "cost": {"total_usd": 0.0}})
    _docx_in(run, "book.docx", "A clean paragraph.",
             "Two  spaces survived here.")
    _docx_in(tmp_path, "source.docx", "A clean paragraph.",
             "Two spaces survived here.")     # clean in the source
    cert = certify_run(run, source=tmp_path / "source.docx")
    check = next(c for c in cert.checks if c.name == "delivered text hygiene")
    assert check.status == "fail"
    assert "double space" in check.detail


def test_certify_text_hygiene_pre_existing_internal_break_space_passes(tmp_path):
    """The exact Redding body-0734 shape: a pre-existing space before a
    paragraph-internal line break, byte-identical in the source."""
    run = _run_dir(tmp_path, {"findings": [], "cost": {"total_usd": 0.0}})
    _docx_with_internal_break(run / "book.docx", "It was gone.", " ",
                              "Next line.")
    _docx_with_internal_break(tmp_path / "source.docx", "It was gone.", " ",
                              "Next line.")
    cert = certify_run(run, source=tmp_path / "source.docx")
    check = next(c for c in cert.checks if c.name == "delivered text hygiene")
    assert check.status == "pass"
    assert "pre-existing" in check.detail


def test_certify_text_hygiene_introduced_internal_break_space_fails(tmp_path):
    """The same shape, but the source paragraph is clean — this run
    introduced the trailing space, so it still fails."""
    run = _run_dir(tmp_path, {"findings": [], "cost": {"total_usd": 0.0}})
    _docx_with_internal_break(run / "book.docx", "It was gone.", " ",
                              "Next line.")
    _docx_with_internal_break(tmp_path / "source.docx", "It was gone.", "",
                              "Next line.")
    cert = certify_run(run, source=tmp_path / "source.docx")
    check = next(c for c in cert.checks if c.name == "delivered text hygiene")
    assert check.status == "fail"


def test_certify_text_hygiene_without_source_still_fails_a_preexisting_fault(
        tmp_path):
    """No --source given means pre-existing can't be proven — every hit still
    fails, exactly like before this fix."""
    run = _run_dir(tmp_path, {"findings": [], "cost": {"total_usd": 0.0}})
    _docx_with_internal_break(run / "book.docx", "It was gone.", " ",
                              "Next line.")
    cert = certify_run(run)
    check = next(c for c in cert.checks if c.name == "delivered text hygiene")
    assert check.status == "fail"


# --- the `ran` flag from `galley verify` -------------------------------------

def test_certify_change_verify_honors_ran_false(tmp_path):
    """`galley verify --walk-only` writes change_verify.json with ran: false and
    no problems — that is a gate that never ran, never a clean read."""
    from galley.manifest import _certify_change_verify, _certify_finished_walk
    (tmp_path / "change_verify.json").write_text(json.dumps(
        {"ran": False, "applied_edits": 12, "problems": []}))
    check = _certify_change_verify(tmp_path)
    assert check.status == "skip"
    assert "--walk-only" in check.detail and "did not run" in check.detail
    (tmp_path / "finished_walk.json").write_text(json.dumps(
        {"ran": False, "residuals": []}))
    walk = _certify_finished_walk(tmp_path)
    assert walk.status == "skip"
    assert "--changes-only" in walk.detail and "did not run" in walk.detail


def test_certify_verify_gates_surface_a_recorded_reason(tmp_path):
    """verify_run that could read no accepted text records ran: false plus a
    reason; the skip repeats it instead of blaming a flag that was not used."""
    from galley.manifest import _certify_change_verify, _certify_finished_walk
    reason = "no accepted text could be read from the deliverable"
    for name in ("change_verify.json", "finished_walk.json"):
        (tmp_path / name).write_text(json.dumps(
            {"ran": False, "reason": reason, "problems": [], "residuals": []}))
    assert reason in _certify_change_verify(tmp_path).detail
    assert reason in _certify_finished_walk(tmp_path).detail


# --- a verify artifact older than the findings it claims to have read --------
#
# Redding Book 1 (2026-09-01): the deliverable was rebuilt after `galley verify`
# ran, and certify reported the PREVIOUS build's clean read over the new build's
# edits. A stale artifact is a skip, never a verdict.

def _clean_verify_records(run: Path) -> None:
    (run / "change_verify.json").write_text(json.dumps(
        {"ran": True, "applied_edits": 3, "problems": []}))
    (run / "finished_walk.json").write_text(json.dumps(
        {"ran": True, "residuals": []}))


def test_certify_skips_verify_artifacts_older_than_the_findings(tmp_path):
    import os
    from galley.manifest import _certify_change_verify, _certify_finished_walk
    _clean_verify_records(tmp_path)
    (tmp_path / "findings.json").write_text(json.dumps({"findings": []}))
    # The rebuild wrote findings.json after both verify artifacts.
    old = 1_700_000_000
    for name in ("change_verify.json", "finished_walk.json"):
        os.utime(tmp_path / name, (old, old))
    os.utime(tmp_path / "findings.json", (old + 3600, old + 3600))
    for check in (_certify_change_verify(tmp_path),
                  _certify_finished_walk(tmp_path)):
        assert check.status == "skip"
        assert "stale" in check.detail and "galley verify" in check.detail


def test_certify_reads_generated_at_over_mtimes(tmp_path):
    """A copied run directory loses its mtimes; the recorded timestamps still
    say which build each artifact belongs to."""
    from galley.manifest import _certify_change_verify, _certify_finished_walk
    (tmp_path / "change_verify.json").write_text(json.dumps(
        {"generated_at": "2026-09-01T10:00:00+00:00", "ran": True,
         "applied_edits": 3, "problems": []}))
    (tmp_path / "finished_walk.json").write_text(json.dumps(
        {"generated_at": "2026-09-01T10:00:00+00:00", "ran": True,
         "residuals": []}))
    (tmp_path / "findings.json").write_text(json.dumps(
        {"generated_at": "2026-09-01T12:00:00+00:00", "findings": []}))
    assert _certify_change_verify(tmp_path).status == "skip"
    assert _certify_finished_walk(tmp_path).status == "skip"


def test_certify_keeps_a_verify_record_written_after_the_findings(tmp_path):
    from galley.manifest import _certify_change_verify, _certify_finished_walk
    (tmp_path / "findings.json").write_text(json.dumps(
        {"generated_at": "2026-09-01T10:00:00+00:00", "findings": []}))
    (tmp_path / "change_verify.json").write_text(json.dumps(
        {"generated_at": "2026-09-01T12:00:00+00:00", "ran": True,
         "applied_edits": 3, "problems": []}))
    (tmp_path / "finished_walk.json").write_text(json.dumps(
        {"generated_at": "2026-09-01T12:00:00+00:00", "ran": True,
         "residuals": []}))
    assert _certify_change_verify(tmp_path).status == "pass"
    assert _certify_finished_walk(tmp_path).status == "pass"


def test_certify_verify_gates_without_ran_field_still_pass_a_clean_record(tmp_path):
    from galley.manifest import _certify_change_verify, _certify_finished_walk
    (tmp_path / "change_verify.json").write_text(json.dumps(
        {"ran": True, "applied_edits": 3, "problems": []}))
    (tmp_path / "finished_walk.json").write_text(json.dumps(
        {"ran": True, "residuals": []}))
    assert _certify_change_verify(tmp_path).status == "pass"
    assert _certify_finished_walk(tmp_path).status == "pass"


# --- the Galley rebuild marker ------------------------------------------------

def _rebuild_envelope(n=2, **extra):
    from galley.deliverable import REBUILD_MARKER
    env = {"findings": [{"finding_id": f"f-{i}", "para_id": "b1",
                         "corrected_text": "x", "status": "validated"}
                        for i in range(n)],
           "cost": {"total_usd": 0.0}, **REBUILD_MARKER}
    env.update(extra)
    return env


def _casefile(run, *charges):
    (run / "casefile.json").write_text(json.dumps({
        "budget": {"charges": [{"label": "wave", "cost_usd": c, "wave": 1}
                               for c in charges], "caps": {}}}))


def test_galley_rebuild_is_not_a_zero_cost_anomaly(tmp_path):
    run = _run_dir(tmp_path, _rebuild_envelope())
    _casefile(run, 1.5, 2.7)
    cert = certify_run(run)
    by = {c.name: c for c in cert.checks}
    assert by["zero-cost anomaly"].status == "pass"
    assert "$4.20" in by["zero-cost anomaly"].detail
    assert by["checkpoint present"].status == "pass"


def test_galley_rebuild_reconciles_budget_against_the_casefile(tmp_path):
    """The envelope says $0; the case file says $4.20 — the approval's ceiling
    is checked against the money actually spent, not the rebuild's own $0."""
    cfg = _mech_cfg()
    src = _source(tmp_path)
    run = _run_dir(tmp_path, _rebuild_envelope())
    _casefile(run, 4.2)
    tight = build_manifest(source=src, config_path=str(CONFIG), cfg=cfg,
                           max_spend_usd=3.0)
    budget = next(c for c in certify_run(run, manifest=tight).checks
                  if c.name == "budget within approval")
    assert budget.status == "fail" and "$4.20" in budget.detail
    roomy = build_manifest(source=src, config_path=str(CONFIG), cfg=cfg,
                           max_spend_usd=20.0)
    budget = next(c for c in certify_run(run, manifest=roomy).checks
                  if c.name == "budget within approval")
    assert budget.status == "pass" and "casefile.json" in budget.detail


def test_galley_rebuild_honors_a_serialized_spent_usd(tmp_path):
    run = _run_dir(tmp_path, _rebuild_envelope())
    (run / "casefile.json").write_text(json.dumps(
        {"budget": {"spent_usd": 7.25, "charges": []}}))
    anomaly = next(c for c in certify_run(run).checks
                   if c.name == "zero-cost anomaly")
    assert anomaly.status == "pass" and "$7.25" in anomaly.detail


def test_galley_rebuild_without_a_casefile_skips_budget_loudly(tmp_path):
    cfg = _mech_cfg()
    src = _source(tmp_path)
    run = _run_dir(tmp_path, _rebuild_envelope())     # no casefile.json
    m = build_manifest(source=src, config_path=str(CONFIG), cfg=cfg,
                       max_spend_usd=3.0)
    by = {c.name: c for c in certify_run(run, manifest=m).checks}
    assert by["zero-cost anomaly"].status == "pass"
    assert by["budget within approval"].status == "skip"
    assert "casefile.json" in by["budget within approval"].detail
    assert by["checkpoint present"].status == "skip"


def test_judge_model_marker_alone_is_enough(tmp_path):
    run = _run_dir(tmp_path, {"findings": [{"a": 1}],
                              "cost": {"total_usd": 0.0},
                              "judge_model": "galley:rebuild"})
    anomaly = next(c for c in certify_run(run).checks
                   if c.name == "zero-cost anomaly")
    assert anomaly.status == "pass"


# --- the reject-all round trip -----------------------------------------------

def _reject_all_check(tmp_path, audit):
    env = {"findings": [], "cost": {"total_usd": 0.0}}
    if audit is not None:
        env["audit"] = audit
    sub = tmp_path / f"case-{len(list(tmp_path.iterdir()))}"
    sub.mkdir()
    return next(c for c in certify_run(_run_dir(sub, env)).checks
                if c.name == "reject-all round trip")


def test_certify_fails_a_failed_reject_all_audit(tmp_path):
    check = _reject_all_check(tmp_path, {
        "ran": True, "passed": False, "checked": 40,
        "mismatches": ["body-0007", "body-0012"], "missing": []})
    assert check.status == "fail"
    assert "2 paragraph(s) differ" in check.detail and "body-0007" in check.detail


def test_certify_skips_loudly_when_the_audit_did_not_run(tmp_path):
    check = _reject_all_check(tmp_path, {"ran": False, "passed": True,
                                         "checked": 0, "mismatches": [],
                                         "missing": []})
    assert check.status == "skip" and "did not run" in check.detail
    absent = _reject_all_check(tmp_path, None)
    assert absent.status == "skip" and "no audit record" in absent.detail


def test_certify_passes_a_passed_reject_all_audit(tmp_path):
    check = _reject_all_check(tmp_path, {"ran": True, "passed": True,
                                         "checked": 40, "mismatches": [],
                                         "missing": []})
    assert check.status == "pass" and "40" in check.detail


# --- artifact regexes + applied rows only -----------------------------------

def _scan(tmp_path, rows):
    run = _run_dir(tmp_path, {"findings": rows, "cost": {"total_usd": 0.0}})
    return next(c for c in certify_run(run).checks if c.name == "artifact scan")


def test_artifact_scan_allows_a_sentence_final_ellipsis_period(tmp_path):
    """An ellipsis directly followed by ONE period is not an artifact — it is
    exactly what `sweep_ellipsis` leaves behind at a genuine sentence end
    (house style: NBSP + … + period), per HARNESS_NOTES item 9. Flagging it
    made a correctly swept book fail its own certificate."""
    check = _scan(tmp_path, [{"para_id": "b1", "corrected_text": "He waited…."}])
    assert check.status == "pass"


def test_artifact_scan_catches_a_doubled_period_after_an_ellipsis(tmp_path):
    """A composed edit appending its own period onto text that already ended
    in an ellipsis+period is a genuine merge artifact."""
    check = _scan(tmp_path, [{"para_id": "b1", "corrected_text": "He waited….."}])
    assert check.status == "fail" and "…" in check.detail


def test_sweep_ellipsis_output_is_clean_under_the_artifact_scan(tmp_path):
    """Run the shipped sweep over sample sentences and confirm the certify
    artifact scan is clean on its output — the sweep and the certificate must
    agree. Mirrors the Redding body-0734-style case: a manuscript sentence
    already ends on the ellipsis GLYPH plus a literal period (a common
    pre-existing style), and the sweep only normalizes the glyph's leading
    space, leaving the sentence-final period exactly where it was."""
    from docproof.sweeps import _sweep_ellipsis, apply_hits
    samples = [
        "She wondered whether it even mattered….",
        "He waited….",
        "“I don't know…” she said.",
    ]
    swept = [apply_hits(t, _sweep_ellipsis(t, None, "nbsp")) for t in samples]
    check = _scan(tmp_path, [{"para_id": f"b{i}", "corrected_text": t}
                             for i, t in enumerate(swept)])
    assert check.status == "pass"


def test_artifact_scan_period_patterns_are_literal_not_any_char(tmp_path):
    check = _scan(tmp_path, [{"para_id": "b1",
                              "corrected_text": "“Go,” she said…and left. ”x"}])
    assert check.status == "pass"


def test_artifact_scan_ignores_rejected_and_skipped_rows(tmp_path):
    """What a rejected row WOULD have written never reaches the author."""
    check = _scan(tmp_path, [
        {"para_id": "b1", "corrected_text": "She paused,, then left.",
         "status": "rejected_overlap"},
        {"para_id": "b2", "corrected_text": "Bad ”. here",
         "status": "skipped_low_confidence"},
        {"para_id": "b3", "corrected_text": "fine ,, but unplaced",
         "status": "validated", "applied": False},
        {"para_id": "b4", "corrected_text": "Clean.", "status": "validated"},
    ])
    assert check.status == "pass"


def test_artifact_scan_still_reads_an_applied_row(tmp_path):
    check = _scan(tmp_path, [{"para_id": "b1", "status": "validated",
                              "applied": True,
                              "corrected_text": "She paused,, then left."}])
    assert check.status == "fail"


def test_insertion_collision_ignores_the_rejected_loser(tmp_path):
    env = {"findings": [
        {"finding_id": "f-1", "para_id": "b1", "original_text": "the cat",
         "occurrence": 1, "corrected_text": "the cats", "status": "validated"},
        {"finding_id": "f-2", "para_id": "b1", "original_text": "the cat",
         "occurrence": 1, "corrected_text": "the CAT",
         "status": "rejected_overlap"},
    ], "cost": {"total_usd": 0.0}}
    coll = next(c for c in certify_run(_run_dir(tmp_path, env)).checks
                if c.name == "no insertion collisions")
    assert coll.status == "pass"


# --- explicit skips for an absent --source / --config ------------------------

def test_certify_records_skips_when_no_approval_is_given(tmp_path):
    run = _run_dir(tmp_path, {"findings": [], "cost": {"total_usd": 0.0}})
    by = {c.name: c for c in certify_run(run).checks}
    assert by["source hash"].status == "skip"
    assert by["config hash"].status == "skip"
    assert by["approved model routes"].status == "skip"
    assert "no approval manifest" in by["source hash"].detail


def test_certify_records_skips_when_source_or_config_is_absent(tmp_path):
    cfg = _mech_cfg()
    src = _source(tmp_path)
    m = build_manifest(source=src, config_path=str(CONFIG), cfg=cfg,
                       max_spend_usd=20.0)
    run = _run_dir(tmp_path, {"findings": [], "cost": {"total_usd": 0.0}})
    by = {c.name: c for c in certify_run(run, manifest=m, cfg=cfg).checks}
    assert by["source hash"].status == "skip" and "--source" in by["source hash"].detail
    assert by["config hash"].status == "pass"
    by = {c.name: c for c in certify_run(run, manifest=m, source=src).checks}
    assert by["source hash"].status == "pass"
    assert by["config hash"].status == "skip" and "--config" in by["config hash"].detail
    assert by["approved model routes"].status == "skip"


# --- two-author attribution, read off the deliverable ------------------------

def _docx_with_revisions(run, name, authors):
    """A .docx whose one paragraph carries a tracked insertion per author."""
    import zipfile
    _docx_in(run, name, "A clean paragraph.")
    path = run / name
    tmp = run / (name + ".tmp")
    ins = "".join(
        f'<w:ins w:id="{i}" w:author="{a}" w:date="2026-01-01T00:00:00Z">'
        f'<w:r><w:t xml:space="preserve"> more</w:t></w:r></w:ins>'
        for i, a in enumerate(authors, 1))
    with zipfile.ZipFile(path) as zin, \
            zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/document.xml":
                data = data.decode("utf-8").replace(
                    "</w:p>", ins + "</w:p>", 1).encode("utf-8")
            zout.writestr(item, data)
    tmp.replace(path)


_TWO_LANES = [
    {"finding_id": "f-1", "para_id": "b1", "error_type": "spelling",
     "lane": "mechanical", "corrected_text": "x", "status": "validated"},
    {"finding_id": "f-2", "para_id": "b2", "error_type": "copyedit",
     "lane": "copyedit", "corrected_text": "y", "status": "validated"},
]


def test_two_author_fails_when_both_lanes_share_one_author(tmp_path):
    run = _run_dir(tmp_path, {"findings": _TWO_LANES, "cost": {"total_usd": 0.0}})
    _docx_with_revisions(run, "book.docx",
                         ["Atmosphere Press Proofreader"] * 2)
    ta = next(c for c in certify_run(run).checks
              if c.name == "two-author attribution")
    assert ta.status == "fail" and "1 author" in ta.detail


def test_two_author_passes_with_two_revision_authors(tmp_path):
    run = _run_dir(tmp_path, {"findings": _TWO_LANES, "cost": {"total_usd": 0.0}})
    _docx_with_revisions(run, "book.docx", ["Atmosphere Press Proofreader",
                                            "Atmosphere Press Copy Editor"])
    ta = next(c for c in certify_run(run).checks
              if c.name == "two-author attribution")
    assert ta.status == "pass" and "Copy Editor" in ta.detail


def test_two_author_skips_without_a_deliverable_and_for_one_lane(tmp_path):
    run = _run_dir(tmp_path, {"findings": _TWO_LANES, "cost": {"total_usd": 0.0}})
    ta = next(c for c in certify_run(run).checks
              if c.name == "two-author attribution")
    assert ta.status == "skip" and "no manuscript .docx" in ta.detail
    # A copy-edit row that was REJECTED is not an applied lane.
    one = [dict(_TWO_LANES[0]), dict(_TWO_LANES[1], status="rejected_overlap")]
    run2 = tmp_path / "run2"
    run2.mkdir()
    (run2 / "findings.json").write_text(json.dumps(
        {"findings": one, "cost": {"total_usd": 0.0}}))
    ta2 = next(c for c in certify_run(run2).checks
               if c.name == "two-author attribution")
    assert ta2.status == "skip" and "single lane" in ta2.detail
