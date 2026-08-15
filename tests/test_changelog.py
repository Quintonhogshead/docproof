"""The change log the press hands to an author.

`summary.md` is for whoever ran the review; this is the deliverable. The
brief's rule for it is "never overstate the depth of the pass", so most of
what is asserted here is that the honesty sections are present and say
something true — a change log that lists corrections and nothing else invites
the reader to assume everything was covered.
"""
from __future__ import annotations

import itertools

import docx
import pytest
from docx.table import Table
from docx.text.paragraph import Paragraph

from docproof.analyzer import MockAnalyzer
from docproof.config import Config, load_config
from docproof.formats import DOCX, IDML
from docproof.models import Usage
from docproof.pipeline import finish, prepare


def _read(path):
    """(all text, [tables as row lists]) — what a reader would see."""
    d = docx.Document(path)
    text, tables = [], []
    for child in d.element.body.iterchildren():
        if child.tag.endswith("}p"):
            text.append(Paragraph(child, d).text)
        elif child.tag.endswith("}tbl"):
            tables.append([[c.text for c in r.cells]
                           for r in Table(child, d).rows])
    return "\n".join(text), tables


def _run(tmp_path, paragraphs, mock=(), *, error_types=("comma_splice",), **kw):
    d = docx.Document()
    for t in paragraphs:
        d.add_paragraph(t)
    src = tmp_path / "Book One.docx"
    d.save(src)

    cfg = load_config("config/default.yaml")
    cfg.error_types = list(error_types)
    for k, v in kw.items():
        setattr(cfg, k, v)
    prepared = prepare(cfg, src, "config/error_types")

    findings, ids = [], itertools.count(1)
    for group in prepared.groups:
        for chunk in prepared.chunks:
            found, _ = MockAnalyzer(group, list(mock), ids).analyze_chunk(
                chunk, Usage())
            findings += found
    return finish(prepared, findings, Usage(), cfg, out_dir=tmp_path / "out",
                  source_path=src)


# --- filenames ----------------------------------------------------------------

def test_the_house_filenames():
    assert DOCX.reviewed_name("/a/Mankin - Book 1.docx") == \
        "Mankin - Book 1 - Atmosphere Press Proofreader.docx"
    assert DOCX.change_log_name("/a/Mankin - Book 1.docx") == \
        "Mankin - Book 1 - Atmosphere Press Proofreader Change Log.docx"


def test_a_layout_keeps_its_extension_but_the_log_is_always_word():
    assert IDML.reviewed_name("/a/Book.idml") == \
        "Book - Atmosphere Press Proofreader.idml"
    assert IDML.change_log_name("/a/Book.idml") == \
        "Book - Atmosphere Press Proofreader Change Log.docx"


def test_the_watcher_knows_these_are_its_own_output():
    """Otherwise the folder watcher would treat a proofread manuscript as a
    new one and pay to review it again."""
    from app.watch.stages import _looks_like_output
    assert _looks_like_output("Book One - Atmosphere Press Proofreader.docx")
    assert _looks_like_output(
        "Book One - Atmosphere Press Proofreader Change Log.docx")
    # Books proofread before the rename still carry the old suffix; the watcher
    # must keep recognising them, or a rename reprocesses the whole back catalogue.
    assert _looks_like_output("Book One - Pre-Proofread.docx")
    assert _looks_like_output("Book One - Pre-Proofread Change Log.docx")
    assert not _looks_like_output("Book One.docx")


# --- the document ------------------------------------------------------------

SPLICE = "The manuscript was finished, nobody wanted to read it."


def test_a_change_log_is_written_beside_the_manuscript(tmp_path):
    out = _run(tmp_path, [SPLICE])
    assert out.change_log is not None and out.change_log.exists()
    assert out.change_log.name == \
        "Book One - Atmosphere Press Proofreader Change Log.docx"


def test_it_can_be_turned_off(tmp_path):
    out = _run(tmp_path, [SPLICE], change_log=False)
    assert out.change_log is None
    assert not list((tmp_path / "out").glob("*Change Log*"))


def test_every_section_the_brief_asks_for_is_present(tmp_path):
    out = _run(tmp_path, [SPLICE])
    text, _ = _read(out.change_log)
    for heading in ("Style basis", "Corrections", "Scripted checks",
                    "Queries", "Deliberately left unchanged",
                    "Footnotes and endnotes", "Final audit",
                    "Limits of this pass", "How this pass was made"):
        assert heading in text, heading


def test_the_style_basis_states_the_variant_and_its_authorities(tmp_path):
    """The brief asks for the variant to be stated, with its authorities, so
    the press can check the pass was made against the right rules."""
    text, _ = _read(_run(tmp_path, [SPLICE]).change_log)
    assert "proofread as U.S. English" in text
    assert "Chicago Manual of Style" in text
    assert "Merriam-Webster" in text


def test_a_hybrid_variant_asks_the_press_to_confirm(tmp_path):
    """Canadian is U.S. punctuation with Canadian spelling, so choosing it is
    a judgment the press should ratify — the brief asks for this by name."""
    text, _ = _read(_run(tmp_path, [SPLICE], variant="ca").change_log)
    assert "proofread as Canadian English" in text
    assert "Canadian Oxford Dictionary" in text
    assert "please confirm" in text


def test_a_non_hybrid_variant_does_not_ask(tmp_path):
    text, _ = _read(_run(tmp_path, [SPLICE], variant="uk").change_log)
    assert "proofread as U.K. English" in text
    assert "Oxford Style Manual" in text
    assert "please confirm" not in text


def test_an_unavailable_dictionary_is_disclosed_as_a_limit(tmp_path):
    """spylls ships en_US only. A U.K. run therefore has no dictionary scan,
    and the change log has to say so rather than letting the reader assume the
    manuscript's vocabulary was protected."""
    text, _ = _read(_run(tmp_path, [SPLICE], variant="uk").change_log)
    assert "dictionary scan did not run" in text
    assert "en_GB is not installed" in text


def test_corrections_are_tabled_with_a_reason(tmp_path):
    out = _run(tmp_path, ["He said what?! The forest---was quiet."])
    _, tables = _read(out.change_log)
    header, *rows = tables[0]
    assert header == ["Where", "Original", "Corrected", "Why"]
    assert rows, "no corrections tabled"
    assert all(r[3].strip() for r in rows), "a correction with no reason"


def test_corrections_are_listed_in_document_order(tmp_path):
    """A reader works through this beside the manuscript. Grouping by which
    sweep found what would send them jumping around the book."""
    out = _run(tmp_path, ["He said what?! Then he waited...",
                          "The forest---was quiet. It was the 20th century."])
    _, tables = _read(out.change_log)
    wheres = [r[0] for r in tables[0][1:]]
    assert wheres == sorted(wheres)


def test_the_scripted_check_counts_are_reported(tmp_path):
    text, _ = _read(_run(tmp_path, ["He said what?!"]).change_log)
    assert "Stacked punctuation: 1 flagged, 1 applied, 0 remaining" in text


def test_the_silent_normalizations_are_disclosed_but_not_itemized(tmp_path):
    """They cannot be rejected in Word, so an author who is not told they
    happened has no way to find out."""
    text, _ = _read(_run(tmp_path, ['"Wait here."  he said.']).change_log)
    assert "Applied without tracked changes" in text
    assert "cannot be rejected in Word" in text


def test_queries_are_listed_and_marked_as_questions(tmp_path):
    out = _run(tmp_path, [SPLICE],
               [{"para_id": "body-0000", "error_type": "comma_splice",
                 "original_text": SPLICE,
                 "corrected_text": "The manuscript was finished; nobody wanted to read it.",
                 "explanation": "Possible splice.", "confidence": "low"}])
    text, _ = _read(out.change_log)
    assert "Possible splice." in text
    assert "a suggestion and not a correction" in text


def test_a_query_is_named_in_the_word_the_format_uses(tmp_path):
    """The log is read beside the manuscript it describes. An author opening an
    .idml is looking for a note; nothing in that file is called a comment.

    The nouns are whole, not a shared "margin " prefix plus a format word: two
    report sites used to prepend the prefix themselves and rendered "a margin
    margin comment" for .docx. So the assertions here are on the exact string
    the log carries, which is what caught that — and the two are asymmetric on
    purpose, because a Word comment sits in the margin and an InDesign Note
    does not. Making them parallel is the tempting edit this guards."""
    from docproof.changelog import write_change_log
    from docproof.models import Anchor, DocumentModel, Finding, ParagraphRef
    from docproof.variants import load_variant

    doc = DocumentModel(source_path="Book.idml", paragraphs=(ParagraphRef(
        "story-ue0-p0000", "Stories/Story_ue0.xml", "body", SPLICE, "Normal"),))
    q = Finding("q-1", "chunk-000", "story-ue0-p0000", "speaker_change",
                SPLICE, 1, SPLICE, "Who is speaking?", "high", status="query",
                anchor=Anchor(0, len(SPLICE), SPLICE, ""))
    cfg = load_config("config/default.yaml")

    write_change_log(tmp_path / "idml.docx", doc=doc, findings=[q], cfg=cfg,
                     applied_ids=(), fmt=IDML, variant=load_variant("us"))
    text, _ = _read(tmp_path / "idml.docx")
    assert "1 question(s) were raised as notes" in text
    assert "each a note in the manuscript" in text
    assert "margin comment" not in text        # never Word's noun, nor the
    assert "margin note" not in text           # default; a Note has no margin

    write_change_log(tmp_path / "docx.docx", doc=doc, findings=[q], cfg=cfg,
                     applied_ids=(), fmt=DOCX, variant=load_variant("us"))
    text, _ = _read(tmp_path / "docx.docx")
    assert "raised as margin comments" in text
    assert "margin margin" not in text         # the prefix is not prepended


def test_the_log_only_claims_the_questions_the_file_actually_carries(tmp_path):
    """With query_comments off, a below-gate finding stays in summary.md. It is
    still listed here — nothing is discarded silently — but the log must not
    say it is in a manuscript that does not hold it."""
    out = _run(tmp_path, [SPLICE],
               [{"para_id": "body-0000", "error_type": "comma_splice",
                 "original_text": SPLICE,
                 "corrected_text": "The manuscript was finished; nobody wanted to read it.",
                 "explanation": "Possible splice.", "confidence": "low"}],
               query_comments=False)
    text, _ = _read(out.change_log)
    assert "0 question(s) were raised as margin comments" in text
    assert "1 question(s), 0 of them a margin comment in the manuscript" in text
    assert "Possible splice." in text          # still reported, just not in the file


def test_the_authors_own_words_are_named_as_protected(tmp_path):
    # Mid-sentence capitalization is what marks a coinage as a name, so it is
    # protected however high the length-scaled repeat floor climbs.
    out = _run(tmp_path, ["The cold marches were where Kaelith crossed at dawn.",
                          "By dusk Kaelith had not returned home."])
    text, _ = _read(out.change_log)
    assert "Kaelith" in text
    assert "protected from correction" in text


def test_the_audit_result_is_stated(tmp_path):
    text, _ = _read(_run(tmp_path, [SPLICE]).change_log)
    assert "Passed." in text
    assert "reproduces all" in text


def test_a_failed_audit_says_so_in_the_change_log(tmp_path):
    """The change log is written before the audit is enforced, so a refused
    run still explains itself in the author's own document."""
    from docproof.utils.xml_helpers import (iter_text_elements, set_text,
                                            walk_package)
    d = docx.Document()
    d.add_paragraph("The forest was quiet and the road went on.")
    src = tmp_path / "Book One.docx"
    d.save(src)

    cfg = load_config("config/default.yaml")
    cfg.audit = "warn"
    prepared = prepare(cfg, src, "config/error_types")
    for wp in walk_package(prepared.pkg):
        t = next(iter_text_elements(wp.element))
        set_text(t, (t.text or "").replace("forest", "woods"))

    out = finish(prepared, [], Usage(), cfg, out_dir=tmp_path / "out",
                 source_path=src)
    text, _ = _read(out.change_log)
    assert "FAILED" in text and "should not be treated as final" in text


def test_notes_coverage_is_stated_either_way(tmp_path):
    text, _ = _read(_run(tmp_path, [SPLICE]).change_log)
    assert "no footnotes or endnotes" in text


def test_a_manuscript_with_footnotes_says_they_were_read():
    """The brief singles this out because the usual failure is silent: most
    tools read only the body."""
    from docproof.changelog import _notes_statement
    from docproof.models import DocumentModel, ParagraphRef
    doc = DocumentModel(source_path="x.docx", paragraphs=(
        ParagraphRef("body-0000", "word/document.xml", "body", "Text.", "Normal"),
        ParagraphRef("footnote-2-p0", "word/footnotes.xml", "footnote",
                     "A note.", "Normal")))
    said = _notes_statement(doc)
    assert "footnotes" in said and "they were read" in said


def test_the_limits_section_admits_what_was_not_read(tmp_path):
    text, _ = _read(_run(tmp_path, [SPLICE]).change_log)
    assert "does not read for plot" in text
    assert "images" in text
    assert "proposed, not imposed" in text


def test_the_default_config_ships_the_change_log_on():
    assert load_config("config/default.yaml").change_log is True
    assert Config().change_log is True
