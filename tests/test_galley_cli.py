"""The `docproof galley` verbs' CLI contract (docproof/__main__.py), where it
meets the approval gate, the $0 promise, and the cost ledger:

- `review --approval` hashes the config AFTER the CLI's own mutations
  (--rounds, --meaning-*, --fix-*) land, so a flag slipped in past approval is
  a refused deviation rather than a free ride;
- `sweep` and `merge` really make no provider call: the two prepare()-time
  spenders (storysheet, candidate screening) are zeroed with the rest;
- `galley score` grades a `docproof review` findings.json directly;
- `galley audit` / `galley verify` write a `cost` field and keep `ran`, and
  verify's exit code draws the same line certify does (flagged edit or a
  HIGH-severity residual);
- `galley calibrate --from-run` works on a bare run with no casefile.json;
- the spend guard refuses a model outside the approval.

No test here makes a network call: providers are scripted fakes handed in
through the same seams the other CLI tests patch (`build_provider`,
`_flights_provider_of`).
"""
from __future__ import annotations

import json
from pathlib import Path

import docx
import yaml

import docproof.__main__ as m
from docproof.__main__ import main
from docproof.config import load_config
from docproof.ingest import build_document_model, preflight
from docproof.providers import NormalizedUsage, ProviderResult

from .conftest import FIXTURES

CONFIG = str(Path("config/default.yaml"))
U = NormalizedUsage(input_tokens=500, output_tokens=50)

PARAGRAPHS = [
    "The lamp on the desk flickered, then finally caught, and Maria sighed.",
    "Its light was thin, but it was enough to read the letter by tonight.",
    "David counted the coins twice, sure the total would not change again.",
    "The garden gate creaked open, and the old dog limped out to greet her.",
]


class _Provider:
    """Replays one scripted parsed body per call; records every call."""

    def __init__(self, *bodies):
        self._bodies = list(bodies)
        self.calls = []

    def complete_structured(self, **kwargs):
        self.calls.append(kwargs)
        body = self._bodies.pop(0) if self._bodies else {}
        return ProviderResult(parsed=body, usage=U, stop_reason="ok")


def _manuscript(tmp_path, name="book.docx"):
    d = docx.Document()
    for p in PARAGRAPHS:
        d.add_paragraph(p)
    src = tmp_path / name
    d.save(src)
    return src


def _mock_review(tmp_path, src, out, *, para_index=0, wrong="teh", right="the"):
    """A $0 review over `src` whose one mock finding lands as a tracked change,
    so the run dir carries a deliverable .docx + a findings.json with a
    validated, anchored row."""
    cfg = load_config(CONFIG)
    doc = build_document_model(preflight(str(src), cfg.tracked_changes_policy),
                               cfg)
    para = doc.paragraphs[para_index]
    assert wrong in para.text
    mocks = [{"para_id": para.para_id, "original_text": wrong,
              "corrected_text": right, "confidence": "high"}]
    mocks_path = tmp_path / "mocks.json"
    mocks_path.write_text(json.dumps(mocks), encoding="utf-8")
    rc = main(["review", str(src), "--error-types", "spelling",
               "--mock-findings", str(mocks_path), "--out", str(out),
               "--config", CONFIG])
    assert rc == 0
    return out


def _seeded_manuscript(tmp_path):
    d = docx.Document()
    for p in PARAGRAPHS:
        d.add_paragraph(p.replace("the", "teh", 1))
    src = tmp_path / "seeded.docx"
    d.save(src)
    return src


# --- 1. review --approval covers the CLI's own config mutations ---------------

def test_review_approval_guard_sees_rounds_and_gate_flags(tmp_path, monkeypatch):
    src = _seeded_manuscript(tmp_path)
    seen = []

    def _capture(args, cfg):
        seen.append(cfg.model_copy(deep=True))
        return 3                                   # abort right after the guard
    monkeypatch.setattr(m, "_approval_guard", _capture)

    rc = main(["review", str(src), "--config", CONFIG, "--out", str(tmp_path / "r"),
               "--approval", str(tmp_path / "unused.json"),
               "--rounds", "2", "--meaning-model", "gpt-5.6-luna",
               "--fix-check"])
    assert rc == 3
    (cfg,) = seen
    # Every flag that changes the effective config landed BEFORE the guard.
    assert cfg.rounds.count == 2
    assert cfg.meaning_check.enabled and cfg.meaning_check.model == "gpt-5.6-luna"
    assert cfg.fix_check.enabled


def test_review_refuses_a_rounds_flag_the_approval_did_not_hash(tmp_path, capsys):
    src = _seeded_manuscript(tmp_path)
    approval = tmp_path / "approval.json"
    rc = main(["galley", "approve", str(src), "--config", CONFIG,
               "--budget", "5", "--out", str(approval)])
    assert rc == 0
    rc = main(["review", str(src), "--config", CONFIG, "--out", str(tmp_path / "r"),
               "--approval", str(approval), "--rounds", "2"])
    assert rc == 5
    err = capsys.readouterr().err
    assert "REFUSED" in err and "config_changed" in err


def test_review_prints_its_dollar_line(tmp_path, capsys):
    src = _seeded_manuscript(tmp_path)
    _mock_review(tmp_path, src, tmp_path / "run")
    out = capsys.readouterr().out
    assert "spent (" in out and "model call(s)" in out


# --- 2. sweep / merge zero the prepare()-time spenders ------------------------

def _spending_config(tmp_path):
    cfg = yaml.safe_load(Path(CONFIG).read_text(encoding="utf-8"))
    cfg["storysheet"] = dict(cfg.get("storysheet") or {}, enabled=True)
    cfg["candidate_screening"] = dict(cfg.get("candidate_screening") or {},
                                      mode="shadow")
    path = tmp_path / "spending.yaml"
    path.write_text(yaml.safe_dump(cfg), encoding="utf-8")
    return path


def _capture_prepare(monkeypatch):
    seen = []
    real = m.prepare

    def _prepare(cfg, *a, **k):
        seen.append((cfg.storysheet.enabled, cfg.candidate_screening.mode))
        # The paid stages are asserted off by the test; the real prepare()
        # then never reaches a provider.
        return real(cfg, *a, **k)
    monkeypatch.setattr(m, "prepare", _prepare)

    def _boom(*a, **k):                       # pragma: no cover - must not run
        raise AssertionError("a $0 verb reached the provider")
    monkeypatch.setattr(m, "build_provider", _boom)
    return seen


def test_sweep_silences_storysheet_and_candidate_screening(tmp_path, monkeypatch):
    cfg_path = _spending_config(tmp_path)
    rule = tmp_path / "rule.yaml"
    rule.write_text("key: bespoke_x\nname: x\npattern: 'zzz'\nreplacement: 'y'\n",
                    encoding="utf-8")
    seen = _capture_prepare(monkeypatch)
    rc = main(["sweep", str(FIXTURES / "simple.docx"), "--rule", str(rule),
               "--config", str(cfg_path), "--out", str(tmp_path / "out")])
    assert rc == 0
    assert seen == [(False, "off")]


def test_merge_silences_storysheet_and_candidate_screening(tmp_path, monkeypatch):
    cfg_path = _spending_config(tmp_path)
    mech = tmp_path / "mech.json"
    mech.write_text(json.dumps([]), encoding="utf-8")
    seen = _capture_prepare(monkeypatch)
    rc = main(["merge", str(FIXTURES / "simple.docx"), "--mechanical", str(mech),
               "--config", str(cfg_path), "--out", str(tmp_path / "out"),
               "--dry-run"])
    assert rc == 0
    assert seen == [(False, "off")]


# --- 3. galley score grades a review findings.json directly -------------------

def test_score_accepts_a_review_findings_json(tmp_path, capsys):
    src = _manuscript(tmp_path)
    seed_dir = tmp_path / "seed"
    rc = main(["galley", "seed", str(src), "-n", "3", "--seed", "0",
               "--out", str(seed_dir)])
    assert rc == 0
    seeded = json.loads((seed_dir / "seeded_manuscript.json").read_text("utf-8"))
    key = json.loads((seed_dir / "answer_key.json").read_text("utf-8"))
    assert key["planted"]

    cfg = load_config(CONFIG)
    doc = build_document_model(preflight(str(src), cfg.tracked_changes_policy),
                               cfg)
    original = {p.para_id: p.text for p in doc.paragraphs}
    mocks = [{"para_id": pe["para_id"],
              "original_text": seeded["paragraphs"][pe["para_id"]],
              "corrected_text": original[pe["para_id"]], "confidence": "high"}
             for pe in key["planted"]]
    mocks_path = tmp_path / "mocks.json"
    mocks_path.write_text(json.dumps(mocks), encoding="utf-8")
    review_dir = tmp_path / "review"
    rc = main(["review", str(seed_dir / "seeded_manuscript.docx"),
               "--error-types", "spelling", "--mock-findings", str(mocks_path),
               "--out", str(review_dir), "--config", CONFIG])
    assert rc == 0

    # The seed hint names this exact path: the run's findings.json, unconverted.
    score_dir = tmp_path / "score"
    rc = main(["galley", "score", str(review_dir / "findings.json"),
               "--answer-key", str(seed_dir / "answer_key.json"),
               "--out", str(score_dir)])
    assert rc == 0
    out = capsys.readouterr().out
    assert "review finding(s) converted" in out
    recall = json.loads((score_dir / "recall.json").read_text("utf-8"))
    assert recall["caught"] == recall["planted"] == len(key["planted"])


def test_score_still_reads_native_gfindings(tmp_path):
    key = {"requested": 1, "rng_seed": 0, "seeded_chapters": [1], "planted": [
        {"para_id": "p1", "error_type": "spelling", "start": 0, "end": 3,
         "original": "the", "mutated": "teh"}]}
    key_path = tmp_path / "answer_key.json"
    key_path.write_text(json.dumps(key), encoding="utf-8")
    findings = tmp_path / "g.json"
    findings.write_text(json.dumps([{
        "id": "g-1", "error_type": "spelling",
        "span": {"para_id": "p1", "start": 0, "end": 3},
        "find": "teh", "replace": "the"}]), encoding="utf-8")
    rc = main(["galley", "score", str(findings), "--answer-key", str(key_path),
               "--out", str(tmp_path)])
    assert rc == 0
    assert (tmp_path / "recall.json").exists()


# --- 4. cost fields on audit / verify; 6. verify's exit policy ----------------

def test_audit_writes_a_cost_field_and_keeps_ran(tmp_path, monkeypatch):
    src = _seeded_manuscript(tmp_path)
    run = _mock_review(tmp_path, src, tmp_path / "run")
    prov = _Provider({"hypotheses": []})
    monkeypatch.setattr(m, "build_provider", lambda cfg, **k: prov)
    rc = main(["galley", "audit", str(run), "--source", str(src),
               "--config", CONFIG, "--model", "gpt-5.6-luna"])
    assert rc == 0
    payload = json.loads((run / "audit.json").read_text("utf-8"))
    assert payload["ran"] is True
    assert payload["cost"]["total_usd"] > 0
    assert "gpt-5.6-luna" in payload["cost"]["by_model"]
    # The experiment record rides beside the hypotheses.
    assert {"control_chapter", "sample_ids", "seed"} <= payload.keys()
    assert len(prov.calls) == 1


def test_verify_writes_cost_per_gate_and_passes_on_low_residuals(tmp_path,
                                                                  monkeypatch):
    src = _seeded_manuscript(tmp_path)
    run = _mock_review(tmp_path, src, tmp_path / "run")
    prov = _Provider(
        {"problems": []},
        {"findings": [{"para_id": "body-0001", "quote": "thin",
                       "problem": "flat word", "suggestion": "faint",
                       "severity": "low"}]})
    monkeypatch.setattr(m, "build_provider", lambda cfg, **k: prov)
    rc = main(["galley", "verify", str(run), "--config", CONFIG,
               "--model", "gpt-5.6-luna"])
    assert rc == 0                      # a low residual is a note, not a defect
    cv = json.loads((run / "change_verify.json").read_text("utf-8"))
    fw = json.loads((run / "finished_walk.json").read_text("utf-8"))
    assert cv["ran"] is True and fw["ran"] is True
    # Each gate bills its own artifact — one call each, no double-counting.
    assert cv["cost"]["total_usd"] > 0 and fw["cost"]["total_usd"] > 0
    assert cv["cost"]["by_model"]["gpt-5.6-luna"] > 0
    assert len(prov.calls) == 2


def test_verify_fails_on_a_high_residual_or_a_flagged_edit(tmp_path, monkeypatch):
    src = _seeded_manuscript(tmp_path)
    run = _mock_review(tmp_path, src, tmp_path / "run")
    high = _Provider({"problems": []},
                     {"findings": [{"para_id": "body-0001", "quote": "thin",
                                    "problem": "wrong word", "suggestion": "x",
                                    "severity": "high"}]})
    monkeypatch.setattr(m, "build_provider", lambda cfg, **k: high)
    assert main(["galley", "verify", str(run), "--config", CONFIG,
                 "--model", "gpt-5.6-luna"]) == 1

    flagged = _Provider({"problems": [{"index": 1, "verdict": "breaks_meaning",
                                       "detail": "no", "fix": "teh"}]},
                        {"findings": []})
    monkeypatch.setattr(m, "build_provider", lambda cfg, **k: flagged)
    assert main(["galley", "verify", str(run), "--config", CONFIG,
                 "--model", "gpt-5.6-luna"]) == 1


def test_verify_dry_run_prices_the_calls_without_a_provider(tmp_path, monkeypatch,
                                                            capsys):
    src = _seeded_manuscript(tmp_path)
    run = _mock_review(tmp_path, src, tmp_path / "run")

    def _boom(*a, **k):                       # pragma: no cover - must not run
        raise AssertionError("dry-run reached the provider")
    monkeypatch.setattr(m, "build_provider", _boom)
    rc = main(["galley", "verify", str(run), "--config", CONFIG,
               "--model", "gpt-5.6-luna", "--dry-run", "--json"])
    assert rc == 0
    out = capsys.readouterr().out
    assert "1 change-verify call(s)" in out and "1 walk read(s)" in out
    line = [l for l in out.splitlines() if l.startswith("{")][-1]
    assert json.loads(line)["calls"] == 2
    assert not (run / "change_verify.json").exists()


def test_spend_guard_refuses_a_model_outside_the_approval(tmp_path, monkeypatch,
                                                          capsys):
    src = _seeded_manuscript(tmp_path)
    run = _mock_review(tmp_path, src, tmp_path / "run")
    approval = tmp_path / "approval.json"
    assert main(["galley", "approve", str(src), "--config", CONFIG,
                 "--budget", "5", "--out", str(approval)]) == 0

    def _boom(*a, **k):                       # pragma: no cover - must not run
        raise AssertionError("a refused verb reached the provider")
    monkeypatch.setattr(m, "build_provider", _boom)
    # The default config approves only its own routes; a Claude judge named on
    # the command line is outside that set.
    rc = main(["galley", "audit", str(run), "--source", str(src),
               "--config", CONFIG, "--model", "claude-fable-5",
               "--approval", str(approval)])
    assert rc == 5
    err = capsys.readouterr().err
    assert "REFUSED" in err and "model_not_approved" in err

    rc = main(["galley", "verify", str(run), "--config", CONFIG,
               "--model", "gpt-5.6-luna", "--budget", "0"])
    assert rc == 5
    assert "budget_over_cap" in capsys.readouterr().err


def test_flights_default_judge_is_the_house_detector_not_claude(tmp_path,
                                                                 monkeypatch,
                                                                 capsys):
    def _boom(cfg):                           # pragma: no cover - must not run
        raise AssertionError("dry-run must not construct a provider")
    monkeypatch.setattr(m, "_flights_provider_of", _boom)
    rc = main(["galley", "flights", str(FIXTURES / "simple.docx"), "--dry-run",
               "--out", str(tmp_path)])
    assert rc == 0
    out = capsys.readouterr().out
    assert "judge   gpt-5.6-luna" in out
    assert "claude" not in out


# --- 7. calibrate --from-run on a bare run ------------------------------------

def test_calibrate_from_run_synthesizes_a_casefile(tmp_path, capsys):
    src = _seeded_manuscript(tmp_path)
    run = _mock_review(tmp_path, src, tmp_path / "run")
    assert not (run / "casefile.json").exists()
    # A mock review bills nothing, and a $0 paid adapter is exactly what the
    # calibration store refuses to learn from (it reads as "didn't run") —
    # stamp the envelope with the spend a real run would carry.
    fj = run / "findings.json"
    payload = json.loads(fj.read_text("utf-8"))
    payload["cost"] = {"total_usd": 0.42, "by_model": {"gpt-5.6-luna": 0.42}}
    fj.write_text(json.dumps(payload), encoding="utf-8")
    store = tmp_path / "calibration.json"
    rc = main(["galley", "calibrate", str(src), "--from-run", str(run),
               "--config", CONFIG, "--calibration", str(store), "--json"])
    assert rc == 0
    captured = capsys.readouterr()
    assert "no casefile.json" in captured.err
    line = [l for l in captured.out.splitlines() if l.startswith("{")][-1]
    payload = json.loads(line)
    # The synthesized wave's one action was recorded over the whole book.
    assert payload["cost"], "the bare run's spend was not recorded"
    (entry,) = payload["cost"].values()
    assert entry["adapter"] == "review" and entry["kwords_total"] > 0
