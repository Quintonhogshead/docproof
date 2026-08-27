"""The query channel: findings that ask instead of correcting.

The house brief calls this "two channels, never blurred" — a tracked change is
a correction the author accepts or rejects, a comment is a question that edits
nothing. The tests here mostly guard the "never blurred" half: a query must
leave the document byte-identical, and must still be there to read.
"""
from __future__ import annotations

import itertools
import json
import re
import zipfile

import pytest

from docproof.analyzer import build_system_prompt
from docproof.config import Config, load_config
from docproof.error_registry import load_error_types
from docproof.models import DocumentModel, ParagraphRef, Usage
from docproof.pipeline import finish, prepare
from docproof.reassembler import paragraph_view_text
from docproof.utils.xml_helpers import DocxPackage, walk_package
from docproof.validator import validate_findings

from .test_error_types import ERROR_DIR

QUERY = frozenset({"speaker_change"})


def _doc(*texts):
    paras = [ParagraphRef(f"body-{i:04d}", "word/document.xml", "body", t,
                          "Normal") for i, t in enumerate(texts)]
    return DocumentModel(source_path="x.docx", paragraphs=tuple(paras))


def _finding(**kw):
    from docproof.models import Finding
    base = dict(finding_id="f-0001", chunk_id="chunk-000", para_id="body-0000",
                error_type="speaker_change", original_text="", occurrence=1,
                corrected_text="", explanation="ask", confidence="high")
    return Finding(**{**base, **kw})


# --- the error type contract --------------------------------------------------

def test_speaker_change_is_a_query_type():
    et = load_error_types(ERROR_DIR, ["speaker_change"])["speaker_change"]
    assert et.is_query and et.channel == "query"


def test_change_types_are_not_queries():
    et = load_error_types(ERROR_DIR, ["comma_splice"])["comma_splice"]
    assert not et.is_query


def test_an_unknown_channel_is_rejected(tmp_path):
    src = (ERROR_DIR / "speaker_change.yaml").read_text()
    (tmp_path / "speaker_change.yaml").write_text(
        src.replace("channel: query", "channel: nonsense"))
    with pytest.raises(ValueError, match="channel must be"):
        load_error_types(tmp_path, ["speaker_change"])


def test_the_prompt_tells_a_query_type_not_to_edit():
    types = list(load_error_types(ERROR_DIR, ["speaker_change"]).values())
    prompt = build_system_prompt(types)
    assert "CHANNEL: QUERY" in prompt
    assert "repeat it unchanged" in prompt
    # A change type gets no such section.
    change = build_system_prompt(
        list(load_error_types(ERROR_DIR, ["comma_splice"]).values()))
    assert "CHANNEL: QUERY" not in change


# --- the validator ------------------------------------------------------------

def test_a_query_anchors_to_its_whole_sentence():
    doc = _doc("He left. She waited a long time. Nobody came.")
    out = validate_findings(
        [_finding(original_text="She waited a long time.",
                  corrected_text="She waited a long time.")],
        doc, "medium", query_types=QUERY)
    assert out[0].status == "query"
    assert out[0].anchor.start == 9
    assert out[0].anchor.end == 9 + len("She waited a long time.")


def test_a_query_is_not_rejected_for_changing_nothing():
    """A change type repeating its sentence is a no-op and gets thrown away.
    For a query that is the whole point."""
    doc = _doc("Two speakers here.")
    same = dict(original_text="Two speakers here.",
                corrected_text="Two speakers here.")
    assert validate_findings([_finding(**same)], doc, "medium",
                             query_types=QUERY)[0].status == "query"
    assert validate_findings([_finding(error_type="comma_splice", **same)],
                             doc, "medium")[0].status == "rejected_noop"


def test_a_query_ignores_the_confidence_gate():
    """A question is the output, not a fallback from a correction, so a
    low-confidence query is still a query."""
    doc = _doc("Two speakers here.")
    out = validate_findings(
        [_finding(original_text="Two speakers here.",
                  corrected_text="Two speakers here.", confidence="low")],
        doc, "high", query_types=QUERY)
    assert out[0].status == "query"


def test_a_query_does_not_block_a_change_on_the_same_text():
    """Different channels: a comment and an edit may sit on the same words."""
    doc = _doc("The gate was open, the dogs were gone.")
    out = validate_findings([
        _finding(original_text="The gate was open, the dogs were gone.",
                 corrected_text="The gate was open, the dogs were gone."),
        _finding(finding_id="f-0002", error_type="comma_splice",
                 original_text="The gate was open, the dogs were gone.",
                 corrected_text="The gate was open; the dogs were gone."),
    ], doc, "medium", query_types=QUERY)
    assert [f.status for f in out] == ["query", "validated"]


def test_duplicate_queries_are_dropped():
    doc = _doc("Two speakers here.")
    kw = dict(original_text="Two speakers here.",
              corrected_text="Two speakers here.")
    out = validate_findings([_finding(**kw),
                             _finding(finding_id="f-0002", **kw)],
                            doc, "medium", query_types=QUERY)
    assert [f.status for f in out] == ["query", "rejected_duplicate"]


def test_two_different_queries_on_one_sentence_both_survive():
    """A speaker-change question and a term-consistency question about the same
    sentence are different questions. The dedupe key includes the error type, so
    the second is not thrown away as a duplicate — the old key collided them and
    silently dropped one."""
    doc = _doc("Two speakers here.")
    kw = dict(original_text="Two speakers here.",
              corrected_text="Two speakers here.")
    out = validate_findings(
        [_finding(**kw),
         _finding(finding_id="f-0002", error_type="term_consistency", **kw)],
        doc, "medium",
        query_types=frozenset({"speaker_change", "term_consistency"}))
    assert [f.status for f in out] == ["query", "query"]


def test_a_query_that_cannot_anchor_is_rejected():
    doc = _doc("Nothing like that here.")
    out = validate_findings([_finding(original_text="Absent sentence.",
                                      corrected_text="Absent sentence.")],
                            doc, "medium", query_types=QUERY)
    assert out[0].status == "rejected_no_anchor"


# --- into the document --------------------------------------------------------

def _run(tmp_path, paragraphs, mock, *, error_types, **cfg_kw):
    import docx
    d = docx.Document()
    for t in paragraphs:
        d.add_paragraph(t)
    src = tmp_path / "q.docx"
    d.save(src)

    cfg = load_config("config/default.yaml")
    cfg.error_types = error_types
    for k, v in cfg_kw.items():
        setattr(cfg, k, v)
    prepared = prepare(cfg, src, "config/error_types")

    from docproof.analyzer import MockAnalyzer
    findings = []
    ids = itertools.count(1)
    for group in prepared.groups:
        for chunk in prepared.chunks:
            found, _ = MockAnalyzer(group, mock, ids).analyze_chunk(chunk, Usage())
            findings += found
    out = finish(prepared, findings, Usage(), cfg, out_dir=tmp_path / "out",
                 source_path=src)
    return out


def _comment_ranges(path) -> dict[str, str]:
    """What each comment actually highlights, read back out of the file."""
    doc = zipfile.ZipFile(path).read("word/document.xml").decode()
    out = {}
    for cid in re.findall(r'commentRangeStart w:id="(\d+)"', doc):
        seg = doc.split(f'commentRangeStart w:id="{cid}"')[1] \
                 .split(f'commentRangeEnd w:id="{cid}"')[0]
        out[cid] = "".join(re.findall(r"<w:t[^>]*>(.*?)</w:t>", seg, re.S))
    return out


def _comment_texts(path) -> list[str]:
    z = zipfile.ZipFile(path)
    if "word/comments.xml" not in z.namelist():
        return []
    body = z.read("word/comments.xml").decode()
    return re.findall(r"<w:t[^>]*>(.*?)</w:t>", body, re.S)


# A JUDGMENT-case speaker query: narration between the quotes, so the
# deterministic speaker-split pass (which owns whitespace-only boundaries and
# restructures the paragraph in prepare) leaves it whole for the query channel.
SPEAKERS = "“I won’t go.” She looked away. “You will.”"


def test_a_query_writes_a_comment_and_changes_nothing(tmp_path):
    out = _run(tmp_path, [f"{SPEAKERS} He said it twice."],
               [{"para_id": "body-0000", "error_type": "speaker_change",
                 "original_text": SPEAKERS, "corrected_text": SPEAKERS,
                 "explanation": "Two speakers — new paragraph for the reply?",
                 "confidence": "high"}],
               error_types=["speaker_change"])
    assert out.applied == 0

    pkg = DocxPackage(out.reviewed_path)
    for wp in walk_package(pkg):
        # No revision markup at all: accept and reject read the same.
        assert (paragraph_view_text(wp.element, "accept")
                == paragraph_view_text(wp.element, "reject"))
    assert list(_comment_ranges(out.reviewed_path).values()) == [SPEAKERS]
    assert any("new paragraph for the reply" in t
               for t in _comment_texts(out.reviewed_path))


def test_a_below_gate_finding_becomes_a_query_comment(tmp_path):
    out = _run(tmp_path, ["The manuscript was finished, nobody wanted it."],
               [{"para_id": "body-0000", "error_type": "comma_splice",
                 "original_text": "The manuscript was finished, nobody wanted it.",
                 "corrected_text": "The manuscript was finished; nobody wanted it.",
                 "explanation": "Two independent clauses joined by a comma.",
                 "confidence": "low"}],
               error_types=["comma_splice"], not_applied_comments=True)
    assert out.applied == 0
    texts = _comment_texts(out.reviewed_path)
    assert any("left as written" in t for t in texts)
    # The suggestion travels with the question, so the author can judge it.
    assert any("Suggested:" in t for t in texts)


def test_an_oversized_edit_becomes_a_comment_not_a_silent_rejection(tmp_path):
    """The edit guard blocks a wholesale re-type from becoming a tracked change,
    but the catch is real — the fix is just too large to apply as a minimal
    edit. It is surfaced as a margin comment so nothing is lost silently."""
    original = ("i had some of that i dont know how much and found myself on "
                "the dance floor beside the bonfire")
    corrected = ("I had some of that; I don't know how much, and found myself "
                 "on the dance floor, beside the bonfire.")
    out = _run(tmp_path, [original],
               [{"para_id": "body-0000", "error_type": "comma_splice",
                 "original_text": original, "corrected_text": corrected,
                 "explanation": "Run-on.", "confidence": "high"}],
               error_types=["comma_splice"], not_applied_comments=True)
    assert out.applied == 0                       # never a tracked change
    texts = _comment_texts(out.reviewed_path)
    assert any("too large to apply" in t for t in texts)
    # The whole sentence is highlighted, not the scattered edit sites.
    assert list(_comment_ranges(out.reviewed_path).values()) == [original]
    # And it has its own summary section, not the terse rejected list.
    assert "too large to auto-correct" in out.summary_md.read_text()


def test_an_oversized_edit_is_reported_even_with_comments_off(tmp_path):
    """With query_comments off there is no margin comment, but the finding is
    still named in the summary — the information is never dropped silently."""
    original = ("i had some of that i dont know how much and found myself on "
                "the dance floor beside the bonfire")
    corrected = ("I had some of that; I don't know how much, and found myself "
                 "on the dance floor, beside the bonfire.")
    out = _run(tmp_path, [original],
               [{"para_id": "body-0000", "error_type": "comma_splice",
                 "original_text": original, "corrected_text": corrected,
                 "explanation": "Run-on.", "confidence": "high"}],
               error_types=["comma_splice"], query_comments=False)
    assert out.applied == 0
    assert _comment_texts(out.reviewed_path) == []
    assert "too large to auto-correct" in out.summary_md.read_text()


def test_a_query_comment_highlights_the_sentence_not_the_edit_site(tmp_path):
    """A below-gate comma splice would have changed one comma. A comment
    hanging off a single comma tells the author nothing."""
    text = "The manuscript was finished, nobody wanted it."
    out = _run(tmp_path, [text],
               [{"para_id": "body-0000", "error_type": "comma_splice",
                 "original_text": text,
                 "corrected_text": "The manuscript was finished; nobody wanted it.",
                 "explanation": "Splice.", "confidence": "low"}],
               error_types=["comma_splice"], not_applied_comments=True)
    assert list(_comment_ranges(out.reviewed_path).values()) == [text]


def test_query_comments_can_be_turned_off(tmp_path):
    out = _run(tmp_path, ["The manuscript was finished, nobody wanted it."],
               [{"para_id": "body-0000", "error_type": "comma_splice",
                 "original_text": "The manuscript was finished, nobody wanted it.",
                 "corrected_text": "The manuscript was finished; nobody wanted it.",
                 "explanation": "Splice.", "confidence": "low"}],
               error_types=["comma_splice"], query_comments=False)
    assert _comment_texts(out.reviewed_path) == []
    # ...but it is still reported, so nothing is lost silently.
    assert "Possibly intentional" in out.summary_md.read_text()


def test_a_below_gate_finding_is_log_only_by_default(tmp_path):
    """not_applied_comments defaults off: a declined correction leaves no margin
    comment on the document, but is still named in summary.md. The author reads
    the changes and the genuine questions, not a commentary on what was not done."""
    out = _run(tmp_path, ["The manuscript was finished, nobody wanted it."],
               [{"para_id": "body-0000", "error_type": "comma_splice",
                 "original_text": "The manuscript was finished, nobody wanted it.",
                 "corrected_text": "The manuscript was finished; nobody wanted it.",
                 "explanation": "Two clauses joined by a comma.",
                 "confidence": "low"}],
               error_types=["comma_splice"])            # default: flag off
    assert out.applied == 0
    assert _comment_texts(out.reviewed_path) == []       # nothing on the page
    assert "Possibly intentional" in out.summary_md.read_text()   # but recorded


def test_a_withheld_edit_is_log_only_by_default(tmp_path):
    """A judge/verifier withdrawal ("Not applied: …") is a declined correction,
    so it too stays out of the document by default while the report keeps it."""
    from docproof.models import Anchor
    from docproof.reassembler import apply_tracked_changes
    from docproof.validator import to_query
    text = "He left. She waited a long time. Nobody came."
    cfg, pkg, doc, pid = _one_para_docx(tmp_path, text)      # default config
    f = _finding(finding_id="w-1", para_id=pid, error_type="comma_splice",
                 original_text="He left.", corrected_text="He departed.",
                 explanation="Withheld by the gate.", status="validated",
                 anchor=Anchor(0, 8, "He left.", "He departed."))
    stats = apply_tracked_changes(pkg, doc, [to_query(f, doc)], cfg)
    assert stats.queried == () and stats.applied == ()      # nothing written
    # Turning the switch on brings it back to the margin.
    cfg.not_applied_comments = True
    cfg2, pkg2, doc2, pid2 = _one_para_docx(tmp_path, text)
    cfg2.not_applied_comments = True
    g = _finding(finding_id="w-2", para_id=pid2, error_type="comma_splice",
                 original_text="He left.", corrected_text="He departed.",
                 explanation="Withheld by the gate.", status="validated",
                 anchor=Anchor(0, 8, "He left.", "He departed."))
    stats2 = apply_tracked_changes(pkg2, doc2, [to_query(g, doc2)], cfg2)
    assert stats2.queried == ("w-2",)


def test_a_query_and_a_change_coexist_in_one_paragraph(tmp_path):
    """The ordering that matters: comments are anchored in canonical-text
    offsets, and applying a deletion moves text out of w:t. Queries therefore
    go on first, and this proves both still land."""
    text = "The gate was open, the dogs were gone. She waited for them."
    out = _run(tmp_path, [text], [
        {"para_id": "body-0000", "error_type": "comma_splice",
         "original_text": "The gate was open, the dogs were gone.",
         "corrected_text": "The gate was open; the dogs were gone.",
         "explanation": "Splice.", "confidence": "high"},
        {"para_id": "body-0000", "error_type": "speaker_change",
         "original_text": "She waited for them.",
         "corrected_text": "She waited for them.",
         "explanation": "Whose line is this?", "confidence": "high"},
    ], error_types=[["comma_splice", "speaker_change"]])

    assert out.applied == 1
    pkg = DocxPackage(out.reviewed_path)
    para = next(iter(walk_package(pkg))).element
    assert paragraph_view_text(para, "accept") == (
        "The gate was open; the dogs were gone. She waited for them.")
    assert paragraph_view_text(para, "reject") == text
    assert "She waited for them." in _comment_ranges(out.reviewed_path).values()


def test_findings_json_records_the_query_status(tmp_path):
    out = _run(tmp_path, [f"{SPEAKERS} He said it twice."],
               [{"para_id": "body-0000", "error_type": "speaker_change",
                 "original_text": SPEAKERS, "corrected_text": SPEAKERS,
                 "explanation": "Two speakers.", "confidence": "high"}],
               error_types=["speaker_change"])
    payload = json.loads(out.findings_json.read_text())
    assert payload["stats"]["query"] == 1
    assert payload["findings"][0]["status"] == "query"
    assert "Queries" in out.summary_md.read_text()


def test_summary_names_the_query_channel_in_the_formats_own_word(tmp_path):
    """Every site interpolates `fmt.comment_noun` bare, so the noun has to be
    complete on its own: prefixing "margin " at a site read "margin margin
    comment" for Word, and folding the prefix into the noun instead read
    "margin note" for InDesign, which has no margin to put a note in."""
    from docproof.formats import DOCX, IDML
    from docproof.reporting import write_summary_md

    cfg = load_config("config/default.yaml")
    # The closing line under test is the comments-on wording; the shipped
    # default now ships with edit explanations off. not_applied_comments on so
    # the below-gate and oversized sections also claim their margin comment —
    # this test is about the noun, not the gating.
    cfg.comments = True
    cfg.not_applied_comments = True
    doc = _doc(SPEAKERS)
    findings = [
        _finding(status="query"),
        _finding(finding_id="f-0002", error_type="comma_splice",
                 confidence="low", status="skipped_low_confidence"),
        _finding(finding_id="f-0003", error_type="run_on_sentence",
                 status="rejected_oversized"),
    ]
    for fmt, noun in ((DOCX, "margin comment"), (IDML, "note")):
        out = tmp_path / f"{fmt.suffix.strip('.')}.md"
        write_summary_md(out, doc=doc, findings=findings, usage=Usage(),
                         cfg=cfg, applied_ids=(), fmt=fmt)
        text = out.read_text("utf-8")
        # One per section that puts a finding in the file: queries, below-gate,
        # oversized. Plus the closing line, which names the noun on its own.
        assert text.count(f"each is a {noun} in the reviewed file") == 3
        assert f"carries a {noun} explaining itself" in text
        assert "margin margin" not in text


def test_the_default_config_ships_the_query_channel_on():
    cfg = load_config("config/default.yaml")
    assert cfg.query_comments is True
    assert "speaker_change" in cfg.error_type_keys
    assert Config().query_comments is True


# -- downgrades: the verifier's and the meaning gate's ------------------------

def test_two_distinct_downgrades_on_one_sentence_both_survive():
    """A downgrade carries a specific correction that was withheld, so two
    different withheld fixes for one sentence are two things to tell the author.
    Keying them only by (paragraph, sentence, type) collapsed them and dropped
    the second — silently losing a finding the pass had promised to keep."""
    from docproof.models import DocumentModel, Finding, ParagraphRef
    from docproof.validator import validate_findings
    text = "He could not have known the road was barred."
    para = ParagraphRef(para_id="body-0000", part="word/document.xml",
                        location="body", text=text, style="Normal")
    doc = DocumentModel(source_path="x.docx", paragraphs=(para,))

    def held(fid, corrected):
        return Finding(finding_id=fid, chunk_id="c0", para_id="body-0000",
                       error_type="spelling", original_text=text, occurrence=1,
                       corrected_text=corrected, explanation="", confidence="high",
                       force_query=True)

    out = validate_findings([held("f-1", text.replace("barred", "blocked")),
                             held("f-2", text.replace("barred", "bolted"))],
                            doc, "medium")
    assert [f.status for f in out] == ["query", "query"]

    # The same withheld fix twice is still one question, not two.
    same = text.replace("barred", "blocked")
    out2 = validate_findings([held("f-1", same), held("f-2", same)],
                             doc, "medium")
    assert [f.status for f in out2] == ["query", "rejected_duplicate"]


def test_to_query_withdraws_a_change_without_re_arbitrating_the_run():
    """The meaning gate withdraws one finding at a time so nothing else moves.
    The result is a margin question anchored on the whole sentence, with the
    span it held still spoken for."""
    from docproof.models import Anchor, DocumentModel, Finding, ParagraphRef
    from docproof.validator import to_query
    text = "He could not have known the road was barred."
    para = ParagraphRef(para_id="body-0000", part="word/document.xml",
                        location="body", text=text, style="Normal")
    doc = DocumentModel(source_path="x.docx", paragraphs=(para,))
    at = text.index("barred")
    f = Finding(finding_id="f-1", chunk_id="c0", para_id="body-0000",
                error_type="spelling", original_text=text, occurrence=1,
                corrected_text=text.replace("barred", "bared"),
                explanation="typo", confidence="high", status="validated",
                anchor=Anchor(start=at, end=at + 6, delete_text="barred",
                              insert_text="bared"))
    q = to_query(f, doc)
    assert q.status == "query" and q.force_query is True
    assert q.anchor.start == 0 and q.anchor.end == len(text)
    assert q.anchor.insert_text == ""          # a question changes no text
    assert q.corrected_text == f.corrected_text  # what was proposed is kept


def test_to_query_on_an_unanchorable_finding_falls_back_to_the_paragraph():
    """A withdrawn change whose quote no longer anchors still has to reach the
    author. Both reassemblers place a comment from the anchor and drop a query
    without one, so no anchor means no margin comment — while summary.md counts
    the question as raised. The paragraph is the coarsest honest span, and it
    still edits nothing."""
    from docproof.models import DocumentModel, Finding, ParagraphRef
    from docproof.validator import to_query
    text = "Some other text."
    para = ParagraphRef(para_id="body-0000", part="word/document.xml",
                        location="body", text=text, style="Normal")
    doc = DocumentModel(source_path="x.docx", paragraphs=(para,))
    f = Finding(finding_id="f-1", chunk_id="c0", para_id="body-0000",
                error_type="spelling", original_text="not in this paragraph",
                occurrence=1, corrected_text="x", explanation="", confidence="high",
                status="validated")
    q = to_query(f, doc)
    assert q.status == "query" and q.force_query is True
    assert q.anchor.start == 0 and q.anchor.end == len(text)
    assert q.anchor.delete_text == text and q.anchor.insert_text == ""


def test_to_query_on_an_unknown_paragraph_has_no_anchor_to_fall_back_to():
    """The one case that genuinely cannot be placed: no paragraph, so nothing
    to hang a comment on. The reassemblers count it as unplaced."""
    from docproof.models import DocumentModel, Finding, ParagraphRef
    from docproof.validator import to_query
    para = ParagraphRef(para_id="body-0000", part="word/document.xml",
                        location="body", text="Some other text.", style="Normal")
    doc = DocumentModel(source_path="x.docx", paragraphs=(para,))
    f = Finding(finding_id="f-1", chunk_id="c0", para_id="body-9999",
                error_type="spelling", original_text="anything", occurrence=1,
                corrected_text="x", explanation="", confidence="high",
                status="validated")
    q = to_query(f, doc)
    assert q.status == "query" and q.anchor is None


def _one_para_docx(tmp_path, text):
    """A one-paragraph .docx, opened as a package and a document model."""
    import docx
    from docproof.ingest import build_document_model, preflight
    d = docx.Document()
    d.add_paragraph(text)
    src = tmp_path / "q.docx"
    d.save(src)
    cfg = load_config("config/default.yaml")
    pkg = preflight(src, "abort")
    doc = build_document_model(pkg, cfg)
    pid = next(p.para_id for p in doc.paragraphs if p.text == text)
    return cfg, pkg, doc, pid


def test_a_fallback_anchored_query_reaches_the_document(tmp_path):
    """to_query's fallback has to be a comment the author can read, not merely
    a non-None anchor: a paragraph-wide span places a real one."""
    from docproof.models import Anchor
    from docproof.reassembler import apply_tracked_changes
    from docproof.validator import to_query
    text = "He left. She waited a long time. Nobody came."
    cfg, pkg, doc, pid = _one_para_docx(tmp_path, text)
    # A withheld edit reaches the margin only with this on; the test is about
    # where the fallback comment lands, not whether it is shown by default.
    cfg.not_applied_comments = True

    # A validated change whose quote no longer matches the paragraph — the
    # belt-and-braces branch of to_query.
    f = _finding(finding_id="q-x", para_id=pid, error_type="comma_splice",
                 original_text="a sentence that is not in the paragraph",
                 corrected_text="x", explanation="Withheld by the gate.",
                 status="validated",
                 anchor=Anchor(0, 8, "He left.", "He departed."))
    stats = apply_tracked_changes(pkg, doc, [to_query(f, doc)], cfg)
    assert stats.queried == ("q-x",) and stats.unplaced == ()
    assert stats.applied == () and stats.skipped == ()

    out = tmp_path / "reviewed.docx"
    pkg.save(out)
    assert list(_comment_ranges(out).values()) == [text]
    assert any("Withheld by the gate." in t for t in _comment_texts(out))


def test_a_query_with_no_anchor_is_counted_as_unplaced_not_dropped(tmp_path):
    """The reassembler places a comment from the anchor, so a query without one
    cannot be written — but summary.md still counts the question as raised. It
    has to be named somewhere, or the loss is invisible: a pass whose every
    suggestion vanished looks exactly like a healthy run that found nothing."""
    from docproof.reassembler import apply_tracked_changes
    text = "He left. She waited a long time. Nobody came."
    cfg, pkg, doc, pid = _one_para_docx(tmp_path, text)

    q = _finding(finding_id="q-x", para_id=pid, original_text="He left.",
                 status="query")
    stats = apply_tracked_changes(pkg, doc, [q], cfg)
    assert stats.queried == () and stats.unplaced == ("q-x",)
    assert stats.applied == () and stats.skipped == ()
    assert _comment_texts(tmp_path / "q.docx") == []


# --- the speaker-split pass end to end ----------------------------------------

def test_speaker_split_lands_tracked_with_a_declarative_comment(tmp_path):
    from docproof.speakersplit import SPLIT_COMMENT
    out = _run(tmp_path,
               ["Narration before.",
                "“I won’t go.” “You will, and you’ll thank me.”",
                "Narration after."],
               [], error_types=["spelling"])
    pkg = DocxPackage(out.reviewed_path)
    texts = [paragraph_view_text(wp.element, "accept")
             for wp in walk_package(pkg)]
    assert "“I won’t go.”" in texts
    assert "“You will, and you’ll thank me.”" in texts
    # The paragraph mark rides as a tracked insertion under the house author.
    doc = zipfile.ZipFile(out.reviewed_path).read("word/document.xml").decode()
    assert re.search(r'<w:pPr>.*?<w:rPr><w:ins [^>]*w:author="Atmosphere Press '
                     r'Proofreader"', doc, re.S)
    # The declarative comment reached the margin, and it states, never asks.
    comments = zipfile.ZipFile(out.reviewed_path).read(
        "word/comments.xml").decode()
    assert "separated the dialogue" in comments
    assert SPLIT_COMMENT.split("—")[0][:40] in comments
    # The reject-all audit held: the run shipped.
    assert out.reviewed_path.exists()
