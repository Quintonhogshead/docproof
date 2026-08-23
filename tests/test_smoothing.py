"""The line-editing pass: suggestions that can never become corrections.

Everything here defends one claim — that a smoothing suggestion reaches the
author as a question and nothing else, whatever the judge said about it. The
rest guards the restraint the pass promises: dialogue left alone, the author's
own words left alone, the volume capped and the cap reported.

See docproof/smoothing.py and docproof.pipeline._smoothing_findings.
"""
from __future__ import annotations

from docproof.config import Config
from docproof.models import DocumentModel, Finding, ParagraphRef, Usage
from docproof.pipeline import Prepared, _smoothing_findings
from docproof.providers.base import ProviderResult
from docproof.smoothing import (cap_for, margin_note, quote_spans,
                                rank_and_cap, touches_lexicon)
from docproof.validator import validate_findings

from .fakes import FakeProvider


def _para(pid: str, text: str) -> ParagraphRef:
    return ParagraphRef(para_id=pid, part="word/document.xml", location="body",
                        text=text, style="Normal")


def _prepared(*paras: ParagraphRef) -> Prepared:
    doc = DocumentModel(source_path="x.docx", paragraphs=tuple(paras))
    return Prepared(pkg=None, doc=doc, chunks=[], groups=[], fmt=None)


def _cfg(**kw) -> Config:
    cfg = Config()
    cfg.smoothing.enabled = True
    for k, v in kw.items():
        setattr(cfg.smoothing, k, v)
    return cfg


def _suggest(*items) -> ProviderResult:
    return ProviderResult(parsed={"suggestions": list(items)})


def _verdicts(*items) -> ProviderResult:
    return ProviderResult(parsed={"verdicts": list(items)})


def _s(para_id, quote, suggestion, category="tighten", rationale="reads tighter"):
    return {"para_id": para_id, "quote": quote, "suggestion": suggestion,
            "category": category, "rationale": rationale}


def _run(monkeypatch, cfg, prepared, results):
    """Run the pass with a scripted provider for both stages."""
    prov = FakeProvider(results)
    import docproof.providers as _p
    monkeypatch.setattr(_p, "build_provider", lambda cfg, **kw: prov)
    findings, report = _smoothing_findings(cfg, prepared, Usage())
    return findings, report, prov


# --- the invariant ------------------------------------------------------------

def test_a_smoothing_suggestion_is_never_a_tracked_change(monkeypatch):
    """The pass's whole reason to exist safely. A smoothing has no right answer,
    so however sure the judge is, it may only ever ask."""
    p = _para("body-0", "She walked over to the door in a very quiet way.")
    findings, _report, _prov = _run(monkeypatch, _cfg(), _prepared(p), [
        _suggest(_s("body-0", "in a very quiet way", "quietly")),
        _verdicts({"index": 1, "is_error": True, "confidence": "high"}),
    ])
    assert len(findings) == 1
    assert findings[0].force_query is True

    # And it survives arbitration as a question, not an edit.
    validated = validate_findings(findings, _prepared(p).doc, "low")
    assert [f.status for f in validated] == ["query"]


def test_even_a_high_confidence_verdict_cannot_reach_validated(monkeypatch):
    """Belt and braces on the same claim, stated the way it would break: the
    routing must not go through a confidence threshold, because a threshold has
    a value that clears it."""
    p = _para("body-0", "He was walking slowly down the long empty road.")
    findings, _r, _prov = _run(monkeypatch, _cfg(min_confidence="low"),
                               _prepared(p), [
        _suggest(_s("body-0", "was walking", "walked", category="aspect")),
        _verdicts({"index": 1, "is_error": True, "confidence": "high"}),
    ])
    validated = validate_findings(findings, _prepared(p).doc, "low")
    assert validated and all(f.status != "validated" for f in validated)


def test_a_suggestion_quotes_its_sentence_not_the_whole_paragraph(monkeypatch):
    """A query is anchored on what it quotes, so quoting the paragraph would
    highlight the paragraph. An author needs to be told where to look."""
    text = ("The rain had stopped. She walked over to the door in a very quiet "
            "way. Outside, the street was empty.")
    p = _para("body-0", text)
    findings, _r, _prov = _run(monkeypatch, _cfg(), _prepared(p), [
        _suggest(_s("body-0", "in a very quiet way", "quietly")),
        _verdicts({"index": 1, "is_error": True, "confidence": "high"}),
    ])
    quoted = findings[0].original_text
    assert "She walked over to the door" in quoted
    assert "The rain had stopped" not in quoted
    assert "the street was empty" not in quoted


def test_the_suggested_wording_reaches_the_margin(monkeypatch):
    """`queries.query_text` renders a query as its explanation ALONE — the
    `Suggested:` line other statuses get is unreachable for a query. So if the
    wording is not in the explanation, the author never sees what to consider."""
    from docproof.queries import query_text
    p = _para("body-0", "She walked over to the door in a very quiet way.")
    findings, _r, _prov = _run(monkeypatch, _cfg(), _prepared(p), [
        _suggest(_s("body-0", "in a very quiet way", "quietly")),
        _verdicts({"index": 1, "is_error": True, "confidence": "high"}),
    ])
    validated = validate_findings(findings, _prepared(p).doc, "low")
    assert "quietly" in query_text(validated[0])


def test_two_suggestions_on_one_sentence_stay_two_questions(monkeypatch):
    """The validator keys a force_query finding partly on corrected_text. A pass
    that emitted the original unchanged would collapse both into one and drop
    the second as a duplicate."""
    p = _para("body-0", "He was walking slowly in a very quiet way indeed.")
    # A generous cap: this test is about the validator's duplicate key, and the
    # default cap would legitimately drop the second on a nine-word document.
    findings, _r, _prov = _run(monkeypatch, _cfg(max_per_1000_words=500),
                               _prepared(p), [
        _suggest(_s("body-0", "was walking", "walked", category="aspect"),
                 _s("body-0", "in a very quiet way", "quietly")),
        _verdicts({"index": 1, "is_error": True, "confidence": "high"},
                  {"index": 2, "is_error": True, "confidence": "high"}),
    ])
    validated = validate_findings(findings, _prepared(p).doc, "low")
    assert [f.status for f in validated] == ["query", "query"]


def test_a_suggestion_reaches_the_document_as_a_comment_and_nothing_else(
        tmp_path, monkeypatch):
    """The invariant through the REAL path — prepare() and finish(), the merge
    order, the validator, the reassembler.

    Every other test here calls _smoothing_findings directly, which means the
    whole pipeline wiring could be deleted with the rest of the suite green.
    This is the one that notices."""
    import docx
    from docproof.config import load_config
    from docproof.pipeline import finish, prepare

    d = docx.Document()
    d.add_paragraph("She walked over to the door in a very quiet way.")
    src = tmp_path / "Book.docx"
    d.save(src)

    cfg = load_config("config/default.yaml")
    cfg.error_types, cfg.sweeps = [], []
    cfg.audit = "off"
    cfg.smoothing.enabled = True
    prov = FakeProvider([
        _suggest(_s("body-0000", "in a very quiet way", "quietly")),
        _verdicts({"index": 1, "is_error": True, "confidence": "high"}),
    ])
    import docproof.providers as _p
    monkeypatch.setattr(_p, "build_provider", lambda cfg, **kw: prov)

    prepared = prepare(cfg, src, "config/error_types")
    out = finish(prepared, [], Usage(), cfg, out_dir=tmp_path / "out",
                 source_path=src)

    import json
    record = json.loads(out.findings_json.read_text(encoding="utf-8"))
    smoothing = [f for f in record["findings"]
                 if f["error_type"] == "smoothing"]
    assert len(smoothing) == 1
    assert smoothing[0]["status"] == "query"
    # Nothing was applied, and the manuscript text is untouched.
    assert out.applied == 0 and out.queried == 1
    reviewed = docx.Document(out.reviewed_path)
    assert reviewed.paragraphs[0].text == \
        "She walked over to the door in a very quiet way."
    # ...but the suggestion did reach the author, wording and all.
    assert "quietly" in (tmp_path / "out" / "summary.md").read_text(
        encoding="utf-8")
    # And the record says what produced it, so a later run can be compared.
    block = record["smoothing"]
    assert block["proposed"] == 1 and block["kept"] == 1
    assert block["unjudged"] == 0
    assert len(block["propose_prompt_sha"]) == 12
    assert block["judge_model"] == cfg.smoothing.judge_model


def test_a_prompt_change_changes_the_recorded_fingerprint():
    """Two runs are only comparable when the prompts match — a prompt change
    moves the output more than any config knob. The fingerprint is what lets an
    eval notice instead of silently scoring a pre-change run against a
    post-change baseline."""
    from docproof.smoothing import PROPOSE_SYSTEM, prompt_sha
    assert prompt_sha(PROPOSE_SYSTEM) == prompt_sha(PROPOSE_SYSTEM)
    assert prompt_sha(PROPOSE_SYSTEM) != prompt_sha(PROPOSE_SYSTEM + " ")
    assert len(prompt_sha("anything")) == 12


# --- restraint ----------------------------------------------------------------

def _no_provider(monkeypatch, why: str):
    """Make BUILDING a provider the failure.

    "Patch nothing and trust it not to call" is not a test: on a machine with a
    vendor key in the environment build_provider succeeds, propose() logs the
    failed call and swallows it, and the assertion passes over a pass that ran."""
    import docproof.providers as _p

    def _never(*a, **kw):
        raise AssertionError(why)
    monkeypatch.setattr(_p, "build_provider", _never)


def test_off_by_default_spends_nothing(monkeypatch):
    cfg = Config()                                  # smoothing defaults off
    p = _para("body-0", "She walked over to the door in a very quiet way.")
    _no_provider(monkeypatch, "smoothing built a provider while disabled")
    findings, report = _smoothing_findings(cfg, _prepared(p), Usage())
    assert findings == [] and report.proposed == 0


def test_a_selected_sections_run_does_not_buy_a_whole_book_line_edit(monkeypatch):
    """Whole-document only, like every pass that reads the book entire. The
    author must not get "Consider: ..." in chapters nobody selected — and must
    not be billed for reading them."""
    import dataclasses
    p = _para("body-0", "She walked over to the door in a very quiet way.")
    prepared = dataclasses.replace(_prepared(p), whole_document=False)
    _no_provider(monkeypatch, "smoothing ran on a selected-sections review")
    findings, report = _smoothing_findings(_cfg(), prepared, Usage())
    assert findings == [] and report.proposed == 0


def test_dialogue_is_left_alone_by_default(monkeypatch):
    """A character's diction is theirs. In dialogue, awkward is often the point."""
    p = _para("body-0", '"I was just sort of going over there," he said.')
    findings, report, prov = _run(monkeypatch, _cfg(), _prepared(p), [
        _suggest(_s("body-0", "just sort of going", "going")),
    ])
    assert findings == [] and report.proposed == 0
    # The judge was never asked — the filter runs before the paid call.
    assert len(prov.calls) == 1


def test_dialogue_can_be_opted_into(monkeypatch):
    p = _para("body-0", '"I was just sort of going over there," he said.')
    findings, _r, _prov = _run(monkeypatch, _cfg(include_dialogue=True),
                               _prepared(p), [
        _suggest(_s("body-0", "just sort of going", "going")),
        _verdicts({"index": 1, "is_error": True, "confidence": "high"}),
    ])
    assert len(findings) == 1


def test_an_author_coinage_is_never_offered_a_smoothing(monkeypatch):
    """The lexicon is what the spell scan agreed to take on trust. A pass that
    offers to tighten the sentence a coined word lives in is one keystroke from
    renaming it."""
    import dataclasses

    from docproof.spellscan import SpellScan
    p = _para("body-0", "The zorrillo moved in a very quiet way across the sand.")
    prepared = dataclasses.replace(_prepared(p),
                                   spell=SpellScan(lexicon=("zorrillo",)))
    findings, report, prov = _run(monkeypatch, _cfg(), prepared, [
        _suggest(_s("body-0", "The zorrillo moved in a very quiet way",
                    "The zorrillo crossed quietly")),
    ])
    assert findings == [] and report.proposed == 0
    assert len(prov.calls) == 1                     # judge never paid for


def test_a_heading_is_neither_read_nor_reworded(monkeypatch):
    """Un-reviewable furniture is out of scope for every other pass, and a
    "Consider: ..." comment on a chapter title is the visible version of the
    same mistake."""
    import dataclasses
    body = _para("body-0", "She walked over to the door in a very quiet way.")
    head = dataclasses.replace(_para("body-1", "Chapter One: The Long Road"),
                               reviewable=False)
    findings, report, prov = _run(monkeypatch, _cfg(), _prepared(body, head), [
        _suggest(_s("body-1", "The Long Road", "A Long Road")),
    ])
    assert findings == [] and report.proposed == 0
    # The heading never even reached the proposer.
    assert "The Long Road" not in str(prov.calls[0])


def test_the_propose_effort_applies_without_a_model_override(monkeypatch):
    """`effort` and `model` are separate knobs. Tying the first to the second
    means the documented default (medium) silently never applies."""
    p = _para("body-0", "She walked over to the door in a very quiet way.")
    seen = {}

    import docproof.providers as _p

    def _build(cfg, **kw):
        seen.setdefault("effort", cfg.api.effort)
        return FakeProvider([_suggest()])
    monkeypatch.setattr(_p, "build_provider", _build)
    _smoothing_findings(_cfg(effort="xhigh"), _prepared(p), Usage())
    assert seen["effort"] == "xhigh"


def test_a_suggestion_that_does_not_anchor_is_dropped_not_guessed(monkeypatch):
    p = _para("body-0", "She walked over to the door.")
    findings, report, _prov = _run(monkeypatch, _cfg(), _prepared(p), [
        _suggest(_s("body-0", "a sentence that is not in the paragraph", "x")),
    ])
    assert findings == [] and report.proposed == 0


def test_an_oversized_rewrite_is_dropped_before_the_judge(monkeypatch):
    """Queries skip the validator's edit guard, so the minimal-edit scale is
    kept here or not at all."""
    long_quote = ("She walked over to the door and stood there for a while "
                  "wondering whether anyone would come at all")
    p = _para("body-0", long_quote + ".")
    findings, report, _prov = _run(monkeypatch, _cfg(), _prepared(p), [
        _suggest(_s("body-0", long_quote,
                    "She waited at the door, wondering if anyone would come")),
    ])
    assert findings == [] and report.proposed == 0


def test_a_long_anchor_with_a_small_change_survives_the_size_guard(monkeypatch):
    """The guard measures the SHRUNK diff, not the raw quote. A smoothing has to
    quote enough of the sentence to anchor uniquely while changing a word or two;
    measuring the raw length dropped exactly those long-anchor/small-change edits.
    Paired with the oversized test above: same guard, opposite verdict."""
    quote = ("The letter that she had been waiting for since the start of the "
             "summer finally came")
    assert len(quote) > 64                       # the raw guard would have dropped it
    p = _para("body-0", quote + ".")
    findings, report, _prov = _run(monkeypatch, _cfg(), _prepared(p), [
        _suggest(_s("body-0", quote,
                    quote.replace("finally came", "finally arrived"),
                    category="idiom")),
        _verdicts({"index": 1, "is_error": True, "confidence": "high"}),
    ])
    assert report.proposed == 1 and len(findings) == 1


def test_two_alternatives_for_one_span_both_reach_the_judge(monkeypatch):
    """The propose dedupe keys on (span, wording), not span alone, so two genuinely
    different rewrites of the same words both reach the judge — which rules on
    each. An exact duplicate would still collapse."""
    p = _para("body-0", "She walked over to the door in a very quiet way.")
    findings, report, _prov = _run(monkeypatch, _cfg(max_per_1000_words=500),
                                   _prepared(p), [
        _suggest(_s("body-0", "in a very quiet way", "quietly"),
                 _s("body-0", "in a very quiet way", "silently")),
        _verdicts({"index": 1, "is_error": True, "confidence": "high"},
                  {"index": 2, "is_error": True, "confidence": "high"}),
    ])
    assert report.proposed == 2
    # And both survive validation as two distinct margin questions.
    validated = validate_findings(findings, _prepared(p).doc, "low")
    assert [f.status for f in validated] == ["query", "query"]


def test_an_exact_duplicate_suggestion_is_still_dropped(monkeypatch):
    """The other half of the (span, wording) key: the same rewrite proposed twice
    is one question, not two."""
    p = _para("body-0", "She walked over to the door in a very quiet way.")
    _findings, report, _prov = _run(monkeypatch, _cfg(max_per_1000_words=500),
                                    _prepared(p), [
        _suggest(_s("body-0", "in a very quiet way", "quietly"),
                 _s("body-0", "in a very quiet way", "quietly")),
        _verdicts({"index": 1, "is_error": True, "confidence": "high"}),
    ])
    assert report.proposed == 1


def test_the_judge_can_refuse_a_suggestion(monkeypatch):
    p = _para("body-0", "She walked over to the door in a very quiet way.")
    findings, report, _prov = _run(monkeypatch, _cfg(), _prepared(p), [
        _suggest(_s("body-0", "in a very quiet way", "quietly")),
        _verdicts({"index": 1, "is_error": False, "confidence": "high"}),
    ])
    assert findings == [] and report.proposed == 1 and report.kept == 0


# --- the cap ------------------------------------------------------------------

def test_cap_scales_with_the_manuscript_and_never_reaches_zero():
    assert cap_for(10_000, 3.0) == 30
    assert cap_for(1_000, 3.0) == 3
    assert cap_for(50, 3.0) == 1              # a short piece still gets one


def _f(conf: str, n: int) -> Finding:
    return Finding(finding_id=f"sm-{n:04d}", chunk_id="smoothing",
                   para_id="body-0", error_type="smoothing",
                   original_text="x", occurrence=1, corrected_text="y",
                   explanation="", confidence=conf, force_query=True)


def test_the_cap_keeps_the_surest_and_counts_the_rest():
    findings = [_f("low", 1), _f("high", 2), _f("medium", 3), _f("high", 4)]
    kept, withheld, below = rank_and_cap(findings, cap=2, min_confidence="low")
    assert [f.finding_id for f in kept] == ["sm-0002", "sm-0004"]
    assert withheld == 2


def test_the_confidence_floor_drops_before_the_cap_counts():
    """A suggestion below the floor was never eligible, so it is not 'withheld' —
    reporting it as withheld would overstate what the cap cost the author."""
    findings = [_f("low", 1), _f("high", 2)]
    kept, withheld, below = rank_and_cap(findings, cap=5, min_confidence="medium")
    assert [f.finding_id for f in kept] == ["sm-0002"] and withheld == 0
    assert below == 1          # the low one was never eligible, and is counted


def test_the_cap_reports_what_it_withheld(monkeypatch):
    """A cap the author cannot see is indistinguishable from a pass that found
    little, and those are very different facts about their manuscript."""
    p = _para("body-0", "He was walking slowly in a very quiet way indeed.")
    findings, report, _prov = _run(
        monkeypatch, _cfg(max_per_1000_words=0.1), _prepared(p), [
            _suggest(_s("body-0", "was walking", "walked", category="aspect"),
                     _s("body-0", "in a very quiet way", "quietly")),
            _verdicts({"index": 1, "is_error": True, "confidence": "high"},
                      {"index": 2, "is_error": True, "confidence": "high"}),
        ])
    assert report.cap == 1
    assert len(findings) == 1 and report.kept == 1 and report.withheld == 1


def test_a_judge_that_never_answered_is_not_reported_as_restraint(monkeypatch):
    """Found on the first real book: both judge batches came back truncated,
    every candidate vanished, and the run reported "0 suggestions from 53
    proposed" — which reads as admirable restraint and was a pass that never
    ran. A failure that flatters the feature is the worst kind to leave
    unreported."""
    p = _para("body-0", "He was walking slowly in a very quiet way indeed.")
    # A reply with no verdicts at all is what truncation produces.
    findings, report, _prov = _run(monkeypatch, _cfg(), _prepared(p), [
        _suggest(_s("body-0", "was walking", "walked", category="aspect"),
                 _s("body-0", "in a very quiet way", "quietly")),
        ProviderResult(parsed={"verdicts": []}),
    ])
    assert findings == []
    assert report.proposed == 2 and report.kept == 0
    # The distinction that matters: nobody ruled on these.
    assert report.unjudged == 2


def test_a_judge_that_refused_everything_is_not_reported_as_failure(monkeypatch):
    """The other side of the same line. An honest rejection IS restraint, and
    must not be tarred as a broken run."""
    p = _para("body-0", "She walked over to the door in a very quiet way.")
    _findings, report, _prov = _run(monkeypatch, _cfg(), _prepared(p), [
        _suggest(_s("body-0", "in a very quiet way", "quietly")),
        _verdicts({"index": 1, "is_error": False, "confidence": "high"}),
    ])
    assert report.proposed == 1 and report.kept == 0 and report.unjudged == 0


def test_a_truncated_reading_pass_is_counted_not_read_as_restraint(monkeypatch):
    """The propose-side twin of the unjudged case, one stage earlier. A read that
    hits the token ceiling returns nothing parsed and drops a whole window of the
    manuscript — counted, so it is never mistaken for a quiet read."""
    p = _para("body-0", "She walked over to the door in a very quiet way.")
    findings, report, prov = _run(monkeypatch, _cfg(), _prepared(p), [
        ProviderResult(stop_reason="max_tokens"),   # the read truncated: nothing back
    ])
    assert findings == [] and report.proposed == 0
    assert report.windows == 1 and report.windows_failed == 1
    # The judge was never reached — there was nothing to rule on.
    assert len(prov.calls) == 1


def test_the_unjudged_count_reaches_summary_md(tmp_path):
    from docproof.formats import DOCX
    from docproof.reporting import write_summary_md
    from docproof.smoothing import SmoothingReport
    doc = DocumentModel(source_path="x.docx",
                        paragraphs=(_para("body-0", "Some prose here."),))
    write_summary_md(tmp_path / "summary.md", doc=doc, findings=[],
                     usage=Usage(), cfg=Config(), applied_ids=(), fmt=DOCX,
                     smoothing=SmoothingReport(proposed=53, kept=0, withheld=0,
                                               cap=3, unjudged=53))
    text = (tmp_path / "summary.md").read_text(encoding="utf-8")
    assert "53 of them were never ruled on" in text
    assert "floor" in text


def test_the_withheld_count_reaches_summary_md(tmp_path):
    from docproof.reporting import write_summary_md
    from docproof.smoothing import SmoothingReport
    doc = DocumentModel(source_path="x.docx",
                        paragraphs=(_para("body-0", "Some prose here."),))

    from docproof.formats import DOCX
    write_summary_md(tmp_path / "summary.md", doc=doc, findings=[],
                     usage=Usage(), cfg=Config(), applied_ids=(), fmt=DOCX,
                     smoothing=SmoothingReport(proposed=9, kept=3, withheld=6,
                                               cap=3))
    text = (tmp_path / "summary.md").read_text(encoding="utf-8")
    assert "6 further suggestion(s) withheld" in text


def test_a_failed_reading_pass_reaches_summary_md(tmp_path):
    """A run whose every window truncated proposes nothing, so the section must
    render on the failure count, not only on `proposed` — otherwise the outage
    hides behind the same silence a clean read produces."""
    from docproof.formats import DOCX
    from docproof.reporting import write_summary_md
    from docproof.smoothing import SmoothingReport
    doc = DocumentModel(source_path="x.docx",
                        paragraphs=(_para("body-0", "Some prose here."),))
    write_summary_md(tmp_path / "summary.md", doc=doc, findings=[],
                     usage=Usage(), cfg=Config(), applied_ids=(), fmt=DOCX,
                     smoothing=SmoothingReport(proposed=0, windows=3,
                                               windows_failed=2))
    text = (tmp_path / "summary.md").read_text(encoding="utf-8")
    assert "Language smoothing" in text
    assert "2 of 3 reading pass(es) did not complete" in text


# --- unit: the deterministic helpers ------------------------------------------

def test_quote_spans_finds_double_quoted_dialogue():
    text = 'She said, "I was going there," and left.'
    assert [text[a:b] for a, b in quote_spans(text, '”"')] == \
        ['"I was going there,"']


def test_quote_spans_does_not_mistake_an_apostrophe_for_a_uk_quote():
    """The U.K. variant's closing mark is the apostrophe character, so a
    possessive or a contraction must not end a quote — or open one."""
    text = "Sarah's brother didn't go, and that was that."
    assert quote_spans(text, "’'") == []


def test_quote_spans_reads_uk_single_quoted_dialogue():
    text = "'I didn't go,' she said."
    spans = quote_spans(text, "’'")
    assert spans and text[spans[0][0]:spans[0][1]] == "'I didn't go,'"


def test_touches_lexicon_matches_a_word_anywhere_in_the_span():
    """Deliberately broader than Sapling's filter, which only compares an edit's
    whole original against the lexicon."""
    assert touches_lexicon("the zorrillo moved", {"zorrillo"}) is True
    assert touches_lexicon("Zorrillo's tail,", {"zorrillo"}) is True
    assert touches_lexicon("the animal moved", {"zorrillo"}) is False


def test_a_coinage_glued_on_by_punctuation_is_still_found():
    """House style here is the UNSPACED em dash, so a coinage with no space
    around it is the common case, not the exotic one. Splitting the span on
    whitespace would hide every one of these from the filter."""
    lex = {"zorrillo", "kaelen"}
    assert touches_lexicon("zorrillo—a small beast—moved", lex) is True
    assert touches_lexicon("Kaelen—who spoke first", lex) is True
    assert touches_lexicon("the zorrillo-hunter came", lex) is True
    assert touches_lexicon("Zorrillo… he thought", lex) is True


def test_uk_dialogue_survives_an_elision_inside_it():
    """The U.K. closing mark is the apostrophe character, so a contraction
    mid-speech must not end the span — that would expose the rest of the line
    to smoothing, which is the failure that matters."""
    text = "'I was just goin' over there,' he said."
    spans = quote_spans(text, "’'")
    assert [text[a:b] for a, b in spans] == ["'I was just goin' over there,'"]


def test_the_margin_note_announces_itself_as_a_suggestion():
    note = margin_note("quietly", "reads tighter")
    assert note.startswith("Consider:")
    assert "quietly" in note and "reads tighter" in note


# --- isolation ----------------------------------------------------------------

def test_a_rejudge_never_pays_for_smoothing_again():
    """A re-judge inherits the original run's config; the suggestions are already
    in the record, so paying again would buy a second, differently-worded set of
    the same questions."""
    from docproof.rejudge import _gates_only
    cfg = Config()
    cfg.smoothing.enabled = True
    assert _gates_only(cfg).smoothing.enabled is False


def test_a_free_rebuild_never_pays_for_smoothing_again():
    from app.jobs import _free_finish
    cfg = Config()
    cfg.smoothing.enabled = True
    _free_finish(cfg)
    assert cfg.smoothing.enabled is False


def test_the_eval_runner_disables_smoothing():
    from docproof.eval.runner import _eval_config
    cfg = Config()
    cfg.smoothing.enabled = True
    assert _eval_config(cfg).smoothing.enabled is False


def test_every_candidate_is_accounted_for(monkeypatch):
    """The identity that makes the pass's numbers reconcilable:

        proposed == kept + withheld + below_floor + refused + unjudged

    Each term is a different thing that happened and has a different fix — a cap
    doing its job, taste, a threshold, and a fault. An unexplained remainder
    would mean a loss path nobody has found, which is exactly how the
    below-floor drop went uncounted until someone tried to make the columns
    add up."""
    p = _para("body-0", "He was walking slowly in a very quiet way indeed, he "
                        "thought, and the long road went on for a while yet.")
    findings, report, _prov = _run(
        monkeypatch, _cfg(max_per_1000_words=500, min_confidence="medium"),
        _prepared(p), [
            _suggest(_s("body-0", "was walking", "walked", category="aspect"),
                     _s("body-0", "in a very quiet way", "quietly"),
                     _s("body-0", "for a while yet", "a while longer"),
                     _s("body-0", "went on", "continued", category="idiom")),
            _verdicts({"index": 1, "is_error": True, "confidence": "high"},
                      {"index": 2, "is_error": True, "confidence": "low"},
                      {"index": 3, "is_error": False, "confidence": "high"}),
        ])
    assert report.proposed == (report.kept + report.withheld
                               + report.below_floor + report.refused
                               + report.unjudged)
    assert report.kept == 1          # the high-confidence one
    assert report.below_floor == 1   # affirmed "low", under the medium floor
    assert report.refused == 1       # the judge said no
    assert report.unjudged == 1      # never ruled on at all
    assert len(findings) == report.kept


# --- Tier C: opt-in recall levers ---------------------------------------------
#
# Every lever below is OFF by default: a run that sets none of them sends the
# shipped prompts, the shipped window size, and the total dialogue skip. The
# tests pin two things — that the levers do what they claim when turned on, and
# that with them off nothing about the pass moves, including its fingerprints
# (the eval keys run-comparability off those, so a flag that changed the output
# without changing the recorded sha would silently fold incomparable runs).

def test_c1_the_open_proposer_keeps_every_voice_rail_and_drops_only_restraint():
    """PROPOSE_SYSTEM_OPEN is a separate constant, not an edit of the shipped
    prompt: every voice-SAFETY line survives verbatim and only the restraint
    FRAMING is lifted. If this drifts, C1 has either weakened a safety rail or
    silently reverted A5's two-independent-touches relaxation."""
    from docproof.smoothing import PROPOSE_SYSTEM, PROPOSE_SYSTEM_OPEN
    assert PROPOSE_SYSTEM_OPEN != PROPOSE_SYSTEM
    # Voice-safety and scope rails — must be present in BOTH, verbatim.
    for rail in [
        "NEVER suggest anything that touches:",
        "dialogue, or any words a character speaks or thinks",
        "invented names, place names, or coined terms",
        "deliberate sentence fragments",
        "stylized, archaic, or poetic diction",
        "repetition that has rhetorical shape",
        "fit inside a single sentence, and change no more than a few words",
        "leave the meaning of the sentence identical",
        "WRONG, it is not yours to make.",   # the mechanical-error wall
        # A5's relaxation must survive the de-restraint rewrite, not revert:
        "addresses a genuinely independent, unrelated spot",
        "do not drop a real smoothing",
    ]:
        assert rail in PROPOSE_SYSTEM, rail
        assert rail in PROPOSE_SYSTEM_OPEN, rail
    # The restraint SATURATION is only in the shipped prompt, never the open one.
    for restraint in ["Your job is RESTRAINT",
                      "Most paragraphs get NOTHING",
                      "an empty result is the correct"]:
        assert restraint in PROPOSE_SYSTEM, restraint
        assert restraint not in PROPOSE_SYSTEM_OPEN, restraint


def test_c1_open_sends_and_fingerprints_the_open_prompt(monkeypatch):
    """The lever swaps the prompt AND records the sha of the prompt it actually
    sent. The failure this guards is subtle and worse than no fingerprint: swap
    the prompt at the call site but fingerprint the shipped one, and two runs
    with different prompts record the same sha — the eval folds them together."""
    from docproof.smoothing import (PROPOSE_SYSTEM, PROPOSE_SYSTEM_OPEN,
                                    prompt_sha)
    p = _para("body-0", "She walked over to the door in a very quiet way.")
    script = lambda: [_suggest(_s("body-0", "in a very quiet way", "quietly")),
                      _verdicts({"index": 1, "is_error": True,
                                 "confidence": "high"})]
    _f, rep_r, prov_r = _run(monkeypatch, _cfg(), _prepared(p), script())
    assert prov_r.calls[0]["system"] == PROPOSE_SYSTEM
    assert rep_r.propose_prompt_sha == prompt_sha(PROPOSE_SYSTEM)

    _f2, rep_o, prov_o = _run(monkeypatch, _cfg(proposer_restraint="open"),
                              _prepared(p), script())
    assert prov_o.calls[0]["system"] == PROPOSE_SYSTEM_OPEN
    assert rep_o.propose_prompt_sha == prompt_sha(PROPOSE_SYSTEM_OPEN)
    assert rep_o.propose_prompt_sha != rep_r.propose_prompt_sha


def test_c1_an_explicit_propose_prompt_still_wins_over_the_flag(monkeypatch):
    """`propose_prompt` is the wholesale override and outranks the restraint
    flag, on both the send and the fingerprint."""
    from docproof.smoothing import prompt_sha
    p = _para("body-0", "She walked over to the door in a very quiet way.")
    custom = "You are a custom proposer. Return nothing."
    _f, rep, prov = _run(monkeypatch,
                         _cfg(proposer_restraint="open", propose_prompt=custom),
                         _prepared(p), [_suggest()])
    assert prov.calls[0]["system"] == custom
    assert rep.propose_prompt_sha == prompt_sha(custom)


def test_c2_smaller_windows_split_more_and_lose_no_paragraph():
    """The satisficing probe: a smaller window makes more reads. Every paragraph
    still lands in exactly one window, both ways — the size only changes where
    the cuts fall, never what is covered."""
    from docproof.smoothing import _windows
    paras = [_para(f"body-{i}", "word " * 200) for i in range(30)]  # ~1000ch ea
    big = _windows(paras, 12_000, 60)
    small = _windows(paras, 5_000, 24)
    assert len(small) > len(big)
    assert sum(len(w) for w in big) == 30 == sum(len(w) for w in small)


def test_c2_the_window_config_reaches_propose(monkeypatch):
    """The config values thread all the way to `_windows`: same three big
    paragraphs, more reads under the smaller setting."""
    paras = [_para(f"body-{i}", "x " * 2500) for i in range(3)]   # ~5000ch each
    _fd, _rd, dprov = _run(monkeypatch, _cfg(), _prepared(*paras),
                           [_suggest(), _suggest(), _suggest()])
    _fs, _rs, sprov = _run(monkeypatch,
                           _cfg(propose_chars=5_000, propose_max_paras=24),
                           _prepared(*paras),
                           [_suggest(), _suggest(), _suggest()])
    assert len(dprov.calls) < len(sprov.calls)   # calls == propose windows here


def test_c4_the_harshness_dial_has_four_distinct_rungs_and_strict_is_the_default():
    """A selector like the effort knob: four levels, and 'strict' IS the shipped
    JUDGE_SYSTEM by identity so the default judge is unchanged."""
    from docproof.smoothing import (JUDGE_SYSTEM, JUDGE_SYSTEMS, prompt_sha)
    assert list(JUDGE_SYSTEMS) == ["lenient", "balanced", "strict", "severe"]
    assert JUDGE_SYSTEMS["strict"] is JUDGE_SYSTEM        # byte-identical default
    assert Config().smoothing.judge_harshness == "strict"
    # every rung is a genuinely different prompt
    shas = {lvl: prompt_sha(txt) for lvl, txt in JUDGE_SYSTEMS.items()}
    assert len(set(shas.values())) == 4


def test_c4_every_harshness_level_keeps_the_voice_safety_vetoes():
    """Leniency buys back the merely-conventional and preference rejects, never
    the voice line: the three voice-SAFETY vetoes appear verbatim at EVERY rung,
    including the most lenient. This is the invariant that keeps the dial safe."""
    from docproof.smoothing import JUDGE_SYSTEMS
    for level, prompt in JUDGE_SYSTEMS.items():
        for veto in [
            "touches dialect, idiolect, a coined term, or a character's voice",
            "alters a deliberate fragment, or repetition with rhetorical shape",
            "changes the meaning, the emphasis, or the rhythm of the sentence",
        ]:
            assert veto in prompt, (level, veto)


def test_c4_the_dial_is_monotonic_in_disposition():
    """Each rung's disposition is where the harshness lives, and they ascend:
    lenient leans to keep, strict defaults to no, severe holds hardest."""
    from docproof.smoothing import JUDGE_SYSTEMS
    assert "Lean toward keeping" in JUDGE_SYSTEMS["lenient"]
    assert "need not reject\nmost" in JUDGE_SYSTEMS["balanced"] \
        or "need not reject" in JUDGE_SYSTEMS["balanced"]
    assert "DEFAULT TO NO." in JUDGE_SYSTEMS["strict"]
    assert "Expect to reject most items." in JUDGE_SYSTEMS["strict"]
    assert "hold that line harder" in JUDGE_SYSTEMS["severe"]
    assert "reject nearly every item" in JUDGE_SYSTEMS["severe"]
    # DEFAULT TO NO belongs to the harsh end only, not the keeping end
    assert "DEFAULT TO NO" not in JUDGE_SYSTEMS["lenient"]
    assert "DEFAULT TO NO" not in JUDGE_SYSTEMS["balanced"]


def test_c4_the_selected_harshness_reaches_the_judge_and_moves_the_fingerprint(
        monkeypatch):
    """With no book context folded in, the recorded judge sha IS the sha of the
    chosen rung — so a sha match proves the selected prompt was the one handed to
    the valve, and default 'strict' records the shipped judge's sha exactly."""
    from docproof.smoothing import JUDGE_SYSTEMS, prompt_sha
    p = _para("body-0", "She walked over to the door in a very quiet way.")
    script = lambda: [_suggest(_s("body-0", "in a very quiet way", "quietly")),
                      _verdicts({"index": 1, "is_error": True,
                                 "confidence": "high"})]
    seen = {}
    for level in JUDGE_SYSTEMS:
        _f, rep, _pv = _run(monkeypatch, _cfg(judge_harshness=level),
                            _prepared(p), script())
        assert rep.judge_prompt_sha == prompt_sha(JUDGE_SYSTEMS[level]), level
        seen[level] = rep.judge_prompt_sha
    # default (no override) matches the strict rung exactly
    _f, rep_default, _pv = _run(monkeypatch, _cfg(), _prepared(p), script())
    assert rep_default.judge_prompt_sha == seen["strict"]
    assert len(set(seen.values())) == 4       # four rungs, four fingerprints


def test_c4_an_explicit_judge_prompt_still_wins_over_the_harshness_dial(
        monkeypatch):
    """`judge_prompt` is the wholesale override and outranks the dial, on the
    fingerprint too."""
    from docproof.smoothing import prompt_sha
    p = _para("body-0", "She walked over to the door in a very quiet way.")
    custom = "You are a custom judge. Keep nothing."
    _f, rep, _pv = _run(
        monkeypatch, _cfg(judge_harshness="lenient", judge_prompt=custom),
        _prepared(p), [_suggest(_s("body-0", "in a very quiet way", "quietly")),
                       _verdicts({"index": 1, "is_error": True,
                                  "confidence": "high"})])
    assert rep.judge_prompt_sha == prompt_sha(custom)


def test_c5_clarity_survives_in_dialogue_when_opted_in(monkeypatch):
    """The middle setting between the total skip and include_dialogue: an
    ambiguous pronoun inside speech reaches the judge, which is the taste gate
    for whether it is a real error or the character's diction."""
    p = _para("body-0", '"He gave it to him," she said.')
    findings, report, prov = _run(
        monkeypatch, _cfg(dialogue_categories=["clarity"]), _prepared(p), [
            _suggest(_s("body-0", "He gave it to him", "He gave it to the boy",
                        category="clarity")),
            _verdicts({"index": 1, "is_error": True, "confidence": "high"}),
        ])
    assert len(findings) == 1 and report.proposed == 1
    assert len(prov.calls) == 2          # the judge WAS paid — it reached it


def test_c5_a_non_clarity_fix_in_dialogue_is_still_dropped_and_counted(
        monkeypatch):
    """Opting clarity in does not open the door to tightening a character's
    line: a tighten candidate in dialogue is still dropped before the judge, and
    still counted in `filtered` so the deterministic tally stays honest."""
    p = _para("body-0", '"I was just sort of going over there," he said.')
    findings, report, prov = _run(
        monkeypatch, _cfg(dialogue_categories=["clarity"]), _prepared(p), [
            _suggest(_s("body-0", "just sort of going", "going",
                        category="tighten")),
        ])
    assert findings == [] and report.proposed == 0
    assert report.filtered == 1          # dropped, and accounted for
    assert len(prov.calls) == 1          # judge never paid


def test_c5_clarity_in_dialogue_is_dropped_by_default(monkeypatch):
    """With the default empty list, dialogue is skipped wholesale — clarity
    included. The opt-in is the only thing that changes this."""
    p = _para("body-0", '"He gave it to him," she said.')
    findings, report, prov = _run(monkeypatch, _cfg(), _prepared(p), [
        _suggest(_s("body-0", "He gave it to him", "He gave it to the boy",
                    category="clarity")),
    ])
    assert findings == [] and report.proposed == 0
    assert len(prov.calls) == 1


def test_c5_dialogue_categories_refuses_anything_but_clarity():
    """Widening dialogue smoothing past clarity is the voice risk the skip
    exists to prevent, so it is refused at config load, not silently accepted."""
    import pytest

    from docproof.config import SmoothingConfig
    SmoothingConfig(dialogue_categories=[])            # the default, fine
    SmoothingConfig(dialogue_categories=["clarity"])   # the one allowed opt-in
    for bad in (["idiom"], ["flow"], ["clarity", "tighten"]):
        with pytest.raises(ValueError):
            SmoothingConfig(dialogue_categories=bad)


def test_every_lever_off_by_default_leaves_the_shipped_pass_untouched():
    """The whole safety story in one assertion: the shipped config sends the
    shipped prompts and the shipped window, and skips dialogue entirely."""
    from docproof.smoothing import (JUDGE_SYSTEM, PROPOSE_SYSTEM,
                                    _PROPOSE_CHARS, _PROPOSE_MAX_PARAS)
    s = Config().smoothing
    assert s.proposer_restraint == "restrained"
    assert s.judge_harshness == "strict"
    assert (s.propose_chars, s.propose_max_paras) == (_PROPOSE_CHARS,
                                                      _PROPOSE_MAX_PARAS)
    assert s.dialogue_categories == []
    # and the chosen constants are the shipped ones (referenced so a rename here
    # can't quietly drift the defaults away from what the pass ships)
    assert PROPOSE_SYSTEM and JUDGE_SYSTEM


# --- the edits switch ---------------------------------------------------------

def test_edits_switch_applies_high_confidence_smoothings_as_tracked_changes(
        monkeypatch):
    """smoothing.edits flips the valve from ask to apply: a judge affirmation
    held at HIGH confidence becomes an ordinary edit finding the author
    accepts or rejects in Word."""
    p = _para("body-0", "She walked over to the door in a very quiet way.")
    findings, _report, _prov = _run(
        monkeypatch, _cfg(edits=True), _prepared(p), [
            _suggest(_s("body-0", "in a very quiet way", "quietly")),
            _verdicts({"index": 1, "is_error": True, "confidence": "high"}),
        ])
    assert len(findings) == 1
    assert findings[0].force_query is False
    assert "quietly" in findings[0].corrected_text


def test_edits_switch_still_queries_below_high_confidence(monkeypatch):
    p = _para("body-0", "She walked over to the door in a very quiet way.")
    findings, _report, _prov = _run(
        monkeypatch, _cfg(edits=True), _prepared(p), [
            _suggest(_s("body-0", "in a very quiet way", "quietly")),
            _verdicts({"index": 1, "is_error": True, "confidence": "medium"}),
        ])
    assert len(findings) == 1
    assert findings[0].force_query is True


def test_edits_defaults_off_and_has_a_feature_switch():
    from app.features import FEATURES_BY_ID
    from docproof.config import Config
    assert Config().smoothing.edits is False
    spec = FEATURES_BY_ID["smoothing_edits"]
    assert spec.path == ("smoothing", "edits")
