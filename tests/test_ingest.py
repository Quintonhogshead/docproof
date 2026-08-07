"""Ingest: which paragraphs the model reviews, which the sweeps reach, and
which nobody touches — decided by paragraph style.

The distinction that matters for recall: a chapter heading is set text the
model has no business rewriting, but it still carries mechanical fixes a sweep
can make ("Chapter Twenty Four" needs a hyphen). So a heading is swept, not
modelled. A table of contents, which regenerates from the headings, is left
alone entirely.
"""
from __future__ import annotations

import docx

from docproof.config import load_config
from docproof.ingest import build_document_model
from docproof.sweeps import run_sweeps
from docproof.utils.xml_helpers import DocxPackage

CONFIG = "config/default.yaml"
BODY = "A body paragraph long enough for the model to actually review it."


def _model(tmp_path, *paras):
    """paras: (text, style-name) pairs -> a built DocumentModel."""
    d = docx.Document()
    for text, style in paras:
        d.add_paragraph(text, style=style)
    path = tmp_path / "doc.docx"
    d.save(path)
    return build_document_model(DocxPackage(path), load_config(CONFIG))


def test_a_heading_is_swept_but_not_modelled(tmp_path):
    doc = _model(tmp_path, ("Chapter Twenty Four", "Heading 1"), (BODY, "Normal"))
    by_text = {p.text: p for p in doc.paragraphs}
    # Present (so the sweeps see it) but not reviewable (so the model does not).
    assert by_text["Chapter Twenty Four"].reviewable is False
    assert by_text[BODY].reviewable is True


def test_a_title_is_skipped_entirely(tmp_path):
    doc = _model(tmp_path, ("The Grand Title", "Title"), (BODY, "Normal"))
    assert "The Grand Title" not in {p.text for p in doc.paragraphs}
    assert any(reason == "style:Title" for _, reason in doc.skipped)


def test_a_heading_now_gets_its_sweep_fix(tmp_path):
    """The whole point of the change: the compound-number fix a chapter heading
    carries is now made. Before, the heading was skipped and it was unreachable."""
    doc = _model(tmp_path, ("Chapter Thirty One", "Heading 1"))
    sweep_only = [p for p in doc.paragraphs if not p.reviewable]
    findings, _ = run_sweeps(sweep_only, ["sweep_compound_number"])
    assert [f.corrected_text for f in findings] == ["Chapter Thirty-One"]


def test_skip_config_classifies_styles():
    skip = load_config(CONFIG).skip
    assert skip.is_sweep_only("Heading1") and skip.is_sweep_only("Heading2")
    assert skip.fully_skipped("Title") and skip.fully_skipped("TOC1")
    # Heading moved out of the full-skip list into sweep-only.
    assert not skip.fully_skipped("Heading1")
    assert not skip.is_sweep_only("Normal")
