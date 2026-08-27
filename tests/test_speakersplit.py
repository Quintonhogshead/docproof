"""The speaker-split pass: the certain two-speaker boundary split as a tracked
paragraph break, with the declarative comment riding the findings channel.

The negative cases matter most: narration between the quotes, tagged speakers,
possessive apostrophes near a boundary, tables — everything the pass must leave
to the speaker_change query rather than restructure.
"""
from __future__ import annotations

import docx
import pytest

from docproof.models import ParagraphRef
from docproof.speakersplit import (SPLIT_COMMENT, find_split_offsets,
                                   split_findings, split_package)
from docproof.utils.xml_helpers import (DocxPackage, paragraph_text, qn,
                                        walk_package)
from docproof.variants import load_variant


def _pkg(tmp_path, *texts):
    d = docx.Document()
    for t in texts:
        d.add_paragraph(t)
    path = tmp_path / "w.docx"
    d.save(path)
    return DocxPackage(path)


def _body_texts(pkg):
    return [paragraph_text(wp.element) for wp in walk_package(pkg)
            if wp.location == "body" and paragraph_text(wp.element)]


# --- the detector -------------------------------------------------------------

@pytest.mark.parametrize("text,n", [
    ("“I won't go.” “You will, and you'll thank me.”", 1),
    ("“Stop!” “No.” “Please.”", 2),                        # three speakers
    ("“Where?” he asked. “Home,” she said.", 0),           # narration between
    ("“I won't go.” She looked away. “You can't make me.”", 0),
    ("He told me “go home” and hung up.", 0),              # one pair
    ("“I won't go.” “You will.", 0),                       # unbalanced: not ours
    ('"I won\'t go." "You will."', 0),                     # straight: not settled
])
def test_find_split_offsets(text, n):
    assert len(find_split_offsets(text)) == n


def test_offsets_point_at_the_opening_quote():
    text = "“I won't go.” “You will.”"
    (off,) = find_split_offsets(text)
    assert text[off] == "“"


def test_single_primary_variant_uses_its_own_marks():
    uk = load_variant("uk")
    text = "‘I won’t go.’ ‘You will.’"
    assert len(find_split_offsets(text, uk)) == 1
    # The apostrophe shape near a boundary never matches: no terminal
    # punctuation before the ’.
    assert find_split_offsets("‘the boys’ ‘ouse was dark.’", uk) == []


# --- the surgery --------------------------------------------------------------

def test_split_makes_two_paragraphs_with_a_tracked_mark(tmp_path):
    pkg = _pkg(tmp_path, "Before.",
               "“I won't go.” “You will, and you'll thank me.”", "After.")
    records = split_package(pkg, None, author="Proofreader")
    texts = _body_texts(pkg)
    assert texts == ["Before.", "“I won't go.” ",
                     "“You will, and you'll thank me.”", "After."]
    # The FIRST fragment's paragraph mark carries the insertion.
    marked = [wp for wp in walk_package(pkg)
              if paragraph_text(wp.element) == "“I won't go.” "]
    ins = marked[0].element.find(
        f"{qn('w:pPr')}/{qn('w:rPr')}/{qn('w:ins')}")
    assert ins is not None and ins.get(qn("w:author")) == "Proofreader"
    # The record names the SECOND paragraph, by its post-split para_id.
    assert [r.text for r in records] == ["“You will, and you'll thank me.”"]
    by_id = {wp.para_id: paragraph_text(wp.element) for wp in walk_package(pkg)}
    assert by_id[records[0].para_id] == records[0].text


def test_three_speakers_split_twice_marks_both_fragments(tmp_path):
    pkg = _pkg(tmp_path, "“Stop!” “No.” “Please.”")
    records = split_package(pkg, None, author="P")
    assert _body_texts(pkg) == ["“Stop!” ", "“No.” ", "“Please.”"]
    assert len(records) == 2
    marks = [wp.element.find(f"{qn('w:pPr')}/{qn('w:rPr')}/{qn('w:ins')}")
             for wp in walk_package(pkg) if paragraph_text(wp.element)]
    assert [m is not None for m in marks] == [True, True, False]
    ids = {m.get(qn("w:id")) for m in marks if m is not None}
    assert len(ids) == 2                       # revision ids stay unique


def test_split_across_runs_preserves_the_boundary(tmp_path):
    d = docx.Document()
    p = d.add_paragraph()
    p.add_run("“I won't go.” “You ")
    p.add_run("will.”")
    path = tmp_path / "w.docx"
    d.save(path)
    pkg = DocxPackage(path)
    split_package(pkg, None, author="P")
    assert _body_texts(pkg) == ["“I won't go.” ", "“You will.”"]


def test_negative_paragraphs_left_whole(tmp_path):
    pkg = _pkg(tmp_path, "“Where?” he asked. “Home,” she said.",
               "“I won't go.” She looked away. “You can't make me.”")
    assert split_package(pkg, None, author="P") == []
    assert len(_body_texts(pkg)) == 2


def test_max_splits_caps_the_pass(tmp_path):
    pkg = _pkg(tmp_path, "“A.” “B.”", "“C.” “D.”")
    records = split_package(pkg, None, author="P", max_splits=1)
    assert len(records) == 1


# --- the comment --------------------------------------------------------------

def test_split_findings_are_declarative_queries(tmp_path):
    pkg = _pkg(tmp_path, "“I won't go.” “You will.”")
    records = split_package(pkg, None, author="P")
    paras = [ParagraphRef(wp.para_id, wp.part, wp.location,
                          paragraph_text(wp.element), "Normal")
             for wp in walk_package(pkg)]
    (f,) = split_findings(records, paras)
    assert f.force_query and f.error_type == "speaker_split"
    assert f.original_text == f.corrected_text        # a query changes nothing
    assert f.explanation == SPLIT_COMMENT
    assert "?" not in f.explanation                   # states, never asks
