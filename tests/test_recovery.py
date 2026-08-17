"""Semantic retry and coverage accounting.

A provider that refuses, truncates, or drops a request used to lose the whole
(pass, chunk) with only a log line — a silent recall hole. run_sync now retries
once (splitting on a truncation, re-asking on a refusal) and records whatever it
still cannot cover, so the report can name the unreviewed paragraphs instead of
letting them read as clean.
"""
from __future__ import annotations

import re
from collections import defaultdict
from pathlib import Path

import docx

from docproof.batch import _assemble, custom_id
from docproof.config import load_config
from docproof.models import CoverageLedger, Usage
from docproof.pipeline import finish, prepare, run_sync
from docproof.providers import NormalizedUsage, ProviderResult

ROOT = Path(__file__).parent.parent
ERROR_DIR = ROOT / "config" / "error_types"
CONFIG = str(ROOT / "config" / "default.yaml")

U = NormalizedUsage(input_tokens=10, output_tokens=5)


def _target_ids(user: str) -> list[str]:
    """The paragraph ids under review — after the read-only context block."""
    return re.findall(r'<paragraph id="([^"]+)"', user.split("</context>")[-1])


def _finding_for(pid: str) -> dict:
    return {"para_id": pid, "error_type": "spelling", "original_text": "teh",
            "corrected_text": "the", "confidence": "high", "explanation": ""}


def _prepared(tmp_path, texts, **chunking):
    d = docx.Document()
    for t in texts:
        d.add_paragraph(t)
    src = tmp_path / "m.docx"
    d.save(src)
    cfg = load_config(CONFIG)
    cfg.error_types = [["spelling"]]
    for k, v in chunking.items():
        setattr(cfg.chunking, k, v)
    return cfg, prepare(cfg, str(src), ERROR_DIR)


# --- the synchronous path -----------------------------------------------------

class _RefuseOnce:
    """Refuses the first call for each paragraph, answers the second."""
    name = "fake"

    def __init__(self):
        self.calls = defaultdict(int)

    def complete_structured(self, *, user, **kw) -> ProviderResult:
        pid = _target_ids(user)[0]
        self.calls[pid] += 1
        if self.calls[pid] == 1:
            return ProviderResult(stop_reason="refusal", error="no", usage=U)
        return ProviderResult(parsed={"findings": [_finding_for(pid)]}, usage=U)


def test_a_refusal_is_recovered_by_one_retry(tmp_path):
    cfg, prepared = _prepared(
        tmp_path,
        [f"Paragraph {i} has teh typo in it." for i in range(3)],
        token_budget=1)                                # one paragraph per chunk
    assert len(prepared.chunks) == 3
    provider = _RefuseOnce()
    coverage = CoverageLedger()

    findings, _ = run_sync(cfg, prepared, provider, coverage=coverage)

    assert len(findings) == 3                          # nothing dropped
    assert all(provider.calls[pid] == 2 for pid in provider.calls)
    assert coverage.complete and coverage.total == 3


class _TruncateThenSplit:
    """Truncates any multi-paragraph request; answers a single paragraph."""
    name = "fake"

    def complete_structured(self, *, user, **kw) -> ProviderResult:
        ids = _target_ids(user)
        if len(ids) > 1:
            return ProviderResult(stop_reason="max_tokens", usage=U)
        return ProviderResult(parsed={"findings": [_finding_for(ids[0])]},
                              usage=U)


def test_a_truncation_is_recovered_by_splitting(tmp_path):
    # Two short paragraphs pack into one chunk under the default budget.
    cfg, prepared = _prepared(tmp_path,
                              ["First line has teh typo here.",
                               "Second line has teh typo too."])
    assert len(prepared.chunks) == 1
    assert len(prepared.chunks[0].paragraphs) == 2
    coverage = CoverageLedger()

    findings, _ = run_sync(cfg, prepared, _TruncateThenSplit(), coverage=coverage)

    # The split reached each paragraph on its own, so both are recovered.
    assert {f.para_id for f in findings} == {"body-0000", "body-0001"}
    assert coverage.complete


class _AlwaysRefuse:
    name = "fake"

    def complete_structured(self, *, user, **kw) -> ProviderResult:
        return ProviderResult(stop_reason="refusal", error="no", usage=U)


def test_an_unrecoverable_chunk_becomes_a_coverage_gap(tmp_path):
    cfg, prepared = _prepared(tmp_path, ["Paragraph with teh typo in it."])
    coverage = CoverageLedger()

    findings, _ = run_sync(cfg, prepared, _AlwaysRefuse(), coverage=coverage)

    assert findings == []                              # still nothing to apply
    assert not coverage.complete
    assert coverage.total == 1 and len(coverage.gaps) == 1
    assert coverage.gaps[0].para_ids == ("body-0000",)
    assert coverage.gaps[0].pass_label == "spelling"


class _AlwaysAnswer:
    name = "fake"

    def complete_structured(self, *, user, **kw) -> ProviderResult:
        return ProviderResult(parsed={"findings": [_finding_for(_target_ids(user)[0])]},
                              usage=U)


def test_a_clean_run_reports_full_coverage(tmp_path):
    cfg, prepared = _prepared(tmp_path,
                              [f"Line {i} has teh typo." for i in range(4)],
                              token_budget=1)
    coverage = CoverageLedger()
    run_sync(cfg, prepared, _AlwaysAnswer(), coverage=coverage)
    assert coverage.complete
    assert coverage.total == 4 and coverage.reviewed == 4


def test_on_phase_reports_the_detector_step(tmp_path):
    """The step callback fires as the run enters the per-chunk loop, so a caller
    can show which stage the document is at."""
    cfg, prepared = _prepared(tmp_path, ["Line has teh typo."])
    cfg.glossary.enabled = False       # keep to the local provider, no key needed
    phases = []
    run_sync(cfg, prepared, _AlwaysAnswer(), on_phase=phases.append)
    assert phases and phases[0] == "reviewing"


# --- the batch path -----------------------------------------------------------

def test_batch_collect_recovers_a_missing_result(tmp_path):
    """A request the batch never returned is recovered with a synchronous call
    at collect time, rather than left as a silent hole."""
    cfg, prepared = _prepared(tmp_path,
                              [f"Line {i} has teh typo." for i in range(3)],
                              token_budget=1)
    # Simulate a batch that returned results for chunks 0 and 2 but dropped 1.
    results = {}
    for c in (prepared.chunks[0], prepared.chunks[2]):
        pid = c.paragraphs[0].para_id
        results[custom_id(0, c.chunk_id)] = ProviderResult(
            parsed={"findings": [_finding_for(pid)]}, usage=U)

    coverage = CoverageLedger()
    findings, _ = _assemble(cfg, prepared, results, _AlwaysAnswer(), coverage)

    # All three paragraphs are present — the missing one recovered inline.
    assert {f.para_id for f in findings} == {"body-0000", "body-0001", "body-0002"}
    assert coverage.complete and coverage.total == 3


def test_batch_collect_without_a_provider_records_the_gap(tmp_path):
    """No provider to recover with (older call sites): the missing request is
    recorded as a gap rather than dropped silently."""
    cfg, prepared = _prepared(tmp_path,
                              [f"Line {i} has teh typo." for i in range(2)],
                              token_budget=1)
    results = {custom_id(0, prepared.chunks[0].chunk_id): ProviderResult(
        parsed={"findings": [_finding_for(prepared.chunks[0].paragraphs[0].para_id)]},
        usage=U)}
    coverage = CoverageLedger()
    findings, _ = _assemble(cfg, prepared, results, None, coverage)

    assert {f.para_id for f in findings} == {"body-0000"}
    assert len(coverage.gaps) == 1 and coverage.gaps[0].para_ids == ("body-0001",)


# --- the report names the gap -------------------------------------------------

def test_summary_names_unreviewed_paragraphs(tmp_path):
    cfg, prepared = _prepared(tmp_path,
                              ["A paragraph with teh typo to review here."])
    coverage = CoverageLedger()
    findings, usage = run_sync(cfg, prepared, _AlwaysRefuse(), coverage=coverage)
    out = finish(prepared, findings, usage, cfg, out_dir=tmp_path / "out",
                 source_path=tmp_path / "m.docx", coverage=coverage)
    text = out.summary_md.read_text()
    assert "## Coverage" in text
    assert "could not be reviewed" in text
    assert "not checked" in text
    assert "body-0000" in text                         # the specific paragraph
    # The same gap is machine-readable in findings.json — both outputs, so a
    # downstream tool sees it too, not just a human reading the summary.
    import json as jsonlib
    cov = jsonlib.loads(out.findings_json.read_text())["coverage"]
    assert cov["reviewed"] < cov["total"]
    assert cov["gaps"][0]["para_ids"] == ["body-0000"]


def test_summary_confirms_full_coverage(tmp_path):
    cfg, prepared = _prepared(tmp_path,
                              ["A paragraph with teh typo to review here."])
    coverage = CoverageLedger()
    findings, usage = run_sync(cfg, prepared, _AlwaysAnswer(), coverage=coverage)
    out = finish(prepared, findings, usage, cfg, out_dir=tmp_path / "out",
                 source_path=tmp_path / "m.docx", coverage=coverage)
    text = out.summary_md.read_text()
    assert "## Coverage" in text
    assert "were reviewed" in text


def test_a_degraded_pass_is_named_everywhere(tmp_path):
    """A paid pass that fell open and produced nothing reaches all three
    surfaces — summary.md, findings.json, and Outputs.warnings (the job card) —
    so a 'done' run that quietly skipped it never reads as a clean one."""
    cfg, prepared = _prepared(tmp_path,
                              ["A paragraph with teh typo to review here."])
    cfg.glossary.enabled = False       # isolate the injected degradation below
    coverage = CoverageLedger()
    findings, usage = run_sync(cfg, prepared, _AlwaysAnswer(), coverage=coverage)
    coverage.record_degraded("continuity read", "provider error 401")
    out = finish(prepared, findings, usage, cfg, out_dir=tmp_path / "out",
                 source_path=tmp_path / "m.docx", coverage=coverage)
    text = out.summary_md.read_text()
    assert "did not run" in text and "continuity read" in text
    import json as jsonlib
    cov = jsonlib.loads(out.findings_json.read_text())["coverage"]
    assert cov["degraded"] == [{"pass": "continuity read",
                                "reason": "provider error 401"}]
    assert out.warnings == ["continuity read: provider error 401"]
