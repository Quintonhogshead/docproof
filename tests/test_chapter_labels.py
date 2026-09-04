"""Chapter/part labels are mechanics, not facts (Quinton, 2026-09-04): a gap,
a mangled number, a style drift is renumbered and restyled as a tracked
heading edit — never an author query. The sequence is the Georgis one.
"""
from __future__ import annotations

import json

import docx

from docproof import editmap as emap
from docproof.chapter_labels import (dominant_style, is_chapter_label,
                                     label_map, parse_number, render,
                                     renumber_rows, spell_number)
from docproof.models import ParagraphRef

GEORGIS = ["Chapter Fifteen", "Chapter Seventeen", "Chapter Eighteen",
           "Chapter Nineteen", "Chapter Twenty-One", "Chapter Twenty-Thirty"]


def _paras(*texts: str, style: str = "Heading 1") -> list[ParagraphRef]:
    return [ParagraphRef(f"h-{i:03d}", "word/document.xml", "body", t, style)
            for i, t in enumerate(texts)]


def test_number_parsing_and_spelling():
    assert parse_number("Fifteen") == 15 and parse_number("twenty-one") == 21
    assert parse_number("Twenty One") == 21 and parse_number("17") == 17
    assert parse_number("XV") == 15 and parse_number("iv") == 4
    assert parse_number("Twenty-Thirty") is None       # not one number
    assert spell_number(20) == "twenty" and spell_number(21) == "twenty-one"
    assert render("Chapter", 21, "spelled", "title") == "Chapter Twenty-One"
    assert render("CHAPTER", 21, "spelled", "upper") == "CHAPTER TWENTY-ONE"
    assert render("Part", 3, "roman", "title") == "Part III"
    assert render("chapter", 4, "numeral", "lower") == "chapter 4"


def test_the_georgis_sequence_is_renumbered_continuously():
    labels = label_map(_paras(*GEORGIS))
    assert [lb.number for lb in labels] == [15, 17, 18, 19, 21, None]
    rows = renumber_rows(labels)
    assert [(r["original_text"], r["corrected_text"]) for r in rows] == [
        ("Chapter Seventeen", "Chapter Sixteen"),
        ("Chapter Eighteen", "Chapter Seventeen"),
        ("Chapter Nineteen", "Chapter Eighteen"),
        ("Chapter Twenty-One", "Chapter Nineteen"),
        ("Chapter Twenty-Thirty", "Chapter Twenty"),
    ]
    assert all(r["error_type"] == "chapter_label" and r["confidence"] == "high"
               for r in rows)
    assert "mechanics" in rows[0]["explanation"]
    # a continuous sequence produces no rows
    assert renumber_rows(label_map(_paras("Chapter One", "Chapter Two"))) == []


def test_style_drift_matches_the_dominant_style_and_titles_survive():
    paras = _paras("PART ONE", "PART TWO", "PART 3", "Chapter One: The Sea",
                   "chapter 2: The Shore", "Chapter Three")
    labels = label_map(paras)
    parts = [lb for lb in labels if lb.kind == "part"]
    assert dominant_style(parts) == ("spelled", "upper")
    rows = renumber_rows(labels)
    by = {r["original_text"]: r["corrected_text"] for r in rows}
    assert by["PART 3"] == "PART THREE"
    assert by["chapter 2"] == "Chapter Two"
    assert len(rows) == 2
    # the row replaces the label span only; the title after the colon stays
    assert rows[1]["original_text"] == "chapter 2"


def test_is_chapter_label_is_a_heading_line_not_prose():
    assert is_chapter_label("Chapter Seventeen")
    assert is_chapter_label("PART 3")
    assert is_chapter_label("Chapter Twenty-One: The Storm")
    assert not is_chapter_label("Chapter three of the report says the levee "
                                "held, which surprised nobody who had "
                                "walked it that spring.")
    assert not is_chapter_label("She closed the chapter on him.")


def test_settle_does_not_demote_a_label_renumbering_to_a_query():
    from galley.settle import Residual, decide
    source = {"h-001": "Chapter Seventeen", "b-001": "He counted 7 coins."}
    em = emap.build_editmap(source, [])
    res = Residual(id="r1", kind="residual", para_id="h-001",
                   quote="Seventeen", problem="numbering gap",
                   suggestion="Sixteen")
    dec = decide(res, em, dict(source), source, {}, None)
    assert dec.action == "add" and dec.replacement == "Sixteen"
    # the same kind of change in prose is still a fact -> query
    res2 = Residual(id="r2", kind="residual", para_id="b-001",
                    quote="7 coins", problem="count", suggestion="6 coins")
    dec2 = decide(res2, em, dict(source), source, {}, None)
    assert dec2.action == "query" and dec2.reason.startswith("fact:")


def test_the_flights_judge_does_not_demote_a_label_change():
    import itertools

    from docproof import flights as fl

    def cluster(cid, pid, text, original, replacement):
        start = text.index(original)
        return fl.Cluster.from_json({
            "para_id": pid, "start": start, "end": start + len(original),
            "original": original, "sentence": text, "para_text": text,
            "options": [{"para_id": pid, "start": start,
                         "end": start + len(original), "original": original,
                         "replacement": replacement, "rationale": "r",
                         "model": "m", "lens": "mechanics"}]})

    def verdict(chosen):
        return fl._Verdict(verdict="pick", chosen_index=0, chosen_text=chosen,
                           confidence="high", reason="numbering")

    label = cluster("c1", "h-001", "Chapter Seventeen", "Seventeen", "Sixteen")
    prose = cluster("c2", "b-001", "He counted 7 coins.", "7 coins",
                    "6 coins")
    out = fl.findings_from_accepted(
        [(label, verdict("Sixteen")), (prose, verdict("6 coins"))],
        itertools.count(1))
    by = {f.para_id: f for f in out}
    assert by["h-001"].force_query is False
    assert by["h-001"].corrected_text == "Chapter Sixteen"
    assert by["b-001"].force_query is True             # prose facts still ask


def test_galley_profile_emits_the_rows(tmp_path, capsys):
    from docproof.__main__ import main
    d = docx.Document()
    for t in GEORGIS:
        d.add_paragraph(t, style="Heading 1")
        d.add_paragraph("A paragraph of the chapter that says enough words to count.")
    src = tmp_path / "Georgis - Book 1.docx"
    d.save(src)
    rows_path = tmp_path / "chapter_rows.json"
    rc = main(["galley", "profile", str(src), "--chapter-rows", str(rows_path)])
    assert rc == 0
    out = capsys.readouterr().out
    assert "5 chapter/part label(s) out of sequence" in out
    rows = json.loads(rows_path.read_text("utf-8"))
    assert [r["corrected_text"] for r in rows] == [
        "Chapter Sixteen", "Chapter Seventeen", "Chapter Eighteen",
        "Chapter Nineteen", "Chapter Twenty"]
