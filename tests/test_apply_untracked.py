"""apply_untracked — building a round's working document (design doc stage C).

Multi-round review corrects the manuscript between rounds and reviews the
corrected text, so each round needs a working .docx whose *text* is the current
render of the edit layer. These edits are applied with no revision markup (the
delivered tracked changes are composed against the original at the very end), so
the checks here are: the text comes out exactly as the layer would render it,
run formatting survives, nothing tracked remains (the next round's preflight
refuses a document that already has tracked changes), and the drift guard holds.
"""
from __future__ import annotations

import docx

from docproof.editlayer import Contribution, EditLayer, RoundEdit
from docproof.models import DocumentModel, ParagraphRef
from docproof.reassembler import (apply_untracked, paragraph_view_text)
from docproof.utils.xml_helpers import (DEL_TAG, INS_TAG, DocxPackage,
                                        paragraph_text, walk_package)


def _pkg(tmp_path, *texts, runs=None):
    """A .docx built by python-docx so the runs are real Word runs. Pass `runs`
    (a list of run-strings) for a single paragraph split across several runs."""
    d = docx.Document()
    if runs is not None:
        p = d.add_paragraph()
        for r in runs:
            p.add_run(r)
    else:
        for t in texts:
            d.add_paragraph(t)
    path = tmp_path / "w.docx"
    d.save(path)
    return DocxPackage(path)


def _doc(pkg):
    paras = tuple(
        ParagraphRef(wp.para_id, wp.part, "body", paragraph_text(wp.element),
                     "Normal")
        for wp in walk_package(pkg))
    return DocumentModel("w.docx", paras)


def _texts(pkg):
    return [paragraph_text(wp.element) for wp in walk_package(pkg)]


def _ids(pkg):
    return [wp.para_id for wp in walk_package(pkg)]


def _no_revisions(pkg):
    for wp in walk_package(pkg):
        if list(wp.element.iter(INS_TAG)) or list(wp.element.iter(DEL_TAG)):
            return False
    return True


def _re(a, b, repl, rnd=1, fid="f-1"):
    return RoundEdit(a, b, repl, Contribution(rnd, fid, "typo", ""))


# --- the basic edit kinds ----------------------------------------------------

def test_replacement(tmp_path):
    pkg = _pkg(tmp_path, "hello world")
    pid = _ids(pkg)[0]
    apply_untracked(pkg, _doc(pkg), {pid: [(6, 11, "there")]})
    assert _texts(pkg) == ["hello there"]
    assert _no_revisions(pkg)


def test_insertion(tmp_path):
    pkg = _pkg(tmp_path, "the cat sat")
    pid = _ids(pkg)[0]
    apply_untracked(pkg, _doc(pkg), {pid: [(7, 7, " quietly")]})
    assert _texts(pkg) == ["the cat quietly sat"]
    assert _no_revisions(pkg)


def test_deletion(tmp_path):
    pkg = _pkg(tmp_path, "a b c")
    pid = _ids(pkg)[0]
    apply_untracked(pkg, _doc(pkg), {pid: [(1, 3, "")]})
    assert _texts(pkg) == ["a c"]
    assert _no_revisions(pkg)


def test_multiple_disjoint_edits_in_one_paragraph(tmp_path):
    pkg = _pkg(tmp_path, "the cat sat")
    pid = _ids(pkg)[0]
    apply_untracked(pkg, _doc(pkg),
                    {pid: [(4, 7, "dog"), (8, 11, "ran"), (3, 3, " big")]})
    assert _texts(pkg) == ["the big dog ran"]
    assert _no_revisions(pkg)


def test_edits_across_several_paragraphs(tmp_path):
    pkg = _pkg(tmp_path, "first line", "second line", "third line")
    p0, p1, p2 = _ids(pkg)
    apply_untracked(pkg, _doc(pkg),
                    {p0: [(0, 5, "worst")], p2: [(6, 10, "row")]})
    assert _texts(pkg) == ["worst line", "second line", "third row"]


# --- run boundaries ----------------------------------------------------------

def test_insertion_at_a_run_boundary(tmp_path):
    # "He said " | "stop" | " now" — insert "!" exactly at the stop/now boundary
    # (offset 12), the case a naive in-place rewrite would drop.
    pkg = _pkg(tmp_path, runs=["He said ", "stop", " now"])
    assert _texts(pkg) == ["He said stop now"]
    pid = _ids(pkg)[0]
    apply_untracked(pkg, _doc(pkg), {pid: [(12, 12, "!")]})
    assert _texts(pkg) == ["He said stop! now"]
    assert _no_revisions(pkg)


def test_replacement_spanning_a_run_boundary(tmp_path):
    pkg = _pkg(tmp_path, runs=["quick ", "brown ", "fox"])
    pid = _ids(pkg)[0]
    # replace "brown fox" (spans two runs) with "red hen"
    apply_untracked(pkg, _doc(pkg), {pid: [(6, 15, "red hen")]})
    assert _texts(pkg) == ["quick red hen"]
    assert _no_revisions(pkg)


def test_run_formatting_survives(tmp_path):
    # the middle run is italic; an edit elsewhere must not disturb it.
    d = docx.Document()
    p = d.add_paragraph()
    p.add_run("plain ")
    p.add_run("italic").italic = True
    p.add_run(" tail")
    path = tmp_path / "fmt.docx"
    d.save(path)
    pkg = DocxPackage(path)
    pid = _ids(pkg)[0]
    apply_untracked(pkg, _doc(pkg), {pid: [(0, 5, "PLAIN")]})
    assert _texts(pkg) == ["PLAIN italic tail"]
    # the italic run is still italic
    italic = [wp for wp in walk_package(pkg)][0].element
    from docproof.utils.xml_helpers import RPR_TAG, qn
    italics = [r for r in italic.iter(qn("w:r"))
               if r.find(RPR_TAG) is not None
               and r.find(RPR_TAG).find(qn("w:i")) is not None]
    assert italics, "italic run lost its formatting"


# --- the invariant that ties stage C to the edit layer -----------------------

def test_matches_editlayer_render(tmp_path):
    orig = "the small grey cat sat still"
    layer = EditLayer()
    # round 1: two edits
    layer = layer.fold_round(orig, [
        _re(4, 9, "tiny", fid="f-1"),
        _re(10, 14, "black", fid="f-2")]).layer            # "the tiny black cat sat still"
    # round 2 (coords of the round-1 render): edit inside round-1's "tiny",
    # insert a word, and delete one
    working = layer.render(orig)
    assert working == "the tiny black cat sat still"
    layer = layer.fold_round(orig, [
        _re(4, 8, "teeny", rnd=2, fid="f-3"),              # inside "tiny" -> composes
        _re(18, 18, " very", rnd=2, fid="f-4"),            # insert before " cat"? position
        _re(23, 27, "lay", rnd=2, fid="f-5")]).layer
    edits = [(e.orig_start, e.orig_end, e.replacement) for e in layer.edits]

    pkg = _pkg(tmp_path, orig)
    pid = _ids(pkg)[0]
    apply_untracked(pkg, _doc(pkg), {pid: edits})
    assert _texts(pkg) == [layer.render(orig)]
    assert _no_revisions(pkg)
    # accept and reject views agree, since there is nothing tracked
    el = [wp for wp in walk_package(pkg)][0].element
    assert paragraph_view_text(el, "accept") == paragraph_view_text(el, "reject")
    assert paragraph_view_text(el, "accept") == layer.render(orig)


# --- guards ------------------------------------------------------------------

def test_drift_refuses_the_edit(tmp_path):
    pkg = _pkg(tmp_path, "hello world")
    pid = _ids(pkg)[0]
    doc = _doc(pkg)
    # a document model that disagrees with the real paragraph text
    bad = DocumentModel("w.docx", (ParagraphRef(pid, doc.paragraphs[0].part,
                                                "body", "different text",
                                                "Normal"),))
    apply_untracked(pkg, bad, {pid: [(0, 5, "HELLO")]})
    assert _texts(pkg) == ["hello world"]                  # untouched


def test_empty_and_unknown_are_noops(tmp_path):
    pkg = _pkg(tmp_path, "hello world")
    pid = _ids(pkg)[0]
    doc = _doc(pkg)
    apply_untracked(pkg, doc, {})                          # empty
    apply_untracked(pkg, doc, {pid: []})                   # empty list
    apply_untracked(pkg, doc, {"body-9999": [(0, 1, "X")]})  # unknown para
    assert _texts(pkg) == ["hello world"]
