"""The scripted sweeps.

Two things are being pinned here, and the second matters more than the first.
One is that each sweep fixes what its house rule says it fixes. The other is
that it stays off everything else: a sweep runs unsupervised over an author's
whole manuscript, so a false positive is not a missed catch, it is an edit
nobody asked for.
"""
from __future__ import annotations

import pytest

from docproof.config import Config
from docproof.models import DocumentModel, ParagraphRef
from docproof.sweeps import (SWEEPS, SWEEPS_BY_KEY, _sweep_ellipsis, apply_hits,
                             ordinal_word, run_sweeps)
from docproof.validator import validate_findings

ALL = [s.key for s in SWEEPS]


def swept(key: str, text: str) -> str:
    """The text as that one sweep would leave it."""
    return apply_hits(text, SWEEPS_BY_KEY[key].scan(text))


def unchanged(key: str, text: str) -> bool:
    return not SWEEPS_BY_KEY[key].scan(text)


# --- ellipsis ----------------------------------------------------------------
#
# House style is a non-breaking space before the ellipsis and a plain space
# after (Bad\u00a0\u2026 she trailed off), so that is the sweep's default. Only
# the leading space is configurable \u2014 see the style test at the end.

@pytest.mark.parametrize("before,after", [
    ("She hesitated... then went on.", "She hesitated\u00a0\u2026 then went on."),
    ("She hesitated . . . then went on.", "She hesitated\u00a0\u2026 then went on."),
    # A run of dots jammed against the next word gets the glyph, the leading
    # non-breaking space, and a plain trailing space.
    ("She hesitated....then went on.", "She hesitated\u00a0\u2026 then went on."),
    # A bare ellipsis missing the required non-breaking space is a target too:
    # the brief says to check pre-existing \u2026 characters.
    ("Bad\u2026 she trailed off.", "Bad\u00a0\u2026 she trailed off."),
    # No space is added before a closing quote, only before the ellipsis.
    ("\u201cI don't know\u2026\u201d she said.", "\u201cI don't know\u00a0\u2026\u201d she said."),
])
def test_ellipsis_normalizes_glyph_and_spacing(before, after):
    assert swept("sweep_ellipsis", before) == after


def test_ellipsis_leaves_a_house_style_one_alone():
    assert unchanged("sweep_ellipsis", "Bad\u00a0\u2026 she trailed off.")


def test_ellipsis_at_the_start_of_a_paragraph_gets_no_leading_space():
    """There is nothing for the non-breaking space to hold on to at the start
    of a paragraph, so only the trailing space is added, in every style."""
    assert swept("sweep_ellipsis", "\u2026and then nothing.") == "\u2026 and then nothing."


@pytest.mark.parametrize("style,after", [
    ("nbsp", "I\u00a0\u2026 guess"),     # house style: a non-breaking space before
    ("closed", "I\u2026 guess"),          # no space before
    ("space", "I \u2026 guess"),          # a plain space before
])
def test_ellipsis_style_chooses_the_leading_space(style, after):
    """style.ellipsis picks the lead; the trailing space before the next word
    is the same in every mode."""
    text = "I... guess"
    assert apply_hits(text, _sweep_ellipsis(text, None, style)) == after


def test_ellipsis_closed_mode_closes_up_a_house_style_ellipsis():
    """The escape hatch: a book that sets ellipses closed up gets the leading
    non-breaking space removed, and an already-closed one is left alone."""
    text = "I\u00a0\u2026 guess"
    assert apply_hits(text, _sweep_ellipsis(text, None, "closed")) == "I\u2026 guess"
    assert not _sweep_ellipsis("I\u2026 guess", None, "closed")


# --- dashes ------------------------------------------------------------------

@pytest.mark.parametrize("before,after", [
    # Between numbers a hyphen run is a range, so it becomes an en dash.
    ("The pages ran 128--129.", "The pages ran 128–129."),
    ("Chapters 1--4 are missing.", "Chapters 1–4 are missing."),
    # Anywhere else it is a sentence break, so it becomes an unspaced em dash.
    ("The forest---which had no name---was quiet.",
     "The forest—which had no name—was quiet."),
    ("He turned -- and stopped.", "He turned—and stopped."),
    # A spaced en dash is an em dash doing the wrong job.
    ("He waited – and waited – for an answer.",
     "He waited—and waited—for an answer."),
    # A pre-existing spaced em dash is house style set the wrong way; close
    # it up. (The proofreader in the field does exactly this by hand.)
    ("Hannah — only nineteen — smiled.", "Hannah—only nineteen—smiled."),
    ("He turned — and stopped.", "He turned—and stopped."),
])
def test_dash_conversions(before, after):
    assert swept("sweep_dash", before) == after


@pytest.mark.parametrize("text", [
    "The range was 1999–2005 exactly.",       # a correct tight en dash
    "She was a well-known author.",           # a single hyphen is a hyphen
    "post–World War II era",                  # open-compound modifier
    "The forest—quiet—waited.",              # a correct tight em dash
    "* * *",                                  # a scene divider, not prose
])
def test_dash_leaves_correct_punctuation_alone(text):
    assert unchanged("sweep_dash", text)


def test_a_line_of_only_dashes_is_a_divider():
    assert unchanged("sweep_dash", "-----")


# --- stacked punctuation -----------------------------------------------------

@pytest.mark.parametrize("before,after", [
    ("You can't be serious!!", "You can't be serious!"),
    ("What?!", "What?"),          # a question with feeling is still a question
    ("What!?", "What?"),
    ("Really???", "Really?"),
    ("He said what‽", "He said what?"),
])
def test_stacked_punctuation_collapses(before, after):
    assert swept("sweep_stacked_punctuation", before) == after


def test_single_marks_are_left_alone():
    assert unchanged("sweep_stacked_punctuation", "Stop! Who goes there?")


# --- doubled words -----------------------------------------------------------

def test_doubled_word_is_removed():
    assert (swept("sweep_doubled_word", "He went to to the the door.")
            == "He went to the door.")


def test_doubled_word_keeps_the_first_capitalization():
    assert swept("sweep_doubled_word", "The the door was open.") == "The door was open."


@pytest.mark.parametrize("text", [
    "She had had enough of it.",              # past perfect
    "He knew that that was the problem.",     # awkward but correct
    "“No, no, no,” she said.",                # punctuation between: not doubled
    "It was very very cold.",                 # deliberate emphasis
    "The door. The door was open.",           # a sentence boundary between
])
def test_doubled_word_leaves_legitimate_repetition_alone(text):
    assert unchanged("sweep_doubled_word", text)


# --- centuries ---------------------------------------------------------------

@pytest.mark.parametrize("n,word", [
    (1, "first"), (9, "ninth"), (12, "twelfth"), (20, "twentieth"),
    (21, "twenty-first"), (45, "forty-fifth"), (99, "ninety-ninth"),
])
def test_ordinal_words(n, word):
    assert ordinal_word(n) == word


def test_ordinal_word_declines_out_of_range():
    assert ordinal_word(0) is None and ordinal_word(100) is None


@pytest.mark.parametrize("before,after", [
    ("He collected 20th century pulp.", "He collected twentieth century pulp."),
    ("It was 21st-century medicine.", "It was twenty-first-century medicine."),
    ("They met in the 19th Century.", "They met in the nineteenth Century."),
])
def test_century_is_spelled_out(before, after):
    assert swept("sweep_century", before) == after


def test_a_century_opening_a_sentence_keeps_its_capital():
    assert (swept("sweep_century", "20th century readers knew better.")
            == "Twentieth century readers knew better.")


def test_century_leaves_other_ordinals_alone():
    assert unchanged("sweep_century", "She came in 4th and never forgot it.")


# --- compound numbers --------------------------------------------------------

@pytest.mark.parametrize("before,after", [
    ("Chapter Twenty Four", "Chapter Twenty-Four"),
    ("seventy five people", "seventy-five people"),
    ("It cost thirty two dollars.", "It cost thirty-two dollars."),
    # Case is taken from the text, on both words.
    ("TWENTY FOUR", "TWENTY-FOUR"),
    ("Twenty four", "Twenty-four"),
])
def test_compound_number_is_hyphenated(before, after):
    assert swept("sweep_compound_number", before) == after


@pytest.mark.parametrize("text", [
    "twenty-four already",                    # already hyphenated
    "one hundred and twenty",                 # a bare ten, no ones word
    "twenty four-year-olds",                  # ambiguous: leave a following hyphen
    "He counted twenty. Four remained.",      # a sentence break, not a compound
    "forty",                                  # a ten on its own
])
def test_compound_number_leaves_others_alone(text):
    assert unchanged("sweep_compound_number", text)


# --- dialogue tags -----------------------------------------------------------
#
# The house brief lays this out as a table and warns that building it from ad
# hoc regexes is how the period + capitalized row gets missed. Every row is
# therefore asserted explicitly, including that one.

@pytest.mark.parametrize("before,after", [
    # period + lowercase → the period should be a comma
    ("“Wait here.” he said.", "“Wait here,” he said."),
    # period + CAPITALIZED → both halves are wrong
    ("“Wait here.” He said, and meant it.", "“Wait here,” he said, and meant it."),
    # comma + CAPITALIZED → lowercase the pronoun
    ("“Not tonight,” She whispered.", "“Not tonight,” she whispered."),
    ("“Who is it?” They asked.", "“Who is it?” they asked."),
    # punctuation outside the quote belongs inside it
    ("“Wait here”, he said.", "“Wait here,” he said."),
    ("“Wait here”. he said.", "“Wait here,” he said."),
    # no punctuation at all before the tag
    ("“Wait here” he said.", "“Wait here,” he said."),
])
def test_dialogue_tag_table(before, after):
    assert swept("sweep_dialogue_tag", before) == after


@pytest.mark.parametrize("text", [
    # correct already: comma / question / exclamation + lowercase pronoun
    "“Wait here,” he said.",
    "“Who's there?” he asked.",
    "“Get out!” she shouted.",
    # an action beat is not a tag, and its period is correct
    "“Wait here.” She turned toward the stairs.",
    "“Wait here.” He was already gone.",
    # a reporting verb taking an object is narration about the speech
    "“I told you already.” She said it again.",
    "“I told you already.” She said nothing more.",
    # "I" is always capitalized
    "“Stop that,” I said.",
    "“Stop that.” I said, and meant it.",
    # a named subject is the error type's judgment call, not the sweep's —
    # lowercasing a name would be worse than missing the comma
    "“Don't move.” Sarah said from the doorway.",
])
def test_dialogue_tag_leaves_these_alone(text):
    assert unchanged("sweep_dialogue_tag", text)


def test_dialogue_tag_never_lowercases_the_pronoun_I():
    assert "i said" not in swept("sweep_dialogue_tag", "“Wait here.” I said, quietly.")


def test_dialogue_tag_leaves_an_action_beat_gerund_alone():
    # A reporting verb run straight into a present participle is an action beat,
    # not a tag: "She continued typing". The period is a sentence break and must
    # not become a comma (the real regression on a manuscript).
    for text in ("“…behind this.” She continued typing, entering commands.",
                 "“Not now.” He kept reading."):
        assert unchanged("sweep_dialogue_tag", text)


# --- times of day ------------------------------------------------------------

@pytest.mark.parametrize("before,after", [
    ("It was 3:40AM.", "It was 3:40 a.m."),
    ("at 4:15PM.", "at 4:15 p.m."),
    ("reached the tree at 1:53AM.", "reached the tree at 1:53 a.m."),
    ("meet at 2:00 AM tonight", "meet at 2:00 a.m. tonight"),
    ("by 2 PM sharp", "by 2 p.m. sharp"),
])
def test_time_of_day_meridiem(before, after):
    assert swept("sweep_time_of_day", before) == after


@pytest.mark.parametrize("text", [
    "already 3:40 a.m. now",          # correct: left alone
    "I AM here and you are not.",     # "AM" not attached to a digit
    "listening to AM radio",
])
def test_time_of_day_leaves_these_alone(text):
    assert unchanged("sweep_time_of_day", text)


# --- the deity capital -------------------------------------------------------

@pytest.mark.parametrize("before,after", [
    ("Oh my god, no!", "Oh my God, no!"),
    ("“Oh my god,” she said.", "“Oh my God,” she said."),
    ("thank god for that", "thank God for that"),
    ("Good god, what now!", "Good God, what now!"),
])
def test_deity_capital(before, after):
    assert swept("sweep_deity_capital", before) == after


@pytest.mark.parametrize("text", [
    "a good god of war",              # common noun: article + "of"
    "the god smiled down",            # article: common noun
    "the gods were angry",            # plural
    "a godforsaken place",            # not a whole word
    "It felt like a godsend.",
    "god damn it",                    # a consistency question, not a cap fix
    "“Oh my God,” she said.",         # already correct: no hit
])
def test_deity_leaves_these_alone(text):
    assert unchanged("sweep_deity_capital", text)


# --- dialogue splice ---------------------------------------------------------

@pytest.mark.parametrize("before,after", [
    # a tag after a sentence-final quote, joined to the next quote by a comma
    ('“Of course!” Raymond said, “Anything for you.”',
     '“Of course!” Raymond said. “Anything for you.”'),
    ('“This is amazing!” she continued, “I called in.”',
     '“This is amazing!” she continued. “I called in.”'),
    ('“Do you?” he asked angrily, “Do you understand?”',
     '“Do you?” he asked angrily. “Do you understand?”'),
    # a physical-action beat mistaken for a tag
    ('“At this point,” Raymond smiled.', '“At this point.” Raymond smiled.'),
    ('“Yeah,” she nodded.', '“Yeah.” She nodded.'),
])
def test_dialogue_splice(before, after):
    assert swept("sweep_dialogue_splice", before) == after


@pytest.mark.parametrize("text", [
    "“I know,” she said, “it just doesn't add up.”",   # continuation: comma right
    "“Yes,” she said.",                                # real tag: leave it
    "“Wait!” he said, hoping she would stop.",         # verb not before a quote
    "“Really?” she asked, tilting her head.",          # tag + participial, not a quote
])
def test_dialogue_splice_leaves_these_alone(text):
    assert unchanged("sweep_dialogue_splice", text)


# --- the engine --------------------------------------------------------------

def _doc(*texts: str) -> tuple[DocumentModel, list[ParagraphRef]]:
    paras = [ParagraphRef(f"body-{i:04d}", "word/document.xml", "body", t,
                          "Normal") for i, t in enumerate(texts)]
    return DocumentModel(source_path="x.docx", paragraphs=tuple(paras)), paras


def test_every_sweep_is_idempotent_so_remaining_is_zero():
    """`remaining` is the number the change log quotes. If a sweep's own
    output still matches its pattern, that claim is false."""
    _, paras = _doc(
        "“Wait here.” he said... The forest---which had no name---was quiet.",
        "He collected 20th century pulp and went to to the door. What?! No!!",
        "“You promised me.” She said, and meant it.",
    )
    _, reports = run_sweeps(paras, ALL)
    assert [r.key for r in reports] == ALL
    assert sum(r.flagged for r in reports) > 0
    for r in reports:
        assert r.remaining == 0, f"{r.key} still matches its own output"


def test_findings_anchor_and_apply_cleanly_through_the_validator():
    """The sweeps' whole reason for producing Findings is to ride the same
    validator as the model. An edit that cannot anchor is worth nothing."""
    doc, paras = _doc(
        "“I don't know…” She said. The forest---which had no name---was quiet. "
        "He went to to the 20th century door!!")
    findings, _ = run_sweeps(paras, ALL)
    validated = validate_findings(findings, doc, "medium")
    assert validated and all(f.status == "validated" for f in validated)

    text = paras[0].text
    edits = sorted((f.anchor for f in validated), key=lambda a: a.start)
    out, last = [], 0
    for a in edits:
        assert text[a.start:a.end] == a.delete_text
        out.append(text[last:a.start])
        out.append(a.insert_text)
        last = a.end
    out.append(text[last:])
    assert "".join(out) == (
        "“I don't know …” she said. The forest—which had no name—was quiet. "
        "He went to the twentieth century door!")


def test_overlaps_treats_a_boundary_touching_insertion_as_a_conflict():
    """VAL-001 unit guard. An insertion at a span's START enters the half-open
    range [start, end) and conflicts (both orders); one that abuts the END sits
    after the span and composes cleanly, so it does not."""
    from docproof.validator import _overlaps

    assert _overlaps(3, 3, 3, 8)          # insertion at the span's start
    assert _overlaps(3, 8, 3, 3)          # ...and the symmetric check order
    assert _overlaps(5, 5, 3, 8)          # insertion strictly inside
    assert not _overlaps(8, 8, 3, 8)      # insertion abutting the span's end
    assert not _overlaps(3, 8, 8, 8)      # ...symmetric
    assert _overlaps(4, 4, 4, 4)          # two insertions at the same point
    assert not _overlaps(4, 4, 5, 5)      # two insertions at different points
    assert _overlaps(3, 6, 5, 8)          # two spans that intersect
    assert not _overlaps(3, 5, 5, 8)      # two spans that only touch end-to-start


def test_a_boundary_touching_insertion_and_retype_do_not_both_apply():
    """VAL-001 regression. A pure insertion at offset X and a replacement whose
    span starts at X used to BOTH validate and compose to duplicated text on
    accept-all ("to to", ",," — seen in the Johnson run's body-0856/body-0445).
    They now conflict: the first claims the span, the second is rejected, and
    the single surviving edit leaves coherent text."""
    from docproof.models import Finding

    doc, _ = _doc("He could go.")
    common = dict(chunk_id="c", para_id="body-0000", original_text="He could go.",
                  occurrence=1, explanation="", confidence="high")
    insert_to = Finding(finding_id="f-0001", error_type="missing_word",
                        corrected_text="He to could go.", **common)   # insert "to " at 3
    retype = Finding(finding_id="r-0001", error_type="rewrite",
                     corrected_text="He to go.", **common)            # replace [3,8] "could"->"to"

    out = validate_findings([insert_to, retype], doc, "medium")
    assert [f.status for f in out] == ["validated", "rejected_overlap"]

    # The one applied edit composes to coherent text — never the "to to" dupe.
    text = "He could go."
    applied = [f for f in out if f.status == "validated"]
    assert len(applied) == 1
    a = applied[0].anchor
    assert text[a.start:a.end] == a.delete_text
    assert text[:a.start] + a.insert_text + text[a.end:] == "He to could go."


def test_an_identical_edit_reported_twice_is_a_duplicate_not_an_overlap():
    """Two copies of one finding used to both come back `rejected_overlap`:
    an identical edit necessarily overlaps the accepted copy of itself, so
    with the overlap check asked first, `rejected_duplicate` could never
    fire on a plain edit."""
    from docproof.models import Finding

    doc, _ = _doc("He went to to the door.")
    twin = dict(chunk_id="c", para_id="body-0000", error_type="repeated_word",
                original_text="to to", occurrence=1, corrected_text="to",
                explanation="", confidence="high")
    out = validate_findings([Finding(finding_id="f-0001", **twin),
                             Finding(finding_id="f-0002", **twin)],
                            doc, "medium")
    assert [f.status for f in out] == ["validated", "rejected_duplicate"]


def test_a_repeated_sentence_anchors_by_occurrence():
    """Two identical sentences in one paragraph must not both anchor to the
    first one — that is what `occurrence` is for."""
    doc, paras = _doc("He waited... She waited... He waited...")
    findings, _ = run_sweeps(paras, ["sweep_ellipsis"])
    validated = validate_findings(findings, doc, "medium")
    assert len(validated) == 3
    assert all(f.status == "validated" for f in validated)
    assert len({f.anchor.start for f in validated}) == 3


def test_sweeps_reach_paragraphs_too_short_for_a_model_pass(tmp_path):
    """When a press sets a min_paragraph_chars floor, short lines skip the model
    pass — but the sweeps must still reach them, because in fiction a short line
    is usually dialogue, exactly where a stray "?!" or a mispunctuated tag
    lives. If the sweeps could not see those, every "zero remaining" would
    quietly mean "zero remaining in the long paragraphs". (The floor is off by
    default now — see test_short_dialogue_reaches_a_model_pass — so this exercises
    the knob explicitly.)"""
    import docx
    from docproof.config import load_config
    from docproof.models import Usage
    from docproof.pipeline import finish, prepare
    from docproof.reassembler import paragraph_view_text
    from docproof.utils.xml_helpers import DocxPackage, walk_package

    d = docx.Document()
    for t in ['"Stop!!"', "He waited...",
              "A paragraph comfortably over the minimum length, for contrast."]:
        d.add_paragraph(t)
    src = tmp_path / "s.docx"
    d.save(src)

    cfg = load_config("config/default.yaml")
    cfg.chunking.min_paragraph_chars = 20      # opt into a floor for this test
    prepared = prepare(cfg, src, "config/error_types")
    # With the floor set, the model still only pays for the long one.
    assert sum(len(c.paragraphs) for c in prepared.chunks) == 1
    assert len(prepared.doc.paragraphs) == 3

    out = finish(prepared, [], Usage(), cfg, out_dir=tmp_path / "out",
                 source_path=src)
    accepted = [paragraph_view_text(wp.element, "accept")
                for wp in walk_package(DocxPackage(out.reviewed_path))]
    assert "“Stop!”" in accepted
    assert any("He waited …" == t for t in accepted)


def test_short_dialogue_reaches_a_model_pass(tmp_path):
    """The default min_paragraph_chars floor is 0, so a short line of dialogue is
    reviewed by a model pass, not merely swept. Under the old 20-char floor a
    homophone in a line like "Their here." never reached a model — a silent
    recall hole."""
    import docx
    from docproof.config import load_config
    from docproof.pipeline import prepare

    d = docx.Document()
    d.add_paragraph('"Their here."')     # ~13 chars; below the old floor
    d.add_paragraph("A paragraph comfortably over any old minimum, for contrast.")
    src = tmp_path / "d.docx"
    d.save(src)

    cfg = load_config("config/default.yaml")
    assert cfg.chunking.min_paragraph_chars == 0        # the new default
    prepared = prepare(cfg, src, "config/error_types")
    reviewable = {p.para_id for c in prepared.chunks for p in c.paragraphs}
    assert len(prepared.doc.paragraphs) == 2
    assert len(reviewable) == 2          # both the short line and the long one


def test_sweep_findings_use_their_own_id_series():
    """`s-` not `f-`: the model's findings are numbered independently, and a
    collision would make two different edits look like one."""
    _, paras = _doc("He said what?!")
    findings, _ = run_sweeps(paras, ALL)
    assert findings and all(f.finding_id.startswith("s-") for f in findings)


def test_disabling_every_sweep_is_allowed():
    _, paras = _doc("He said what?!")
    findings, reports = run_sweeps(paras, [])
    assert findings == [] and reports == []


def test_an_unknown_sweep_is_rejected_by_config():
    with pytest.raises(Exception, match="unknown sweep"):
        Config(sweeps=["sweep_nonexistent"])


def test_a_duplicated_sweep_is_rejected_by_config():
    with pytest.raises(Exception, match="more than once"):
        Config(sweeps=["sweep_ellipsis", "sweep_ellipsis"])


def test_shipped_config_enables_sweeps():
    from app.settings import CONFIG_PATH
    from docproof.config import load_config
    cfg = load_config(CONFIG_PATH)
    assert cfg.sweeps, "the shipped config ships with sweeps turned off"
    assert set(cfg.sweeps) <= set(SWEEPS_BY_KEY)


# --- terminal period ---------------------------------------------------------

from docproof.sweeps import _sweep_terminal_period


@pytest.mark.parametrize("text", [
    "Annie managed a curtsy. Alex just crossed his arms",
    "“It might be sick,” Victor warned",
    "“Stay still, damn it!” She shouted",
    "He followed the others down the long corridor toward the far light ahead",  # 13 words, no internal mark
])
def test_terminal_period_fixes_a_run_off_narrative_sentence(text):
    hits = _sweep_terminal_period(text)
    assert len(hits) == 1
    assert apply_hits(text, hits) == text + "."


def test_terminal_period_lands_inside_a_closing_quote():
    text = 'He turned away and muttered, “I never wanted any of this”'
    assert apply_hits(text, _sweep_terminal_period(text)) == \
        'He turned away and muttered, “I never wanted any of this.”'


@pytest.mark.parametrize("text", [
    "Tides of The Squall",                       # a song/poem title
    "The Grey Elegy, Stanza One",                # a title
    "Mortal: A regular person of any species",   # a glossary definition (colon)
    "He said",                                   # too short to be sure it is prose
    "She turned away and left—",                 # an interruption, deliberate
    "I don’t know what to say…",                 # a trailing-off, deliberate
    "The room was cold.",                        # already punctuated
    "Chapter Twenty-Four",                       # a heading (title case, short)
    "The Star Thief and the Winter Court",       # title case, no prose shape
])
def test_terminal_period_leaves_these_alone(text):
    assert _sweep_terminal_period(text) == []


def test_terminal_period_is_idempotent():
    text = "The horses bolted at the sudden crack of the whip behind them"
    once = apply_hits(text, _sweep_terminal_period(text))
    assert _sweep_terminal_period(once) == []
