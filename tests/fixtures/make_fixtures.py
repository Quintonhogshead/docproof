"""Generate test fixtures. Run directly or let conftest.py do it."""
from pathlib import Path
import zipfile

import docx

HERE = Path(__file__).parent

SPLICE_A = "The manuscript was finished, nobody wanted to read it."
CLEAN_B  = "Although it was late, we kept driving toward the coast."
SPLICE_C = "She opened the letter, her hands would not stay still."


def simple() -> None:
    d = docx.Document()
    for text in (SPLICE_A, CLEAN_B, SPLICE_C):
        d.add_paragraph(text)
    d.save(HERE / "simple.docx")


def styled() -> None:
    """Bold span mid-sentence — exercises run splitting across formatting."""
    d = docx.Document()
    d.add_heading("Chapter One", level=1)          # must be skipped by style
    p = d.add_paragraph("It was late, ")
    p.add_run("we were tired").bold = True
    p.add_run(", the road went on.")
    d.save(HERE / "styled.docx")


def table() -> None:
    d = docx.Document()
    d.add_paragraph("An introductory paragraph long enough to be reviewed.")
    t = d.add_table(rows=1, cols=2)
    t.cell(0, 0).text = "The gate was open, the dogs were gone."
    t.cell(0, 1).text = "The gate was open, and the dogs were gone."
    d.save(HERE / "table.docx")


_CONTENT_TYPES = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
  <Override PartName="/word/footnotes.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"/>
</Types>"""

_ROOT_RELS = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>"""

_DOC_RELS = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes" Target="footnotes.xml"/>
</Relationships>"""

_DOCUMENT = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>
    <w:p>
      <w:r><w:t xml:space="preserve">A body paragraph with a footnote reference.</w:t></w:r>
      <w:r><w:footnoteReference w:id="2"/></w:r>
    </w:p>
  </w:body>
</w:document>"""

_FOOTNOTES = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:footnote w:type="separator" w:id="0">
    <w:p><w:r><w:separator/></w:r></w:p>
  </w:footnote>
  <w:footnote w:type="continuationSeparator" w:id="1">
    <w:p><w:r><w:continuationSeparator/></w:r></w:p>
  </w:footnote>
  <w:footnote w:id="2">
    <w:p><w:r><w:t>The note was short, the implications were not.</w:t></w:r></w:p>
  </w:footnote>
</w:footnotes>"""


def footnotes() -> None:
    with zipfile.ZipFile(HERE / "footnotes.docx", "w",
                         zipfile.ZIP_DEFLATED) as z:
        z.writestr("[Content_Types].xml", _CONTENT_TYPES)
        z.writestr("_rels/.rels", _ROOT_RELS)
        z.writestr("word/document.xml", _DOCUMENT)
        z.writestr("word/_rels/document.xml.rels", _DOC_RELS)
        z.writestr("word/footnotes.xml", _FOOTNOTES)


_GOOGLE_CONTENT_TYPES = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>"""

# A Google Docs export as they actually arrive: no styles.xml at all, every
# paragraph "Normal" with direct formatting, blank lines doing the work of
# scene breaks and spacing, typed leading whitespace, a leftover goog_rdk
# content control, one italic run, one URL, one tab indent.
_GOOGLE_DOC = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>
    <w:p><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr><w:b/><w:sz w:val="48"/><w:rFonts w:ascii="Arial"/></w:rPr><w:t>Wildflower</w:t></w:r></w:p>
    <w:p/>
    <w:p><w:r><w:rPr><w:sz w:val="18"/></w:rPr><w:t>Copyright 2026 by A. Author. All rights reserved.</w:t></w:r></w:p>
    <w:p/>
    <w:sdt><w:sdtPr><w:tag w:val="goog_rdk_0"/></w:sdtPr><w:sdtContent/></w:sdt>
    <w:p><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr><w:b/></w:rPr><w:t>Chapter One</w:t></w:r></w:p>
    <w:p/>
    <w:p><w:r><w:t xml:space="preserve">   The road out of town was longer than she remembered, and </w:t></w:r><w:r><w:rPr><w:i/></w:rPr><w:t>much</w:t></w:r><w:r><w:t xml:space="preserve"> darker.  </w:t></w:r></w:p>
    <w:p><w:r><w:tab/></w:r><w:r><w:t>She had not driven it since the winter her mother died.</w:t></w:r></w:p>
    <w:p/>
    <w:p><w:r><w:t>Morning came grey and slow over the fields.</w:t></w:r></w:p>
    <w:p><w:r><w:t xml:space="preserve">Her notes are still online at www.example.com/wildflower for anyone who wants them.</w:t></w:r></w:p>
    <w:p/>
    <w:p><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr><w:b/></w:rPr><w:t>Acknowledgments</w:t></w:r></w:p>
    <w:p><w:r><w:t>Thanks to everyone who read this before it was any good.</w:t></w:r></w:p>
  </w:body>
</w:document>"""

_GOOGLE_ROOT_RELS = _ROOT_RELS


def googledoc() -> None:
    """The manuscript-prep fixture: a styleless export, the shape prep exists
    for. Written as raw XML rather than through python-docx so the content
    control, the typed whitespace and the missing styles.xml are all exactly
    what a real export contains."""
    with zipfile.ZipFile(HERE / "googledoc.docx", "w",
                         zipfile.ZIP_DEFLATED) as z:
        z.writestr("[Content_Types].xml", _GOOGLE_CONTENT_TYPES)
        z.writestr("_rels/.rels", _GOOGLE_ROOT_RELS)
        z.writestr("word/document.xml", _GOOGLE_DOC)


# A manuscript whose front matter carries a real Word-generated table of
# contents, written the three ways Word and its imitators actually write one:
# entries styled TOC1/TOC2, entries linked to a _Toc bookmark, and entries
# carrying their own PAGEREF field. The "Contents" title above them is
# TOCHeading, which is a front-matter title and must NOT be read as an entry —
# and "Chapter One" appears twice, once as a pointer and once as the heading it
# points at, which is the whole difficulty this fixture exists for.
_TOC_DOC = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>
    <w:p><w:pPr><w:pStyle w:val="TOCHeading"/></w:pPr><w:r><w:t>Contents</w:t></w:r></w:p>
    <w:p><w:pPr><w:pStyle w:val="TOC1"/></w:pPr><w:r><w:fldChar w:fldCharType="begin"/></w:r><w:r><w:instrText xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText></w:r><w:r><w:fldChar w:fldCharType="separate"/></w:r><w:hyperlink w:anchor="_Toc0001"><w:r><w:t>Chapter One</w:t></w:r><w:r><w:tab/><w:t>1</w:t></w:r></w:hyperlink></w:p>
    <w:p><w:pPr><w:pStyle w:val="TOC1"/></w:pPr><w:hyperlink w:anchor="_Toc0002"><w:r><w:t>Chapter Two</w:t></w:r><w:r><w:tab/><w:t>14</w:t></w:r></w:hyperlink></w:p>
    <w:p><w:hyperlink w:anchor="_Toc0003"><w:r><w:t>Chapter Three</w:t></w:r></w:hyperlink><w:r><w:tab/></w:r><w:r><w:instrText xml:space="preserve"> PAGEREF _Toc0003 \\h </w:instrText></w:r><w:r><w:t>27</w:t></w:r><w:r><w:fldChar w:fldCharType="end"/></w:r></w:p>
    <w:p><w:pPr><w:pStyle w:val="Contents2"/></w:pPr><w:r><w:t>Acknowledgments</w:t></w:r><w:r><w:tab/><w:t>40</w:t></w:r></w:p>
    <w:p/>
    <w:p><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:rPr><w:b/></w:rPr><w:t>Chapter One</w:t></w:r></w:p>
    <w:p/>
    <w:p><w:r><w:t>The road out of town was longer than she remembered.</w:t></w:r></w:p>
    <w:p><w:r><w:t>She had not driven it since the winter her mother died.</w:t></w:r></w:p>
  </w:body>
</w:document>"""


def toc() -> None:
    """A manuscript with a generated table of contents in its front matter."""
    with zipfile.ZipFile(HERE / "toc.docx", "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("[Content_Types].xml", _GOOGLE_CONTENT_TYPES)
        z.writestr("_rels/.rels", _GOOGLE_ROOT_RELS)
        z.writestr("word/document.xml", _TOC_DOC)


# A manuscript with text boxes — the floating frames Word anchors in a run.
# Two shapes, the two ways they arrive:
#   * a modern DrawingML box wrapped in mc:AlternateContent, whose mc:Choice
#     holds the real story and whose mc:Fallback repeats it verbatim as legacy
#     VML — the reader must take the Choice once and never the Fallback, or the
#     line is read, edited and counted twice;
#   * a plain legacy VML box with no AlternateContent around it, read once.
# Each box carries a comma splice, so the review pipeline has a real error to
# find inside it, and the anchoring paragraphs carry their own separate text
# that must stay clear of the box's.
_TEXTBOX_DOC = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape" xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:v="urn:schemas-microsoft-com:vml">
  <w:body>
    <w:p><w:r><w:t>An introductory paragraph long enough to be reviewed on its own.</w:t></w:r></w:p>
    <w:p>
      <w:r><w:t xml:space="preserve">The sign by the road read as follows.</w:t></w:r>
      <w:r>
        <mc:AlternateContent>
          <mc:Choice Requires="wps">
            <w:drawing><wp:anchor><a:graphic><a:graphicData><wps:wsp><wps:txbx><w:txbxContent>
              <w:p><w:r><w:t>Beware the dog, it bites without warning.</w:t></w:r></w:p>
            </w:txbxContent></wps:txbx></wps:wsp></a:graphicData></a:graphic></wp:anchor></w:drawing>
          </mc:Choice>
          <mc:Fallback>
            <w:pict><v:shape><v:textbox><w:txbxContent>
              <w:p><w:r><w:t>Beware the dog, it bites without warning.</w:t></w:r></w:p>
            </w:txbxContent></v:textbox></v:shape></w:pict>
          </mc:Fallback>
        </mc:AlternateContent>
      </w:r>
    </w:p>
    <w:p>
      <w:r><w:pict><v:shape><v:textbox><w:txbxContent>
        <w:p><w:r><w:t>The gate stood open, the yard was empty.</w:t></w:r></w:p>
      </w:txbxContent></v:textbox></v:shape></w:pict></w:r>
    </w:p>
    <w:p><w:r><w:t>A closing paragraph, also comfortably long enough to review.</w:t></w:r></w:p>
  </w:body>
</w:document>"""


def textbox() -> None:
    """A manuscript whose prose lives partly inside floating text boxes —
    modern (mc:Choice + mc:Fallback duplicate) and legacy VML."""
    with zipfile.ZipFile(HERE / "textbox.docx", "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("[Content_Types].xml", _GOOGLE_CONTENT_TYPES)
        z.writestr("_rels/.rels", _ROOT_RELS)
        z.writestr("word/document.xml", _TEXTBOX_DOC)


def tracked() -> None:
    """A manuscript that still has an unresolved edit in it. Prep refuses this
    outright — it restyles every paragraph, and doing that around somebody
    else's undecided change would bury it."""
    d = docx.Document()
    d.add_paragraph("A perfectly ordinary paragraph of manuscript text.")
    p = d.add_paragraph()
    run = p.add_run("An inserted phrase.")
    ins = run._r.makeelement(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ins",
        {"{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id": "1",
         "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author": "Someone",
         "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}date":
             "2026-01-01T00:00:00Z"})
    run._r.addprevious(ins)
    ins.append(run._r)
    d.save(HERE / "tracked.docx")


if __name__ == "__main__":
    simple(); styled(); table(); footnotes(); googledoc(); toc(); tracked()
    textbox()
    print("fixtures written to", HERE)