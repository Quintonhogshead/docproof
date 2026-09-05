"""Genre posture presets (docproof/genre.py).

The load-bearing guarantee under test: a genre preset can only ever touch the
stylistic lane (ALLOWED_PRESET_KEYS) — never mechanics/proofreading, and never
by way of a typo either, since a bad key is refused at load. The second thing
under test is precedence: a review PROFILE (docproof.profiles) is a stricter
boundary than a genre, and must win when both are applied, in either order the
caller happens to apply them in code, and specifically in the order
`docproof.__main__._configure` actually uses (genre, then profile).
"""
from __future__ import annotations

from argparse import Namespace
from pathlib import Path

import pytest
import yaml

from docproof.__main__ import _configure
from docproof.config import Config, load_config
from docproof.genre import (ALLOWED_PRESET_KEYS, apply_genre,
                            available_genres, load_genre_preset,
                            materialize_genre_pack, write_genre_pack)
from docproof.profiles import DETECTOR_ONLY, apply_profile

CONFIG = Path(__file__).parent.parent / "config" / "default.yaml"
GENRES_DIR = Path(__file__).parent.parent / "config" / "genres"

SHIPPED_GENRES = ("academic", "fantasy_sf", "general_fiction",
                  "general_nonfiction", "historical", "literary_memoir",
                  "poetry", "religious", "self_help_business")

# The non-fiction posture presets added by the config/genre taxonomy work.
# House rule for every one of them: a genre sets POSTURE, never a workflow
# stage's lane switches — so none may turn on edits-mode smoothing or the
# rewrite recall lever. That contamination (religious non-fiction run under
# self_help_business, which flips both on) is exactly what these presets fix.
NONFICTION_GENRES = ("academic", "general_nonfiction", "historical",
                     "religious")

# Every knob a mechanics/proofreading section owns — none of these may ever
# appear in a genre preset's `overlay`. Not exhaustive of every Config field,
# but exhaustive of the sections house policy names as hammered-every-genre:
# normalization, style mechanics, the edit guard, and the spell scan's own
# enforcement (allowlist/denylist are runtime seeding, not a posture toggle).
_MECHANICS_SECTIONS = ("normalize", "style", "edit_guard", "spellcheck",
                       "audit", "tracked_changes_policy", "error_types",
                       "sweeps")


# --- the shipped presets ------------------------------------------------------

def test_available_genres_lists_every_shipped_preset():
    assert available_genres() == tuple(sorted(SHIPPED_GENRES))


@pytest.mark.parametrize("genre", NONFICTION_GENRES)
def test_nonfiction_preset_never_enables_a_copyedit_wave_lane(genre):
    """A genre is a posture, not a workflow stage: no non-fiction preset may
    turn on edits-mode smoothing or the rewrite lever. Those belong to the
    copy-edit wave (config/stages/), and letting a genre enable them is the
    contamination the taxonomy work removes."""
    cfg = load_config(CONFIG)
    cfg, _pending = apply_genre(cfg, genre)
    assert cfg.smoothing.edits is False, f"{genre} enabled edits-mode smoothing"
    assert cfg.rewrite.enabled is False, f"{genre} enabled the rewrite lever"


def test_religious_preset_protects_quoted_text_and_raises_the_name_bar():
    """The preset that gives theological non-fiction a home: no automatic
    line-edit source (Scripture must not be reworded) and a name-spelling bar
    well above house default (transliteration near-matches like Deut/Deute)."""
    cfg = load_config(CONFIG)
    cfg, _ = apply_genre(cfg, "religious")
    assert cfg.smoothing.enabled is False
    assert cfg.rewrite.enabled is False
    assert cfg.consistency.name_dominance > 5      # above the house default
    assert cfg.consistency.name_min_count > 20


def test_historical_pack_wires_era_and_period_continuity():
    cfg, summary = materialize_genre_pack(CONFIG, "historical", era=1863)
    assert cfg.genre_scans.anachronism.enabled is True
    assert cfg.genre_scans.anachronism.era == 1863
    assert cfg.genre_scans.citation_format.enabled is True
    assert "period consistency" in cfg.continuity.prompt
    # Without an era the scan is a deliberate no-op (never guessed).
    cfg2, summary2 = materialize_genre_pack(CONFIG, "historical")
    assert cfg2.genre_scans.anachronism.era is None


@pytest.mark.parametrize("genre", SHIPPED_GENRES)
def test_shipped_preset_loads_and_only_uses_allowed_keys(genre):
    preset = load_genre_preset(genre)
    overlay = preset.get("overlay") or {}
    assert overlay, f"{genre} preset has an empty overlay"
    assert set(overlay) <= ALLOWED_PRESET_KEYS
    # And the negative half of the guarantee: no key even LOOKS like it
    # touches a mechanics section, whatever ALLOWED_PRESET_KEYS says today.
    for key in overlay:
        section = key.split(".", 1)[0]
        assert section not in _MECHANICS_SECTIONS


@pytest.mark.parametrize("genre", SHIPPED_GENRES)
def test_shipped_preset_applies_cleanly_to_a_loaded_config(genre):
    cfg = load_config(CONFIG)
    cfg, pending = apply_genre(cfg, genre)
    assert isinstance(cfg, Config)
    assert pending == {}
    assert cfg.flights.posture in ("strict", "lenient")


def test_apply_genre_none_is_a_noop():
    cfg = load_config(CONFIG)
    before = cfg.model_dump()
    cfg, pending = apply_genre(cfg, None)
    assert cfg.model_dump() == before
    assert pending == {}


def test_unknown_genre_raises_with_the_available_list():
    with pytest.raises(ValueError, match="steampunk_romance"):
        load_genre_preset("steampunk_romance")


# --- the whitelist is actually enforced, not just documented -----------------

def test_preset_overlay_key_outside_whitelist_is_refused(tmp_path):
    bad = tmp_path / "rogue.yaml"
    bad.write_text(yaml.safe_dump({
        "name": "rogue",
        "overlay": {"normalize.quotes": False,          # a mechanics key
                   "smoothing.enabled": True},
    }))
    with pytest.raises(ValueError, match="normalize.quotes"):
        load_genre_preset("rogue", genres_dir=tmp_path)


def test_preset_overlay_bad_value_fails_at_apply_not_silently(tmp_path):
    """A preset can only name whitelisted KEYS; a bad VALUE for an allowed key
    still has to clear that section's own validation, exactly like a
    hand-edited config file would."""
    bad = tmp_path / "malformed.yaml"
    bad.write_text(yaml.safe_dump({
        "name": "malformed",
        "overlay": {"smoothing.judge_harshness": "furious"},   # not a Literal
    }))
    cfg = Config()
    with pytest.raises(Exception):
        apply_genre(cfg, "malformed", genres_dir=tmp_path)


def test_preset_overlay_must_be_a_mapping(tmp_path):
    bad = tmp_path / "notmap.yaml"
    bad.write_text(yaml.safe_dump({"name": "notmap", "overlay": ["nope"]}))
    with pytest.raises(ValueError, match="mapping"):
        load_genre_preset("notmap", genres_dir=tmp_path)


# --- genre x profile precedence ----------------------------------------------

def test_genre_then_profile_lets_the_profile_win():
    """fantasy_sf turns smoothing ON; detector-only turns it back OFF. Genre
    applied first, profile second (the order docproof/genre.py's docstring
    prescribes) must leave smoothing off — a profile's reproducibility
    boundary is not something a genre preset should be able to reopen."""
    cfg = load_config(CONFIG)
    cfg, _pending = apply_genre(cfg, "fantasy_sf")
    assert cfg.smoothing.enabled is True          # the preset alone did this
    apply_profile(cfg, DETECTOR_ONLY)
    assert cfg.smoothing.enabled is False          # the profile wins


def test_configure_cli_applies_genre_before_profile():
    """The actual CLI path: `docproof review --genre fantasy_sf --profile
    detector-only` must produce the same strict, profile-governed config as
    plain `--profile detector-only` — the genre must not leak past it."""
    args = Namespace(config=str(CONFIG), genre="fantasy_sf",
                     profile=DETECTOR_ONLY, error_types=None,
                     min_confidence=None, variant=None, dictionary=None,
                     no_comments=False, out=None, model=None)
    cfg, _error_dir = _configure(args)
    assert cfg.smoothing.enabled is False
    assert cfg.rewrite.enabled is False
    assert cfg.consistency.enabled is False


def test_configure_cli_genre_alone_applies_the_posture():
    args = Namespace(config=str(CONFIG), genre="self_help_business",
                     profile=None, error_types=None, min_confidence=None,
                     variant=None, dictionary=None, no_comments=False,
                     out=None, model=None)
    cfg, _error_dir = _configure(args)
    assert cfg.smoothing.enabled is True
    assert cfg.smoothing.edits is True
    assert cfg.rewrite.enabled is True


# --- genre-pack materialization ----------------------------------------------

def test_materialize_genre_pack_applies_overlay_and_continuity_prompt():
    cfg, summary = materialize_genre_pack(CONFIG, "fantasy_sf")
    assert cfg.smoothing.judge_harshness == "lenient"
    assert cfg.consistency.name_dominance == 10
    assert "world rule" in cfg.continuity.prompt
    assert summary["genre"] == "fantasy_sf"
    assert summary["pending"] == {}
    assert cfg.flights.posture == "lenient"
    assert "continuity_prompt" in summary


def test_materialize_genre_pack_without_continuity_prompt_leaves_it_blank():
    cfg, _summary = materialize_genre_pack(CONFIG, "general_fiction")
    assert cfg.continuity.prompt == ""


def test_materialize_genre_pack_self_help_turns_on_genre_scans():
    cfg, summary = materialize_genre_pack(CONFIG, "self_help_business")
    assert cfg.genre_scans.citation_format.enabled is True
    assert cfg.genre_scans.reading_level.enabled is True
    assert cfg.genre_scans.anachronism.enabled is False   # untouched
    assert "genre_scans_applied" in summary


def test_materialize_genre_pack_era_is_never_inferred_only_stated():
    cfg, summary = materialize_genre_pack(CONFIG, "fantasy_sf")
    assert cfg.genre_scans.anachronism.era is None
    assert "anachronism_era" not in summary

    cfg, summary = materialize_genre_pack(CONFIG, "fantasy_sf", era=1350)
    assert cfg.genre_scans.anachronism.era == 1350
    assert summary["anachronism_era"] == 1350


def test_materialize_genre_pack_seeds_names_and_reading_band_from_a_profile():
    class _Noun:
        def __init__(self, name, count):
            self.name, self.count = name, count

    class _ReadingLevel:
        ari = 42.5

    class _Profile:
        proper_nouns = [_Noun("Zylandria", 6), _Noun("Kestrion", 3)]
        reading_level = _ReadingLevel()

    cfg, summary = materialize_genre_pack(
        CONFIG, "fantasy_sf", profile=_Profile())
    assert set(cfg.consistency.seeded_names) >= {"Zylandria", "Kestrion"}
    assert set(cfg.spellcheck.allowlist) >= {"Zylandria", "Kestrion"}
    assert cfg.genre_scans.reading_level.target_ari == 42.5
    assert summary["seeded_names_count"] == 2
    assert summary["reading_level_target_ari"] == 42.5


def test_materialize_genre_pack_without_a_profile_leaves_seeding_untouched():
    cfg, summary = materialize_genre_pack(CONFIG, "fantasy_sf")
    assert cfg.consistency.seeded_names == []
    assert summary["seeded_names_count"] == 0
    assert "reading_level_target_ari" not in summary


def test_materialize_genre_pack_seeding_unions_with_an_existing_allowlist(
        tmp_path):
    """Seeding must never REPLACE whatever the base config already protects —
    only add to it."""
    base = tmp_path / "base.yaml"
    base.write_text(yaml.safe_dump(
        {"spellcheck": {"allowlist": ["Eldergate"]},
         "consistency": {"seeded_names": ["Marrowdeep"]}}))

    class _Noun:
        def __init__(self, name):
            self.name, self.count = name, 1

    class _Profile:
        proper_nouns = [_Noun("Zylandria")]
        reading_level = None

    cfg, _summary = materialize_genre_pack(base, "fantasy_sf",
                                           profile=_Profile())
    assert set(cfg.spellcheck.allowlist) == {"Eldergate", "Zylandria"}
    assert set(cfg.consistency.seeded_names) == {"Marrowdeep", "Zylandria"}


def test_write_genre_pack_round_trips_through_load_config(tmp_path):
    out = tmp_path / "run.yaml"
    summary = write_genre_pack(CONFIG, "literary_memoir", out)
    assert out.is_file()
    reloaded = load_config(out)
    assert reloaded.smoothing.enabled is False
    assert reloaded.rewrite.enabled is False
    assert summary["out_path"] == str(out)


def test_materialized_config_is_self_contained_for_error_types(tmp_path):
    """A genre-pack config written into a book workspace has no sibling
    error_types/ directory. `_resolve_error_dir` must fall back to the packaged
    shipped prompts so the relocated config still runs `inventory` and every
    typed pass — the failure the portability fix removes."""
    from docproof.__main__ import _resolve_error_dir
    out = tmp_path / "bookws" / "run.yaml"
    write_genre_pack(CONFIG, "religious", out)
    assert not (out.parent / "error_types").exists()
    resolved = _resolve_error_dir(out)
    assert resolved.is_dir()
    assert (resolved / "spelling.yaml").is_file()


def test_resolve_error_dir_prefers_a_sibling_directory(tmp_path):
    """When a config DOES have its own error_types/ beside it (an edited
    workspace copy), that wins over the packaged fallback."""
    from docproof.__main__ import _resolve_error_dir
    (tmp_path / "error_types").mkdir()
    cfg = tmp_path / "run.yaml"
    cfg.write_text("{}")
    assert _resolve_error_dir(cfg) == tmp_path / "error_types"


# --- name seeding actually reaches the consistency checks (not just config) --

def _para(text, i):
    from docproof.models import ParagraphRef
    return ParagraphRef(f"body-{i:04d}", "word/document.xml", "body", text,
                        "Normal")


def _write_toward_docx(path):
    import docx
    d = docx.Document()
    d.add_paragraph(
        "She walked toward the door and paused before opening it.",
        style="Normal")
    d.add_paragraph(
        "He drifted toward sleep long before the fire went out.",
        style="Normal")
    d.add_paragraph(
        "They moved towards the light at the end of the hall.",
        style="Normal")
    d.add_paragraph(
        "The quiet town had nothing special about it at all today.",
        style="Normal")
    d.save(path)


ERROR_DIR = Path(__file__).parent.parent / "config" / "error_types"


def test_seeded_names_reach_the_consistency_scan(tmp_path):
    """`consistency.seeded_names` is not just a config field: pipeline.py
    unions it into the `protected` set the mechanical scans read, exactly
    like the spell scan's own lexicon. "toward"/"towards" is a real
    spelling-variant pair with NO respell entry for the US variant (unlike
    grey/gray), so the finding here is the seeding's doing, not something
    else already suppressing it."""
    from docproof.pipeline import prepare

    docx_path = tmp_path / "toward.docx"
    _write_toward_docx(docx_path)

    cfg = load_config(CONFIG)
    cfg.error_types = [["spelling"]]
    prepared = prepare(cfg, docx_path, ERROR_DIR)
    kinds = [f.error_type for f in prepared.consistency_findings]
    assert "term_consistency" in kinds

    cfg2 = load_config(CONFIG)
    cfg2.error_types = [["spelling"]]
    cfg2.consistency.seeded_names = ["towards"]
    prepared2 = prepare(cfg2, docx_path, ERROR_DIR)
    kinds2 = [f.error_type for f in prepared2.consistency_findings]
    assert "term_consistency" not in kinds2
