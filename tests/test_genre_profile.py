"""The deterministic ($0) manuscript profile (docproof/genre_profile.py).

Builds a tiny, self-contained fixture .docx (two chapter headings, a coined
proper noun, dialogue, a doubled-punctuation tic, and a scene-break glyph
line) so every field on Profile has something real to find, rather than
reusing tests/fixtures/tiny_novel.docx — that fixture's names are all
dictionary-known first names (Kathryn/Katherine), so it cannot exercise the
proper-noun-candidate path at all.
"""
from __future__ import annotations

import docx
import pytest

from docproof.config import Config, load_config
from docproof.genre_profile import Profile, build_profile, confirm_with_model

CONFIG = __import__("pathlib").Path(__file__).parent.parent / "config" / "default.yaml"


@pytest.fixture(scope="module")
def fixture_docx(tmp_path_factory):
    d = docx.Document()
    d.add_paragraph("Chapter One: The Gate", style="Heading 1")
    d.add_paragraph(
        "Zylandria stood at the gate long after the others had gone "
        "inside, watching the last light fade over the wall.", style="Normal")
    d.add_paragraph(
        '"Are you coming, Zylandria?" called a voice from the courtyard, '
        'sharp against the quiet.', style="Normal")
    d.add_paragraph(
        "She sat by the the window and watched the stars come out one "
        "by one over the ruined tower.", style="Normal")
    d.add_paragraph("* * *", style="Normal")
    d.add_paragraph(
        "By morning Zylandria had already crossed the outer field, her "
        "boots dark with dew and her thoughts far ahead of her feet.",
        style="Normal")
    d.add_paragraph("Chapter Two: The Road", style="Heading 1")
    d.add_paragraph(
        "The road forward was long, and Zylandria walked it without "
        "looking back even once, not even at the sound of her own name.",
        style="Normal")
    d.add_paragraph(
        '"Wait!!" someone shouted behind her, but she did not stop.',
        style="Normal")
    path = tmp_path_factory.mktemp("genre_profile") / "fixture.docx"
    d.save(path)
    return path


@pytest.fixture(scope="module")
def cfg():
    return load_config(CONFIG)


@pytest.fixture(scope="module")
def profile(fixture_docx, cfg):
    return build_profile(fixture_docx, cfg)


# --- basic counts --------------------------------------------------------

def test_word_and_paragraph_counts_are_positive(profile):
    assert profile.word_count > 50
    assert profile.paragraph_count > 5


def test_chapters_are_detected_from_headings(profile):
    titles = [c.title for c in profile.chapters]
    assert titles == ["Chapter One: The Gate", "Chapter Two: The Road"]
    assert all(c.word_count > 0 for c in profile.chapters)


def test_dialogue_density_is_nonzero_when_dialogue_is_present(profile):
    assert 0.0 < profile.dialogue_density < 1.0


def test_dialogue_density_is_zero_with_no_quoted_speech():
    d = docx.Document()
    d.add_paragraph("The field stretched on and on with nothing in it.",
                    style="Normal")
    import tempfile
    from pathlib import Path
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "no_dialogue.docx"
        d.save(path)
        p = build_profile(path, Config())
    assert p.dialogue_density == 0.0


# --- proper nouns ----------------------------------------------------------

def test_proper_noun_candidates_include_the_coined_name(profile):
    names = {c.name for c in profile.proper_nouns}
    assert "Zylandria" in names
    zy = next(c for c in profile.proper_nouns if c.name == "Zylandria")
    assert zy.count >= 3


# --- author tics -------------------------------------------------------------

def test_tics_include_the_doubled_word(profile):
    kinds = {t.kind for t in profile.tics}
    assert "sweep_doubled_word" in kinds


def test_tics_include_the_doubled_punctuation(profile):
    kinds = {t.kind for t in profile.tics}
    assert "sweep_stacked_punctuation" in kinds


def test_tics_include_the_scene_break_glyph(profile):
    glyph_tics = [t for t in profile.tics if t.kind == "scene_break_glyph"]
    assert len(glyph_tics) == 1
    assert glyph_tics[0].count == 1
    assert glyph_tics[0].samples[0].before.strip() == "* * *"


def test_no_scene_break_glyph_when_none_present():
    d = docx.Document()
    d.add_paragraph("A perfectly ordinary paragraph with no glyph line at "
                    "all in it, just words.", style="Normal")
    import tempfile
    from pathlib import Path
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "no_glyph.docx"
        d.save(path)
        p = build_profile(path, Config())
    assert all(t.kind != "scene_break_glyph" for t in p.tics)


# --- reading level -------------------------------------------------------

def test_reading_level_metrics_are_computed(profile):
    rl = profile.reading_level
    assert rl.ari is not None
    assert rl.avg_sentence_words is not None and rl.avg_sentence_words > 0
    assert rl.avg_word_chars is not None and rl.avg_word_chars > 0
    # mean_zipf is best-effort (wordfreq); when it IS available it should be
    # a plausible zipf score, not left at its unset default.
    if rl.mean_zipf is not None:
        assert 0 < rl.mean_zipf < 8


# --- genre guesses -------------------------------------------------------

def test_genre_guesses_are_a_normalized_ranking(profile):
    assert len(profile.genre_guesses) == 4
    scores = [g.score for g in profile.genre_guesses]
    assert scores == sorted(scores, reverse=True)
    assert pytest.approx(sum(scores), abs=1e-6) == 1.0
    assert profile.recommended_preset == profile.genre_guesses[0].genre


def test_recommended_preset_is_one_of_the_four_shipped_genres(profile):
    from docproof.genre import available_genres
    assert profile.recommended_preset in available_genres()


# --- bespoke sweep candidates -------------------------------------------

def test_bespoke_candidates_only_for_tics_seen_at_least_three_times():
    """The fixture's doubled word and doubled punctuation each occur once —
    below the >=3 bar — so neither should mint a bespoke-sweep candidate."""
    d = docx.Document()
    for i in range(4):
        d.add_paragraph(f"The gate creaked open again on day {i}.",
                        style="Normal")
        d.add_paragraph("* * *", style="Normal")
    import tempfile
    from pathlib import Path
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "recurring.docx"
        d.save(path)
        p = build_profile(path, Config())
    glyph_candidates = [c for c in p.bespoke_sweep_candidates
                        if c.sample.strip() == "* * *"]
    assert glyph_candidates, "a glyph line repeated 4x should be a candidate"
    assert glyph_candidates[0].count >= 3


def test_fixture_tics_below_the_bar_mint_no_bespoke_candidates(profile):
    # The fixture's own doubled word/punctuation each occur exactly once.
    assert profile.bespoke_sweep_candidates == []


# --- JSON round-trip -------------------------------------------------------

def test_profile_serializes_and_round_trips_as_json(profile):
    raw = profile.model_dump_json()
    reloaded = Profile.model_validate_json(raw)
    assert reloaded.word_count == profile.word_count
    assert [c.title for c in reloaded.chapters] == \
        [c.title for c in profile.chapters]


# --- optional model confirmation is additive and best-effort ----------------

def test_confirm_with_model_falls_back_cleanly_with_no_provider_available(
        profile, fixture_docx, cfg):
    """No API key is configured in the test environment (and the socket
    guard blocks any real network call regardless), so this must degrade to
    the deterministic profile, not raise."""
    from docproof.formats import get_format
    fmt = get_format(fixture_docx)
    pkg = fmt.preflight(str(fixture_docx), cfg.tracked_changes_policy)
    doc = fmt.build_document_model(pkg, cfg)
    result = confirm_with_model(profile, doc.paragraphs,
                                model="gpt-5.6-luna", cfg=cfg)
    assert result.model_confirmed is False
    assert result.word_count == profile.word_count
    assert result.recommended_preset == profile.recommended_preset
