"""Punctuation-tolerant anchoring (Phase 1.5).

The manuscript normalizer curls quotes before ingest, so the canonical text
holds “ ” ‘ ’. A weaker model routinely re-types the sentence it quotes with
"straight" punctuation; an exact substring anchor then fails and the finding is
discarded as unanchorable even though it is the same sentence character for
character. The baseline made this concrete — claude-haiku-4-5 lost ~24% of its
findings this way while gpt-5.6-luna lost none. These tests pin the fix: fold
curly quotes, dashes and nbsp (one character for one) so the quote anchors, and
take the edit's delete_text from the REAL paragraph so it still applies.
"""
from __future__ import annotations

import docx

from docproof.config import load_config
from docproof.models import DocumentModel, Finding, ParagraphRef, Usage
from docproof.pipeline import finish, prepare
from docproof.reassembler import paragraph_view_text
from docproof.utils.xml_helpers import DocxPackage, walk_package
from docproof.validator import anchor_offset, fold_punct, validate_findings

CURLY_APOS = "The grocer chalked fresh apple’s on the board outside."


def _doc(text: str) -> DocumentModel:
    return DocumentModel(
        source_path="x.docx",
        paragraphs=(ParagraphRef("body-0000", "word/document.xml", "body",
                                 text, "Normal"),))


def _finding(original: str, corrected: str, etype: str = "apostrophe_error"
             ) -> Finding:
    return Finding("f-0001", "chunk-000", "body-0000", etype, original, 1,
                   corrected, "x", "high")


# --- anchor_offset -----------------------------------------------------------

def test_exact_match_still_works():
    assert anchor_offset("hello world", "world", 1) == 6


def test_absent_text_returns_minus_one():
    assert anchor_offset("hello world", "goodbye", 1) == -1


def test_straight_quotes_match_curly_text():
    hay = "She said “yes” and left."         # curly “ ”
    s = anchor_offset(hay, 'said "yes"', 1)             # straight " "
    assert s != -1
    assert fold_punct(hay)[s:s + len('said "yes"')] == 'said "yes"'


def test_em_dash_matches_hyphen_and_nbsp_matches_space():
    assert anchor_offset("Wait — I said so.", "Wait - I said", 1) != -1
    assert anchor_offset("a b c", "a b c", 1) != -1


def test_the_fold_is_length_preserving():
    # Offsets only stay valid because no character changes width.
    for s in ["“x”", "a—b", "it’s", "a b c"]:
        assert len(fold_punct(s)) == len(s)


def test_a_genuinely_different_quote_still_fails():
    assert anchor_offset(CURLY_APOS, "a sentence that is not present", 1) == -1


# --- validate_findings: the edit still applies to the real text --------------

def test_a_straight_quoted_finding_anchors_against_curly_text():
    doc = _doc(CURLY_APOS)
    # Model quoted the sentence with a straight apostrophe and fixed apple's.
    out = validate_findings(
        [_finding("The grocer chalked fresh apple's on the board outside.",
                  "The grocer chalked fresh apples on the board outside.")],
        doc, "medium")
    assert out[0].status == "validated"
    a = out[0].anchor
    # The slice the reassembler re-asserts before editing is the REAL text...
    assert doc.paragraphs[0].text[a.start:a.end] == a.delete_text
    # ...which is the manuscript's curly apostrophe, not the model's straight one.
    assert a.delete_text == "’"
    assert a.insert_text == ""


def test_a_straight_quoted_query_anchors_against_curly_text():
    doc = _doc("He said “wait” then “run” without a breath.")
    out = validate_findings(
        [_finding('He said "wait" then "run" without a breath.',
                  'He said "wait" then "run" without a breath.',
                  etype="speaker_change")],
        doc, "medium", query_types=frozenset({"speaker_change"}))
    assert out[0].status == "query"
    a = out[0].anchor
    assert doc.paragraphs[0].text[a.start:a.end] == a.delete_text   # real text


def test_exact_match_delete_text_is_unchanged():
    """Regression: on the exact-match path (what gpt-5.6-luna hits every time),
    delete_text is byte-identical to before — the real slice equals the model's
    quote, so nothing about that path changes."""
    doc = _doc("It was, late.")
    out = validate_findings(
        [_finding("It was, late.", "It was late.", etype="comma_splice")],
        doc, "medium")
    assert out[0].status == "validated"
    a = out[0].anchor
    assert a.delete_text == "," and a.insert_text == ""
    assert doc.paragraphs[0].text[a.start:a.end] == a.delete_text


# --- end to end: the tracked change actually lands ---------------------------

def test_a_fold_recovered_edit_applies_through_the_reassembler(tmp_path):
    d = docx.Document()
    d.add_paragraph(CURLY_APOS)                    # curly apostrophe in the file
    src = tmp_path / "m.docx"
    d.save(src)

    cfg = load_config("config/default.yaml")
    cfg.error_types = ["apostrophe_error"]
    prepared = prepare(cfg, src, "config/error_types")

    from docproof.analyzer import MockAnalyzer
    import itertools
    ids = itertools.count(1)
    findings = []
    for group in prepared.groups:
        for chunk in prepared.chunks:
            found, _ = MockAnalyzer(group, [{
                "para_id": "body-0000", "error_type": "apostrophe_error",
                # straight apostrophe — would not anchor before the fold
                "original_text": "The grocer chalked fresh apple's on the board outside.",
                "corrected_text": "The grocer chalked fresh apples on the board outside.",
                "explanation": "possessive, not plural", "confidence": "high",
            }], ids).analyze_chunk(chunk, Usage())
            findings += found

    out = finish(prepared, findings, Usage(), cfg, out_dir=tmp_path / "out",
                 source_path=src)
    assert out.applied == 1
    pkg = DocxPackage(out.reviewed_path)
    para = next(iter(walk_package(pkg))).element
    # Accepting the change removes the (curly) apostrophe; rejecting restores it.
    assert paragraph_view_text(para, "accept") == \
        "The grocer chalked fresh apples on the board outside."
    assert paragraph_view_text(para, "reject") == CURLY_APOS
