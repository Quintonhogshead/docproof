"""`docproof review`'s intra-run checkpoint (docproof/checkpoint.py, wired into
the sync CLI's run_sync call in __main__.py's cmd_review) and the fingerprint
fix that goes with it: a checkpoint is keyed off the config's CONTENTS, not
its path, so editing the config file a run points at can never cause a stale
checkpoint to replay findings produced under the config that came before.

Composes with the existing findings-boundary checkpoint (run_checkpoint.py,
tests/test_run_checkpoint.py): this one survives a crash DURING run_sync, that
one survives a crash AFTER it but before finish() finishes writing.
"""
from __future__ import annotations

import re
from pathlib import Path

import docx
import pytest
import yaml

from docproof.__main__ import main
from docproof.config import load_config
from docproof.providers import NormalizedUsage, ProviderResult

U = NormalizedUsage(input_tokens=5, output_tokens=2)
REPO_ERROR_TYPES = (Path(__file__).parent.parent / "config" / "error_types").resolve()


class _CrashingProvider:
    """Answers "teh" -> "the" for every chunk except `crash_on`'s paragraph,
    where it raises instead — a stand-in for the provider call that never
    comes back mid run_sync."""

    def __init__(self, calls: list, crash_on: str | None = None):
        self.calls = calls
        self.crash_on = crash_on

    def complete_structured(self, *, user, **kw) -> ProviderResult:
        pid = re.findall(r'<paragraph id="([^"]+)"',
                         user.split("</context>")[-1])[0]
        self.calls.append(pid)
        if pid == self.crash_on:
            raise RuntimeError("simulated crash mid run_sync")
        return ProviderResult(parsed={"findings": [{
            "para_id": pid, "error_type": "spelling", "original_text": "teh",
            "corrected_text": "the", "confidence": "high", "explanation": ""}]},
            usage=U)


def _manuscript(tmp_path):
    d = docx.Document()
    d.add_paragraph("This has a teh typo right here in the first line today.")
    d.add_paragraph("Another line with a teh typo somewhere in it as well.")
    src = tmp_path / "m.docx"
    d.save(src)
    return src


def _write_config(path, **overrides):
    """A real config/default.yaml, isolated to one cheap pass and no other
    provider-backed stage, written out so --config can point at it. `overrides`
    are dotted-path tweaks applied on top before serializing.

    `_configure()` resolves error types from `<config's dir>/error_types`, so
    a symlink to the repo's real one is planted beside `path` the first time
    it is written."""
    error_types_link = path.parent / "error_types"
    if not error_types_link.exists():
        error_types_link.symlink_to(REPO_ERROR_TYPES, target_is_directory=True)
    cfg = load_config("config/default.yaml")
    cfg.error_types = [["spelling"]]
    cfg.chunking.token_budget = 1          # one paragraph per chunk
    cfg.glossary.enabled = False
    cfg.adjudicate.enabled = False
    cfg.consistency.enabled = False
    cfg.spellcheck.enabled = False
    for dotted, value in overrides.items():
        obj = cfg
        parts = dotted.split(".")
        for p in parts[:-1]:
            obj = getattr(obj, p)
        setattr(obj, parts[-1], value)
    path.write_text(yaml.safe_dump(cfg.model_dump(mode="json")),
                    encoding="utf-8")
    return cfg


def test_a_crash_mid_run_sync_is_resumed_without_recalling(tmp_path, monkeypatch):
    src = _manuscript(tmp_path)
    cfg_path = tmp_path / "config.yaml"
    _write_config(cfg_path)
    out = tmp_path / "out"

    first_calls: list[str] = []
    monkeypatch.setattr("docproof.__main__.build_provider",
                        lambda cfg: _CrashingProvider(first_calls,
                                                       crash_on="body-0001"))
    with pytest.raises(RuntimeError, match="simulated crash"):
        main(["review", str(src), "--config", str(cfg_path), "--out", str(out)])

    assert first_calls == ["body-0000", "body-0001"]
    checkpoint_path = out / "checkpoint.json"
    assert checkpoint_path.exists(), (
        "the intra-run checkpoint must be on disk the moment the crash hits — "
        "that is the whole point of wiring it into run_sync")
    # findings.checkpoint.json (the coarser, findings-boundary one) must NOT
    # exist yet: run_sync itself never returned.
    assert not (out / "findings.checkpoint.json").exists()

    second_calls: list[str] = []
    monkeypatch.setattr("docproof.__main__.build_provider",
                        lambda cfg: _CrashingProvider(second_calls))
    # No --resume flag: the per-call checkpoint applies on its own, the same
    # way the app path always does (see app/jobs.py's _checkpoint) — presence
    # plus a matching fingerprint is what makes it apply.
    rc = main(["review", str(src), "--config", str(cfg_path), "--out", str(out)])
    assert rc == 0

    # body-0000's call was already paid for and checkpointed; only the one
    # that crashed is paid for again.
    assert second_calls == ["body-0001"]
    assert not checkpoint_path.exists()               # cleared on success
    assert not (out / "findings.checkpoint.json").exists()

    payload = (out / "findings.json").read_text("utf-8")
    assert payload.count('"status": "validated"') == 2


def test_editing_the_config_between_runs_invalidates_the_checkpoint(
        tmp_path, monkeypatch):
    """The hazard this fixes: the old fingerprint hashed the config's PATH,
    not its contents, so --out reused across two runs of the SAME path but
    DIFFERENT config content would replay findings produced under the config
    that came before — a correctness bug, not just a wasted-money one."""
    src = _manuscript(tmp_path)
    cfg_path = tmp_path / "config.yaml"
    _write_config(cfg_path)                            # min_confidence: medium
    out = tmp_path / "out"

    first_calls: list[str] = []
    monkeypatch.setattr("docproof.__main__.build_provider",
                        lambda cfg: _CrashingProvider(first_calls,
                                                       crash_on="body-0001"))
    with pytest.raises(RuntimeError):
        main(["review", str(src), "--config", str(cfg_path), "--out", str(out)])
    assert first_calls == ["body-0000", "body-0001"]
    assert (out / "checkpoint.json").exists()

    # Same PATH, different CONTENT — min_confidence never touches a system
    # prompt, so this isolates the config-dict comparison itself, not
    # pass_prompts() noticing a changed detection prompt.
    _write_config(cfg_path, **{"min_confidence": "high"})

    second_calls: list[str] = []
    monkeypatch.setattr("docproof.__main__.build_provider",
                        lambda cfg: _CrashingProvider(second_calls))
    rc = main(["review", str(src), "--config", str(cfg_path), "--out", str(out)])
    assert rc == 0

    # Both paragraphs are read fresh: the stale, differently-configured
    # checkpoint was discarded rather than trusted.
    assert second_calls == ["body-0000", "body-0001"]
