"""The two silent edits, and the audit that guards everything else.

These are the only changes docproof makes that a person cannot reject in Word,
so the tests here lean hard on the negative cases: what the normalizer must
*not* touch, and what it must leave straight rather than guess at. A mark
curled the wrong way ships.
"""
from __future__ import annotations

import pytest

from docproof.audit import AuditError, AuditReport, enforce, run_audit
from docproof.config import Config
from docproof.normalize import (NormalizationReport, _quote_edits,
                                normalize_package, normalize_text)
from docproof.utils.xml_helpers import (DocxPackage, iter_text_elements,
                                        paragraph_text, set_text, walk_package)

from .conftest import FIXTURES


def left_straight(text: str) -> int:
    return _quote_edits(text)[1]


# --- double quotes ------------------------------------------------------------

@pytest.mark.parametrize("before,after", [
    ('"Wait here," he said.', '“Wait here,” he said.'),
    ('He said "stop" and left.', 'He said “stop” and left.'),
    ('("nested") and [also "this"]', '(“nested”) and [also “this”]'),
    # A dash before a quote opens it: interrupted speech resuming.
    ('—"and then he left"', '—“and then he left”'),
])
def test_double_quotes_curl_by_context(before, after):
    assert normalize_text(before) == after


# --- apostrophes and single quotes --------------------------------------------

@pytest.mark.parametrize("before,after", [
    # Contractions and possessives: the overwhelming majority, and never
    # ambiguous, because a letter precedes the mark.
    ("It's fine.", "It’s fine."),
    ("The dogs' bowls.", "The dogs’ bowls."),
    ("He was runnin' late.", "He was runnin’ late."),
    ("don't, won't, shouldn't", "don’t, won’t, shouldn’t"),
    # Decades take a right-curling mark, not an opening quote.
    ("Back in the '60s.", "Back in the ’60s."),
    # Dropped leading letters curl right too. This is the case a naive
    # open/close rule gets backwards.
    ("'Tis the season.", "’Tis the season."),
    ("'twas ever thus", "’twas ever thus"),
    ("'ere now, guv", "’ere now, guv"),
    ("rock 'n' roll", "rock ’n’ roll"),
])
def test_apostrophes_curl_right(before, after):
    assert normalize_text(before) == after


def test_a_nested_quotation_is_curled_as_a_pair():
    """Both halves or neither. One straight and one curled reads worse than
    leaving the pair alone."""
    assert (normalize_text('"He said \'hello\' and left," she reported.')
            == '“He said ‘hello’ and left,” she reported.')


def test_an_unmatched_single_mark_is_left_straight_and_counted():
    text = "unmatched 'quote with nothing to close it"
    assert normalize_text(text) == text
    assert left_straight(text) == 1


def test_nothing_is_left_straight_in_ordinary_prose():
    text = "\"It's the dogs' fault,\" he said. \"'Tis always so.\""
    assert left_straight(text) == 0


# --- spaces -------------------------------------------------------------------

@pytest.mark.parametrize("before,after", [
    ("Some  double  spaces.", "Some double spaces."),
    ("Three   in   a   row.", "Three in a row."),
    ("After a sentence.  Another one.", "After a sentence. Another one."),
])
def test_space_runs_collapse(before, after):
    assert normalize_text(before) == after


def test_leading_indentation_is_preserved():
    assert normalize_text("    Indented line.") == "    Indented line."


def test_a_divider_keeps_its_spacing():
    """"*   *   *" is spaced that way on purpose, and it is not prose."""
    assert normalize_text("*   *   *") == "*   *   *"


def test_tabs_are_left_alone():
    assert normalize_text("Name\t\tValue") == "Name\t\tValue"


# --- switches -----------------------------------------------------------------

def test_each_half_can_be_turned_off():
    text = 'He said  "stop".'
    assert normalize_text(text, spaces=False) == 'He said  “stop”.'
    assert normalize_text(text, quotes=False) == 'He said "stop".'
    assert normalize_text(text, quotes=False, spaces=False) == text


# --- against a real package ---------------------------------------------------

def _pkg_with(tmp_path, *texts):
    """A .docx built by editing a fixture's text in place, so the runs are
    real Word runs rather than something this test invented."""
    import docx
    d = docx.Document()
    for t in texts:
        d.add_paragraph(t)
    path = tmp_path / "n.docx"
    d.save(path)
    return DocxPackage(path)


def test_normalizing_a_package_rewrites_the_text(tmp_path):
    pkg = _pkg_with(tmp_path, '"Wait here."  he said.', "It's the '60s.")
    report = normalize_package(pkg)
    assert report.ran and report.quotes == 4 and report.spaces == 1
    assert report.paragraphs == 2

    texts = [paragraph_text(wp.element) for wp in walk_package(pkg)]
    assert texts == ['“Wait here.” he said.', "It’s the ’60s."]


def test_normalization_survives_text_split_across_runs(tmp_path):
    """Word splits a sentence across runs for its own reasons. A quote whose
    context lives in the previous run still has to curl correctly."""
    import docx
    d = docx.Document()
    p = d.add_paragraph()
    p.add_run('He said ')
    p.add_run('"stop"')          # the opening quote's context is the run before
    p.add_run(' and  left.')
    path = tmp_path / "split.docx"
    d.save(path)

    pkg = DocxPackage(path)
    normalize_package(pkg)
    assert [paragraph_text(wp.element) for wp in walk_package(pkg)] == [
        'He said “stop” and left.']


def test_turning_both_off_touches_nothing(tmp_path):
    pkg = _pkg_with(tmp_path, '"Wait here."  he said.')
    report = normalize_package(pkg, quotes=False, spaces=False)
    assert not report.ran
    assert [paragraph_text(wp.element) for wp in walk_package(pkg)] == [
        '"Wait here."  he said.']


# --- the audit ----------------------------------------------------------------

def test_audit_passes_when_the_reject_view_matches():
    baseline = {"body-0000": "The road went on.", "body-0001": "It rained."}
    report = run_audit(baseline, dict(baseline))
    assert report.passed and report.checked == 2
    assert "passed" in report.summary()


def test_audit_fails_on_an_untracked_change():
    baseline = {"body-0000": "The forest was quiet."}
    report = run_audit(baseline, {"body-0000": "The woods was quiet."})
    assert not report.passed
    assert report.mismatches[0].para_id == "body-0000"
    # The message points at where they diverge, not at two whole paragraphs.
    assert "character 4" in report.mismatches[0].describe()


def test_audit_fails_on_a_vanished_paragraph():
    report = run_audit({"body-0000": "Here."}, {})
    assert not report.passed and report.missing == ("body-0000",)


def test_strict_refuses_and_warn_does_not():
    failed = run_audit({"body-0000": "a"}, {"body-0000": "b"})
    with pytest.raises(AuditError, match="No file was written"):
        enforce(failed, "strict")
    enforce(failed, "warn")        # returns quietly
    enforce(failed, "off")


def test_a_passing_audit_never_raises():
    enforce(run_audit({"body-0000": "a"}, {"body-0000": "a"}), "strict")


def test_an_audit_that_did_not_run_never_raises():
    enforce(AuditReport(), "strict")


# --- the two together, end to end ---------------------------------------------

def test_a_real_run_normalizes_and_passes_its_own_audit(tmp_path):
    """The invariant that matters: normalization happens before the baseline
    is taken, so the silent edits are outside the audit by construction, and
    everything else must be inside a tracked change."""
    from docproof.config import load_config
    from docproof.formats import DOCX
    from docproof.pipeline import finish, prepare
    from docproof.models import Usage

    import docx
    d = docx.Document()
    d.add_heading("Chapter One", level=1)
    d.add_paragraph('"Wait here." he said.  She didn\'t move.')
    d.add_paragraph("The forest---which had no name---was quiet.")
    src = tmp_path / "run.docx"
    d.save(src)

    cfg = load_config("config/default.yaml")
    prepared = prepare(cfg, src, "config/error_types")
    assert prepared.normalization.quotes > 0
    assert prepared.normalization.spaces > 0

    out = finish(prepared, [], Usage(), cfg, out_dir=tmp_path / "out",
                 source_path=src)
    assert out.reviewed_path.exists()

    # The sweeps did land, as tracked changes...
    reviewed = DocxPackage(out.reviewed_path)
    from docproof.reassembler import paragraph_view_text
    accepted = [paragraph_view_text(wp.element, "accept")
                for wp in walk_package(reviewed)]
    assert any("—which had no name—" in t for t in accepted)
    assert any("“Wait here,” he said." in t for t in accepted)

    # ...and rejecting them returns every paragraph to the ingested text.
    rejected = DOCX.snapshot(reviewed, "reject")
    assert rejected == prepared.baseline


def test_an_untracked_edit_makes_a_real_run_refuse_to_ship(tmp_path):
    """The audit is only worth having if it fails. Sneak a change past the
    tracked-changes machinery and the run must produce no file at all."""
    from docproof.config import load_config
    from docproof.models import Usage
    from docproof.pipeline import finish, prepare

    import docx
    d = docx.Document()
    d.add_paragraph("The forest was quiet and the road went on.")
    src = tmp_path / "sneak.docx"
    d.save(src)

    cfg = load_config("config/default.yaml")
    prepared = prepare(cfg, src, "config/error_types")
    for wp in walk_package(prepared.pkg):
        t = next(iter_text_elements(wp.element))
        set_text(t, (t.text or "").replace("forest", "woods"))

    out_dir = tmp_path / "out"
    with pytest.raises(AuditError, match="without wrapping it in a tracked"):
        finish(prepared, [], Usage(), cfg, out_dir=out_dir, source_path=src)
    assert not list(out_dir.glob("*.docx")), "a failed audit still shipped a file"
    # The diagnosis survives, though: refusing to ship should not also mean
    # refusing to explain.
    assert (out_dir / "summary.md").exists()
    assert "FAILED" in (out_dir / "summary.md").read_text()


def test_warn_writes_the_file_and_records_the_failure(tmp_path):
    from docproof.config import load_config
    from docproof.models import Usage
    from docproof.pipeline import finish, prepare
    import json

    import docx
    d = docx.Document()
    d.add_paragraph("The forest was quiet and the road went on.")
    src = tmp_path / "warn.docx"
    d.save(src)

    cfg = load_config("config/default.yaml")
    cfg.audit = "warn"
    prepared = prepare(cfg, src, "config/error_types")
    for wp in walk_package(prepared.pkg):
        t = next(iter_text_elements(wp.element))
        set_text(t, (t.text or "").replace("forest", "woods"))

    out = finish(prepared, [], Usage(), cfg, out_dir=tmp_path / "out",
                 source_path=src)
    assert out.reviewed_path.exists()
    payload = json.loads(out.findings_json.read_text())
    assert payload["audit"]["passed"] is False
    assert payload["audit"]["mismatches"] == ["body-0000"]


def test_config_defaults_to_refusing():
    assert Config().audit == "strict"
    assert Config().normalize.quotes and Config().normalize.spaces
