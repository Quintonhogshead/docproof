import json

import pytest

from docproof.candidate_generators import (
    INITIAL_CANDIDATE_TYPES, generate_initial_candidates)
from docproof.candidate_judge import judge_screening_packet
from docproof.candidate_ledger import (
    CandidateLedger, IncompleteCandidateVerdicts,
    validate_candidate_verdict_coverage)
from docproof.candidate_models import (
    Candidate, CandidateAnchor, CandidateStatus, Verdict,
    deterministic_candidate_id)
from docproof.candidate_packet import build_screening_packets
from docproof.candidate_screening import (
    prepare_candidate_screening, validate_candidate_anchor)
from docproof.config import Config, load_config
from docproof.context_service import ContextService
from docproof.models import Anchor, DocumentModel, Finding, ParagraphRef
from docproof.providers import ProviderResult

from .fakes import FakeProvider, USAGE


def _doc(*paragraphs):
    return DocumentModel("book.docx", tuple(paragraphs))


def _para(para_id, text, style="Normal", reviewable=True):
    return ParagraphRef(
        para_id, "word/document.xml", "body", text, style, reviewable)


def _candidate(candidate_id="C-one", start=0, end=3):
    return Candidate(
        candidate_id=candidate_id, candidate_type="number_style",
        generator_id="test.generator",
        anchors=(CandidateAnchor(
            document_part="word/document.xml", paragraph_id="body-0000",
            start_offset=start, end_offset=end),),
        observed_text="123", evidence={"source": "test"})


def test_feature_has_off_shadow_and_apply_modes_with_a_kill_switch(monkeypatch):
    assert Config().candidate_screening.mode == "off"
    assert load_config("config/default.yaml").candidate_screening.mode == "off"
    cfg = Config(candidate_screening={"mode": "shadow"})
    assert cfg.candidate_screening.mode == "shadow"
    assert cfg.candidate_screening.judgment_enabled is True
    monkeypatch.setenv("DOCPROOF_CANDIDATE_SCREENING", "0")
    doc = _doc(_para("body-0000", "This is is wrong."))
    assert prepare_candidate_screening(
        cfg, doc, paragraphs=doc.paragraphs) is None


def test_candidate_ids_are_deterministic_and_ledger_merges_provenance():
    first = _candidate()
    duplicate = first.model_copy(update={
        "generator_id": "other.generator",
        "generator_ids": ("other.generator",),
    })
    expected = deterministic_candidate_id(
        first.candidate_type, first.anchors, first.observed_text,
        first.candidate_correction)
    assert expected.startswith("C-") and len(expected) == 26

    ledger = CandidateLedger()
    candidate_id, created = ledger.register(first)
    merged_id, duplicate_created = ledger.register(duplicate)
    assert created is True and duplicate_created is False
    assert merged_id == candidate_id
    assert ledger.candidate(candidate_id).generator_ids == (
        "test.generator", "other.generator")
    assert len(ledger) == 1
    assert ledger.events[-1].reason_code == "canonical_duplicate_merged"


def test_verdict_coverage_requires_an_exact_partition():
    validate_candidate_verdict_coverage(["a", "b"], ["b", "a"])
    with pytest.raises(IncompleteCandidateVerdicts) as missing:
        validate_candidate_verdict_coverage(["a", "b"], ["a"])
    assert missing.value.missing == ("b",)
    with pytest.raises(IncompleteCandidateVerdicts) as duplicate:
        validate_candidate_verdict_coverage(["a"], ["a", "a"])
    assert duplicate.value.duplicate == ("a",)
    with pytest.raises(IncompleteCandidateVerdicts) as unknown:
        validate_candidate_verdict_coverage(["a"], ["other"])
    assert unknown.value.unknown == ("other",)


def test_initial_rollout_generates_every_locally_generatable_type():
    # term_consistency and grammar are analyzer-sourced (fed via finding_sources
    # from the reused deterministic scans), not produced from raw paragraphs.
    analyzer_sourced = {"term_consistency", "grammar"}
    paragraphs = (
        _para("body-0000", (
            '“Wait.” he said. However we stayed. Hello John. '
            'This is is 5 and the lantern shone while the lantern dimmed. $5. '
            'They took there bags ; the ( stayed open all evening and '
            'it never closed again.')),
        _para("body-0001", "Chapter 1", "Heading1", False),
        _para("body-0002", "Chapter 3", "Heading1", False),
        _para("body-0003", "First item", "ListParagraph", False),
        _para("body-0004", "Second item.", "ListParagraph", False),
    )
    candidates = generate_initial_candidates(_doc(*paragraphs), paragraphs)
    produced = {candidate.candidate_type for candidate in candidates}
    assert (set(INITIAL_CANDIDATE_TYPES) - analyzer_sourced) <= produced
    assert all(candidate.candidate_id.startswith("C-")
               for candidate in candidates)


def test_analyzer_sourced_types_arrive_through_finding_sources():
    from docproof.models import Finding
    paras = (_para("body-0000", "He drank the the potion."),)
    doc = _doc(*paras)
    finding = Finding(
        finding_id="c-0001", chunk_id="consistency", para_id="body-0000",
        error_type="term_consistency", original_text="He drank the the potion.",
        occurrence=1, corrected_text="He drank the the potion.",
        explanation="Written more than one way.", confidence="high")
    candidates = generate_initial_candidates(
        doc, paras, finding_sources={"term_consistency": [finding]})
    assert any(c.candidate_type == "term_consistency" for c in candidates)


def test_direct_address_generator_does_not_treat_please_wait_as_a_name():
    para = _para("body-0000", "Please wait while I find John.")
    candidates = generate_initial_candidates(
        _doc(para), [para], candidate_types=["direct_address_comma"])
    assert candidates == []


def test_existing_deterministic_sweep_hits_are_adapted_with_provenance():
    para = _para("body-0000", "This is is wrong.")
    finding = Finding(
        "s-1", "sweep", para.para_id, "sweep_doubled_word", para.text, 1,
        "This is wrong.", "Doubled word.", "high")
    candidates = generate_initial_candidates(
        _doc(para), [para], candidate_types=["repeated_word"],
        sweep_findings=[finding])
    adapted = [candidate for candidate in candidates
               if candidate.generator_id.startswith("candidate.adapter")]
    assert len(adapted) == 1
    assert adapted[0].evidence["finding_id"] == "s-1"
    assert adapted[0].evidence["local_screening"]["decision"] == "error"
    assert all(candidate.anchors[0].document_part == "word/document.xml"
               for candidate in candidates)


def test_anchor_validation_rejects_stale_text_and_allows_registered_ghosts():
    doc = _doc(_para("body-0000", "123"))
    assert validate_candidate_anchor(_candidate(), doc) is None
    stale = _candidate().model_copy(update={"observed_text": "999"})
    assert validate_candidate_anchor(stale, doc).reason_code \
        == "anchored_text_mismatch"

    anchor = CandidateAnchor(
        document_part="word/document.xml", paragraph_id="body-0000",
        virtual_location="expected closing quotation mark")
    ghost = Candidate(
        candidate_id="C-ghost", candidate_type="quote_balance",
        generator_id="test", anchors=(anchor,), observed_text=None)
    assert validate_candidate_anchor(ghost, doc) is None
    unregistered = ghost.model_copy(update={"candidate_type": "number_style"})
    assert validate_candidate_anchor(unregistered, doc).reason_code \
        == "unregistered_ghost_site"


def test_prepare_resolves_strong_local_cases_and_keeps_ambiguity_explicit(
        tmp_path):
    para = _para("body-0000", "This is is 5. However we waited.")
    doc = _doc(para)
    cfg = Config(candidate_screening={"mode": "shadow"})
    run = prepare_candidate_screening(cfg, doc, paragraphs=[para])
    assert run is not None
    states = run.ledger.state_counts()
    assert states["error"] >= 2
    assert states["packet_ready"] >= 1
    assert len(run.packets) >= 1

    report, ledger_path, markdown_path = run.write(
        tmp_path, cfg.candidate_screening, source=doc.source_path)
    assert ledger_path.is_file() and markdown_path.is_file()
    assert (tmp_path / "candidate-screening-report.json").is_file()
    assert report["scope"]["shadow_only"] is True
    assert report["scope"]["may_modify_documents"] is False
    assert report["accounting"]["all_candidates_have_state"] is True
    stored = json.loads(
        (tmp_path / "candidate-screening-report.json").read_text())
    assert stored["accounting"]["generated_candidates"] == len(run.ledger)


def test_apply_mode_emits_only_safe_errors_through_the_shared_validator(
        monkeypatch):
    from docproof.validator import validate_findings

    monkeypatch.setenv("DOCPROOF_CANDIDATE_APPLY", "1")
    para = _para("body-0000", "This is is wrong.")
    doc = _doc(para)
    cfg = Config(candidate_screening={
        "mode": "apply", "candidate_types": ["repeated_word"],
        "judgment_enabled": False,
    })
    run = prepare_candidate_screening(cfg, doc, paragraphs=[para])

    proposed = run.production_findings()
    assert len(proposed) == 1
    assert proposed[0].finding_id.startswith("cs-")
    assert proposed[0].error_type == "repeated_word"
    validated = validate_findings(
        proposed, doc, cfg.min_confidence, edit_guard=cfg.edit_guard)
    assert validated[0].status == "validated"
    assert validated[0].anchor is not None

    run.observe_findings(
        validated, applied_ids={validated[0].finding_id})
    report = run.report(cfg.candidate_screening, source=doc.source_path)
    assert report["scope"]["may_modify_documents"] is True
    assert report["application"]["applied_tracked_changes"] == 1


def test_report_surfaces_generation_coverage(tmp_path):
    from docproof.candidate_screening import _markdown_report
    para = _para("body-0000", "However he left. This is is wrong.")
    doc = _doc(para)
    cfg = Config(candidate_screening={"mode": "shadow"})
    run = prepare_candidate_screening(cfg, doc, paragraphs=[para])
    report = run.report(cfg.candidate_screening, source=doc.source_path)
    assert "coverage" in report
    assert set(report["coverage"]["counts"]) == {"covered", "partial", "gap"}
    markdown = _markdown_report(report)
    assert "Generation coverage" in markdown


def test_queries_become_comments_not_edits(monkeypatch):
    # P3-05: an uncertain candidate becomes a force_query comment (never an edit)
    # once Apply is released, and nothing when it is not.
    import docproof.candidate_screening as cs
    monkeypatch.setenv("DOCPROOF_CANDIDATE_APPLY", "1")
    para = _para("body-0000", "They went there today.")
    doc = _doc(para)
    run = cs.CandidateScreeningRun(cs.CandidateLedger(), doc, mode="apply")
    cand = Candidate(
        candidate_id="C-q", candidate_type="homophone",
        generator_id="candidate.homophone",
        anchors=(CandidateAnchor(
            document_part="word/document.xml", paragraph_id="body-0000",
            start_offset=10, end_offset=15),),
        observed_text="there", channel_preference="either")
    run.ledger.register(cand)
    run.ledger.apply_verdict(Verdict(
        candidate_id="C-q", decision="uncertain", confidence="low",
        reason_code="confusable_word_needs_context",
        explanation="Confusable with 'their'.", judge_id="test"),
        stage="test")

    queries = run.production_queries()
    assert len(queries) == 1
    assert queries[0].force_query is True
    assert queries[0].corrected_text == queries[0].original_text  # not an edit
    assert run.production_findings() == []  # an uncertain is never an edit

    monkeypatch.delenv("DOCPROOF_CANDIDATE_APPLY", raising=False)
    run2 = cs.CandidateScreeningRun(cs.CandidateLedger(), doc, mode="apply")
    run2.ledger.register(cand)
    run2.ledger.apply_verdict(Verdict(
        candidate_id="C-q", decision="uncertain", confidence="low",
        reason_code="x", explanation="y", judge_id="t"), stage="test")
    assert run2.production_queries() == []  # contained until release


def test_insufficient_context_defers_instead_of_judging(monkeypatch):
    # P3-02: a candidate whose recipe needs context ContextService cannot build
    # is deferred with an explicit reason, never packaged and judged on partial
    # context.
    import docproof.candidate_screening as cs
    from docproof.context_service import unsupported_recipes

    assert unsupported_recipes(("current paragraph",)) == ()
    assert unsupported_recipes(("scene summary", "current paragraph")) == (
        "scene summary",)

    para = _para("body-0000", "The entity moved.")
    doc = _doc(para)
    needy = Candidate(
        candidate_id="C-needy", candidate_type="number_style",
        generator_id="test.generator",
        anchors=(CandidateAnchor(
            document_part="word/document.xml", paragraph_id="body-0000",
            start_offset=0, end_offset=3),),
        observed_text="The", evidence={"source": "test"},
        context_recipe=("scene summary",),
        status=CandidateStatus.GENERATED)
    monkeypatch.setattr(cs, "generate_initial_candidates",
                        lambda *a, **k: [needy])
    cfg = Config(candidate_screening={"mode": "shadow"}).candidate_screening
    run = cs.CandidateScreeningRun.prepare(cfg, doc, paragraphs=[para])
    assert run.ledger.status("C-needy") == CandidateStatus.DEFERRED
    events = [e for e in run.ledger.events if e.candidate_id == "C-needy"]
    assert any(e.reason_code == "insufficient_context" for e in events)
    assert not run.packets


def test_standalone_mode_reuses_local_analyzers_as_candidates():
    # P2-01/02: the candidate-only profile must not fall back to the regex
    # generators alone — the free deterministic sweeps feed the ledger too.
    from docproof.profiles import CANDIDATE_ONLY, apply_profile

    paras = tuple(
        ParagraphRef(f"body-{i:04d}", "word/document.xml", "body", t, "Normal", True)
        for i, t in enumerate([
            "He slammed the door and left the room!!",
            "She took there bags from the hall table.",
        ]))
    doc = DocumentModel("book.docx", paras)
    cfg = apply_profile(load_config("config/default.yaml"), CANDIDATE_ONLY)
    run = prepare_candidate_screening(cfg, doc, paragraphs=list(paras))
    generators = {c.generator_id for c in run.ledger.candidates}
    types = {c.candidate_type for c in run.ledger.candidates}
    # Reused deterministic ERROR sweeps became candidates (not direct findings)...
    assert any(g.startswith("candidate.adapter.sweep_") for g in generators)
    # ...alongside the signal-gated lexical/punctuation generators.
    assert {"homophone", "punctuation_style"} <= types
    # Containment still holds: standalone apply is clamped to shadow.
    assert run.mode == "shadow"


def test_reuse_local_analyzers_can_be_disabled():
    from docproof.profiles import CANDIDATE_ONLY, apply_profile

    paras = (ParagraphRef("body-0000", "word/document.xml", "body",
                          "He said . . . nothing -- then left.", "Normal", True),)
    doc = DocumentModel("book.docx", paras)
    cfg = apply_profile(load_config("config/default.yaml"), CANDIDATE_ONLY)
    cfg.candidate_screening.reuse_local_analyzers = False
    run = prepare_candidate_screening(cfg, doc, paragraphs=list(paras))
    generators = {c.generator_id for c in run.ledger.candidates}
    assert not any(g.startswith("candidate.adapter.sweep_") for g in generators)


def test_apply_is_contained_to_shadow_until_the_release_gate_opens(monkeypatch):
    # P0-01: without the release gate, a requested apply mode must not be able to
    # modify a document. It is clamped to shadow, the run reports itself as
    # non-mutating, and no production findings are emitted.
    monkeypatch.delenv("DOCPROOF_CANDIDATE_APPLY", raising=False)
    para = _para("body-0000", "This is is wrong.")
    doc = _doc(para)
    cfg = Config(candidate_screening={
        "mode": "apply", "candidate_types": ["repeated_word"],
        "judgment_enabled": False,
    })
    run = prepare_candidate_screening(cfg, doc, paragraphs=[para])
    assert run.mode == "shadow"
    assert run.production_findings() == []
    report = run.report(cfg.candidate_screening, source=doc.source_path)
    assert report["scope"]["may_modify_documents"] is False


def test_release_gate_env_opt_in_restores_apply(monkeypatch):
    monkeypatch.setenv("DOCPROOF_CANDIDATE_APPLY", "1")
    para = _para("body-0000", "This is is wrong.")
    doc = _doc(para)
    cfg = Config(candidate_screening={
        "mode": "apply", "candidate_types": ["repeated_word"],
        "judgment_enabled": False,
    })
    run = prepare_candidate_screening(cfg, doc, paragraphs=[para])
    assert run.mode == "apply"
    assert len(run.production_findings()) == 1


def test_model_packet_retries_only_missing_candidate_ids():
    doc = _doc(_para("body-0000", "123 456"))
    first = _candidate("C-a", 0, 3)
    second = _candidate("C-b", 4, 7).model_copy(update={
        "observed_text": "456"})
    packet = build_screening_packets(
        [first, second], ContextService(doc), batch_size=100)[0]
    provider = FakeProvider([
        ProviderResult(parsed={
            "pass_ids": ["C-a"], "errors": [], "uncertain": [],
            "deferred": [],
        }, usage=USAGE),
        ProviderResult(parsed={
            "pass_ids": ["C-b"], "errors": [], "uncertain": [],
            "deferred": [],
        }, usage=USAGE),
    ])

    verdicts = judge_screening_packet(
        packet, provider, model="gpt-5.6-luna", max_tokens=500,
        retry_log=(retries := []))
    assert [verdict.candidate_id for verdict in verdicts] == ["C-a", "C-b"]
    assert len(provider.calls) == 2
    assert '"candidate_id":"C-a"' in provider.calls[0]["user"]
    assert '"candidate_id":"C-a"' not in provider.calls[1]["user"]
    assert '"candidate_id":"C-b"' in provider.calls[1]["user"]
    assert retries == [{
        "attempt": 1, "missing_ids": ["C-b"], "invalid_ids": [],
        "retried_ids": ["C-b"],
    }]


def test_model_packet_retries_an_invalid_correction_only():
    doc = _doc(_para("body-0000", "123 456"))
    first = _candidate("C-a", 0, 3)
    second = _candidate("C-b", 4, 7).model_copy(update={
        "observed_text": "456"})
    packet = build_screening_packets(
        [first, second], ContextService(doc), batch_size=100)[0]
    provider = FakeProvider([
        ProviderResult(parsed={
            "pass_ids": ["C-a"],
            "errors": [{
                "candidate_id": "C-b", "corrected_text": "456",
                "reason_code": "number_style", "explanation": "No change.",
                "confidence": "high",
            }],
            "uncertain": [], "deferred": [],
        }, usage=USAGE),
        ProviderResult(parsed={
            "pass_ids": ["C-b"], "errors": [], "uncertain": [],
            "deferred": [],
        }, usage=USAGE),
    ])

    verdicts = judge_screening_packet(
        packet, provider, model="gpt-5.6-luna", max_tokens=500)

    assert [row.decision for row in verdicts] == ["pass", "pass"]
    assert '"candidate_id":"C-a"' not in provider.calls[1]["user"]
    assert '"candidate_id":"C-b"' in provider.calls[1]["user"]


def test_model_screening_is_ledger_only_and_records_packet_ids():
    para = _para("body-0000", "5")
    doc = _doc(para)
    cfg = Config(candidate_screening={
        "mode": "shadow", "candidate_types": ["number_style"],
        "judgment_enabled": True, "max_output_tokens": 500,
        "max_cost_usd": 2.0,
    })
    run = prepare_candidate_screening(cfg, doc, paragraphs=[para])
    candidate_id = run.packets[0].candidate_ids[0]
    provider = FakeProvider([ProviderResult(parsed={
        "pass_ids": [candidate_id], "errors": [], "uncertain": [],
        "deferred": [],
    }, usage=USAGE)])

    usage = run.screen(cfg, provider_factory=lambda _cfg: provider)

    assert usage.api_calls == 1
    assert run.ledger.status(candidate_id) == CandidateStatus.PASS
    assert run.packet_records[0]["submitted_ids"] == [candidate_id]
    assert run.packet_records[0]["returned_ids"] == [candidate_id]
    # The shadow object owns no Finding list and exposes no document writer.
    assert not hasattr(run, "findings")
    assert not hasattr(run, "apply")


def test_open_discovery_findings_enter_the_same_ledger():
    para = _para("body-0000", "The teh cat slept.")
    doc = _doc(para)
    cfg = Config(candidate_screening={
        "mode": "shadow", "candidate_types": ["number_style"]})
    run = prepare_candidate_screening(cfg, doc, paragraphs=[para])
    finding = Finding(
        "f-1", "chunk-000", para.para_id, "spelling", para.text, 1,
        "The the cat slept.", "Teh is misspelled.", "high",
        status="validated", anchor=Anchor(4, 7, "teh", "the"))

    run.observe_findings([finding])

    assert len(run.discovery_ids) == 1
    candidate_id = next(iter(run.discovery_ids))
    assert run.ledger.candidate(candidate_id).generator_id \
        == "open_discovery.existing_reviewer"
    assert run.ledger.status(candidate_id) == CandidateStatus.ERROR
    assert run.production_matches[candidate_id][0]["finding_id"] == "f-1"


def test_verdict_schema_requires_explanations_for_nontrivial_decisions():
    with pytest.raises(ValueError, match="require an explanation"):
        Verdict(
            candidate_id="C-one", decision="error", corrected_text="124",
            reason_code="bad_number", judge_id="test")


def test_prepare_and_finish_write_shadow_artifacts_without_adding_findings(
        tmp_path):
    from docproof.__main__ import _run_mock
    from docproof.pipeline import finish, prepare
    from .conftest import FIXTURES

    cfg = load_config("config/default.yaml")
    cfg.examination_graph.enabled = False
    cfg.candidate_screening.mode = "shadow"
    cfg.candidate_screening.judgment_enabled = False
    prepared = prepare(
        cfg, FIXTURES / "simple.docx", FIXTURES.parent.parent /
        "config" / "error_types")
    findings, usage = _run_mock(cfg, prepared, [])
    model_finding_count = len(findings)

    outputs = finish(
        prepared, findings, usage, cfg, out_dir=tmp_path,
        source_path=FIXTURES / "simple.docx")

    assert len(findings) == model_finding_count
    assert outputs.candidate_screening_ledger.is_file()
    assert outputs.candidate_screening_report.is_file()
    assert outputs.reviewed_path.is_file()
    report = json.loads(
        (tmp_path / "candidate-screening-report.json").read_text())
    assert report["scope"]["may_create_findings"] is False


def test_candidate_only_profile_writes_a_guarded_tracked_change(
        tmp_path, monkeypatch):
    from docx import Document
    from zipfile import ZipFile

    from docproof.__main__ import _run_mock
    from docproof.pipeline import finish, prepare
    from docproof.profiles import CANDIDATE_ONLY, apply_profile
    from .test_error_types import ERROR_DIR

    monkeypatch.setenv("DOCPROOF_CANDIDATE_APPLY", "1")
    source = tmp_path / "candidate-only.docx"
    document = Document()
    document.add_paragraph("This is is wrong.")
    document.save(source)

    cfg = apply_profile(load_config("config/default.yaml"), CANDIDATE_ONLY)
    prepared = prepare(cfg, source, ERROR_DIR)
    findings, usage = _run_mock(cfg, prepared, [])
    outputs = finish(
        prepared, findings, usage, cfg, out_dir=tmp_path / "out",
        source_path=source)

    assert outputs.applied == 1
    with ZipFile(outputs.reviewed_path) as archive:
        xml = archive.read("word/document.xml")
    assert b"<w:del" in xml
    report = json.loads(
        (tmp_path / "out" / "candidate-screening-report.json").read_text())
    assert report["application"]["applied_tracked_changes"] == 1


def test_candidate_only_is_valid_batch_work_without_detector_requests(
        tmp_path, monkeypatch):
    from docx import Document

    from docproof import batch
    from docproof.profiles import CANDIDATE_ONLY, apply_profile
    from .fakes import FakeProvider
    from .test_error_types import ERROR_DIR

    monkeypatch.setenv("DOCPROOF_CANDIDATE_APPLY", "1")
    source = tmp_path / "candidate-batch.docx"
    document = Document()
    document.add_paragraph("This is is wrong.")
    document.save(source)
    cfg = apply_profile(load_config("config/default.yaml"), CANDIDATE_ONLY)
    cfg.candidate_screening.judgment_enabled = False
    provider = FakeProvider()

    job = batch.submit(
        cfg, source, ERROR_DIR, provider, tmp_path / "batch-workspace")
    assert job.batch_id == "" and job.request_count == 0
    assert not provider.calls
    batch.poll(job, provider, tmp_path / "batch-workspace")
    assert job.state == "ready"

    result = batch.collect_findings(job, provider, ERROR_DIR)
    assert any(f.finding_id.startswith("cs-") for f in result.findings)


def test_multi_round_reviews_run_candidate_generation_once():
    from docproof.rounds import _round_config

    cfg = Config(candidate_screening={"mode": "apply"})
    assert _round_config(cfg, 1, True).candidate_screening.mode == "apply"
    assert _round_config(cfg, 2, True).candidate_screening.mode == "off"
