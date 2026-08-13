"""Promo engine, end to end, without a model.

The whole pipeline runs offline here: a canned answer through finish for the
happy path, a FakeProvider through run to exercise schema validation, and the
grounding check on a manuscript with one invented name planted in the copy.
"""
from __future__ import annotations

import json
from pathlib import Path

import docx
import pytest

from docproof import promo
from docproof.ingest import IngestError
from docproof.models import Usage
from docproof.promo import (ClaimCheck, MarketingPlan, PlanMeta, PromoError,
                            PromoResult, SocialPost, verify_claims)
from tests.fakes import FakeProvider, USAGE
from docproof.providers import ProviderResult

CONFIG_DIR = Path(__file__).parent.parent / "config"

# "Mira" and "Aldous" are in the book below; "Zorbon" is not.
_PARAS = [
    "Mira crossed the Aldous bridge at first light, the river low beneath it.",
    "She had not slept. The letter from her brother lay folded in her coat.",
    "By noon the market at Aldous square was loud, and no one looked at her.",
    "Whatever waited on the far shore, Mira meant to reach it before dark.",
]


def _make_novel(path: Path) -> None:
    d = docx.Document()
    d.add_heading("Chapter One", level=1)
    for text in _PARAS:
        d.add_paragraph(text)
    d.save(str(path))


def _canned() -> PromoResult:
    return PromoResult(
        teaser="When Mira crosses the Aldous bridge, the Zorbon empire stirs.",
        posts=[SocialPost(text=f"Will Mira reach the far shore? ({i + 1})",
                          hashtags=["Fiction"]) for i in range(12)])


def test_prepare_reads_whole_manuscript(tmp_path, cfg):
    src = tmp_path / "Rowan - book 1.docx"
    _make_novel(src)
    prepared = promo.prepare(cfg, src, config_dir=CONFIG_DIR)
    assert prepared.post_count == 12
    assert prepared.manuscript.title == "Rowan - book 1"
    # The whole book is in the user turn, headings and all.
    assert "Aldous bridge" in prepared.user
    assert "{post_count}" not in prepared.system   # placeholder was filled
    assert "12" in prepared.system


def test_finish_writes_two_docx_and_json(tmp_path, cfg):
    src = tmp_path / "Rowan - book 1.docx"
    _make_novel(src)
    prepared = promo.prepare(cfg, src, config_dir=CONFIG_DIR)
    result, usage = promo.run_mock(prepared, canned=_canned())
    out = promo.finish(prepared, result, usage, cfg,
                       out_dir=tmp_path / "out", source_path=src)

    assert set(out.documents) == {"teaser", "posts"}
    assert out.documents["teaser"].name == "Rowan - book 1 - teaser.docx"
    assert out.documents["posts"].name == "Rowan - book 1 - social posts.docx"
    assert all(p.exists() for p in out.documents.values())

    saved = json.loads(out.data_json.read_text(encoding="utf-8"))
    assert saved["teaser"].startswith("When Mira")
    assert len(saved["posts"]) == 12


def test_grounding_flags_invented_names_only(tmp_path, cfg):
    src = tmp_path / "Rowan - book 1.docx"
    _make_novel(src)
    prepared = promo.prepare(cfg, src, config_dir=CONFIG_DIR)
    result, usage = promo.run_mock(prepared, canned=_canned())
    out = promo.finish(prepared, result, usage, cfg,
                       out_dir=tmp_path / "out", source_path=src)

    # "Zorbon" appears nowhere in the manuscript → flagged.
    assert "Zorbon" in out.unverified
    # Names that ARE in the book do not flag; nor does a sentence opener.
    assert "Mira" not in out.unverified
    assert "Aldous" not in out.unverified
    assert "When" not in out.unverified


def test_run_validates_provider_answer(tmp_path, cfg):
    src = tmp_path / "Rowan - book 1.docx"
    _make_novel(src)
    prepared = promo.prepare(cfg, src, config_dir=CONFIG_DIR)

    from tests.fakes import FakeProvider
    provider = FakeProvider(results=[
        ProviderResult(parsed=_canned().model_dump(), usage=USAGE)])
    result, usage = promo.run(cfg, prepared, provider)

    assert len(result.posts) == 12
    assert result.teaser.startswith("When Mira")
    assert usage.api_calls == 1


def test_run_raises_on_unusable_answer(tmp_path, cfg):
    src = tmp_path / "Rowan - book 1.docx"
    _make_novel(src)
    prepared = promo.prepare(cfg, src, config_dir=CONFIG_DIR)

    from tests.fakes import FakeProvider
    provider = FakeProvider(results=[
        ProviderResult(parsed=None, stop_reason="refusal",
                       error="declined", usage=USAGE)])
    with pytest.raises(PromoError):
        promo.run(cfg, prepared, provider)


def test_prepare_refuses_a_book_over_the_limit(tmp_path, cfg):
    src = tmp_path / "Rowan - book 1.docx"
    _make_novel(src)
    cfg.promo.max_input_tokens = 5           # smaller than any real manuscript
    from docproof.promo import PromoTooLarge
    with pytest.raises(PromoTooLarge) as excinfo:
        promo.prepare(cfg, src, config_dir=CONFIG_DIR)
    # It is a PromoError too, so every existing handler still catches it.
    assert isinstance(excinfo.value, PromoError)
    # And it carries the numbers a caller needs to act (email, or the override).
    assert excinfo.value.limit == 5
    assert excinfo.value.tokens > 5
    assert excinfo.value.words > 0


def test_allow_oversize_runs_a_book_over_the_limit(tmp_path, cfg):
    """The human override: past the single-pass limit, prepare goes through
    instead of raising, so a person who chose to run anyway can."""
    src = tmp_path / "Rowan - book 1.docx"
    _make_novel(src)
    cfg.promo.max_input_tokens = 5
    prepared = promo.prepare(cfg, src, config_dir=CONFIG_DIR, allow_oversize=True)
    assert prepared.est_input_tokens > cfg.promo.max_input_tokens
    assert "Aldous bridge" in prepared.user   # the whole book still went in


def test_prepare_refuses_non_docx(tmp_path, cfg):
    src = tmp_path / "Rowan - book 1.txt"
    src.write_text("not a docx", encoding="utf-8")
    with pytest.raises(IngestError):
        promo.prepare(cfg, src, config_dir=CONFIG_DIR)


# --- the entailment (claim) grounding pass ------------------------------------

def _prepared(tmp_path, cfg):
    src = tmp_path / "Rowan - book 1.docx"
    _make_novel(src)
    return promo.prepare(cfg, src, config_dir=CONFIG_DIR)


def test_verify_claims_flags_the_unsupported_ones(tmp_path, cfg):
    prepared = _prepared(tmp_path, cfg)
    report = {"claims": [
        {"claim": "Mira crosses the Aldous bridge", "supported": True, "note": ""},
        {"claim": "The Zorbon empire falls", "supported": False,
         "note": "the book never mentions it"}]}
    provider = FakeProvider(results=[ProviderResult(parsed=report, usage=USAGE)])
    usage = Usage()

    claims = verify_claims(cfg, prepared, _canned(), provider, usage,
                           config_dir=CONFIG_DIR)

    assert [c.claim for c in claims if not c.supported] == \
        ["The Zorbon empire falls"]
    assert usage.api_calls == 1           # its cost folded into the run's


def test_verify_claims_never_fails_the_run(tmp_path, cfg):
    """A grounding check that broke must not sink the copy it was checking."""
    prepared = _prepared(tmp_path, cfg)
    provider = FakeProvider(results=[
        ProviderResult(parsed=None, stop_reason="error", error="down",
                       usage=USAGE)])
    assert verify_claims(cfg, prepared, _canned(), provider, Usage(),
                         config_dir=CONFIG_DIR) == []


def test_finish_records_unsupported_claims(tmp_path, cfg):
    prepared = _prepared(tmp_path, cfg)
    result, usage = promo.run_mock(prepared, canned=_canned())
    claims = [ClaimCheck(claim="Made up", supported=False, note="not in book"),
              ClaimCheck(claim="Real", supported=True)]

    out = promo.finish(prepared, result, usage, cfg, out_dir=tmp_path / "out",
                       source_path=prepared.manuscript.source_path, claims=claims)

    assert out.unsupported_claims == ("Made up",)
    # flag_count combines the term check ("Zephyr"? no — canned uses Zorbon) and
    # the unsupported claim.
    assert out.flag_count == len(out.unverified) + 1
    saved = json.loads((tmp_path / "out" / "promo_verify.json")
                       .read_text("utf-8"))
    assert len(saved["claims"]) == 2 and saved["terms"] == list(out.unverified)


# --- the marketing plan -------------------------------------------------------

_PLAN_MD = (
    "# Rowan - book 1 — Rowan Ash\n\n"
    "This is a resource for you, the author, to consider some publicity "
    "efforts and endeavors to pursue!\n\n"
    "## Target Audience\n\n"
    "- **Bridge-crossers:** readers who like a quiet, tense journey.\n\n"
    "## Marketing Strategy\n\n"
    "- **Riverside readings:** stage events by water to echo the setting.\n\n"
    "## Comparable Titles\n\n"
    "- **A Real Novel by Someone Else:** shares the lone-traveller mood.\n\n"
    "## Key Message\n\n"
    "One woman's walk to the far shore is a walk toward herself.")


def test_prepare_plan_fills_metadata_and_falls_back_to_filename(tmp_path, cfg):
    src = tmp_path / "Rowan - book 1.docx"
    _make_novel(src)
    # A blank title falls back to the manuscript name; the rest is threaded in.
    meta = PlanMeta(title="", author_name="Rowan Ash", keywords="rivers, journey",
                    back_cover="A brave, quiet book.", author_city="Austin, TX")
    prepared = promo.prepare_plan(cfg, src, meta, config_dir=CONFIG_DIR)

    assert prepared.meta.title == "Rowan - book 1"
    assert "Aldous bridge" in prepared.user        # the whole book went in
    assert "Rowan Ash" in prepared.user
    assert "rivers, journey" in prepared.user
    assert "Austin, TX" in prepared.user


def test_prepare_plan_blank_optionals_become_dashes_not_the_name(tmp_path, cfg):
    src = tmp_path / "Rowan - book 1.docx"
    _make_novel(src)
    prepared = promo.prepare_plan(cfg, src, PlanMeta(title="Book"),
                                  config_dir=CONFIG_DIR)
    # Optional fields read as "none given"; the pen name is left empty (not a
    # dash) so the title heading is never "Book — —".
    assert "Keywords: —" in prepared.user
    assert "Author's city: —" in prepared.user
    assert "printed as-is): —" not in prepared.user


def test_run_plan_validates_and_finish_writes_docx(tmp_path, cfg):
    src = tmp_path / "Rowan - book 1.docx"
    _make_novel(src)
    prepared = promo.prepare_plan(cfg, src, PlanMeta(title="", author_name="Rowan Ash"),
                                  config_dir=CONFIG_DIR)
    provider = FakeProvider(results=[
        ProviderResult(parsed=MarketingPlan(plan=_PLAN_MD).model_dump(),
                       usage=USAGE)])
    plan, usage = promo.run_plan(cfg, prepared, provider)
    assert usage.api_calls == 1
    assert plan.plan.startswith("# Rowan")

    out = promo.finish_plan(prepared, plan, cfg, out_dir=tmp_path / "out",
                            source_path=src)
    assert out.document.name == "Rowan - book 1 - marketing plan.docx"
    assert out.document.exists()
    saved = json.loads(out.data_json.read_text("utf-8"))
    assert saved["plan"].startswith("# Rowan")

    # The Markdown parsed into headings and bullets, and the bold label lost its
    # asterisks (a bold run, not literal **text**).
    doc = docx.Document(str(out.document))
    heads = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert "Target Audience" in heads and "Comparable Titles" in heads
    assert all("**" not in p.text for p in doc.paragraphs)


def test_run_plan_raises_on_empty_plan(tmp_path, cfg):
    src = tmp_path / "Rowan - book 1.docx"
    _make_novel(src)
    prepared = promo.prepare_plan(cfg, src, PlanMeta(title="Book"),
                                  config_dir=CONFIG_DIR)
    provider = FakeProvider(results=[
        ProviderResult(parsed=MarketingPlan(plan="   ").model_dump(),
                       usage=USAGE)])
    with pytest.raises(PromoError):
        promo.run_plan(cfg, prepared, provider)


def test_write_plan_survives_plain_prose(tmp_path, cfg):
    """A plan that arrives with no Markdown marks still writes — headingless
    paragraphs under a fallback title — rather than losing the copy."""
    from docproof.promo.writers import write_plan
    path = write_plan(tmp_path / "p.docx",
                      "Just some prose with no marks at all.", title="Book")
    doc = docx.Document(str(path))
    assert any("Book — Marketing Plan" in p.text for p in doc.paragraphs)
    assert any("Just some prose" in p.text for p in doc.paragraphs)
