"""galley/settle.py — residual settlement end to end, on real .docx builds.

Every test builds a run the way a Galley run does (a $0 mock review or an
import-findings replay), plants verify artifacts by hand (what `galley verify`
would have written), runs `galley settle`, and checks the deliverable, the
settlement record, and certify. No model call is made: the deterministic lane
settles what it can and a scripted provider stands in for the judge / delta
verifier where a test needs one.
"""
from __future__ import annotations

import json
from pathlib import Path

import docx
import yaml

import docproof.__main__ as m
from docproof.__main__ import main
from docproof.config import load_config
from docproof.ingest import build_document_model, preflight
from docproof.providers import NormalizedUsage, ProviderResult
from galley.manifest import certify_run
from galley.settle import Settlement, open_items, terminal_state
from galley.verify import paragraph_views, problem_id, residual_id

CONFIG = str(Path("config/default.yaml"))

PARAGRAPHS = [
    "The lamp on teh desk flickered, then finally caught, and Maria sighed.",
    "Its light was thin, but it was enough to recieve the letter by tonight.",
    "David counted the coins twice, sure the total would not change again.",
    "The garden gate creaked open, and the old dog limped out to greet her.",
    "She wrote Mom a note and left it under the lamp before she slept.",
]


class _Provider:
    """Replays one scripted parsed body per call; records every call."""

    def __init__(self, *bodies):
        self._bodies = list(bodies)
        self.calls = []

    def complete_structured(self, **kwargs):
        self.calls.append(kwargs)
        body = self._bodies.pop(0) if self._bodies else {}
        return ProviderResult(parsed=body, usage=NormalizedUsage(10, 5),
                              stop_reason="ok")


def _manuscript(tmp_path, paragraphs=PARAGRAPHS, name="book.docx"):
    d = docx.Document()
    for p in paragraphs:
        d.add_paragraph(p)
    src = tmp_path / name
    d.save(src)
    return src


def _para_ids(src):
    cfg = load_config(CONFIG)
    doc = build_document_model(preflight(str(src), cfg.tracked_changes_policy),
                               cfg)
    return [p.para_id for p in doc.paragraphs], doc


def _replay_config(tmp_path, **extra):
    """A $0 rebuild config: the shipped defaults with no typed passes and
    every paid stage off (what final-replay materializes)."""
    cfg = yaml.safe_load(Path(CONFIG).read_text("utf-8"))
    cfg["error_types"] = []
    for k in ("languagetool", "spellcheck", "consistency", "glossary",
              "residuals", "meaning_check", "fix_check", "factcheck",
              "continuity", "chapter_continuity", "chapter_sweep",
              "smoothing", "rewrite", "repair", "storysheet"):
        cfg.setdefault(k, {})
        if isinstance(cfg[k], dict):
            cfg[k]["enabled"] = False
    cfg["ensemble"] = {"detectors": [], "verify_policy": "none"}
    cfg["candidate_screening"] = {"mode": "off"}
    cfg.update(extra)
    path = tmp_path / "replay.yaml"
    path.write_text(yaml.safe_dump(cfg), encoding="utf-8")
    return str(path)


def _build(tmp_path, src, rows, out="run"):
    """A finished build from import-findings rows: findings.json, the docx,
    editmap.json."""
    rows_path = tmp_path / "rows.json"
    rows_path.write_text(json.dumps(rows), encoding="utf-8")
    run = tmp_path / out
    rc = main(["import-findings", str(rows_path), str(src), "--config",
               _replay_config(tmp_path), "--out", str(run)])
    assert rc == 0
    assert (run / "editmap.json").exists()
    return run


def _walk(run, residuals, problems=()):
    (run / "finished_walk.json").write_text(json.dumps({
        "generated_at": "2099-01-01T00:00:00+00:00", "ran": True,
        "residuals": list(residuals)}), encoding="utf-8")
    (run / "change_verify.json").write_text(json.dumps({
        "generated_at": "2099-01-01T00:00:00+00:00", "ran": True,
        "applied_edits": 0, "problems": list(problems)}), encoding="utf-8")


def _settle(tmp_path, run, src, *extra):
    return main(["galley", "settle", str(run), "--source", str(src),
                 "--config", _replay_config(tmp_path), "--engine", "none",
                 "--no-verify", *extra])


def _accepted(run):
    return paragraph_views(run)[1]


def _records(run):
    st = Settlement.load(run)
    assert st is not None
    return {r.residual_id: r for r in st.latest().values()}, st


# --- T1: every residual class settles to a terminal state --------------------

def test_settle_closes_owned_unowned_zone_note_and_fact_residuals(tmp_path):
    src = _manuscript(tmp_path)
    ids, doc = _para_ids(src)
    p0, p1, p2, p4 = ids[0], ids[1], ids[2], ids[4]
    # An applied edit that is itself wrong (teh -> thee) owns a span in p0.
    run = _build(tmp_path, src, [
        {"para_id": p0, "original_text": "teh", "corrected_text": "thee",
         "explanation": "typo", "confidence": "high"},
    ])
    zones = tmp_path / "zones.json"
    zones.write_text(json.dumps({"zones": [
        {"label": "Mom caps", "category": "declared_caps",
         "permission": "locked", "terms": ["Mom"]}]}), encoding="utf-8")
    acc = _accepted(run)
    assert "thee desk" in acc[p0]
    _walk(run, [
        # (a) inside an owned span: the owner's replacement is revised
        {"para_id": p0, "quote": "thee", "problem": "wrong word",
         "suggestion": "the", "severity": "high"},
        # (b) untouched text: a new edit, with an editorial note stripped
        {"para_id": p1, "quote": "recieve", "problem": "misspelling",
         "suggestion": "receive (spell correctly)", "severity": "high"},
        # (c) a fact change: a query, never an edit
        {"para_id": p2, "quote": "coins twice", "problem": "count?",
         "suggestion": "coins 3 times", "severity": "low"},
        # (d) inside a locked intent zone: dropped
        {"para_id": p4, "quote": "Mom", "problem": "capitalization",
         "suggestion": "mom", "severity": "medium"},
        # (e) a quote the accepted text does not hold: dropped(unanchorable)
        {"para_id": p1, "quote": "zzzz", "problem": "?", "suggestion": "y",
         "severity": "low"},
    ])
    cfgpath = _replay_config(tmp_path, intent_zones_file=str(zones))
    rc = main(["galley", "settle", str(run), "--source", str(src), "--config",
               cfgpath, "--engine", "none", "--no-verify"])
    assert rc == 0

    recs, st = _records(run)
    assert recs[residual_id(p0, "thee")].action == "absorb"
    assert recs[residual_id(p1, "recieve")].action == "add"
    q = recs[residual_id(p2, "coins twice")]
    assert q.action == "query" and q.reason.startswith("fact:")
    z = recs[residual_id(p4, "Mom")]
    assert z.action == "drop" and z.reason.startswith("intent_zone")
    u = recs[residual_id(p1, "zzzz")]
    assert u.action == "drop" and u.reason == "unanchorable"
    assert st.open == [] and st.rounds == 1

    # the deliverable: composite revised, new edit landed, zone untouched
    acc = _accepted(run)
    assert "the desk" in acc[p0] and "thee" not in acc[p0]
    assert "receive" in acc[p1] and "spell correctly" not in acc[p1]
    assert "coins twice" in acc[p2]
    assert "Mom" in acc[p4]
    # I6: reject-all is the source, byte for byte
    original = paragraph_views(run)[0]
    for p in doc.paragraphs:
        assert original[p.para_id] == p.text
    # the fact query reached the margin
    env = json.loads((run / "findings.json").read_text("utf-8"))
    queries = [r for r in env["findings"] if r.get("queried")]
    assert any(r["para_id"] == p2 and "3 times" in r["explanation"]
               for r in queries)
    # I1: every row terminal, and stamped
    assert all(r.get("state") in ("applied", "dropped", "query")
               for r in env["findings"])
    composite = [r for r in env["findings"]
                 if r.get("chunk_id") == f"settle:{residual_id(p0, 'thee')}"
                 and r.get("applied")]
    assert composite and composite[0]["error_type"] == "galley_settle"
    assert composite[0]["absorbed"]
    # nothing is open; certify's settlement checks pass
    assert open_items(run) == []
    cert = certify_run(run)
    by = {c.name: c for c in cert.checks}
    assert by["residual settlement"].status == "pass"
    assert by["finished-text walk"].status == "pass"
    assert by["change verifier"].status == "pass"
    assert by["terminal states"].status == "pass"
    assert by["outcome"].status == "pass" and "done" in by["outcome"].detail
    assert (run / "outcome.json").exists()
    oc = json.loads((run / "outcome.json").read_text("utf-8"))
    assert oc["outcome"] == "done" and oc["hubspot"]["property"] == "docproof"


# --- edit damage: the change verifier's own items ------------------------------

def test_settle_reverts_a_flagged_edit_and_revises_another(tmp_path):
    src = _manuscript(tmp_path)
    ids, _doc = _para_ids(src)
    p0, p1 = ids[0], ids[1]
    run = _build(tmp_path, src, [
        {"para_id": p0, "original_text": "teh", "corrected_text": "tea",
         "confidence": "high"},                                # voice damage
        {"para_id": p1, "original_text": "recieve", "corrected_text": "recieves",
         "confidence": "high"},                                # still wrong
    ])
    _walk(run, [], [
        {"para_id": p0, "original_text": "teh", "corrected_text": "tea",
         "verdict": "breaks_meaning", "detail": "not a drink", "fix": "teh"},
        {"para_id": p1, "original_text": "recieve", "corrected_text": "recieves",
         "verdict": "breaks_grammar", "detail": "misspelled", "fix": "receive"},
    ])
    assert _settle(tmp_path, run, src) == 0
    recs, st = _records(run)
    a = recs[problem_id(p0, "teh", "tea")]
    assert a.action == "drop" and a.reason.startswith("edit_damage")
    b = recs[problem_id(p1, "recieve", "recieves")]
    assert b.action == "revise" and b.after_replacement == "receive"
    acc = _accepted(run)
    assert "teh desk" in acc[p0]                # reverted to the source
    assert "receive the letter" in acc[p1]      # revised in place
    cert = certify_run(run)
    by = {c.name: c for c in cert.checks}
    assert by["change verifier"].status == "pass"
    assert by["residual settlement"].status == "pass"


# --- T3: a residual inside a repair cluster member keeps the cluster ----------

def test_settle_keeps_a_cluster_applied_when_one_member_is_revised(tmp_path):
    src = _manuscript(tmp_path)
    ids, _doc = _para_ids(src)
    p0 = ids[0]
    run = _build(tmp_path, src, [
        {"para_id": p0, "original_text": "teh", "corrected_text": "thee",
         "confidence": "high", "cluster_id": "rep-1"},
        {"para_id": p0, "original_text": "sighed", "corrected_text": "sighs",
         "confidence": "high", "cluster_id": "rep-1"},
    ])
    acc = _accepted(run)
    assert "thee desk" in acc[p0] and "sighs" in acc[p0]
    _walk(run, [{"para_id": p0, "quote": "thee", "problem": "wrong word",
                 "suggestion": "the", "severity": "high"}])
    assert _settle(tmp_path, run, src) == 0
    recs, _st = _records(run)
    assert recs[residual_id(p0, "thee")].action == "absorb"
    acc = _accepted(run)
    assert "the desk" in acc[p0] and "sighs" in acc[p0]    # sibling stayed
    env = json.loads((run / "findings.json").read_text("utf-8"))
    applied = [r for r in env["findings"] if r.get("applied")]
    assert {r["cluster_id"] for r in applied} == {"rep-1"}
    assert len(applied) == 2


# --- T4: bounded rounds — what cannot be decided ships as a question ---------

def test_settle_turns_an_undecidable_residual_into_a_query(tmp_path):
    src = _manuscript(tmp_path)
    ids, _doc = _para_ids(src)
    p2 = ids[2]
    run = _build(tmp_path, src, [
        {"para_id": ids[0], "original_text": "teh", "corrected_text": "the",
         "confidence": "high"}])
    _walk(run, [{"para_id": p2, "quote": "sure the total", "problem":
                 "garbled — two plausible repairs", "suggestion": "",
                 "severity": "medium"}])
    assert _settle(tmp_path, run, src, "--rounds", "1") == 0
    recs, st = _records(run)
    r = recs[residual_id(p2, "sure the total")]
    assert r.action == "query" and r.question
    assert st.open == []
    env = json.loads((run / "findings.json").read_text("utf-8"))
    assert any(row.get("queried") and row["para_id"] == p2
               and "two plausible repairs" in row["explanation"]
               for row in env["findings"])
    # state machine: settled is reachable now, and refuses while open
    ws = tmp_path / "ws"
    ws.mkdir()
    assert main(["galley", "state", str(ws), "--advance", "intake"]) == 0
    assert main(["galley", "state", str(ws), "--advance", "settled",
                 "--results", str(run)]) == 0


def test_state_refuses_settled_while_items_are_open(tmp_path):
    src = _manuscript(tmp_path)
    ids, _doc = _para_ids(src)
    run = _build(tmp_path, src, [
        {"para_id": ids[0], "original_text": "teh", "corrected_text": "the",
         "confidence": "high"}])
    _walk(run, [{"para_id": ids[1], "quote": "recieve", "problem": "sp",
                 "suggestion": "receive", "severity": "high"}])
    ws = tmp_path / "ws"
    ws.mkdir()
    assert main(["galley", "state", str(ws), "--advance", "intake"]) == 0
    assert main(["galley", "state", str(ws), "--advance", "settled",
                 "--results", str(run)]) == 7
    cert = certify_run(run)
    by = {c.name: c for c in cert.checks}
    assert by["residual settlement"].status == "fail"
    assert by["finished-text walk"].status == "fail"


# --- T5: an unanchorable walker row is recorded, never silently lost ---------

def test_unanchorable_rows_are_recorded_as_dropped(tmp_path):
    src = _manuscript(tmp_path)
    ids, _doc = _para_ids(src)
    run = _build(tmp_path, src, [
        {"para_id": ids[0], "original_text": "teh", "corrected_text": "the",
         "confidence": "high"}])
    _walk(run, [{"para_id": ids[3], "quote": "not in the text", "problem": "?",
                 "suggestion": "x", "severity": "low"},
                {"para_id": "body-9999", "quote": "gate", "problem": "?",
                 "suggestion": "x", "severity": "low"}])
    assert _settle(tmp_path, run, src) == 0
    recs, _st = _records(run)
    assert recs[residual_id(ids[3], "not in the text")].reason == "unanchorable"
    assert recs[residual_id("body-9999", "gate")].action == "drop"
    fw = json.loads((run / "finished_walk.json").read_text("utf-8"))
    assert {r["settled"] for r in fw["residuals"]} == {"drop"}
    assert fw["unsettled"] == []


# --- T2: the delta verifier flags a composite -> revert + query ---------------

def test_a_verifier_flag_on_a_composite_reverts_it_to_a_query(tmp_path,
                                                             monkeypatch):
    src = _manuscript(tmp_path)
    ids, _doc = _para_ids(src)
    p0 = ids[0]
    run = _build(tmp_path, src, [
        {"para_id": p0, "original_text": "teh", "corrected_text": "thee",
         "confidence": "high"}])
    _walk(run, [{"para_id": p0, "quote": "thee", "problem": "wrong word",
                 "suggestion": "the", "severity": "high"}])
    # round 1 delta verify: the change verifier flags the composite (index 1
    # of the paragraph's applied edits), the walk is clean; the judge gets a
    # second look and CONFIRMS the verifier; nothing after.
    prov = _Provider(
        {"problems": [{"index": 1, "verdict": "voice_damage",
                       "detail": "author wrote thee on purpose",
                       "fix": "thee"}]},
        {"findings": []},
        {"answer": "revert", "reason": "archaic voice, deliberate"})
    monkeypatch.setattr(m, "_resolve_engine",
                        lambda args, cfg, default_model=None:
                        ("provider", prov, "fake-model"))
    rc = main(["galley", "settle", str(run), "--source", str(src),
               "--config", _replay_config(tmp_path), "--engine", "provider"])
    assert rc == 0
    recs, st = _records(run)
    r = recs[residual_id(p0, "thee")]
    assert r.action == "query" and r.reason == "verifier_confirmed"
    acc = _accepted(run)
    assert "thee desk" in acc[p0]               # the owner's edit stands
    env = json.loads((run / "findings.json").read_text("utf-8"))
    assert any(row.get("queried") and row["para_id"] == p0
               for row in env["findings"])
    assert len(prov.calls) == 3
    look = prov.calls[2]
    assert "VERIFIER SAYS (voice_damage)" in look["user"]
    assert "HOUSE STYLE" in look["system"]
    assert certify_run(run).checks and {
        c.name: c.status for c in certify_run(run).checks
    }["residual settlement"] == "pass"


# --- the judge: a residual with no suggestion, decided by the model ----------

def test_the_judge_settles_a_residual_with_no_suggestion(tmp_path, monkeypatch):
    src = _manuscript(tmp_path)
    ids, _doc = _para_ids(src)
    p1 = ids[1]
    run = _build(tmp_path, src, [
        {"para_id": ids[0], "original_text": "teh", "corrected_text": "the",
         "confidence": "high"}])
    _walk(run, [{"para_id": p1, "quote": "recieve", "problem": "misspelling",
                 "suggestion": "", "severity": "high"}])
    prov = _Provider(
        {"action": "add", "replacement": "receive", "reason": "", "question": ""},
        {"problems": []},                                    # delta verify
        {"findings": []})
    monkeypatch.setattr(m, "_resolve_engine",
                        lambda args, cfg, default_model=None:
                        ("provider", prov, "fake-model"))
    rc = main(["galley", "settle", str(run), "--source", str(src),
               "--config", _replay_config(tmp_path), "--engine", "provider"])
    assert rc == 0
    recs, _st = _records(run)
    assert recs[residual_id(p1, "recieve")].action == "add"
    assert "receive the letter" in _accepted(run)[p1]
    assert "FLAGGED SPAN: 'recieve'" in prov.calls[0]["user"]


# --- residuals / outcome verbs -------------------------------------------------

def test_residuals_verb_lists_open_items_with_owners(tmp_path, capsys):
    src = _manuscript(tmp_path)
    ids, _doc = _para_ids(src)
    run = _build(tmp_path, src, [
        {"para_id": ids[0], "original_text": "teh", "corrected_text": "thee",
         "confidence": "high"}])
    _walk(run, [{"para_id": ids[0], "quote": "thee", "problem": "w",
                 "suggestion": "the", "severity": "high"}])
    assert main(["galley", "residuals", str(run), "--source", str(src),
                 "--config", _replay_config(tmp_path), "--json"]) == 0
    out = capsys.readouterr().out
    assert "1 open item(s)" in out
    line = [l for l in out.splitlines() if l.startswith("{")][-1]
    item = json.loads(line)["open"][0]
    assert item["owner_finding_id"] and item["resolution"] == "resolvable"


def test_outcome_needs_human_when_most_paragraphs_need_rewrite(tmp_path, capsys):
    src = _manuscript(tmp_path)
    ids, _doc = _para_ids(src)
    # three edits in four of five paragraphs -> rewrite share 80%
    rows = []
    for pid, words in ((ids[0], ("lamp", "desk", "caught")),
                       (ids[1], ("light", "thin", "letter")),
                       (ids[2], ("coins", "total", "change")),
                       (ids[3], ("gate", "dog", "greet"))):
        for w in words:
            rows.append({"para_id": pid, "original_text": w,
                         "corrected_text": w + "s", "confidence": "high"})
    run = _build(tmp_path, src, rows)
    assert main(["galley", "outcome", str(run), "--json"]) == 0
    out = capsys.readouterr().out
    oc = json.loads([l for l in out.splitlines() if l.startswith("{")][-1])
    assert oc["outcome"] == "needs_human"
    assert "rewrite-class" in oc["reason"]
    # needs_human names its own HubSpot value — the option that puts the book
    # in front of a human proofreader. Both verdicts move the book off "Ready
    # for Proofing"; a book left sitting at ready is one nobody would notice.
    assert oc["hubspot"] == {"object": "0-970", "property": "docproof",
                             "value": "Needs Human PR"}
    # a human overrule is recorded as such
    assert main(["galley", "outcome", str(run), "--set", "done",
                 "--reason", "reviewed by hand"]) == 0
    assert json.loads((run / "outcome.json").read_text("utf-8"))["set_by"] \
        == "human"


def test_hubspot_fields_names_a_value_for_both_verdicts():
    """The whole HubSpot contract, in one place. Both verdicts name a real
    option on the `docproof` property — "Proofing Complete" when the loop
    finished the book, "Needs Human PR" when a human proofreader has to take it
    on — so a book leaves "Ready for Proofing" either way and nothing sits in a
    queue nobody is reading.

    A blanked value is the one case that yields nothing: PATCHing "" would blank
    the status property rather than move it, so the answer is "write nothing"
    instead."""
    from galley.outcome import (DEFAULT_DONE_VALUE, DEFAULT_NEEDS_HUMAN_VALUE,
                                hubspot_fields)

    assert DEFAULT_DONE_VALUE == "Proofing Complete"
    assert DEFAULT_NEEDS_HUMAN_VALUE == "Needs Human PR"
    assert hubspot_fields("done") == {"object": "0-970", "property": "docproof",
                                      "value": "Proofing Complete"}
    assert hubspot_fields("needs_human") == {
        "object": "0-970", "property": "docproof", "value": "Needs Human PR"}
    assert hubspot_fields("done", done_value="") == {}
    assert hubspot_fields("needs_human", needs_human_value="") == {}


# --- import-findings --anchor accepted -----------------------------------------

def test_import_findings_anchor_accepted_revises_the_owner(tmp_path):
    src = _manuscript(tmp_path)
    ids, _doc = _para_ids(src)
    p0, p1 = ids[0], ids[1]
    run = _build(tmp_path, src, [
        {"para_id": p0, "original_text": "teh", "corrected_text": "thee",
         "confidence": "high"}])
    rows = tmp_path / "accepted_rows.json"
    rows.write_text(json.dumps([
        {"para_id": p0, "quote": "thee", "replacement": "the",
         "explanation": "wrong word"},
        {"para_id": p1, "quote": "recieve", "replacement": "receive"},
    ]), encoding="utf-8")
    out = tmp_path / "run2"
    rc = main(["import-findings", str(rows), str(src), "--anchor", "accepted",
               "--run", str(run), "--config", _replay_config(tmp_path),
               "--out", str(out)])
    assert rc == 0
    acc = _accepted(out)
    assert "the desk" in acc[p0] and "receive the letter" in acc[p1]
    original = paragraph_views(out)[0]
    assert "teh desk" in original[p0]           # reject-all is the source


# --- guards learned on the Redding trial ---------------------------------------

def test_instruction_shaped_suggestions_never_become_text(tmp_path):
    src = _manuscript(tmp_path)
    ids, _doc = _para_ids(src)
    run = _build(tmp_path, src, [
        {"para_id": ids[0], "original_text": "teh", "corrected_text": "the",
         "confidence": "high"}])
    _walk(run, [{"para_id": ids[3], "quote": "The garden gate",
                 "problem": "leading space", "suggestion":
                 "Delete the trailing whitespace.", "severity": "low"}])
    assert _settle(tmp_path, run, src) == 0
    recs, _st = _records(run)
    r = recs[residual_id(ids[3], "The garden gate")]
    assert r.action == "query" and r.reason == "instruction"
    assert "Delete the trailing" not in _accepted(run)[ids[3]]


def test_a_whole_sentence_suggestion_replaces_the_sentence_not_the_fragment(
        tmp_path):
    src = _manuscript(tmp_path)
    ids, _doc = _para_ids(src)
    p2 = ids[2]
    run = _build(tmp_path, src, [
        {"para_id": ids[0], "original_text": "teh", "corrected_text": "the",
         "confidence": "high"}])
    # quote is a fragment; the suggestion rewrites the whole sentence
    _walk(run, [{"para_id": p2, "quote": "sure the total would not change",
                 "problem": "wordy", "suggestion":
                 "David counted the coins twice, certain the total would not "
                 "change.", "severity": "low"}])
    assert _settle(tmp_path, run, src) == 0
    acc = _accepted(run)[p2]
    assert acc.startswith("David counted the coins twice, certain the total")
    assert acc.count("David counted") == 1


def test_settle_queries_are_never_collapsed_into_one_comment(tmp_path):
    src = _manuscript(tmp_path)
    ids, _doc = _para_ids(src)
    run = _build(tmp_path, src, [
        {"para_id": ids[0], "original_text": "teh", "corrected_text": "the",
         "confidence": "high"}])
    # three undecidable residuals in three paragraphs -> three questions
    _walk(run, [{"para_id": ids[i], "quote": q, "problem": f"q{i}",
                 "suggestion": "", "severity": "medium"}
                for i, q in ((1, "Its light"), (2, "the coins"),
                             (3, "old dog"))])
    assert _settle(tmp_path, run, src) == 0
    env = json.loads((run / "findings.json").read_text("utf-8"))
    assert sum(1 for r in env["findings"] if r.get("queried")
               and r["error_type"] == "galley_settle") == 3


def test_case_only_change_inside_quoted_speech_is_not_a_fact_change():
    from galley.settle import _fact
    assert _fact('“just be careful,” she said', '“Just be careful,” she said') \
        is None
    assert _fact("my friend's luck", "Ava's luck") is not None
    assert _fact("one 12 minute talk", "one 21 minute talk") is not None


def test_a_suggestion_that_ships_a_house_artifact_is_queried(tmp_path):
    src = _manuscript(tmp_path)
    ids, _doc = _para_ids(src)
    run = _build(tmp_path, src, [
        {"para_id": ids[0], "original_text": "teh", "corrected_text": "the",
         "confidence": "high"}])
    # British punctuation: the period outside the closing quote
    _walk(run, [{"para_id": ids[1], "quote": "the letter by tonight.",
                 "problem": "quote it", "suggestion": "“the letter by tonight”.",
                 "severity": "low"}])
    assert _settle(tmp_path, run, src) == 0
    recs, _st = _records(run)
    r = recs[residual_id(ids[1], "the letter by tonight.")]
    assert r.action == "query" and r.reason.startswith("artifact:")
    assert "”." not in _accepted(run)[ids[1]]


def test_typographic_parentheticals_are_notes_not_text():
    from galley.settle import strip_note
    assert strip_note("the Rocky theme (with Rocky in italics)", "the Rocky theme") \
        == ("the Rocky theme", True)
    # the author's own parenthetical, present in the quote, survives
    q = "my husband (a comma fan) laughed"
    assert strip_note(q, q) == (q, False)


def test_deleting_a_leaked_editorial_aside_is_not_a_fact_change():
    from galley.settle import _fact, deletes_an_aside
    before = ('actresses who we were going to be when we grew up (parallel '
              'with "were going to be"; breaks tense agreement as written)')
    after = 'actresses who we were going to be when we grew up'
    assert deletes_an_aside(before, after)
    assert _fact(before, after) is None
    # the author's own parenthetical is not an aside
    assert not deletes_an_aside("we left (all three of us) at noon",
                                "we left at noon")


def test_restore_rows_never_overwrites_a_current_row():
    from galley.settle import restore_rows
    # after a rebuild "f-0003" names an unrelated row; the old owner that
    # was absorbed also used to be "f-0003"
    working = {"f-0003": {"para_id": "p9", "original_text": "keep me"}}
    removed = {"f-0003": {"para_id": "p1", "original_text": "old owner"}}
    assert restore_rows(working, "r-abc", removed.values()) == 1
    assert working["f-0003"]["original_text"] == "keep me"
    assert any(v["original_text"] == "old owner" for v in working.values())


# --- propagation: the same fix at identical untouched sites nearby ----------

def test_a_settled_fix_propagates_to_the_same_surface_nearby(tmp_path):
    paras = [
        "I will recieve the letter, and you will recieve the reply.",
        "Nobody could recieve it in time, though formal receipts came later.",
        "The end.",
    ]
    src = _manuscript(tmp_path, paras)
    ids, _doc = _para_ids(src)
    run = _build(tmp_path, src, [
        {"para_id": ids[2], "original_text": "end", "corrected_text": "End",
         "confidence": "high"}])
    _walk(run, [{"para_id": ids[0], "quote": "recieve", "problem": "sp",
                 "suggestion": "receive", "severity": "high"}])
    assert _settle(tmp_path, run, src) == 0
    acc = _accepted(run)
    assert acc[ids[0]].count("receive") == 2 and "recieve" not in acc[ids[0]]
    assert "could receive it" in acc[ids[1]]            # the neighbour too
    assert "formal receipts" in acc[ids[1]]             # word boundary held
    recs, _st = _records(run)
    prop = [r for r in recs.values() if r.reason.startswith("propagated:")]
    assert len(prop) == 2 and {r.action for r in prop} == {"add"}


# --- --until-clean: keep sweeping while rounds find work, stop when quiet --

def _fake_engine(monkeypatch, prov):
    monkeypatch.setattr(m, "_resolve_engine",
                        lambda args, cfg, default_model=None:
                        ("provider", prov, "fake-model"))


def test_until_clean_stops_after_a_quiet_round(tmp_path, monkeypatch):
    src = _manuscript(tmp_path)
    ids, _doc = _para_ids(src)
    run = _build(tmp_path, src, [
        {"para_id": ids[0], "original_text": "teh", "corrected_text": "the",
         "confidence": "high"}])
    _walk(run, [{"para_id": ids[1], "quote": "recieve", "problem": "sp",
                 "suggestion": "receive", "severity": "high"}])
    # round 1 verify: clean changes, walk raises ONE new item (quiet) ->
    # round 2 settles it, its verify is clean -> stop. Never a third round.
    prov = _Provider(
        {"problems": []},
        {"findings": [{"para_id": ids[1], "quote": "was thin",
                       "problem": "x", "suggestion": "was faint",
                       "severity": "low"}]},
        {"problems": []}, {"findings": []})
    _fake_engine(monkeypatch, prov)
    rc = main(["galley", "settle", str(run), "--source", str(src), "--config",
               _replay_config(tmp_path), "--engine", "provider",
               "--until-clean"])
    assert rc == 0
    _recs, st = _records(run)
    assert st.rounds == 2 and st.open == []
    assert any("quiet" in n for n in st.notes)


def test_until_clean_respects_the_turn_budget(tmp_path, monkeypatch):
    src = _manuscript(tmp_path)
    ids, _doc = _para_ids(src)
    run = _build(tmp_path, src, [
        {"para_id": ids[0], "original_text": "teh", "corrected_text": "the",
         "confidence": "high"}])
    _walk(run, [{"para_id": ids[1], "quote": "recieve", "problem": "sp",
                 "suggestion": "receive", "severity": "high"}])
    # every verify keeps raising a fresh item; the budget of 2 turns ends it
    # after round 1 and the leftover ships as a question.
    prov = _Provider(
        {"problems": []},
        {"findings": [{"para_id": ids[1], "quote": "Its light",
                       "problem": "x", "suggestion": "Its glow",
                       "severity": "low"},
                      {"para_id": ids[1], "quote": "was thin",
                       "problem": "y", "suggestion": "was faint",
                       "severity": "low"},
                      {"para_id": ids[1], "quote": "was enough",
                       "problem": "z", "suggestion": "was sufficient",
                       "severity": "low"},
                      {"para_id": ids[1], "quote": "by tonight",
                       "problem": "w", "suggestion": "by nightfall",
                       "severity": "low"}]})
    _fake_engine(monkeypatch, prov)
    rc = main(["galley", "settle", str(run), "--source", str(src), "--config",
               _replay_config(tmp_path), "--engine", "provider",
               "--until-clean", "--max-turns", "2"])
    assert rc == 0
    recs, st = _records(run)
    assert st.rounds == 2 and st.open == []
    assert any("turn budget" in n for n in st.notes)
    assert sum(1 for r in recs.values()
               if r.reason.startswith("unresolved_after_")) == 4


def test_a_sweep_that_will_not_converge_flags_needs_human(tmp_path, monkeypatch):
    src = _manuscript(tmp_path)
    ids, _doc = _para_ids(src)
    run = _build(tmp_path, src, [
        {"para_id": ids[0], "original_text": "teh", "corrected_text": "the",
         "confidence": "high"}])
    _walk(run, [{"para_id": ids[1], "quote": "recieve", "problem": "sp",
                 "suggestion": "receive", "severity": "high"}])
    # the walk raises four new items every round; with a 2-turn budget the
    # sweep stops noisy -> the outcome says a human proofreader is needed
    def noisy(pairs):
        return {"findings": [{"para_id": ids[1], "quote": q, "problem": "x",
                              "suggestion": r, "severity": "low"}
                             for q, r in pairs]}
    prov = _Provider(
        {"problems": []},
        noisy((("Its light", "Its glow"), ("was thin", "was faint"),
               ("was enough", "was sufficient"), ("by tonight", "by nightfall"))),
        {"problems": []},
        noisy((("Its glow", "Its shine"), ("was faint", "was dim"),
               ("was sufficient", "was ample"), ("by nightfall", "by dusk"))))
    _fake_engine(monkeypatch, prov)
    # two noisy rounds fit in a 4-turn budget; the sweep stops still noisy
    assert main(["galley", "settle", str(run), "--source", str(src),
                 "--config", _replay_config(tmp_path), "--engine", "provider",
                 "--until-clean", "--max-turns", "4"]) == 0
    oc = json.loads((run / "outcome.json").read_text("utf-8"))
    assert oc["outcome"] == "needs_human"
    assert "still finding" in oc["reason"]
    _recs, st = _records(run)
    assert st.convergence["stopped"] == "turn_budget"


def test_until_clean_takes_an_explicit_rounds_as_its_ceiling(tmp_path,
                                                             monkeypatch):
    """`--until-clean --rounds 2` sweeps AT MOST two rounds. The sweep used to
    ignore --rounds entirely and run to the hard cap of 12, which is a trap for
    an unattended driver that names a round budget and means it."""
    src = _manuscript(tmp_path)
    ids, _doc = _para_ids(src)
    run = _build(tmp_path, src, [
        {"para_id": ids[0], "original_text": "teh", "corrected_text": "the",
         "confidence": "high"}])
    _walk(run, [{"para_id": ids[1], "quote": "recieve", "problem": "sp",
                 "suggestion": "receive", "severity": "high"}])

    def noisy(pairs):
        return {"findings": [{"para_id": ids[1], "quote": q, "problem": "x",
                              "suggestion": r, "severity": "low"}
                             for q, r in pairs]}
    prov = _Provider(
        {"problems": []},
        noisy((("Its light", "Its glow"), ("was thin", "was faint"),
               ("was enough", "was sufficient"), ("by tonight", "by nightfall"))),
        {"problems": []},
        noisy((("Its glow", "Its shine"), ("was faint", "was dim"),
               ("was sufficient", "was ample"), ("by nightfall", "by dusk"))),
        {"problems": []}, noisy((("Its shine", "Its blaze"),)))
    _fake_engine(monkeypatch, prov)
    assert main(["galley", "settle", str(run), "--source", str(src),
                 "--config", _replay_config(tmp_path), "--engine", "provider",
                 "--until-clean", "--rounds", "2", "--quiet-floor", "0",
                 "--quiet-share", "0"]) == 0
    _recs, st = _records(run)
    # Two settling rounds, then the leftovers shipped as questions — never a
    # third read, though the provider had one more noisy answer scripted.
    assert st.convergence["rounds"] == 2
    assert st.convergence["stopped"] == "round_cap"
    assert st.convergence["quiet"] is False
    assert prov.calls and len(prov.calls) == 4


def test_until_clean_without_rounds_keeps_its_old_reach(tmp_path, monkeypatch):
    """…and an unset --rounds still means "sweep to a quiet round", not the
    verb's default of 3: the ceiling only exists when a caller names one."""
    from galley.settle import HARD_MAX_ROUNDS, SettleOptions
    src = _manuscript(tmp_path)
    ids, _doc = _para_ids(src)
    run = _build(tmp_path, src, [
        {"para_id": ids[0], "original_text": "teh", "corrected_text": "the",
         "confidence": "high"}])
    _walk(run, [{"para_id": ids[1], "quote": "recieve", "problem": "sp",
                 "suggestion": "receive", "severity": "high"}])
    prov = _Provider({"problems": []}, {"findings": []})
    _fake_engine(monkeypatch, prov)
    seen = {}
    from galley import settle as settle_mod
    real = settle_mod.Settler.__init__

    def spy(self, *a, **kw):
        real(self, *a, **kw)
        seen["rounds"] = self.opt.rounds
        seen["until_clean"] = self.opt.until_clean
    monkeypatch.setattr(settle_mod.Settler, "__init__", spy)
    assert main(["galley", "settle", str(run), "--source", str(src),
                 "--config", _replay_config(tmp_path), "--engine", "provider",
                 "--until-clean"]) == 0
    # rounds 0 = "no ceiling of your own"; the loop falls back to the hard cap.
    assert seen == {"rounds": 0, "until_clean": True}
    assert HARD_MAX_ROUNDS == 12
    assert SettleOptions().rounds == 3          # the non-until-clean default


def test_until_clean_with_nothing_open_starts_with_a_fresh_sweep(tmp_path,
                                                                  monkeypatch):
    src = _manuscript(tmp_path)
    ids, _doc = _para_ids(src)
    run = _build(tmp_path, src, [
        {"para_id": ids[0], "original_text": "teh", "corrected_text": "the",
         "confidence": "high"}])
    _walk(run, [])                                   # nothing open
    # the fresh sweep: 1 change batch (clean) + 1 walk read raising one item;
    # round 1 settles it; its delta verify is clean -> quiet, done.
    prov = _Provider(
        {"problems": []},
        {"findings": [{"para_id": ids[1], "quote": "recieve", "problem": "sp",
                       "suggestion": "receive", "severity": "high"}]},
        {"problems": []}, {"findings": []})
    _fake_engine(monkeypatch, prov)
    assert main(["galley", "settle", str(run), "--source", str(src),
                 "--config", _replay_config(tmp_path), "--engine", "provider",
                 "--until-clean"]) == 0
    assert "receive the letter" in _accepted(run)[ids[1]]
    _recs, st = _records(run)
    assert any(n.startswith("fresh sweep") for n in st.notes)
    assert len(prov.calls) == 4
    fw = json.loads((run / "finished_walk.json").read_text("utf-8"))
    assert fw["ran"] is True and fw["engine"] == "provider"


def test_outcome_density_counts_wording_edits_not_mechanics(tmp_path, capsys):
    src = _manuscript(tmp_path)
    ids, _doc = _para_ids(src)
    # lots of punctuation/case edits, few wording edits: a sound book with
    # a heavy house style, not one that needs rewriting
    rows = []
    for pid, text in zip(ids, PARAGRAPHS):
        for w in ("lamp", "light", "coins", "gate", "note"):
            if w in text:
                rows.append({"para_id": pid, "original_text": w,
                             "corrected_text": w.capitalize(),
                             "confidence": "high"})
    run = _build(tmp_path, src, rows)
    assert main(["galley", "outcome", str(run), "--json"]) == 0
    out = capsys.readouterr().out
    oc = json.loads([l for l in out.splitlines() if l.startswith("{")][-1])
    assert oc["outcome"] == "done"
    assert oc["evidence"]["mechanical_edits"] >= 4
    assert oc["evidence"]["wording_edits"] == 0


# --- Georgis (2026-09-04): the verifier's second look ---------------------------

def test_the_judge_can_overrule_the_verifier_and_the_composite_stands(
        tmp_path, monkeypatch):
    """The verifier is Chicago-trained; the house is not. When the judge,
    told the house rules, answers `keep`, the settlement stays applied and
    the record says `verifier_overruled` — no revert, no query."""
    src = _manuscript(tmp_path)
    ids, _doc = _para_ids(src)
    p0 = ids[0]
    run = _build(tmp_path, src, [
        {"para_id": p0, "original_text": "teh", "corrected_text": "thee",
         "confidence": "high"}])
    _walk(run, [{"para_id": p0, "quote": "thee", "problem": "wrong word",
                 "suggestion": "the", "severity": "high"}])
    prov = _Provider(
        {"problems": [{"index": 1, "verdict": "wrong_rule",
                       "detail": "should be thee", "fix": "thee"}]},
        {"findings": []},
        {"answer": "keep", "reason": "the is the plain typo fix"})
    monkeypatch.setattr(m, "_resolve_engine",
                        lambda args, cfg, default_model=None:
                        ("provider", prov, "fake-model"))
    rc = main(["galley", "settle", str(run), "--source", str(src),
               "--config", _replay_config(tmp_path), "--engine", "provider"])
    assert rc == 0
    recs, st = _records(run)
    r = recs[residual_id(p0, "thee")]
    assert r.action == "absorb" and r.reason.startswith("verifier_overruled")
    assert "the desk" in _accepted(run)[p0]
    assert len(prov.calls) == 3
    assert certify_run(run).checks and {
        c.name: c.status for c in certify_run(run).checks
    }["residual settlement"] == "pass"


def test_a_failed_second_look_keeps_todays_revert(tmp_path, monkeypatch):
    src = _manuscript(tmp_path)
    ids, _doc = _para_ids(src)
    p0 = ids[0]
    run = _build(tmp_path, src, [
        {"para_id": p0, "original_text": "teh", "corrected_text": "thee",
         "confidence": "high"}])
    _walk(run, [{"para_id": p0, "quote": "thee", "problem": "wrong word",
                 "suggestion": "the", "severity": "high"}])
    prov = _Provider(
        {"problems": [{"index": 1, "verdict": "voice_damage",
                       "detail": "deliberate", "fix": "thee"}]},
        {"findings": []})                 # the second look gets {} -> a loss
    monkeypatch.setattr(m, "_resolve_engine",
                        lambda args, cfg, default_model=None:
                        ("provider", prov, "fake-model"))
    rc = main(["galley", "settle", str(run), "--source", str(src),
               "--config", _replay_config(tmp_path), "--engine", "provider"])
    assert rc == 0
    recs, _st = _records(run)
    r = recs[residual_id(p0, "thee")]
    assert r.action == "query" and r.reason == "verifier_reverted"
    assert "thee desk" in _accepted(run)[p0]


# --- Georgis: house style is not up for settlement ------------------------------

def test_a_settlement_the_sweeps_would_undo_is_dropped(tmp_path):
    """The walk, prompted Chicago, flagged the house "4:00 AM" and settle
    wrote "4:00 a.m."; `sweep_time_of_day` fired straight back. Now the
    candidate paragraph is swept before the settlement lands."""
    paras = list(PARAGRAPHS) + ["The train left at 4:00 AM and we slept."]
    src = _manuscript(tmp_path, paras)
    ids, _doc = _para_ids(src)
    p5 = ids[5]
    run = _build(tmp_path, src, [
        {"para_id": ids[0], "original_text": "teh", "corrected_text": "the",
         "confidence": "high"}])
    _walk(run, [{"para_id": p5, "quote": "4:00 AM",
                 "problem": "Chicago uses lowercase periods for a.m.",
                 "suggestion": "4:00 a.m.", "severity": "medium"},
                {"para_id": p5, "quote": "we slept", "problem": "tense",
                 "suggestion": "we sleep", "severity": "low"}])
    assert _settle(tmp_path, run, src) == 0
    recs, st = _records(run)
    r = recs[residual_id(p5, "4:00 AM")]
    assert r.action == "drop"
    assert r.reason == "undoes_house_style:sweep_time_of_day"
    assert "4:00 AM" in _accepted(run)[p5]
    # an ordinary settlement in the same paragraph still lands
    assert recs[residual_id(p5, "we slept")].action == "add"
    assert "we sleep" in _accepted(run)[p5]
    assert st.open == []


def test_the_judge_prompt_carries_the_house_rules():
    from galley.house_style import HOUSE_RULES
    from galley.settle import _judge_system
    system = _judge_system("Keep the narrator's slang.")
    for rule in HOUSE_RULES:
        assert rule in system
    assert system.index("HOUSE STYLE") < system.index("VOICE NOTES")
    assert "Keep the narrator's slang." in system


# --- Georgis: the mechanical-only guard ------------------------------------------

def test_rewrite_class_names_the_georgis_examples():
    from galley.settle import rewrite_class
    # must become queries
    assert rewrite_class("No job wasn't good enough.",
                         "No job was too menial for me.")
    assert rewrite_class("became more resolved with the fact",
                         "became more resigned to the fact")
    assert rewrite_class("virulent male", "virile male")
    assert rewrite_class("fresh grave", "fresh gravesite")
    assert rewrite_class("and told to join", "and join")
    assert rewrite_class("Chewing rapidly.", "I chewed rapidly.")
    # may apply
    assert rewrite_class("at that moment Emily's age",
                         "at that moment that Emily's age") is None
    assert rewrite_class("borne from", "born from") is None
    assert rewrite_class("Ju-Jitsu", "ju-jitsu") is None
    assert rewrite_class("dark sunken eyes", "dark, sunken eyes") is None
    assert rewrite_class("hellbent", "hell-bent") is None
    assert rewrite_class("recieve the letter", "receive the letter") is None
    assert rewrite_class("teh desk", "the desk") is None
    assert rewrite_class("in a moonless night", "on a moonless night") is None
    assert rewrite_class("ordered I be taken", "ordered that I be taken") is None
    assert rewrite_class("as a man and wife", "as man and wife") is None


def test_mechanical_only_turns_a_rewrite_into_a_query_with_the_suggestion(
        tmp_path):
    src = _manuscript(tmp_path)
    ids, _doc = _para_ids(src)
    p2, p1 = ids[2], ids[1]
    run = _build(tmp_path, src, [
        {"para_id": ids[0], "original_text": "teh", "corrected_text": "the",
         "confidence": "high"}])
    _walk(run, [
        {"para_id": p2, "quote": "sure the total would not change again",
         "problem": "awkward", "suggestion": "certain the sum was fixed",
         "severity": "low"},
        {"para_id": p2, "quote": "sure the total",
         "problem": "missing that", "suggestion": "sure that the total",
         "severity": "medium"},
        {"para_id": p1, "quote": "recieve", "problem": "misspelling",
         "suggestion": "receive", "severity": "high"},
    ])
    assert _settle(tmp_path, run, src, "--mechanical-only") == 0
    recs, st = _records(run)
    q = recs[residual_id(p2, "sure the total would not change again")]
    assert q.action == "query" and q.reason.startswith("rewrite_class:")
    assert "certain the sum was fixed" in q.question
    assert recs[residual_id(p2, "sure the total")].action == "add"
    assert recs[residual_id(p1, "recieve")].action == "add"
    acc = _accepted(run)
    assert "sure that the total would not change again" in acc[p2]
    assert "receive" in acc[p1]
    env = json.loads((run / "findings.json").read_text("utf-8"))
    assert any(r.get("queried") and "certain the sum" in r["explanation"]
               for r in env["findings"])
    assert st.open == []


def test_an_approval_that_says_mechanical_only_implies_the_guard(tmp_path):
    src = _manuscript(tmp_path)
    ids, _doc = _para_ids(src)
    p2 = ids[2]
    run = _build(tmp_path, src, [
        {"para_id": ids[0], "original_text": "teh", "corrected_text": "the",
         "confidence": "high"}])
    _walk(run, [{"para_id": p2, "quote": "sure the total would not change",
                 "problem": "awkward", "suggestion": "certain the sum was fixed",
                 "severity": "low"}])
    approval = tmp_path / "approval.json"
    approval.write_text(json.dumps({"mechanical_only": True}), "utf-8")
    assert _settle(tmp_path, run, src, "--approval", str(approval)) == 0
    recs, _st = _records(run)
    r = recs[residual_id(p2, "sure the total would not change")]
    assert r.action == "query" and r.reason.startswith("rewrite_class:")
    assert "coins twice, sure the total would not change" in _accepted(run)[p2]


# --- Georgis: the composite self-check and duplicated fragments ------------------

def test_duplicated_fragments_finds_the_georgis_splice():
    from galley.settle import duplicated_fragments, introduced_fragments
    bad = ("trot, communicating to urge the horse into a trot ,communicating "
           "his frustration and worry.")
    assert duplicated_fragments(bad) == []          # only 4 words repeat
    worse = ("Mr. Nestor jerked the reins to urge the horse into a trot, "
             "communicating; jerked the reins to urge the horse into a trot "
             ",communicating his frustration and worry.")
    frags = duplicated_fragments(worse)
    assert "jerked the reins to urge" in frags
    src = "The rain, the rain, the rain, the rain, the rain fell."
    assert duplicated_fragments(src)               # the author's own refrain
    assert introduced_fragments(src, src) == []
    assert introduced_fragments("Plain text here.", worse)


def test_self_check_reverts_a_composite_that_did_not_compose_as_planned(
        tmp_path, monkeypatch):
    """A composite whose rebuilt paragraph does not read as the decision said
    it should (the Georgis splice) is reverted to a query with reason
    `composite_mismatch`; an honest one in another paragraph stands."""
    import galley.verify as gv
    src = _manuscript(tmp_path)
    ids, _doc = _para_ids(src)
    p0, p1 = ids[0], ids[1]
    run = _build(tmp_path, src, [
        {"para_id": p0, "original_text": "teh", "corrected_text": "thee",
         "confidence": "high"}])
    _walk(run, [{"para_id": p0, "quote": "thee", "problem": "wrong word",
                 "suggestion": "the", "severity": "high"},
                {"para_id": p1, "quote": "recieve", "problem": "misspelling",
                 "suggestion": "receive", "severity": "high"}])
    real = gv.paragraph_views
    state = {"n": 0}

    def views(run_dir):
        original, accepted = real(run_dir)
        state["n"] += 1
        # Corrupt p0 ONCE, at the self-check's read of the rebuilt docx (the
        # loads before it must see the true text or nothing resolves).
        if state["n"] == 2 and "the desk" in accepted.get(p0, ""):
            accepted = dict(accepted)
            accepted[p0] = accepted[p0].replace(
                "the desk", "the desk flickered, then the desk")
        return original, accepted
    monkeypatch.setattr(gv, "paragraph_views", views)
    assert _settle(tmp_path, run, src) == 0
    recs, st = _records(run)
    r = recs[residual_id(p0, "thee")]
    assert r.action == "query" and r.reason == "composite_mismatch"
    assert r.before_replacement == "the"
    assert recs[residual_id(p1, "recieve")].action == "add"
    acc = real(run)[1]
    assert "thee desk" in acc[p0]                  # the owner's edit restored
    assert "receive" in acc[p1]
    env = json.loads((run / "findings.json").read_text("utf-8"))
    assert any(row.get("queried") and row["para_id"] == p0
               for row in env["findings"])
    assert st.open == []
    assert any("failed the self-check" in n for n in st.notes)


def test_certify_artifact_scan_fails_an_introduced_duplicate_fragment(tmp_path):
    src = _manuscript(tmp_path)
    ids, _doc = _para_ids(src)
    p3 = ids[3]
    run = _build(tmp_path, src, [
        {"para_id": p3, "original_text": "The garden gate creaked open, and",
         "corrected_text": "The garden gate creaked open, and the garden "
                           "gate creaked open, and", "confidence": "high"}])
    scan = next(c for c in certify_run(run).checks if c.name == "artifact scan")
    assert scan.status == "fail"
    assert "duplicated fragment" in scan.detail and p3 in scan.detail


def test_certify_artifact_scan_forgives_a_repeat_the_source_already_had(
        tmp_path):
    paras = list(PARAGRAPHS) + [
        "Row, row, row your boat; row, row, row your boat, gently."]
    src = _manuscript(tmp_path, paras)
    ids, _doc = _para_ids(src)
    run = _build(tmp_path, src, [
        {"para_id": ids[0], "original_text": "teh", "corrected_text": "the",
         "confidence": "high"}])
    scan = next(c for c in certify_run(run).checks if c.name == "artifact scan")
    assert scan.status == "pass", scan.detail


# --- Georgis final settle/certify: XML-safe text, local repeats only -----------

def test_xml_safe_strips_what_ooxml_cannot_carry_and_nothing_else():
    """A settle rebuild died with "All strings must be XML compatible" after a
    judge reply smuggled a control character into a replacement."""
    from galley.settle import (Decision, Residual, _question, judge_decision,
                               xml_safe)
    assert xml_safe("a\x00b\x08c\x0bd\x1fe\x7ff\x85g\ufffeh\uffffi") == \
        "abcdefghi"
    assert xml_safe("tab\tnew\nline\r ok — “curly” ’ é") == \
        "tab\tnew\nline\r ok — “curly” ’ é"
    assert xml_safe("") == "" and xml_safe(None) == ""
    # every path a model string reaches the document by is covered
    res = Residual.from_walk({"para_id": "p", "quote": "teh",
                              "problem": "typo\x00", "suggestion": "the\x0b"})
    assert res.suggestion == "the"
    assert "\x00" not in _question(res) and _question(res).endswith("'the'")
    fix = Residual.from_problem({"para_id": "p", "original_text": "a",
                                 "corrected_text": "b", "fix": "c\x1f"})
    assert fix.suggestion == "c"
    parsed = {"action": "add", "replacement": "recei\x00ve", "reason": "",
              "question": "why\x0c?"}
    dec = judge_decision(res, ProviderResult(parsed=parsed, stop_reason="ok"))
    assert isinstance(dec, Decision) and dec.replacement == "receive"
    parsed["action"] = "query"
    assert judge_decision(res, ProviderResult(parsed=parsed,
                                              stop_reason="ok")).question == "why?"


def test_introduced_fragments_counts_a_local_splice_not_a_far_repeat():
    from galley.settle import introduced_fragments
    # Georgis: the same clause fixed the same way twice in one paragraph —
    # "for Peter and I to" -> "for Peter and me to" — repeats far apart.
    far = ("It was time for Peter and me to leave the house, and the long "
           "road to the harbor took us past the market, the church, and the "
           "school before it was time for Peter and me to say goodbye.")
    assert introduced_fragments("", far) == []
    # A splice duplicates locally (the two copies within a clause).
    local = ("Mr. Nestor jerked the reins to urge the horse into a trot, "
             "jerked the reins to urge the horse into a trot, communicating "
             "his frustration.")
    frags = introduced_fragments("", local)
    assert "jerked the reins to urge" in frags
    # ...unless the source paragraph already repeated it.
    assert introduced_fragments(local, local) == []
    # the window is measured between the two copies' starts
    tight = "one two three four five, one two three four five."
    assert introduced_fragments("", tight) == ["one two three four five"]
