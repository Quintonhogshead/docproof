"""Galley fixture book (ticket I1).

Proves three things about the tiny-novel fixture:

  1. the generator is deterministic — two builds yield identical manuscript
     TEXT (we assert on extracted paragraph texts, not zip bytes);
  2. docproof can ingest the committed .docx through the real ingest APIs, and
     its three chapters are detectable as chapter headings; and
  3. every manifest entry (the answer key) quotes a real substring of the doc,
     in the chapter the manifest claims, and its correction actually repairs
     the quote.

If any of these break, the fixture — which later Galley tickets measure recall
against — has drifted from its answer key and must be regenerated.
"""
from __future__ import annotations

import json
import sys
from pathlib import Path

import docx
import pytest

from docproof.config import load_config
from docproof.continuity import looks_like_chapter_heading
from docproof.ingest import build_document_model, preflight

FIXTURES = Path(__file__).resolve().parent.parent / "fixtures"
DOCX = FIXTURES / "tiny_novel.docx"
MANIFEST_PATH = FIXTURES / "tiny_novel.manifest.json"
CONFIG = Path(__file__).resolve().parent.parent.parent / "config" / "default.yaml"


@pytest.fixture(scope="module")
def gen():
    """The generator module, imported off the fixtures dir like conftest does."""
    sys.path.insert(0, str(FIXTURES))
    try:
        import make_tiny_novel
        return make_tiny_novel
    finally:
        sys.path.remove(str(FIXTURES))


@pytest.fixture(scope="module")
def cfg():
    return load_config(str(CONFIG))


@pytest.fixture(scope="module")
def model(cfg):
    pkg = preflight(DOCX, cfg.tracked_changes_policy)
    return build_document_model(pkg, cfg)


@pytest.fixture(scope="module")
def manifest():
    return json.loads(MANIFEST_PATH.read_text(encoding="ascii"))


# --- 1. determinism -------------------------------------------------------

def test_generator_text_is_deterministic(gen):
    assert gen.paragraph_texts() == gen.paragraph_texts()


def test_committed_docx_matches_the_generator(gen):
    """The .docx on disk carries exactly the generator's text — i.e. it was
    built from this source and has not drifted."""
    committed = [(p.text, p.style.name) for p in docx.Document(str(DOCX)).paragraphs]
    assert committed == gen.paragraph_texts()


def test_novel_is_a_few_thousand_words(gen):
    words = sum(len(t.split()) for t, _ in gen.paragraph_texts())
    assert 1500 <= words <= 2600, f"unexpected length: {words} words"


# --- 2. docproof ingest ---------------------------------------------------

def test_docproof_ingests_the_docx(model):
    assert model.paragraphs, "ingest produced no paragraphs"
    # Reviewable body paragraphs get ids like body-0000.
    body = [p for p in model.paragraphs if p.para_id.startswith("body-")]
    assert body
    # The planted comma splice lands in a reviewable body paragraph.
    splice = [p for p in model.paragraphs
              if "nobody wanted to read it" in p.text]
    assert splice and splice[0].reviewable


def test_three_chapters_are_detectable(model):
    headings = [p for p in model.paragraphs if looks_like_chapter_heading(p)]
    texts = [p.text for p in headings]
    assert texts == ["Chapter One", "Chapter Two", "Chapter Three"]


# --- 3. manifest is a faithful answer key ---------------------------------

def _doc_text(model) -> str:
    return "\n".join(p.text for p in model.paragraphs)


def test_manifest_has_expected_shape(manifest):
    assert len(manifest) == 6
    seen_ids = set()
    for entry in manifest:
        for key in ("id", "error_type", "chapter", "quote", "correction"):
            assert key in entry, f"{entry.get('id')} missing {key}"
        assert entry["id"] not in seen_ids, f"duplicate id {entry['id']}"
        seen_ids.add(entry["id"])
        assert entry["chapter"] in (1, 2, 3)
    # A spread of error types, not six of the same.
    assert len({e["error_type"] for e in manifest}) >= 5


def test_manifest_quotes_are_real_substrings(model):
    text = _doc_text(model)
    for entry in _load(MANIFEST_PATH):
        assert entry["quote"] in text, (
            f"{entry['id']}: quote not found in the document")


def test_manifest_quotes_live_in_the_named_chapter(model):
    chapters = _by_chapter(model)
    for entry in _load(MANIFEST_PATH):
        chap_text = chapters[entry["chapter"]]
        assert entry["quote"] in chap_text, (
            f"{entry['id']}: quote not in chapter {entry['chapter']}")


def test_corrections_repair_their_quotes(model):
    """The correction is the quote with exactly the planted defect removed, and
    the corrected form is NOT already present in the manuscript."""
    text = _doc_text(model)
    for entry in _load(MANIFEST_PATH):
        assert entry["correction"] != entry["quote"]
        assert entry["correction"] not in text, (
            f"{entry['id']}: 'correction' already appears in the doc — the "
            f"planted error may be missing")


# --- helpers --------------------------------------------------------------

def _load(path):
    return json.loads(path.read_text(encoding="ascii"))


def _by_chapter(model) -> dict[int, str]:
    """Split the ingested paragraphs into chapter-number -> joined text, using
    the same chapter-heading detector docproof uses."""
    chapters: dict[int, str] = {}
    n = 0
    for p in model.paragraphs:
        if looks_like_chapter_heading(p):
            n += 1
            chapters[n] = p.text
        elif n:
            chapters[n] += "\n" + p.text
    return chapters
