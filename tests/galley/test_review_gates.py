"""The six findings of the 2026-09-04 external review (REVIEW_2026-09-04.json)
and their acceptance criteria: certification needs evidence (GALLEY-001),
verification is bound to the build and coverage is merged, never restamped
(002), the turn-cap detector ignores the driver's own header (003), the app
path runs the same delivery gate (004), a failed hand-off upload is a durable
pending delivery retried without rerunning the book (005), and a workspace is
identified by the Drive file id and its source hash (006).
"""
from __future__ import annotations

import json
from pathlib import Path

import docx
import pytest

import galley.agent as ga
import galley.driver as gd
from galley.manifest import MANDATORY_CHECKS, certify_run
from galley.state_machine import RunStateMachine
from galley.verify import (VerifyRunResult, accepted_fingerprint,
                           build_fingerprints, write_artifacts)
from docproof.models import Usage


# --- GALLEY-001: no evidence, no certificate ----------------------------------

def test_an_empty_run_cannot_pass_certification(tmp_path):
    run = tmp_path / "run"
    run.mkdir()
    cert = certify_run(run)
    assert not cert.passed
    assert {c.name for c in cert.missing} >= set(MANDATORY_CHECKS)
    assert cert.to_json()["missing"]


def test_missing_verify_artifacts_block_delivery_but_optional_skips_do_not(
        tmp_path):
    run = tmp_path / "run"
    run.mkdir()
    (run / "findings.json").write_text(json.dumps(
        {"findings": [{"a": 1}], "cost": {"total_usd": 1.0},
         "checkpoint": {"x": 1}}), "utf-8")
    cert = certify_run(run)
    by = {c.name: c for c in cert.checks}
    # required and absent: not a pass
    assert by["change verifier"].required and by["change verifier"].status != "pass"
    assert by["finished-text walk"].required
    assert by["residual settlement"].required
    assert by["outcome"].required and by["run state"].required
    # optional skips stay skips and are NOT what blocks delivery
    assert by["two-author attribution"].required is False
    assert by["intent-zone collisions"].required is False
    assert not cert.passed
    assert {c.name for c in cert.missing} == set(MANDATORY_CHECKS)


# --- GALLEY-002: binding and coverage ----------------------------------------------

def _build(tmp_path, paragraphs):
    d = docx.Document()
    for p in paragraphs:
        d.add_paragraph(p)
    run = tmp_path / "run"
    run.mkdir(exist_ok=True)
    d.save(run / "book - Atmosphere Press Proofreader.docx")
    (run / "findings.json").write_text(json.dumps(
        {"findings": [], "cost": {"total_usd": 0.0},
         "generated_at": "2000-01-01T00:00:00+00:00"}), "utf-8")
    return run


def _ok(problems=(), residuals=(), *, ran=True, reason=""):
    return VerifyRunResult(list(problems), list(residuals), ran_changes=ran,
                           ran_walk=ran, reason=reason)


def test_verify_artifacts_are_bound_to_the_build_and_a_changed_build_fails(
        tmp_path):
    run = _build(tmp_path, ["One teh two.", "Three four."])
    write_artifacts(run, _ok(), _ok(), model="m", engine="none",
                    usage_changes=Usage(), usage_walk=Usage(), applied=0,
                    paragraphs=2)
    cv = json.loads((run / "change_verify.json").read_text("utf-8"))
    fp = build_fingerprints(run)
    assert cv["accepted_sha256"] == fp["accepted_sha256"]
    assert cv["build_sha256"] == fp["build_sha256"]
    assert set(cv["paragraph_sha256"]) == set(fp["paragraph_sha256"])
    by = {c.name: c for c in certify_run(run).checks}
    assert by["change verifier"].status == "pass"
    assert by["finished-text walk"].status == "pass"
    # the deliverable changes underneath the evidence
    _build(tmp_path, ["One the two.", "Three four."])
    by = {c.name: c for c in certify_run(run).checks}
    assert by["change verifier"].status == "fail"
    assert "different accepted text" in by["change verifier"].detail
    assert by["finished-text walk"].status == "fail"


def test_unread_paragraphs_and_batches_block_certification(tmp_path):
    from galley import verify
    run = _build(tmp_path, ["One.", "Two."])
    verify.UNREAD[:] = ["body-0001"]
    verify.UNREAD_BATCHES[:] = [{"index": 1, "edits": 3, "para_ids": ["body-0000"]}]
    try:
        write_artifacts(run, _ok(), _ok(), model="m", engine="none",
                        usage_changes=Usage(), usage_walk=Usage(),
                        applied=3, paragraphs=2)
    finally:
        verify.UNREAD.clear()
        verify.UNREAD_BATCHES.clear()
    fw = json.loads((run / "finished_walk.json").read_text("utf-8"))
    cv = json.loads((run / "change_verify.json").read_text("utf-8"))
    assert fw["unread_paragraphs"] == ["body-0001"]
    assert cv["unread_batches"][0]["edits"] == 3
    assert "accepted_sha256" not in fw and "accepted_sha256" not in cv
    by = {c.name: c for c in certify_run(run).checks}
    assert by["finished-text walk"].status == "fail"
    assert "unread" in by["finished-text walk"].detail
    assert by["change verifier"].status == "fail"
    assert "batch" in by["change verifier"].detail


def test_a_partial_reread_merges_coverage_without_discarding_findings(tmp_path):
    from galley import verify
    from galley.verify import ResidualFinding
    run = _build(tmp_path, ["One.", "Two.", "Three."])
    first = ResidualFinding("body-0000", "One", "x", "y", "low")
    write_artifacts(run, _ok(), _ok(residuals=[first]), model="m",
                    engine="none", usage_changes=Usage(), usage_walk=Usage(),
                    applied=0, paragraphs=3)
    # a delta over body-0001 only, which finds something there
    second = ResidualFinding("body-0001", "Two", "p", "q", "low")
    write_artifacts(run, _ok(), _ok(residuals=[second]), model="m",
                    engine="none", usage_changes=Usage(), usage_walk=Usage(),
                    applied=0, paragraphs=3, para_ids=["body-0001"], merge=True)
    fw = json.loads((run / "finished_walk.json").read_text("utf-8"))
    assert {r["para_id"] for r in fw["residuals"]} == {"body-0000", "body-0001"}
    assert fw["ran"] is True and fw["unread_paragraphs"] == []
    # a delta whose read failed cannot upgrade or erase anything
    lost = VerifyRunResult([], [], ran_changes=False, ran_walk=False,
                           reason="every read failed")
    write_artifacts(run, lost, lost, model="m", engine="none",
                    usage_changes=Usage(), usage_walk=Usage(), applied=0,
                    paragraphs=3, para_ids=["body-0002"], merge=True)
    fw = json.loads((run / "finished_walk.json").read_text("utf-8"))
    assert fw["ran"] is True                    # the earlier read still stands
    assert {r["para_id"] for r in fw["residuals"]} == {"body-0000", "body-0001"}


def test_settlement_merges_the_read_outcome_and_never_restamps_ran(tmp_path):
    from galley.settle import Settlement, SettlementRecord, rewrite_verify_artifacts
    run = _build(tmp_path, ["One."])
    (run / "finished_walk.json").write_text(json.dumps({
        "generated_at": "2001-01-01T00:00:00+00:00", "ran": False,
        "reason": "every read failed", "residuals": [],
        "unread_paragraphs": ["body-0000"]}), "utf-8")
    (run / "change_verify.json").write_text(json.dumps({
        "generated_at": "2001-01-01T00:00:00+00:00", "ran": False,
        "reason": "every read failed", "problems": [],
        "unread_batches": [{"index": 1, "edits": 2, "para_ids": ["body-0000"]}]}),
        "utf-8")
    st = Settlement(records=[SettlementRecord("r-1", 1, "drop", None, "", "",
                                              "walker_wrong", "deterministic",
                                              para_id="body-0000")])
    rewrite_verify_artifacts(run, st, unverified=["body-0000"])
    fw = json.loads((run / "finished_walk.json").read_text("utf-8"))
    cv = json.loads((run / "change_verify.json").read_text("utf-8"))
    assert fw["ran"] is False and fw["reason"] == "every read failed"
    assert fw["unread_paragraphs"] == ["body-0000"]
    assert fw["unverified_paragraphs"] == ["body-0000"]
    assert cv["ran"] is False and cv["unread_batches"]
    by = {c.name: c for c in certify_run(run).checks}
    assert by["finished-text walk"].status == "fail"
    assert by["change verifier"].status == "fail"
    assert not certify_run(run).passed


def test_a_deterministic_settle_leaves_changed_paragraphs_unverified(tmp_path):
    """The Georgis corroboration: certify passed a build settle had changed
    without an engine to re-read it. Now the changed paragraphs are recorded
    unverified and the verify checks fail until a re-read covers them."""
    from tests.galley.test_settle import (_build as build_run, _manuscript,
                                          _para_ids, _settle, _walk)
    from galley.verify import residual_id
    src = _manuscript(tmp_path)
    ids, _doc = _para_ids(src)
    run = build_run(tmp_path, src, [
        {"para_id": ids[0], "original_text": "teh", "corrected_text": "the",
         "confidence": "high"}])
    # bind the (hand-planted) artifacts to the pre-settle build
    _walk(run, [{"para_id": ids[1], "quote": "recieve", "problem": "sp",
                 "suggestion": "receive", "severity": "high"}])
    fp = build_fingerprints(run)
    for name in ("finished_walk.json", "change_verify.json"):
        payload = json.loads((run / name).read_text("utf-8"))
        payload.update(accepted_sha256=fp["accepted_sha256"],
                       build_sha256=fp["build_sha256"])
        (run / name).write_text(json.dumps(payload), "utf-8")
    assert _settle(tmp_path, run, src) == 0
    fw = json.loads((run / "finished_walk.json").read_text("utf-8"))
    assert fw["unverified_paragraphs"] == [ids[1]]
    assert fw["ran"] is True
    cert = certify_run(run)
    by = {c.name: c for c in cert.checks}
    assert by["residual settlement"].status == "pass"
    assert by["finished-text walk"].status == "fail"
    assert "changed after their last read" in by["finished-text walk"].detail
    assert not cert.passed


# --- GALLEY-003: the turn-cap detector ---------------------------------------------

def test_the_drivers_own_header_is_never_turn_exhaustion():
    log = ("# galley driver: phase sweeps at 2026-09-04T12:00:00 "
           "(max-turns 120, timeout 2.0h)\nwrote runs/sweeps.txt\n")
    assert not gd.detect_turn_cap(log)
    assert gd.session_limit(None, gd.transcript_tail(log)) is None
    # a session that merely MENTIONS its configured cap is not exhausted
    assert gd.session_limit({"type": "result", "subtype": "success",
                             "num_turns": 3},
                            "I was given max turns 120 and used three.") is None


def test_actual_turn_exhaustion_is_still_detected():
    assert gd.session_limit({"type": "result", "subtype": "error_max_turns",
                             "num_turns": 120}, "") == "max_turns"
    assert gd.detect_turn_cap("Reached max turns (100)")
    assert gd.session_limit(None, "error: maximum turns exceeded") == "max_turns"
    stream = ('{"type":"assistant","message":{"content":[{"type":"text",'
              '"text":"working"}]}}\n'
              '{"type":"result","subtype":"error_max_turns","num_turns":5}\n')
    assert gd.parse_session_result(stream)["subtype"] == "error_max_turns"


def test_spawn_claude_reads_the_structured_result(tmp_path, monkeypatch):
    import subprocess

    def fake_run(argv, cwd, env, stdout, stderr, text, timeout):
        stdout.write('{"type":"assistant","message":{"content":[{"type":"text",'
                     '"text":"Done: wrote runs/sweeps.txt"}]}}\n')
        stdout.write('{"type":"result","subtype":"success","num_turns":2,'
                     '"is_error":false}\n')
        return subprocess.CompletedProcess(argv, 0)
    monkeypatch.setattr(gd.subprocess, "run", fake_run)
    spec = gd.PhaseSpec("sweeps", "p", tmp_path, tmp_path / "sweeps.log",
                        ["claude"], {}, max_turns=120, timeout_s=10)
    out = gd.spawn_claude(spec)
    assert out.ok and out.limit is None and out.subtype == "success"
    assert out.num_turns == 2
    log = (tmp_path / "sweeps.log").read_text("utf-8")
    assert "max-turns 120" in log                     # the header is written…
    assert "Done: wrote runs/sweeps.txt" in log       # …the transcript rendered…
    assert "max-turns" not in out.tail                # …and excluded from the tail


# --- GALLEY-004: the app path runs the same gate ------------------------------------

def test_the_app_gate_never_marks_done_without_evidence_or_a_manuscript(tmp_path):
    from app.jobs import JobRunner
    out = tmp_path / "out"
    out.mkdir()
    (out / "findings.json").write_text(json.dumps(
        {"findings": [], "cost": {"total_usd": 0.0}}), "utf-8")
    state, notes = JobRunner._galley_gate(out, built=False)
    assert state == "needs_human" and any("not built" in n for n in notes)
    assert (out / "certificate.json").is_file()
    state, notes = JobRunner._galley_gate(out, built=True)
    assert state == "needs_human"
    assert any("required evidence missing" in n for n in notes)
    cert = json.loads((out / "certificate.json").read_text("utf-8"))
    assert cert["passed"] is False and cert["missing"]


# --- GALLEY-005: a failed upload is a durable pending delivery ----------------------

from tests.galley.test_agent import (BOOK, ENV_TEXT, FakeApp,  # noqa: E402
                                     FakeResult, _agent, _downloader)


@pytest.fixture()
def env(tmp_path) -> ga.AgentEnv:
    path = tmp_path / "agent.env"
    path.write_text(ENV_TEXT, encoding="utf-8")
    path.chmod(0o600)
    return ga.read_env(path)


def test_a_temporary_upload_failure_leaves_a_durable_pending_delivery(env, tmp_path):
    calls: list[list[str]] = []
    fail = {"on": True}

    def upload(files, _folder):
        calls.append([p.name for p in files])
        if fail["on"]:
            raise RuntimeError("Drive is down")
        return [f"id-{p.name}" for p in files]

    ran = []
    agent = _agent(env, tmp_path, opener=FakeApp([BOOK]),
                   download=_downloader(tmp_path), poll_interval_s=0,
                   run_driver=lambda **kw: (ran.append(kw), (_ for _ in ()).throw(
                       RuntimeError("boom")))[1],
                   upload=upload)
    agent.poll_once()
    ledger = ga.Ledger.load(agent.ledger_path)
    entry = ledger.books["drive-1"]
    assert entry["state"] == ga.PENDING_DELIVERY
    assert entry["delivery_attempts"] == 1 and entry["handoff_files"]
    assert "drive-1" in ledger.pending_deliveries()
    # a later poll retries the delivery — and does not rerun the book
    fail["on"] = False
    report = agent.poll_once()
    assert len(ran) == 1
    entry = ga.Ledger.load(agent.ledger_path).books["drive-1"]
    assert entry["state"] == ga.FAILED and entry["delivery"] == "delivered"
    assert entry["uploaded"]
    assert report.delivered == ["Test - Book 1.docx"]


def test_a_retried_delivery_uploads_only_what_is_missing(env, tmp_path):
    uploaded: list[str] = []
    state = {"calls": 0}

    def upload(files, _folder):
        state["calls"] += 1
        name = files[0].name
        if state["calls"] == 2:               # the second file fails first time
            raise RuntimeError("blip")
        uploaded.append(name)
        return [f"id-{name}"]

    agent = _agent(env, tmp_path, opener=FakeApp([BOOK]),
                   download=_downloader(tmp_path), poll_interval_s=0,
                   run_driver=lambda **kw: (_ for _ in ()).throw(
                       RuntimeError("boom")),
                   upload=upload)
    agent.poll_once()
    entry = ga.Ledger.load(agent.ledger_path).books["drive-1"]
    assert entry["state"] == ga.PENDING_DELIVERY
    first_batch = list(uploaded)
    assert len(first_batch) >= 1
    agent.poll_once()
    entry = ga.Ledger.load(agent.ledger_path).books["drive-1"]
    assert entry["state"] == ga.FAILED
    assert sorted(uploaded) == sorted(set(uploaded))      # no file twice
    assert set(entry["uploaded_names"]) == {Path(p).name
                                            for p in entry["handoff_files"]}


def test_delivery_is_abandoned_after_the_bounded_retries(env, tmp_path):
    agent = _agent(env, tmp_path, opener=FakeApp([BOOK]),
                   download=_downloader(tmp_path), poll_interval_s=60,
                   run_driver=lambda **kw: (_ for _ in ()).throw(
                       RuntimeError("boom")),
                   upload=lambda f, d: (_ for _ in ()).throw(
                       RuntimeError("down")))
    agent.poll_once()
    # each retry waits out its own (doubling) backoff before the next
    clock = 1e12
    for _ in range(ga.MAX_DELIVERY_ATTEMPTS + 1):
        ledger = ga.Ledger.load(agent.ledger_path)
        clock = max(clock, float(ledger.books["drive-1"].get(
            "next_delivery_at") or 0) + 1)
        agent.retry_deliveries(ledger, ga.RunReport(), now=clock)
    entry = ga.Ledger.load(agent.ledger_path).books["drive-1"]
    assert entry["state"] == ga.FAILED and entry["delivery"] == "abandoned"


@pytest.mark.parametrize("files_present", ["none", "some", "all"])
def test_missing_handoff_files_keep_delivery_pending(env, tmp_path, files_present):
    files = [tmp_path / "book.docx", tmp_path / "outcome.json"]
    count = {"none": 0, "some": 1, "all": 2}[files_present]
    for path in files[:count]:
        path.write_text("artifact")
    uploaded = []

    def upload(paths, folder):
        uploaded.extend(p.name for p in paths)
        return [f"id-{p.name}" for p in paths]

    agent = _agent(env, tmp_path, upload=upload)
    ledger = agent.ledger()
    ledger.record("drive-1", ga.PENDING_DELIVERY, name="Test - Book 1.docx",
                  folder_id="folder-A", outcome="done", delivery_attempts=1,
                  handoff_files=[str(p) for p in files], next_delivery_at=0)
    report = ga.RunReport()
    agent.retry_deliveries(ledger, report, now=1)
    entry = ga.Ledger.load(agent.ledger_path).claimed("drive-1")
    assert uploaded == [p.name for p in files[:count]]
    if count < len(files):
        assert entry["state"] == ga.PENDING_DELIVERY
        assert entry["delivery_attempts"] == 2
        assert "missing" in entry["delivery_error"]
        assert report.delivered == []
        for path in files[:count]:
            path.unlink()  # Previously confirmed uploads need no local copy.
        for path in files[count:]:
            path.write_text("restored artifact")
        agent.retry_deliveries(agent.ledger(), report, now=1e12)
        entry = ga.Ledger.load(agent.ledger_path).claimed("drive-1")
        assert entry["state"] == ga.FINISHED
        assert entry["delivery"] == "delivered"
        assert uploaded == [p.name for p in files]
    else:
        assert entry["state"] == ga.FINISHED
        assert entry["delivery"] == "delivered"


@pytest.mark.parametrize("ids", [None, [], [""], ["  "], [None]])
def test_upload_without_confirmation_stays_pending(env, tmp_path, ids):
    path = tmp_path / "book.docx"
    path.write_text("artifact")
    agent = _agent(env, tmp_path, upload=lambda paths, folder: ids)
    ledger = agent.ledger()
    agent.owe_delivery(ga.AwaitingBook.from_json(BOOK), ledger, "test", "folder-A",
                      "done", "", [path], why="")
    entry = ga.Ledger.load(agent.ledger_path).claimed("drive-1")
    assert entry["state"] == ga.PENDING_DELIVERY
    assert entry["uploaded_names"] == {}
    assert "No upload id" in entry["delivery_error"]


def test_empty_handoff_is_not_delivered(env, tmp_path):
    agent = _agent(env, tmp_path)
    assert agent._upload_missing([], "folder-A", {}) is False


# --- GALLEY-006: workspace identity ------------------------------------------------

def test_workspaces_are_named_by_surname_and_drive_id():
    assert ga.slug_for("Test - Book 1.docx", "Test", "drive-1") == "test-drive-1"
    assert ga.slug_for("Test - Book 1.docx", "Test", "1AbCdEfGhIjKlMnOp") == \
        "test-ijklmnop"
    assert ga.slug_for("Test - Book 1.docx", "Other", "drive-2") == \
        "other-drive-2"
    assert ga.slug_for("Test - Book 1.docx", "Test", "drive-1") != \
        ga.slug_for("Test - Book 1.docx", "Test", "drive-2")
    assert ga.slug_for("Test - Book 1.docx") == "test-book-1"   # no id: as before


def _book(tmp_path, text="Once upon a time."):
    d = docx.Document()
    d.add_paragraph(text)
    path = tmp_path / "Ford - Book 1.docx"
    d.save(path)
    return path


def test_seed_records_the_source_hash_and_preserves_unchanged_progress(tmp_path):
    book = _book(tmp_path)
    ws = gd.seed_workspace(book, "ford-abc12345", workspace_root=tmp_path / "ws",
                           source_id="abc12345")
    machine = RunStateMachine.load(ws / "state.json")
    assert machine.source_sha256 and machine.source_id == "abc12345"
    assert machine.revision == 1
    (ws / "runs" / "ladder.log").write_text("progress", "utf-8")
    machine.advance("intake", by="test")
    machine.save(ws / "state.json")
    # the SAME manuscript again: everything preserved
    gd.seed_workspace(book, "ford-abc12345", workspace_root=tmp_path / "ws",
                      source_id="abc12345")
    assert (ws / "runs" / "ladder.log").read_text("utf-8") == "progress"
    assert RunStateMachine.load(ws / "state.json").current == "intake"


def test_a_changed_source_needs_an_explicit_revision_transition(tmp_path):
    book = _book(tmp_path)
    ws = gd.seed_workspace(book, "ford-abc12345", workspace_root=tmp_path / "ws")
    (ws / "runs" / "final").mkdir()
    (ws / "runs" / "final" / "outcome.json").write_text("{}", "utf-8")
    changed = _book(tmp_path, "Once upon a very different time.")
    with pytest.raises(gd.SourceChanged, match="different manuscript"):
        gd.seed_workspace(changed, "ford-abc12345",
                          workspace_root=tmp_path / "ws")
    # nothing was touched by the refusal
    assert (ws / "runs" / "final" / "outcome.json").is_file()
    # the explicit transition: previous results set aside, revision 2 fresh
    gd.seed_workspace(changed, "ford-abc12345", workspace_root=tmp_path / "ws",
                      on_source_change="revise")
    assert (ws / "runs.rev1" / "final" / "outcome.json").is_file()
    assert not (ws / "runs" / "final").exists()
    machine = RunStateMachine.load(ws / "state.json")
    assert machine.revision == 2 and machine.current == ""
    assert (ws / "state.rev1.json").is_file()
