"""galley/outcome.py — the `needs_human` verdict, proven through the real $0
path over real run directories, on a genuinely bad manuscript.

`assess()` has only ever been exercised on one clean book (Redding Book 1)
plus two synthetic unit tests in test_settle.py. Nothing has shown the verdict
actually flips on a bad book that went through `import-findings` ->
`rebuild_from_rows` -> `galley outcome`, the same $0 rebuild path a real
Galley run uses (see docproof/replay.py's module docstring). These tests do
that, on two fixtures:

  * `bad_novel.docx` (tests/fixtures/make_bad_novel.py) — Chapter One of the
    tiny-novel fixture with ten of its twelve body paragraphs carrying
    rewrite-class damage (subject-verb disagreement, a dropped word, a
    run-on, garbled word order). The corrections a proofread would make are
    the manifest's `quote` -> `correction` pairs, each landed as a
    whole-sentence rewrite in its own repair cluster — the same shape the
    real repair channel (Avenue D) produces for a broken sentence.

  * `tiny_novel.docx` (tests/fixtures/make_tiny_novel.py) — the existing
    clean fixture with its six planted, mostly-mechanical errors, run the
    same way, as the mirror: a sound book must still come back `done`.

A third test proves the flip side of the same fixture: a MECHANICAL-ONLY
correction pass over the bad book (punctuation/case only, no wording touched
— the owner has scoped go-live to mechanical proofreading only) must not be
flagged `needs_human` by edit density, because density is computed over
WORDING edits only. The evidence split (`wording_edits` vs `mechanical_edits`)
is asserted directly so the report stays honest about what a mechanical pass
did and did not fix.

No model call is made anywhere in this file: `import-findings` zeroes every
paid pass before it builds (docproof.replay.zero_paid_passes, invoked inside
rebuild_from_rows), and no `galley settle` is run.
"""
from __future__ import annotations

import json
import sys
from pathlib import Path

import docx
import pytest
import yaml

from docproof.__main__ import main
from docproof.config import load_config
from docproof.ingest import build_document_model, preflight

CONFIG = str(Path("config/default.yaml"))
FIXTURES = Path(__file__).resolve().parent.parent / "fixtures"
TINY_DOCX = FIXTURES / "tiny_novel.docx"
TINY_MANIFEST = FIXTURES / "tiny_novel.manifest.json"


@pytest.fixture(scope="module")
def bad_gen():
    """The bad-novel generator module, imported off the fixtures dir the same
    way tests/galley/test_fixture_book.py imports the tiny-novel one."""
    sys.path.insert(0, str(FIXTURES))
    try:
        import make_bad_novel
        return make_bad_novel
    finally:
        sys.path.remove(str(FIXTURES))


# --- shared $0-rebuild plumbing (mirrors tests/galley/test_settle.py) --------

def _replay_config(tmp_path, **extra):
    """A $0 rebuild config: the shipped defaults with no typed passes and
    every paid stage off (what final-replay materializes)."""
    cfg = yaml.safe_load(Path(CONFIG).read_text("utf-8"))
    cfg["error_types"] = []
    for k in ("languagetool", "spellcheck", "consistency", "glossary",
              "residuals", "meaning_check", "fix_check", "factcheck",
              "continuity", "chapter_continuity", "chapter_sweep",
              "smoothing", "rewrite", "repair", "storysheet"):
        cfg.setdefault(k, {})
        if isinstance(cfg[k], dict):
            cfg[k]["enabled"] = False
    cfg["ensemble"] = {"detectors": [], "verify_policy": "none"}
    cfg["candidate_screening"] = {"mode": "off"}
    cfg.update(extra)
    path = tmp_path / "replay.yaml"
    path.write_text(yaml.safe_dump(cfg), encoding="utf-8")
    return str(path)


def _para_ids(src):
    cfg = load_config(CONFIG)
    doc = build_document_model(preflight(str(src), cfg.tracked_changes_policy),
                               cfg)
    return [p.para_id for p in doc.paragraphs], doc


def _build(tmp_path, src, rows, out="run"):
    """A finished, $0 build from import-findings rows — findings.json, the
    docx, editmap.json — via the exact CLI path a real replay uses."""
    rows_path = tmp_path / "rows.json"
    rows_path.write_text(json.dumps(rows), encoding="utf-8")
    run = tmp_path / out
    rc = main(["import-findings", str(rows_path), str(src), "--config",
               _replay_config(tmp_path), "--out", str(run)])
    assert rc == 0
    assert (run / "editmap.json").exists()
    return run


def _outcome(run, capsys, *extra_args):
    """Runs `galley outcome --json` (the real CLI verb) over `run` and
    returns the parsed Outcome dict."""
    assert main(["galley", "outcome", str(run), "--json", *extra_args]) == 0
    out = capsys.readouterr().out
    return json.loads([ln for ln in out.splitlines() if ln.startswith("{")][-1])


# --- the bad book: needs_human ------------------------------------------------

def test_bad_novel_needs_human_through_the_real_run_path(tmp_path, capsys,
                                                          bad_gen):
    src = tmp_path / "bad_novel.docx"
    bad_gen.build().save(src)
    ids, _doc = _para_ids(src)
    manifest = bad_gen.manifest()
    assert len(manifest) == 10

    # The corrections a proofread would make: a whole-sentence rewrite per
    # damaged paragraph, landed as a single-member repair cluster — the shape
    # the real repair channel (Avenue D) produces for a broken sentence, and
    # exactly what makes a paragraph "rewrite-class" to assess()'s own
    # definition (a repair cluster, 3+ edits, or 2+ unsettled residuals).
    rows = [{"para_id": ids[entry["paragraph"]],
             "original_text": entry["quote"],
             "corrected_text": entry["correction"],
             "confidence": "high",
             "cluster_id": f"repair-{entry['id']}"} for entry in manifest]
    run = _build(tmp_path, src, rows)

    oc = _outcome(run, capsys)
    assert oc["outcome"] == "needs_human"
    assert "rewrite-class" in oc["reason"]
    ev = oc["evidence"]
    # measured margins, for the record: 13 paragraphs (1 heading + 12 body),
    # 10 of them rewrite-class -> ~77% rewrite share, well past the 50%
    # threshold; a real proofread's per-sentence-level correction volume
    # would push edit density up further still.
    assert ev["paragraphs"] == 13
    assert ev["rewrite_paragraphs"] == 10
    assert ev["rewrite_share"] == pytest.approx(10 / 13, abs=1e-6)
    assert ev["rewrite_share"] >= 0.50
    assert ev["wording_edits"] == 10
    # outcome.json is on disk with a verdict a human/DocWatch can read, no
    # matter what hubspot_fields("needs_human") maps to today.
    on_disk = json.loads((run / "outcome.json").read_text("utf-8"))
    assert on_disk["outcome"] == "needs_human"
    assert on_disk["reason"]


# --- the mirror: a clean book with only its planted, mostly-mechanical -------
# --- errors stays done --------------------------------------------------------

def test_tiny_novel_stays_done_through_the_real_run_path(tmp_path, capsys):
    src = TINY_DOCX
    ids, doc = _para_ids(src)
    by_text = {p.para_id: p.text for p in doc.paragraphs}
    manifest = json.loads(TINY_MANIFEST.read_text("ascii"))
    assert len(manifest) == 6

    rows = []
    for entry in manifest:
        owners = [pid for pid, text in by_text.items() if entry["quote"] in text]
        assert owners, f"{entry['id']}: quote not found in any paragraph"
        rows.append({"para_id": owners[0], "original_text": entry["quote"],
                     "corrected_text": entry["correction"],
                     "confidence": "high"})
    run = _build(tmp_path, src, rows)

    oc = _outcome(run, capsys)
    assert oc["outcome"] == "done"
    ev = oc["evidence"]
    # measured margins: six planted edits over ~2,000 words is nowhere near
    # the 60/1,000 density floor, and none of them share a cluster or stack
    # 3+ to a paragraph, so rewrite share stays near zero too.
    assert ev["wording_edits"] + ev["mechanical_edits"] == 6
    assert ev["edit_density_per_kword"] < 60.0
    assert ev["rewrite_share"] < 0.50
    on_disk = json.loads((run / "outcome.json").read_text("utf-8"))
    assert on_disk["outcome"] == "done"


# --- mechanical-only pass over the SAME bad book: honest, not needs_human ---
# --- by density ---------------------------------------------------------------

def test_bad_novel_mechanical_only_pass_is_not_flagged_by_density(tmp_path,
                                                                   capsys,
                                                                   bad_gen):
    """The owner has scoped go-live to mechanical proofreading only: a run
    that touches punctuation/case/spacing and none of the wording damage must
    not have its density number inflated by those mechanics (they are
    excluded from edit_density_per_kword by design — see Thresholds'
    docstring), and the evidence it writes must say plainly that no wording
    was touched, not silently claim the book is clean."""
    src = tmp_path / "bad_novel.docx"
    bad_gen.build().save(src)
    ids, doc = _para_ids(src)
    texts = {p.para_id: p.text for p in doc.paragraphs}

    # Five case-only fixes, scattered across damaged AND undamaged
    # paragraphs — punctuation/case, never a word — while the real
    # subject-verb/dropped-word/run-on/garbled-order damage from `bad_gen`
    # goes untouched.
    case_fixes = [(1, "valley"), (2, "rain"), (3, "years"), (5, "world"),
                  (6, "crow")]
    rows = []
    for para_no, word in case_fixes:
        pid = ids[para_no]
        assert texts[pid].count(word) == 1, (
            f"fixture drift: {word!r} not a unique anchor in paragraph "
            f"{para_no}")
        rows.append({"para_id": pid, "original_text": word,
                     "corrected_text": word.capitalize(),
                     "confidence": "high"})
    run = _build(tmp_path, src, rows)

    oc = _outcome(run, capsys)
    ev = oc["evidence"]
    # the honest split: five mechanical edits landed, zero counted as wording
    assert ev["mechanical_edits"] == 5
    assert ev["wording_edits"] == 0
    assert ev["edit_density_per_kword"] == 0.0
    # no cluster, no 3-edit paragraph, no settlement residuals -> rewrite
    # share reads zero too, even though the underlying book is the same bad
    # one as the first test above.
    assert ev["rewrite_paragraphs"] == 0
    assert ev["rewrite_share"] == 0.0
    # so nothing here trips needs_human — a mechanical pass correctly reports
    # what it did, not what it left undone.
    assert oc["outcome"] == "done"
    assert "per 1,000" not in oc["reason"]
